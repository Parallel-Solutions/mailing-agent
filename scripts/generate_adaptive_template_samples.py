from __future__ import annotations

import argparse
import html
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from pypdf import PdfReader, PdfWriter
from pypdf.generic import (
    ArrayObject,
    BooleanObject,
    DictionaryObject,
    FloatObject,
    NameObject,
    NumberObject,
    TextStringObject,
)

from src.generator.generation.pdf_converter import convert_html_to_pdf
from src.generator.templates.adapters.docx import DocxTemplateAdapter
from src.generator.templates.adapters.pdf import PdfAcroFormAdapter


A4_WIDTH_PT = 595.28
A4_HEIGHT_PT = 841.89
MM_TO_PT = 72.0 / 25.4

FIELDS = (
    "OUTGOING_NUMBER",
    "DATE",
    "ADM_NAME",
    "HEAD_FIO_SHORT",
    "HEAD_GREETING",
    "WORK_TITLE",
    "MUN_NAME_1",
    "SUB_RF_1",
)

PREVIEW_CONTEXT = {
    "OUTGOING_NUMBER": "247-КП",
    "DATE": "13.07.2026",
    "ADM_NAME": "Администрация муниципального образования городской округ Северный Берег",
    "HEAD_FIO_SHORT": "А. В. Александрову",
    "HEAD_GREETING": "Уважаемый Алексей Викторович!",
    "WORK_TITLE": "разработка документов территориального планирования и градостроительного зонирования",
    "MUN_NAME_1": "городского округа Северный Берег",
    "SUB_RF_1": "Ленинградской области",
}


@dataclass(frozen=True)
class Brand:
    slug: str
    company: str
    short_name: str
    slogan: str
    primary: str
    secondary: str
    accent: str
    style: str
    director: str
    phone: str
    email: str
    intro: str
    benefits: tuple[str, str, str]
    price: str
    duration: str


BRANDS = (
    Brand(
        slug="severnaya_geodeziya",
        company="ООО «Северная Геодезия»",
        short_name="СГ",
        slogan="Пространственные данные для точных решений",
        primary="#123B63",
        secondary="#EAF2F8",
        accent="#2F80C0",
        style="technical",
        director="И. П. Воронцов",
        phone="+7 812 555-18-40",
        email="project@north-geo.test",
        intro="Команда геодезистов, кадастровых инженеров и градостроителей выполняет проекты полного цикла.",
        benefits=("ГИС-анализ территории", "Юридическая проверка", "Сопровождение согласования"),
        price="180 000 ₽",
        duration="45 рабочих дней",
    ),
    Brand(
        slug="gorodskie_sistemy",
        company="АО «Городские Системы»",
        short_name="ГС",
        slogan="Проектируем устойчивое развитие территорий",
        primary="#681F2B",
        secondary="#F7F0E8",
        accent="#B98A52",
        style="formal",
        director="Е. М. Романова",
        phone="+7 495 555-02-71",
        email="office@city-systems.test",
        intro="Экспертное бюро объединяет градостроительное, экономическое и правовое проектирование.",
        benefits=("Экспертная методология", "Прозрачный календарный план", "Персональный руководитель"),
        price="225 000 ₽",
        duration="60 рабочих дней",
    ),
    Brand(
        slug="terraproekt",
        company="ООО «ТерраПроект»",
        short_name="ТП",
        slogan="Понятные документы. Живые территории.",
        primary="#1F5B45",
        secondary="#EDF5EF",
        accent="#76A56F",
        style="minimal",
        director="М. С. Беляев",
        phone="+7 343 555-44-09",
        email="hello@terra-project.test",
        intro="Разрабатываем практичные градостроительные документы и передаём заказчику готовую цифровую модель.",
        benefits=("Открытые рабочие встречи", "Контроль качества данных", "Передача исходных материалов"),
        price="165 000 ₽",
        duration="40 рабочих дней",
    ),
)


def _hex_rgb(value: str) -> RGBColor:
    clean = value.lstrip("#")
    return RGBColor(int(clean[0:2], 16), int(clean[2:4], 16), int(clean[4:6], 16))


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill.lstrip("#"))


def _set_cell_border(cell, **edges: dict[str, Any]) -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge_name, options in edges.items():
        tag = f"w:{edge_name}"
        edge = borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)
        for key, value in options.items():
            edge.set(qn(f"w:{key}"), str(value))


def _set_cell_margins(cell, *, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _set_font(run, *, name: str = "Arial", size: float = 9.5, bold: bool = False, color: str = "#202830") -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = _hex_rgb(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _style_paragraph(paragraph, *, before: float = 0, after: float = 3, line: float = 1.05) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.widow_control = True


def _add_text(paragraph, text: str, *, size: float = 9.5, bold: bool = False, color: str = "#202830"):
    run = paragraph.add_run(text)
    _set_font(run, size=size, bold=bold, color=color)
    return run


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    section.left_margin = Mm(13)
    section.right_margin = Mm(13)
    section.header_distance = Mm(4)
    section.footer_distance = Mm(4)
    section.start_type = WD_SECTION.NEW_PAGE
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def _add_header(document: Document, brand: Brand) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(3.0)
    table.columns[1].width = Cm(15.2)
    logo_cell, info_cell = table.rows[0].cells
    logo_cell.width = Cm(3.0)
    info_cell.width = Cm(15.2)
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    info_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_shading(logo_cell, brand.primary)
    _set_cell_shading(info_cell, brand.secondary)
    for cell in (logo_cell, info_cell):
        _set_cell_margins(cell, top=110, bottom=110, start=130, end=130)
        _set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})

    paragraph = logo_cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_paragraph(paragraph, after=0)
    _add_text(paragraph, brand.short_name, size=20, bold=True, color="#FFFFFF")

    title = info_cell.paragraphs[0]
    _style_paragraph(title, after=1)
    _add_text(title, brand.company, size=13.5, bold=True, color=brand.primary)
    slogan = info_cell.add_paragraph()
    _style_paragraph(slogan, after=0)
    _add_text(slogan, brand.slogan, size=8.5, color=brand.primary)


def _add_meta_and_recipient(document: Document, brand: Brand) -> None:
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8.5)
    table.columns[1].width = Cm(9.7)
    left, right = table.rows[0].cells
    for cell in (left, right):
        _set_cell_margins(cell, top=70, bottom=70, start=0, end=0)
        _set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})

    paragraph = left.paragraphs[0]
    _style_paragraph(paragraph, after=1)
    _add_text(paragraph, "Исх. № ", size=8.7, color="#5E6872")
    _add_text(paragraph, "{{OUTGOING_NUMBER}}", size=8.7, bold=True, color=brand.primary)
    _add_text(paragraph, " от ", size=8.7, color="#5E6872")
    _add_text(paragraph, "{{DATE}}", size=8.7, bold=True, color=brand.primary)

    recipient = right.paragraphs[0]
    recipient.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _style_paragraph(recipient, after=1)
    _add_text(recipient, "{{ADM_NAME}}", size=9.2, bold=True, color=brand.primary)
    person = right.add_paragraph()
    person.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _style_paragraph(person, after=0)
    _add_text(person, "{{HEAD_FIO_SHORT}}", size=9.2, color="#30383F")


def _add_body(document: Document, brand: Brand) -> None:
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if brand.style == "formal" else WD_ALIGN_PARAGRAPH.LEFT
    _style_paragraph(heading, before=3, after=4)
    _add_text(heading, "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ", size=14, bold=True, color=brand.primary)

    greeting = document.add_paragraph()
    _style_paragraph(greeting, after=4)
    _add_text(greeting, "{{HEAD_GREETING}}", size=10.2, bold=True, color=brand.primary)

    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _style_paragraph(intro, after=4, line=1.08)
    _add_text(intro, brand.intro + " Предлагаем выполнить ", size=9.5)
    _add_text(intro, "{{WORK_TITLE}}", size=9.5, bold=True, color=brand.primary)
    _add_text(intro, " для ", size=9.5)
    _add_text(intro, "{{MUN_NAME_1}}", size=9.5, bold=True, color=brand.primary)
    _add_text(intro, " ", size=9.5)
    _add_text(intro, "{{SUB_RF_1}}", size=9.5, bold=True, color=brand.primary)
    _add_text(intro, ".", size=9.5)

    benefits = document.add_table(rows=1, cols=3)
    benefits.alignment = WD_TABLE_ALIGNMENT.CENTER
    benefits.autofit = False
    for index, (cell, benefit) in enumerate(zip(benefits.rows[0].cells, brand.benefits, strict=True), start=1):
        cell.width = Cm(6.05)
        _set_cell_shading(cell, brand.secondary if index % 2 else "#FFFFFF")
        _set_cell_margins(cell, top=110, bottom=110, start=100, end=100)
        _set_cell_border(
            cell,
            top={"val": "single", "sz": "5", "color": brand.accent.lstrip("#")},
            bottom={"val": "single", "sz": "5", "color": brand.accent.lstrip("#")},
            start={"val": "single", "sz": "5", "color": brand.accent.lstrip("#")},
            end={"val": "single", "sz": "5", "color": brand.accent.lstrip("#")},
        )
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_paragraph(paragraph, after=0)
        _add_text(paragraph, f"{index:02d}\n", size=8, bold=True, color=brand.accent)
        _add_text(paragraph, benefit, size=8.5, bold=True, color=brand.primary)

    scope_title = document.add_paragraph()
    _style_paragraph(scope_title, before=4, after=2)
    _add_text(scope_title, "Состав результата", size=10.2, bold=True, color=brand.primary)
    scope = (
        "• аналитическая записка и исходная модель данных;\n"
        "• комплект проектных материалов в редактируемых форматах;\n"
        "• сопровождение замечаний до получения согласованного результата."
    )
    scope_paragraph = document.add_paragraph()
    _style_paragraph(scope_paragraph, after=4, line=1.05)
    _add_text(scope_paragraph, scope, size=9.1)

    offer = document.add_table(rows=3, cols=2)
    offer.alignment = WD_TABLE_ALIGNMENT.CENTER
    offer.autofit = False
    rows = (("Предмет", "Комплекс работ по техническому заданию"), ("Срок", brand.duration), ("Стоимость", brand.price))
    for row_index, (label, value) in enumerate(rows):
        label_cell, value_cell = offer.rows[row_index].cells
        label_cell.width = Cm(4.4)
        value_cell.width = Cm(13.8)
        _set_cell_shading(label_cell, brand.primary)
        _set_cell_shading(value_cell, brand.secondary if row_index != 2 else "#FFFFFF")
        for cell in (label_cell, value_cell):
            _set_cell_margins(cell, top=75, bottom=75, start=120, end=120)
            _set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": brand.primary.lstrip("#")},
                bottom={"val": "single", "sz": "4", "color": brand.primary.lstrip("#")},
                start={"val": "single", "sz": "4", "color": brand.primary.lstrip("#")},
                end={"val": "single", "sz": "4", "color": brand.primary.lstrip("#")},
            )
        left = label_cell.paragraphs[0]
        _style_paragraph(left, after=0)
        _add_text(left, label, size=8.8, bold=True, color="#FFFFFF")
        right = value_cell.paragraphs[0]
        _style_paragraph(right, after=0)
        _add_text(right, value, size=9.2, bold=row_index == 2, color=brand.primary)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _style_paragraph(note, before=4, after=4, line=1.05)
    _add_text(
        note,
        "Предложение включает организацию работ, контроль качества и передачу заказчику полного цифрового архива.",
        size=8.8,
        color="#4D5861",
    )


def _add_signature(document: Document, brand: Brand) -> None:
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(9.2)
    table.columns[1].width = Cm(9.0)
    left, right = table.rows[0].cells
    for cell in (left, right):
        _set_cell_margins(cell, top=50, bottom=40, start=0, end=0)
        _set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})

    paragraph = left.paragraphs[0]
    _style_paragraph(paragraph, after=1)
    _add_text(paragraph, "С уважением,\nгенеральный директор", size=9.2, bold=True, color=brand.primary)
    contacts = left.add_paragraph()
    _style_paragraph(contacts, after=0)
    _add_text(contacts, f"{brand.phone}  •  {brand.email}", size=8.2, color="#5A646C")

    signature = right.paragraphs[0]
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _style_paragraph(signature, before=8, after=0)
    _add_text(signature, brand.director, size=10, bold=True, color=brand.primary)


def build_docx_template(brand: Brand, output_path: Path) -> None:
    document = Document()
    _configure_document(document)
    _add_header(document, brand)
    _add_meta_and_recipient(document, brand)
    _add_body(document, brand)
    _add_signature(document, brand)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _field_box(name: str, label: str, x: float, y: float, width: float, height: float) -> str:
    return (
        f'<div class="field-label" style="left:{x}mm;top:{y - 4}mm">{html.escape(label)}</div>'
        f'<div class="field-box" data-field="{name}" style="left:{x}mm;top:{y}mm;'
        f'width:{width}mm;height:{height}mm"></div>'
    )


def build_pdf_background_html(brand: Brand) -> str:
    benefits = "".join(
        f'<div class="benefit"><b>{index:02d}</b><span>{html.escape(item)}</span></div>'
        for index, item in enumerate(brand.benefits, start=1)
    )
    fields = "".join(
        (
            _field_box("OUTGOING_NUMBER", "Исх. №", 27, 42, 33, 7),
            _field_box("DATE", "Дата", 66, 42, 28, 7),
            _field_box("ADM_NAME", "Получатель", 108, 50, 87, 14),
            _field_box("HEAD_FIO_SHORT", "Руководитель", 108, 69, 87, 8),
            _field_box("HEAD_GREETING", "Обращение", 15, 87, 180, 9),
            _field_box("WORK_TITLE", "Предмет предложения", 15, 106, 180, 15),
            _field_box("MUN_NAME_1", "Муниципальное образование", 15, 131, 112, 10),
            _field_box("SUB_RF_1", "Субъект РФ", 132, 131, 63, 10),
        )
    )
    return f"""<!doctype html>
<html lang="ru"><head><meta charset="utf-8"><style>
@page {{ size: A4; margin: 0; }}
* {{ box-sizing: border-box; }}
html, body {{ margin: 0; width: 210mm; height: 297mm; font-family: Arial, sans-serif; color: #27313a; }}
.page {{ position: relative; width: 210mm; height: 297mm; overflow: hidden; background: white; }}
.top {{ position: absolute; left: 0; top: 0; width: 210mm; height: 34mm; background: {brand.secondary}; border-top: 6mm solid {brand.primary}; }}
.logo {{ position: absolute; left: 15mm; top: 11mm; width: 23mm; height: 17mm; display: flex; align-items: center; justify-content: center; background: {brand.primary}; color: white; font-size: 18pt; font-weight: 800; letter-spacing: .5pt; }}
.company {{ position: absolute; left: 44mm; top: 12mm; color: {brand.primary}; font-size: 17pt; font-weight: 800; }}
.slogan {{ position: absolute; left: 44mm; top: 22mm; color: {brand.primary}; font-size: 8.5pt; letter-spacing: .2pt; }}
.title {{ position: absolute; left: 15mm; top: 52mm; width: 86mm; color: {brand.primary}; font-size: 18pt; font-weight: 800; line-height: 1.05; }}
.field-label {{ position: absolute; color: #69747d; font-size: 6.5pt; text-transform: uppercase; letter-spacing: .35pt; }}
.field-box {{ position: absolute; border-bottom: .45mm solid {brand.accent}; background: rgba(255,255,255,.68); }}
.benefits {{ position: absolute; left: 15mm; top: 151mm; width: 180mm; display: flex; gap: 4mm; }}
.benefit {{ width: 57.3mm; min-height: 25mm; padding: 4mm; border: .35mm solid {brand.accent}; background: {brand.secondary}; }}
.benefit b {{ display: block; color: {brand.accent}; font-size: 8pt; margin-bottom: 2mm; }}
.benefit span {{ color: {brand.primary}; font-size: 9pt; font-weight: 700; line-height: 1.2; }}
.scope {{ position: absolute; left: 15mm; top: 184mm; width: 180mm; font-size: 9pt; line-height: 1.45; }}
.scope h3 {{ margin: 0 0 2mm; color: {brand.primary}; font-size: 12pt; }}
.offer {{ position: absolute; left: 15mm; top: 220mm; width: 180mm; border-collapse: collapse; font-size: 9pt; }}
.offer td {{ border: .3mm solid {brand.primary}; padding: 2.4mm 3mm; }}
.offer td:first-child {{ width: 42mm; color: white; background: {brand.primary}; font-weight: 700; }}
.offer td:last-child {{ color: {brand.primary}; background: {brand.secondary}; }}
.footer {{ position: absolute; left: 15mm; bottom: 12mm; width: 180mm; padding-top: 3mm; border-top: .5mm solid {brand.primary}; color: {brand.primary}; font-size: 8.5pt; }}
.footer strong {{ font-size: 10pt; }}
.director {{ float: right; font-weight: 700; font-size: 10pt; }}
.glyphs {{ position: absolute; left: 1mm; bottom: 1mm; color: #fff; font-size: .1pt; font-weight: 400; }}
</style></head><body><div class="page">
<div class="top"></div><div class="logo">{html.escape(brand.short_name)}</div>
<div class="company">{html.escape(brand.company)}</div><div class="slogan">{html.escape(brand.slogan)}</div>
<div class="title">КОММЕРЧЕСКОЕ<br>ПРЕДЛОЖЕНИЕ</div>{fields}
<div class="benefits">{benefits}</div>
<div class="scope"><h3>Что входит в результат</h3>Аналитика исходных данных • проектные материалы в редактируемых форматах • цифровая модель • сопровождение замечаний до согласования.</div>
<table class="offer"><tr><td>Срок</td><td>{html.escape(brand.duration)}</td></tr><tr><td>Стоимость</td><td><b>{html.escape(brand.price)}</b></td></tr></table>
<div class="footer"><strong>С уважением, генеральный директор</strong><span class="director">{html.escape(brand.director)}</span><br>{html.escape(brand.phone)} &nbsp;•&nbsp; {html.escape(brand.email)}</div>
<div class="glyphs">АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ абвгдеёжзийклмнопрстуфхцчшщъыьэюя 0123456789 №.,:;!?()—–-«»₽ расширенное проверочное значение максимальной ожидаемой длины</div>
</div></body></html>"""


PDF_FIELDS = {
    "OUTGOING_NUMBER": (27, 42, 33, 7, False),
    "DATE": (66, 42, 28, 7, False),
    "ADM_NAME": (108, 50, 87, 14, True),
    "HEAD_FIO_SHORT": (108, 69, 87, 8, False),
    "HEAD_GREETING": (15, 87, 180, 9, False),
    "WORK_TITLE": (15, 106, 180, 15, True),
    "MUN_NAME_1": (15, 131, 112, 10, False),
    "SUB_RF_1": (132, 131, 63, 10, False),
}


def _pdf_rect(x_mm: float, y_mm: float, width_mm: float, height_mm: float) -> ArrayObject:
    left = x_mm * MM_TO_PT
    right = (x_mm + width_mm) * MM_TO_PT
    top = A4_HEIGHT_PT - y_mm * MM_TO_PT
    bottom = A4_HEIGHT_PT - (y_mm + height_mm) * MM_TO_PT
    return ArrayObject([FloatObject(left), FloatObject(bottom), FloatObject(right), FloatObject(top)])


def add_acroform_fields(background_pdf: Path, output_pdf: Path) -> None:
    reader = PdfReader(str(background_pdf))
    writer = PdfWriter()
    writer.clone_document_from_reader(reader)
    page = writer.pages[0]

    resources = page[NameObject("/Resources")].get_object()
    font_resources = resources.get(NameObject("/Font"))
    fonts = font_resources.get_object() if font_resources is not None else DictionaryObject()
    if not fonts:
        raise RuntimeError("Converted PDF does not expose a font resource for AcroForm appearances")
    def font_score(item) -> tuple[int, int, int]:
        _, reference = item
        font = reference.get_object()
        base_name = str(font.get("/BaseFont") or "")
        return (
            1 if font.get("/Subtype") == "/Type0" else 0,
            1 if font.get("/ToUnicode") is not None else 0,
            1 if "Bold" not in base_name else 0,
        )

    font_key, font_ref = max(fonts.items(), key=font_score)
    fields = ArrayObject()
    acro_form = DictionaryObject(
        {
            NameObject("/Fields"): fields,
            NameObject("/NeedAppearances"): BooleanObject(False),
            NameObject("/DA"): TextStringObject(f"{font_key} 0 Tf 0 g"),
            NameObject("/DR"): DictionaryObject(
                {NameObject("/Font"): DictionaryObject({NameObject(str(font_key)): font_ref})}
            ),
        }
    )
    writer._root_object[NameObject("/AcroForm")] = writer._add_object(acro_form)
    annotations = page.get(NameObject("/Annots"), ArrayObject())
    annotations = annotations.get_object() if hasattr(annotations, "get_object") else annotations
    page[NameObject("/Annots")] = annotations

    for field_name, (x, y, width, height, multiline) in PDF_FIELDS.items():
        widget = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Annot"),
                NameObject("/Subtype"): NameObject("/Widget"),
                NameObject("/FT"): NameObject("/Tx"),
                NameObject("/T"): TextStringObject(field_name),
                NameObject("/V"): TextStringObject(""),
                NameObject("/Rect"): _pdf_rect(x, y, width, height),
                NameObject("/DA"): TextStringObject(f"{font_key} 0 Tf 0 g"),
                NameObject("/F"): NumberObject(4),
                NameObject("/Ff"): NumberObject(4096 if multiline else 0),
                NameObject("/Q"): NumberObject(0),
                NameObject("/P"): page.indirect_reference,
                NameObject("/BS"): DictionaryObject(
                    {NameObject("/W"): NumberObject(0), NameObject("/S"): NameObject("/S")}
                ),
            }
        )
        widget_ref = writer._add_object(widget)
        fields.append(widget_ref)
        annotations.append(widget_ref)

    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    with output_pdf.open("wb") as handle:
        writer.write(handle)


def render_preview_png(pdf_path: Path, png_path: Path) -> None:
    try:
        import fitz
    except ImportError:
        return
    with fitz.open(pdf_path) as document:
        pixmap = document[0].get_pixmap(matrix=fitz.Matrix(1.35, 1.35), alpha=False)
        pixmap.save(png_path)


def build_pdf_template(brand: Brand, template_path: Path, preview_path: Path, preview_png: Path) -> None:
    work_dir = template_path.parent / ".build"
    work_dir.mkdir(parents=True, exist_ok=True)
    background_path = work_dir / "background.pdf"
    converted = convert_html_to_pdf(
        build_pdf_background_html(brand),
        background_path,
        filename=f"{brand.slug}.html",
    )
    if converted is None or not converted.exists():
        raise RuntimeError(f"Could not convert PDF background for {brand.company}")
    add_acroform_fields(converted, template_path)
    PdfAcroFormAdapter().render(template_path, PREVIEW_CONTEXT, preview_path)
    render_preview_png(preview_path, preview_png)
    background_path.unlink(missing_ok=True)
    work_dir.rmdir()


def generate(output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    manifest: dict[str, Any] = {
        "purpose": "Test-only adaptive commercial proposal templates",
        "fields": list(FIELDS),
        "companies": [],
    }
    for brand in BRANDS:
        brand_dir = output_dir / brand.slug
        brand_dir.mkdir(parents=True, exist_ok=True)
        docx_template = brand_dir / "kp_template.docx"
        docx_preview = brand_dir / "kp_preview.docx"
        pdf_template = brand_dir / "kp_template.pdf"
        pdf_preview = brand_dir / "kp_preview.pdf"
        preview_png = brand_dir / "kp_preview.png"
        build_docx_template(brand, docx_template)
        DocxTemplateAdapter().render(docx_template, PREVIEW_CONTEXT, docx_preview)
        build_pdf_template(brand, pdf_template, pdf_preview, preview_png)
        manifest["companies"].append(
            {
                "slug": brand.slug,
                "company": brand.company,
                "style": brand.style,
                "files": [
                    "kp_template.docx",
                    "kp_preview.docx",
                    "kp_template.pdf",
                    "kp_preview.pdf",
                    "kp_preview.png",
                ],
            }
        )

    (output_dir / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    (output_dir / "README.md").write_text(
        "# Тестовые адаптивные шаблоны КП\n\n"
        "Все компании и реквизиты вымышлены. Файлы предназначены только для проверки движка.\n\n"
        "- `kp_template.docx` — DOCX с маркерами `{{FIELD_NAME}}`.\n"
        "- `kp_preview.docx` — заполненный пример DOCX.\n"
        "- `kp_template.pdf` — PDF с именованными AcroForm-полями.\n"
        "- `kp_preview.pdf` и `kp_preview.png` — заполненное превью.\n\n"
        "Поддерживаемые поля перечислены в `manifest.json`.\n",
        encoding="utf-8",
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    generate(args.output.resolve())


if __name__ == "__main__":
    main()

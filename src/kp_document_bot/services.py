from __future__ import annotations

import json
import re
import shutil
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document

try:
    import httpx  # type: ignore
except ImportError:  # pragma: no cover
    httpx = None

try:
    from src.generator.ai_case_agent import (
        OpenAI,
        _resolve_openai_api_key,
        _resolve_openai_base_url,
        apply_case_agent_result,
        run_case_validation_agent,
    )
    from src.generator.config_generator import (
        BASE_DIR,
        CASE_AGENT_MODEL,
        DATA_XLSX_PATH,
        PDF_CHUNK_SIZE,
        START_OUTGOING_NUMBER,
    )
    from src.generator.document_review_agent import review_docx
    from src.generator.document_builder import (
        CONTRACT_TEMPLATE_PATH,
        KP_TEMPLATE_PATH,
        build_contract_filename,
        build_contract_replacements,
        build_kp_filename,
        build_kp_replacements,
        ensure_output_folder,
        render_docx,
    )
    from src.generator.excel_io import load_rows
    from src.generator.pdf_converter import convert_docx_batch
    from src.generator.transforms import build_document_context
except ImportError:  # pragma: no cover
    from generator.ai_case_agent import (
        OpenAI,
        _resolve_openai_api_key,
        _resolve_openai_base_url,
        apply_case_agent_result,
        run_case_validation_agent,
    )
    from generator.config_generator import (
        BASE_DIR,
        CASE_AGENT_MODEL,
        DATA_XLSX_PATH,
        PDF_CHUNK_SIZE,
        START_OUTGOING_NUMBER,
    )
    from generator.document_review_agent import review_docx
    from generator.document_builder import (
        CONTRACT_TEMPLATE_PATH,
        KP_TEMPLATE_PATH,
        build_contract_filename,
        build_contract_replacements,
        build_kp_filename,
        build_kp_replacements,
        ensure_output_folder,
        render_docx,
    )
    from generator.excel_io import load_rows
    from generator.pdf_converter import convert_docx_batch
    from generator.transforms import build_document_context


TMP_ROOT = BASE_DIR / "data" / "tmp"
TMP_ROOT.mkdir(parents=True, exist_ok=True)


def _json_default(value: Any) -> Any:
    if isinstance(value, Path):
        return str(value)
    raise TypeError(f"Object of type {type(value).__name__} is not JSON serializable")


def _extract_docx_text(docx_path: Path) -> str:
    doc = Document(docx_path)
    chunks: List[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        chunks.append(text)

    return "\n".join(chunks)


def _build_llm_client():
    if not OpenAI:
        return None
    api_key = _resolve_openai_api_key()
    if not api_key:
        return None
    base_url = _resolve_openai_base_url()
    kwargs = {"api_key": api_key}
    if base_url:
        kwargs["base_url"] = base_url
    if httpx:
        kwargs["http_client"] = httpx.Client(
            http2=False,
            timeout=httpx.Timeout(connect=10, read=60, write=60, pool=60),
            trust_env=False,
        )
    return OpenAI(**kwargs)


def _call_json_llm(prompt: str, model: str) -> dict:
    client = _build_llm_client()
    if not client:
        return {"status": "unavailable", "issues": [], "comment": "LLM client unavailable"}

    request_kwargs: Dict[str, Any] = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
    }
    if not _resolve_openai_base_url():
        request_kwargs["response_format"] = {"type": "json_object"}

    response = client.chat.completions.create(**request_kwargs)
    content = response.choices[0].message.content or "{}"
    parsed = json.loads(_extract_json_payload(content))
    if isinstance(parsed, list):
        return {"status": "ok", "issues": parsed}
    if isinstance(parsed, dict):
        return parsed
    return {"status": "unavailable", "issues": [], "comment": "Unsupported LLM response"}


def _extract_json_payload(content: str) -> str:
    text = (content or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?", "", text, flags=re.IGNORECASE).strip()
        text = re.sub(r"```$", "", text).strip()
    start_object = text.find("{")
    end_object = text.rfind("}")
    if start_object >= 0 and end_object > start_object:
        return text[start_object : end_object + 1]
    return text or "{}"


def review_text_content(text: str, model: Optional[str] = None) -> dict:
    prompt = (
        "Проверь текст официального документа на корректность русского языка, падежей, "
        "согласования, канцелярского стиля и странных сокращений. "
        "Верни JSON-объект вида "
        '{"status":"ok|needs_fix","issues":[{"severity":"low|medium|high","fragment":"...","comment":"...","suggestion":"..."}],'
        '"summary":"...","corrected_excerpt":"..."}.\n\n'
        f"Текст для проверки:\n{text}"
    )
    selected_model = model or "openai/gpt-4o-mini"
    return _call_json_llm(prompt, selected_model)


def review_generated_documents(docx_paths: List[Path], model: Optional[str] = None) -> dict:
    documents = []
    for path in docx_paths:
        if not path.exists():
            continue
        documents.append(
            {
                "name": path.name,
                "path": str(path),
                "review": review_text_content(_extract_docx_text(path), model=model),
            }
        )
    return {"documents": documents}


def review_uploaded_document(file_name: str, content: bytes, model: Optional[str] = None) -> dict:
    suffix = Path(file_name or "document").suffix.lower()
    temp_root = Path(tempfile.mkdtemp(prefix="kpbot_review_", dir=str(TMP_ROOT)))
    try:
        temp_path = temp_root / (Path(file_name).name or f"document{suffix or '.bin'}")
        temp_path.write_bytes(content)
        if suffix == ".docx":
            review = review_docx(temp_path)
            return {
                "status": "ok",
                "file_name": file_name,
                "file_type": "docx",
                "review": review,
            }
        if suffix in {".txt", ".md", ".csv", ".json"}:
            text = temp_path.read_text(encoding="utf-8", errors="ignore")
            return {
                "status": "ok",
                "file_name": file_name,
                "file_type": suffix.lstrip("."),
                "review": review_text_content(text, model=model),
            }
        return {
            "status": "error",
            "message": f"Неподдерживаемый формат файла: {suffix or 'unknown'}",
            "supported_types": [".docx", ".txt", ".md", ".csv", ".json"],
        }
    finally:
        shutil.rmtree(temp_root, ignore_errors=True)


def _load_rows_from_excel(xlsx_path: Optional[Path] = None) -> list[dict]:
    source_path = Path(xlsx_path) if xlsx_path else DATA_XLSX_PATH
    _, _, rows = load_rows(source_path)
    return rows


def _select_rows_by_range(rows: list[dict], start: int, end: int) -> list[dict]:
    start_index = max(1, int(start))
    end_index = max(start_index, int(end))
    return rows[start_index - 1 : end_index]


def generate_documents_batch(
    start_row: int,
    end_row: int,
    review_final_text: bool = False,
    review_model: Optional[str] = None,
    xlsx_path: Optional[Path] = None,
) -> dict:
    source_path = Path(xlsx_path) if xlsx_path else DATA_XLSX_PATH
    rows = _load_rows_from_excel(source_path)
    selected_rows = _select_rows_by_range(rows, start_row, end_row)
    results = []
    for offset, row in enumerate(selected_rows, start=start_row):
        payload = generate_document_package(
            row=row,
            outgoing_number=START_OUTGOING_NUMBER + offset - 1,
            generate_pdf=True,
            review_final_text=review_final_text,
            review_model=review_model,
        )
        results.append(payload)

    return {
        "status": "ok",
        "start_row": start_row,
        "end_row": end_row,
        "count": len(results),
        "source_xlsx": str(source_path),
        "documents": results,
    }


def iter_generate_documents_batch(
    start_row: int,
    end_row: int,
    review_final_text: bool = False,
    review_model: Optional[str] = None,
    xlsx_path: Optional[Path] = None,
):
    source_path = Path(xlsx_path) if xlsx_path else DATA_XLSX_PATH
    rows = _load_rows_from_excel(source_path)
    selected_rows = _select_rows_by_range(rows, start_row, end_row)
    total = len(selected_rows)
    for index, row in enumerate(selected_rows, start=1):
        offset = start_row + index - 1
        payload = generate_document_package(
            row=row,
            outgoing_number=START_OUTGOING_NUMBER + offset - 1,
            generate_pdf=True,
            review_final_text=review_final_text,
            review_model=review_model,
        )
        yield {
            "index": index,
            "total": total,
            "row_index": offset,
            "payload": payload,
            "source_xlsx": str(source_path),
        }


def _parse_natural_action(text: str) -> Optional[dict]:
    normalized = text.strip().lower()
    numbers = [int(match) for match in re.findall(r"\d+", normalized)]
    has_generate_intent = False
    for token in ("сгенер", "собери", "подготов", "сформируй"):
        if normalized.find(token) >= 0:
            has_generate_intent = True
            break

    if has_generate_intent:
        if len(numbers) >= 2:
            return {"action": "generate_batch", "start": max(1, numbers[0]), "end": max(numbers[0], numbers[1])}
        if len(numbers) == 1:
            return {"action": "generate_batch", "start": 1, "end": max(1, numbers[0])}
        return {"action": "ask_generate_range"}

    has_review_batch_intent = False
    for token in ("проверь эти документ", "проверь документы", "проверь док", "проверь их", "проверь последние"):
        if normalized.find(token) >= 0:
            has_review_batch_intent = True
            break
    if has_review_batch_intent:
        if len(numbers) >= 2:
            return {"action": "review_batch", "start": max(1, numbers[0]), "end": max(numbers[0], numbers[1])}
        if len(numbers) == 1:
            return {"action": "review_batch", "start": 1, "end": max(1, numbers[0])}
        return {"action": "review_last"}

    for token in ("проверь текст", "проверь ошибки", "проверь граммат"):
        if normalized.find(token) >= 0:
            return {"action": "review_text"}

    return None


def _extract_number_range(text: str) -> Optional[tuple[int, int]]:
    numbers = [int(match) for match in re.findall(r"\d+", text or "")]
    if not numbers:
        return None
    if len(numbers) == 1:
        count = max(1, numbers[0])
        return (1, count)
    start = max(1, numbers[0])
    end = max(start, numbers[1])
    return (start, end)


def _is_affirmative(text: str) -> bool:
    return text.strip().lower() in {"да", "ага", "угу", "ок", "окей", "хорошо", "давай", "подтверждаю"}


def _default_help_reply() -> str:
    return (
        "Могу помочь с документами: сгенерировать их по таблице, проверить уже готовые файлы "
        "или посмотреть загруженный документ. Можно сначала прислать Excel-таблицу `.xlsx`, "
        "и тогда я буду работать именно по ней. "
        "Можно писать свободно, например: `сгенерируй 5 документов`, "
        "`сгенерируй документы для строк 10 20`, `проверь эти документы`, `проверь этот текст`."
    )


def _looks_like_greeting(text: str) -> bool:
    normalized = (text or "").strip().lower()
    return normalized in {
        "привет",
        "здравствуйте",
        "добрый день",
        "доброе утро",
        "добрый вечер",
        "ты тут",
        "ты на месте",
    }


def _classify_chat_intent(
    message: str,
    session: dict,
    has_row: bool,
    has_text: bool,
    has_uploaded_document: bool,
    model: Optional[str] = None,
) -> dict:
    prompt = (
        "Ты маршрутизатор намерений для чат-агента по документам. "
        "Определи только ближайшее действие пользователя и верни JSON без пояснений. "
        "Доступные intent: "
        "generate_batch, review_generated, review_uploaded_document, review_text, generate_documents, send_archive, answer_question, ask_clarification, help. "
        "Если пользователь просит сгенерировать несколько документов, используй generate_batch. "
        "Если просит проверить уже созданные документы, используй review_generated. "
        "Если пользователь прислал документ и просит проверить его, используй review_uploaded_document. "
        "Если пользователь задает вопрос по документу или процессу, используй answer_question. "
        "Если информации не хватает, используй ask_clarification. "
        "Верни JSON вида "
        '{"intent":"...","arguments":{"start_row":1,"end_row":5},"reply":"..."}.\n\n'
        f"message={json.dumps(message, ensure_ascii=False)}\n"
        f"session={json.dumps(session, ensure_ascii=False)}\n"
        f"has_row={json.dumps(has_row)}\n"
        f"has_text={json.dumps(has_text)}\n"
        f"has_uploaded_document={json.dumps(has_uploaded_document)}\n"
    )
    parsed = _call_json_llm(prompt, model or CASE_AGENT_MODEL)
    if not isinstance(parsed, dict):
        return {}
    return parsed


def handle_agent_message(
    message: str,
    session: Optional[dict] = None,
    row: Optional[dict] = None,
    text: Optional[str] = None,
    uploaded_document: Optional[dict] = None,
    outgoing_number: Optional[int] = None,
    review_final_text: bool = True,
    review_model: Optional[str] = None,
) -> dict:
    session_state = dict(session or {})
    spreadsheet_path_raw = session_state.get("spreadsheet_path") or session_state.get("last_uploaded_spreadsheet_path")
    spreadsheet_path = Path(spreadsheet_path_raw) if spreadsheet_path_raw else None
    normalized = (message or "").strip()
    if not normalized:
        return {
            "action": "help",
            "reply": _default_help_reply(),
            "payload": {},
            "session": session_state,
        }

    if _looks_like_greeting(normalized):
        return {
            "action": "greeting",
            "reply": (
                "Привет! Я на месте. Могу сгенерировать документы по таблице, "
                "проверить уже готовые документы или посмотреть файл, который ты отправишь. "
                "Если пришлёшь Excel-таблицу, буду работать по ней."
            ),
            "payload": {},
            "session": session_state,
        }

    last_range = session_state.get("last_generated_range") or {}
    explicit_last_docs_request = any(
        token in normalized.lower()
        for token in (
            "проверь эти документы",
            "проверь эти док",
            "проверь их",
            "проверь последние документы",
            "проверь последние док",
        )
    )
    if explicit_last_docs_request and last_range:
        start_row = int(last_range.get("start", 1))
        end_row = int(last_range.get("end", start_row))
        payload = review_generated_batch(
            start_row,
            end_row,
            xlsx_path=spreadsheet_path,
            use_cached_only=True,
        )
        session_state["pending_action"] = None
        return {
            "action": "review_batch",
            "reply": f"Проверила последние сгенерированные документы для строк {start_row}–{end_row}.",
            "payload": payload,
            "session": session_state,
        }

    pending = session_state.get("pending_action")
    if pending == "await_generate_range":
        range_values = _extract_number_range(normalized)
        if range_values:
            start_row, end_row = range_values
            payload = generate_documents_batch(
                start_row,
                end_row,
                review_final_text=False,
                review_model=review_model,
                xlsx_path=spreadsheet_path,
            )
            session_state["pending_action"] = None
            session_state["last_generated_range"] = {"start": start_row, "end": end_row}
            return {
                "action": "generate_batch",
                "reply": f"Сгенерировала документы для строк {start_row}–{end_row}.",
                "payload": payload,
                "session": session_state,
            }
        if _is_affirmative(normalized):
            payload = generate_documents_batch(
                1,
                1,
                review_final_text=False,
                review_model=review_model,
                xlsx_path=spreadsheet_path,
            )
            session_state["pending_action"] = None
            session_state["last_generated_range"] = {"start": 1, "end": 1}
            return {
                "action": "generate_batch",
                "reply": "Сгенерировала первый комплект.",
                "payload": payload,
                "session": session_state,
            }
        return {
            "action": "ask_generate_range",
            "reply": "Напиши количество или диапазон. Например: `5` или `10 20`.",
            "payload": {},
            "session": session_state,
        }

    if pending == "await_review_range":
        range_values = _extract_number_range(normalized)
        if range_values:
            start_row, end_row = range_values
            payload = review_generated_batch(
                start_row,
                end_row,
                xlsx_path=spreadsheet_path,
                use_cached_only=True,
            )
            session_state["pending_action"] = None
            return {
                "action": "review_batch",
                "reply": f"Проверила документы для строк {start_row}–{end_row}.",
                "payload": payload,
                "session": session_state,
            }
        last_range = session_state.get("last_generated_range")
        if _is_affirmative(normalized) and last_range:
            payload = review_generated_batch(
                int(last_range["start"]),
                int(last_range["end"]),
                xlsx_path=spreadsheet_path,
                use_cached_only=True,
            )
            session_state["pending_action"] = None
            return {
                "action": "review_batch",
                "reply": "Проверила последние сгенерированные документы.",
                "payload": payload,
                "session": session_state,
            }
        return {
            "action": "ask_review_range",
            "reply": "Напиши количество или диапазон для проверки. Например: `5` или `10 20`.",
            "payload": {},
            "session": session_state,
        }

    llm_route = _classify_chat_intent(
        message=normalized,
        session=session_state,
        has_row=row is not None,
        has_text=bool(text),
        has_uploaded_document=uploaded_document is not None,
        model=review_model,
    )
    intent = str(llm_route.get("intent", "")).strip()
    arguments = llm_route.get("arguments") or {}
    if intent == "generate_batch":
        start_row = int(arguments.get("start_row", 1))
        end_row = int(arguments.get("end_row", start_row))
        payload = generate_documents_batch(
            start_row,
            end_row,
            review_final_text=False,
            review_model=review_model,
            xlsx_path=spreadsheet_path,
        )
        session_state["last_generated_range"] = {"start": start_row, "end": end_row}
        session_state["pending_action"] = None
        return {
            "action": "generate_batch",
            "reply": str(llm_route.get("reply") or f"Сгенерировала документы для строк {start_row}–{end_row}."),
            "payload": payload,
            "session": session_state,
        }
    if intent == "review_generated":
        last_range = session_state.get("last_generated_range") or {}
        start_row = int(arguments.get("start_row", last_range.get("start", 1)))
        end_row = int(arguments.get("end_row", last_range.get("end", start_row)))
        payload = review_generated_batch(
            start_row,
            end_row,
            xlsx_path=spreadsheet_path,
            use_cached_only=True,
        )
        session_state["pending_action"] = None
        return {
            "action": "review_batch",
            "reply": str(llm_route.get("reply") or f"Проверила документы для строк {start_row}–{end_row}."),
            "payload": payload,
            "session": session_state,
        }
    if intent == "review_uploaded_document" and uploaded_document:
        payload = review_uploaded_document(
            file_name=str(uploaded_document.get("file_name", "document.docx")),
            content=uploaded_document.get("content", b""),
            model=review_model,
        )
        return {
            "action": "review_uploaded_document",
            "reply": str(llm_route.get("reply") or "Проверила загруженный документ."),
            "payload": payload,
            "session": session_state,
        }
    if intent == "review_text" and text:
        payload = review_text_content(text, model=review_model)
        return {
            "action": "review_text",
            "reply": str(llm_route.get("reply") or "Проверила текст."),
            "payload": payload,
            "session": session_state,
        }
    if intent == "generate_documents" and row:
        payload = generate_document_package(
            row=row,
            outgoing_number=outgoing_number,
            generate_pdf=True,
            review_final_text=review_final_text,
            review_model=review_model,
        )
        return {
            "action": "generate_documents",
            "reply": str(llm_route.get("reply") or "Документы собраны, проверка выполнена."),
            "payload": payload,
            "session": session_state,
        }
    if intent == "send_archive":
        return {
            "action": "send_archive",
            "reply": str(llm_route.get("reply") or "Могу отправить последний архив, если он уже есть в сессии."),
            "payload": {},
            "session": session_state,
        }
    if intent == "ask_clarification":
        return {
            "action": "ask_clarification",
            "reply": str(llm_route.get("reply") or "Уточни, что именно нужно сделать: сгенерировать документы, проверить их или проверить загруженный файл."),
            "payload": {},
            "session": session_state,
        }
    if intent == "answer_question":
        return {
            "action": "answer_question",
            "reply": str(llm_route.get("reply") or _default_help_reply()),
            "payload": {},
            "session": session_state,
        }

    parsed = _parse_natural_action(normalized)
    if parsed:
        action = parsed["action"]
        if action == "generate_batch":
            payload = generate_documents_batch(
                parsed["start"],
                parsed["end"],
                review_final_text=False,
                review_model=review_model,
                xlsx_path=spreadsheet_path,
            )
            session_state["last_generated_range"] = {"start": parsed["start"], "end": parsed["end"]}
            session_state["pending_action"] = None
            return {
                "action": "generate_batch",
                "reply": f"Сгенерировала документы для строк {parsed['start']}–{parsed['end']}.",
                "payload": payload,
                "session": session_state,
            }
        if action == "ask_generate_range":
            session_state["pending_action"] = "await_generate_range"
            return {
                "action": "ask_generate_range",
                "reply": "Сколько документов нужно сгенерировать? Можно написать `5` или диапазон `10 20`.",
                "payload": {},
                "session": session_state,
            }
        if action == "review_batch":
            payload = review_generated_batch(
                parsed["start"],
                parsed["end"],
                xlsx_path=spreadsheet_path,
                use_cached_only=True,
            )
            session_state["pending_action"] = None
            return {
                "action": "review_batch",
                "reply": f"Проверила документы для строк {parsed['start']}–{parsed['end']}.",
                "payload": payload,
                "session": session_state,
            }
        if action == "review_last":
            last_range = session_state.get("last_generated_range")
            if last_range:
                payload = review_generated_batch(
                    int(last_range["start"]),
                    int(last_range["end"]),
                    xlsx_path=spreadsheet_path,
                    use_cached_only=True,
                )
                return {
                    "action": "review_batch",
                    "reply": "Проверила последние сгенерированные документы.",
                    "payload": payload,
                    "session": session_state,
                }
            session_state["pending_action"] = "await_review_range"
            return {
                "action": "ask_review_range",
                "reply": "У меня пока нет последних сгенерированных документов. Напиши диапазон, например `1 5`.",
                "payload": {},
                "session": session_state,
            }
        if action == "review_text":
            if text:
                payload = review_text_content(text, model=review_model)
                return {
                    "action": "review_text",
                    "reply": "Проверила текст.",
                    "payload": payload,
                    "session": session_state,
                }

    has_row_generate_intent = False
    for token in ("сгенер", "собери", "подготов", "сделай"):
        if normalized.lower().find(token) >= 0:
            has_row_generate_intent = True
            break
    if row and has_row_generate_intent:
        payload = generate_document_package(
            row=row,
            outgoing_number=outgoing_number,
            generate_pdf=True,
            review_final_text=review_final_text,
            review_model=review_model,
        )
        return {
            "action": "generate_documents",
            "reply": "Документы собраны, проверка выполнена.",
            "payload": payload,
            "session": session_state,
        }

    has_text_review_intent = False
    for token in ("проверь", "ошиб", "текст", "граммат"):
        if normalized.lower().find(token) >= 0:
            has_text_review_intent = True
            break
    if text and has_text_review_intent:
        payload = review_text_content(text, model=review_model)
        return {
            "action": "review_text",
            "reply": "Проверка текста выполнена.",
            "payload": payload,
            "session": session_state,
        }

    return {
        "action": "help",
        "reply": _default_help_reply(),
        "payload": {
            "supported_actions": [
                "generate_documents",
                "generate_batch",
                "review_generated_documents",
                "review_text",
                "review_uploaded_document",
            ]
        },
        "session": session_state,
    }


def review_generated_batch(
    start_row: int,
    end_row: int,
    xlsx_path: Optional[Path] = None,
    use_cached_only: bool = False,
) -> dict:
    rows = _load_rows_from_excel(xlsx_path)
    selected_rows = _select_rows_by_range(rows, start_row, end_row)
    documents = []
    for row in selected_rows:
        output_folder = ensure_output_folder(row)
        row_documents = []
        cached_review_path = output_folder / "document_review.json"
        if cached_review_path.exists():
            try:
                cached = json.loads(cached_review_path.read_text(encoding="utf-8"))
                row_documents = list(cached.get("documents", []))
            except (OSError, json.JSONDecodeError):
                row_documents = []
        if not row_documents and not use_cached_only:
            for docx_path in sorted(output_folder.glob("*.docx")):
                row_documents.append(
                    {
                        "name": docx_path.name,
                        "path": str(docx_path),
                        "review": review_docx(docx_path),
                    }
                )
        documents.append(
            {
                "row_id": row.get("ID"),
                "mun_name": row.get("MUN_NAME"),
                "output_folder": str(output_folder),
                "used_cached_review": bool(row_documents),
                "documents": row_documents,
            }
        )
    return {
        "status": "ok",
        "start_row": start_row,
        "end_row": end_row,
        "count": len(documents),
        "rows": documents,
    }


def generate_document_package(
    row: dict,
    outgoing_number: Optional[int] = None,
    generate_pdf: bool = True,
    review_final_text: bool = False,
    review_model: Optional[str] = None,
    perform_document_review: bool = True,
    document_review_ai: bool = False,
) -> dict:
    if not isinstance(row, dict):
        raise ValueError("row must be a dict")

    outgoing_number = int(outgoing_number or START_OUTGOING_NUMBER)
    context = build_document_context(row, outgoing_number=outgoing_number)
    agent_result = run_case_validation_agent(row, context)
    context = apply_case_agent_result(context, agent_result)

    output_folder = ensure_output_folder(row)
    temp_root = Path(tempfile.mkdtemp(prefix="kpbot_", dir=str(TMP_ROOT)))
    temp_docx_dir = temp_root / "docx"
    temp_pdf_dir = temp_root / "pdf"
    temp_profiles_dir = temp_root / "profiles"
    temp_docx_dir.mkdir(parents=True, exist_ok=True)
    temp_pdf_dir.mkdir(parents=True, exist_ok=True)

    rendered_docx: List[Path] = []
    final_files: Dict[str, Optional[str]] = {
        "kp_docx": None,
        "kp_pdf": None,
        "contract_docx": None,
        "contract_pdf": None,
    }

    try:
        if KP_TEMPLATE_PATH.exists():
            kp_temp_path = temp_docx_dir / build_kp_filename(row)
            render_docx(KP_TEMPLATE_PATH, build_kp_replacements(context), kp_temp_path, context)
            kp_final_docx = output_folder / build_kp_filename(row)
            shutil.copy2(str(kp_temp_path), str(kp_final_docx))
            rendered_docx.append(kp_temp_path)
            final_files["kp_docx"] = str(kp_final_docx)

        if CONTRACT_TEMPLATE_PATH.exists():
            contract_temp_path = temp_docx_dir / build_contract_filename(row)
            render_docx(
                CONTRACT_TEMPLATE_PATH,
                build_contract_replacements(context),
                contract_temp_path,
                context,
            )
            contract_final_docx = output_folder / build_contract_filename(row)
            shutil.copy2(str(contract_temp_path), str(contract_final_docx))
            rendered_docx.append(contract_temp_path)
            final_files["contract_docx"] = str(contract_final_docx)

        document_reviews = []
        if perform_document_review:
            for temp_docx_path in rendered_docx:
                final_docx_path = output_folder / temp_docx_path.name
                document_reviews.append(
                    {
                        "name": temp_docx_path.name,
                        "path": str(final_docx_path),
                        "review": review_docx(temp_docx_path, ai_enabled=document_review_ai),
                    }
                )

        if generate_pdf and rendered_docx:
            pdf_map = convert_docx_batch(
                rendered_docx,
                temp_pdf_dir,
                chunk_size=PDF_CHUNK_SIZE,
                worker_count=1,
                profiles_root=temp_profiles_dir,
            )
            for temp_docx_path in rendered_docx:
                pdf_path = pdf_map.get(temp_docx_path)
                if not pdf_path or not pdf_path.exists():
                    continue
                if temp_docx_path.name.startswith("КП_"):
                    final_pdf = output_folder / temp_docx_path.name.replace(".docx", ".pdf")
                    shutil.move(str(pdf_path), str(final_pdf))
                    final_files["kp_pdf"] = str(final_pdf)
                elif temp_docx_path.name.startswith("Договор_"):
                    final_pdf = output_folder / temp_docx_path.name.replace(".docx", ".pdf")
                    shutil.move(str(pdf_path), str(final_pdf))
                    final_files["contract_pdf"] = str(final_pdf)

        process_result = {
            "case_agent_status": context.get("CASE_AGENT_STATUS"),
            "case_agent_items": context.get("CASE_AGENT_ITEMS", []),
            "case_agent_summary": context.get("CASE_AGENT_SUMMARY", {}),
            "case_agent_mode": agent_result.get("mode"),
            "case_agent_enabled": agent_result.get("enabled", False),
            "case_agent_error": agent_result.get("error"),
        }
        text_review = None
        if review_final_text:
            final_docx_paths = [Path(path) for key, path in final_files.items() if key.endswith("_docx") and path]
            text_review = review_generated_documents(final_docx_paths, model=review_model)

        return {
            "status": "ok",
            "row_id": row.get("ID"),
            "mun_name": row.get("MUN_NAME"),
            "output_folder": str(output_folder),
            "files": final_files,
            "case_agent": {
                "status": process_result["case_agent_status"],
                "summary": process_result["case_agent_summary"],
                "items": process_result["case_agent_items"],
                "error": process_result["case_agent_error"],
            },
            "document_review": {
                "documents": document_reviews,
            },
            "text_review": text_review,
        }
    finally:
        shutil.rmtree(temp_root, ignore_errors=True)


def handle_chat_message(
    message: str,
    row: Optional[dict] = None,
    text: Optional[str] = None,
    outgoing_number: Optional[int] = None,
    review_final_text: bool = True,
) -> dict:
    normalized = (message or "").strip().lower()

    if row and any(token in normalized for token in ("сгенер", "подготов", "собери", "сделай")):
        payload = generate_document_package(
            row=row,
            outgoing_number=outgoing_number,
            generate_pdf=True,
            review_final_text=review_final_text,
        )
        return {
            "action": "generate_documents",
            "reply": "Документы собраны, AI-проверка и дополнительная проверка текста выполнены.",
            "payload": payload,
        }

    if text and any(token in normalized for token in ("проверь", "ошиб", "текст", "граммат")):
        payload = review_text_content(text)
        return {
            "action": "review_text",
            "reply": "Проверка текста выполнена.",
            "payload": payload,
        }

    return {
        "action": "help",
        "reply": (
            "Я могу собрать документы по строке данных, проверить спорные формулировки, "
            "а также сделать дополнительную проверку готового текста."
        ),
        "payload": {
            "supported_actions": [
                "generate_documents",
                "review_text",
            ]
        },
    }

from __future__ import annotations

import base64
import html
import mimetypes
import posixpath
import re
from dataclasses import dataclass
from pathlib import Path
from time import perf_counter
from typing import Any
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree

from src.generator.generation.structured_kp import build_structured_kp_model, format_money_decimal
from src.generator.generation.work_types import WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES

WORD_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "asvg": "http://schemas.microsoft.com/office/drawing/2016/SVG/main",
}
REL_NS = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
EMU_PER_MM = 36000


@dataclass(frozen=True)
class KPTemplateAssets:
    logo: str = ""
    stamp: str = ""
    phone_icon: str = ""
    mail_icon: str = ""
    background_right: str = ""
    background_left: str = ""


@dataclass(frozen=True)
class KPSignatureText:
    director_title: str = "С уважением,\nисполнительный директор"
    signer: str = "К.И. Крашенинников"
    executor: str = "Исп. Черкашина Наталья"
    phone: str = "тел. +7 993 079-45-61"
    email: str = "ks@parresh.ru"


def render_kp_pdf(
    source_docx: Path,
    output_pdf: Path,
    *,
    template_docx: Path | None = None,
    context: dict[str, Any] | None = None,
    timing_callback: Any | None = None,
) -> Path:
    """Render a KP DOCX into PDF through a controlled HTML/CSS renderer.

    This intentionally avoids DOCX-to-PDF layout engines for KP files. The DOCX
    remains the editable source, while the PDF uses stable coordinates for the
    signature assets that were unreliable in LibreOffice/Gotenberg conversion.
    """

    started = perf_counter()
    source_docx = Path(source_docx)
    output_pdf = Path(output_pdf)
    output_pdf.parent.mkdir(parents=True, exist_ok=True)

    template_source = template_docx if template_docx and Path(template_docx).exists() else source_docx
    assets = extract_kp_template_assets(Path(template_source))
    blocks = extract_kp_body_blocks(source_docx)
    if not blocks and context:
        blocks = fallback_blocks_from_context(context)
    signature = extract_kp_signature(source_docx)
    html_text = build_kp_html(blocks, signature, assets)
    render_html_pdf(html_text, output_pdf)

    if timing_callback is not None:
        timing_callback(
            {
                "backend": "custom_kp_html",
                "total": 1,
                "success_count": 1,
                "failed_count": 0,
                "seconds": round(perf_counter() - started, 3),
                "source": str(source_docx),
            }
        )
    return output_pdf


def extract_kp_template_assets(template_docx: Path) -> KPTemplateAssets:
    if not template_docx.exists():
        return KPTemplateAssets()

    logo = stamp = phone_icon = mail_icon = background_right = background_left = ""
    try:
        with ZipFile(template_docx) as archive:
            doc_rels = _rels_for_part(archive, "word/document.xml")
            document_root = etree.fromstring(archive.read("word/document.xml"))

            for header_part in sorted(name for name in archive.namelist() if re.fullmatch(r"word/header\d+\.xml", name)):
                header_rels = _rels_for_part(archive, header_part)
                header_root = etree.fromstring(archive.read(header_part))
                image_node = _first_image_node(header_root)
                if image_node is not None:
                    logo = _data_uri_for_image_node(archive, header_part, header_rels, image_node)
                    if logo:
                        break

            contact_icons: list[tuple[int, str]] = []
            for anchor in document_root.xpath("//wp:anchor|//wp:inline", namespaces=WORD_NS):
                uri = _data_uri_for_image_node(archive, "word/document.xml", doc_rels, anchor)
                if not uri:
                    continue
                extent = anchor.xpath("./wp:extent", namespaces=WORD_NS)
                cx = _safe_int(extent[0].get("cx") if extent else "0")
                cy = _safe_int(extent[0].get("cy") if extent else "0")
                position_h = _safe_int(_first_text(anchor.xpath("./wp:positionH/wp:posOffset/text()", namespaces=WORD_NS)))
                position_v = _safe_int(_first_text(anchor.xpath("./wp:positionV/wp:posOffset/text()", namespaces=WORD_NS)))
                behind = anchor.get("behindDoc") == "1"
                paragraph_text = _paragraph_text_for_node(anchor)

                if behind:
                    if position_h < 0:
                        background_left = background_left or uri
                    elif cx >= 4_000_000:
                        background_right = background_right or uri
                    elif not background_left:
                        background_left = uri
                    continue

                if cx >= 1_000_000 and cy >= 1_000_000:
                    stamp = stamp or uri
                    continue

                if "тел" in paragraph_text.casefold() or "parresh" in paragraph_text.casefold():
                    contact_icons.append((position_v, uri))

            contact_icons.sort(key=lambda item: item[0])
            if contact_icons:
                phone_icon = contact_icons[0][1]
            if len(contact_icons) > 1:
                mail_icon = contact_icons[1][1]
    except Exception:
        return KPTemplateAssets()

    return KPTemplateAssets(
        logo=logo,
        stamp=stamp,
        phone_icon=phone_icon,
        mail_icon=mail_icon,
        background_right=background_right,
        background_left=background_left,
    )


def extract_kp_body_blocks(source_docx: Path) -> list[dict[str, Any]]:
    try:
        document = Document(source_docx)
    except Exception:
        return []

    blocks: list[dict[str, Any]] = []
    for child in document.element.body:
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            paragraph = Paragraph(child, document)
            block = _paragraph_block(paragraph)
            if block:
                blocks.append(block)
        elif tag == "tbl":
            table = Table(child, document)
            table_text = _table_text(table).casefold()
            if "с уважением" in table_text or "крашенинников" in table_text:
                break
            block = _table_block(table)
            if block:
                blocks.append(block)
    return _trim_body_blocks(blocks)


def fallback_blocks_from_context(context: dict[str, Any]) -> list[dict[str, Any]]:
    model = build_structured_kp_model(context)
    return [
        {"type": "meta", "left": f"№ {model.outgoing_number}-КП от {model.date}", "right": model.recipient},
        {"type": "paragraph", "class": "title", "html": html.escape(model.title)},
        {"type": "paragraph", "class": "body", "html": html.escape(model.intro)},
        {"type": "paragraph", "class": "body", "html": html.escape(model.included_services)},
        {
            "type": "table",
            "class": "price-table",
            "rows": [
                ["Вид работ", "Стоимость,\nруб."],
                [model.work_table_title, format_money_decimal(model.price.amount_rubles)],
                ["ИТОГО:", format_money_decimal(model.price.amount_rubles)],
            ],
        },
        {"type": "paragraph", "class": "body", "html": html.escape(model.price_note)},
        {"type": "paragraph", "class": "body", "html": html.escape(_company_paragraph())},
        {"type": "paragraph", "class": "body", "html": html.escape(_closing_paragraph(context))},
        {"type": "paragraph", "class": "body", "html": html.escape(f"Срок действия коммерческого предложения: до {model.validity_date}.")},
    ]


def extract_kp_signature(source_docx: Path) -> KPSignatureText:
    try:
        document = Document(source_docx)
    except Exception:
        return KPSignatureText()

    for table in reversed(document.tables):
        table_text = _table_text(table)
        if "С уважением" not in table_text and "Крашенинников" not in table_text:
            continue
        lines = [line.strip() for line in re.split(r"[\n\r]+", table_text) if line.strip()]
        phone_match = re.search(r"тел\.\s*\+?\s*[\d\s()\-\u00a0]+", table_text, flags=re.I)
        email_match = re.search(r"[\w.\-+]+\s*@\s*[\w.\-]+", table_text)
        executor = next((line for line in lines if line.startswith("Исп.")), KPSignatureText.executor)
        signer = next((line for line in lines if "Крашенинников" in line), KPSignatureText.signer)
        title_lines = []
        for line in lines:
            if line.startswith("С уважением") or "исполнительный директор" in line:
                title_lines.append(line)
        director_title = "\n".join(title_lines[:2]) or KPSignatureText.director_title
        return KPSignatureText(
            director_title=director_title,
            signer=signer,
            executor=executor,
            phone=_normalize_contact(phone_match.group(0)) if phone_match else KPSignatureText.phone,
            email=_normalize_contact(email_match.group(0)) if email_match else KPSignatureText.email,
        )
    return KPSignatureText()


def build_kp_html(blocks: list[dict[str, Any]], signature: KPSignatureText, assets: KPTemplateAssets) -> str:
    body_html = "\n".join(_render_block(block) for block in blocks)
    signature_html = _render_signature(signature, assets)
    logo_html = f'<img class="logo" src="{assets.logo}" alt="">' if assets.logo else '<div class="logo-text">ПАРАЛЛЕЛЬНЫЕ РЕШЕНИЯ</div>'
    return f"""<!doctype html>
<html lang="ru">
<head>
<meta charset="utf-8">
<style>
{_css()}
</style>
</head>
<body>
  <main class="page">
    <header class="header">
      <div class="header-logo">{logo_html}</div>
      <div class="company">
        <strong>Общество с ограниченной ответственностью<br>«Параллельные Решения» (ООО «ПР»)</strong><br><br>
        195220, г. Санкт-Петербург,<br>
        Кушелевская дор, д. 12 литера А, помещ. 2-н<br>
        ИНН 5038110107, КПП 780401001,<br>
        ОГРН 1145038110458,<br>
        т. +7 (812) 242-93-12, parresh.ru
      </div>
    </header>
    <section class="content">{body_html}</section>
    {signature_html}
  </main>
</body>
</html>"""


def render_html_pdf(html_text: str, output_pdf: Path) -> None:
    from playwright.sync_api import Error as PlaywrightError
    from playwright.sync_api import sync_playwright

    last_error: Exception | None = None
    with sync_playwright() as playwright:
        launch_attempts: list[dict[str, Any]] = [
            {},
            {"channel": "chrome"},
            {"channel": "msedge"},
        ]
        for executable in _browser_executable_candidates():
            launch_attempts.append({"executable_path": str(executable)})

        for options in launch_attempts:
            browser = None
            try:
                browser = playwright.chromium.launch(headless=True, **options)
                page = browser.new_page(viewport={"width": 1240, "height": 1754}, device_scale_factor=1)
                page.set_content(html_text, wait_until="load")
                page.pdf(path=str(output_pdf), format="A4", print_background=True, prefer_css_page_size=True)
                return
            except PlaywrightError as exc:
                last_error = exc
            finally:
                if browser is not None:
                    browser.close()

    raise RuntimeError(f"Не удалось запустить Chromium/Chrome для рендера КП PDF: {last_error}")


def _browser_executable_candidates() -> list[Path]:
    candidates = [
        Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe"),
        Path(r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe"),
        Path(r"C:\Program Files\Microsoft\Edge\Application\msedge.exe"),
        Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"),
    ]
    return [path for path in candidates if path.exists()]


def _render_block(block: dict[str, Any]) -> str:
    kind = block.get("type")
    if kind == "meta":
        return f'<div class="meta"><div>{html.escape(str(block.get("left") or ""))}</div><div>{html.escape(str(block.get("right") or ""))}</div></div>'
    if kind == "table":
        rows = block.get("rows") or []
        table_class = html.escape(str(block.get("class") or "doc-table"))
        rendered_rows = []
        for row_index, row in enumerate(rows):
            cells = []
            tag = "th" if row_index == 0 and table_class == "price-table" else "td"
            for cell in row:
                cell_html = html.escape(str(cell or "")).replace("\n", "<br>")
                cells.append(f"<{tag}>{cell_html}</{tag}>")
            rendered_rows.append("<tr>" + "".join(cells) + "</tr>")
        return f'<table class="{table_class}"><tbody>{"".join(rendered_rows)}</tbody></table>'
    paragraph_class = html.escape(str(block.get("class") or "body"))
    return f'<p class="{paragraph_class}">{block.get("html") or ""}</p>'


def _render_signature(signature: KPSignatureText, assets: KPTemplateAssets) -> str:
    bg_right = f'<img class="sig-bg sig-bg-right" src="{assets.background_right}" alt="">' if assets.background_right else ""
    bg_left = f'<img class="sig-bg sig-bg-left" src="{assets.background_left}" alt="">' if assets.background_left else ""
    stamp = f'<img class="stamp" src="{assets.stamp}" alt="">' if assets.stamp else ""
    phone_icon = f'<img class="contact-icon" src="{assets.phone_icon}" alt="">' if assets.phone_icon else '<span class="contact-symbol">☎</span>'
    mail_icon = f'<img class="contact-icon mail" src="{assets.mail_icon}" alt="">' if assets.mail_icon else '<span class="contact-symbol">✉</span>'
    title_html = "<br>".join(html.escape(line) for line in signature.director_title.splitlines() if line.strip())
    return f"""
    <section class="signature">
      {bg_left}
      {bg_right}
      <div class="director-title">{title_html}</div>
      {stamp}
      <div class="signer">{html.escape(signature.signer)}</div>
      <div class="contacts">
        <div class="executor">{html.escape(signature.executor)}</div>
        <div class="contact-row">{phone_icon}<span>{html.escape(signature.phone)}</span></div>
        <div class="contact-row">{mail_icon}<span>{html.escape(signature.email)}</span></div>
      </div>
    </section>"""


def _paragraph_block(paragraph: Paragraph) -> dict[str, Any] | None:
    text = _normalize_text(paragraph.text)
    if not text:
        return None
    paragraph_html = "".join(_run_html(run) for run in paragraph.runs) or html.escape(text)
    paragraph_class = "body"
    if "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ" in text:
        paragraph_class = "title"
    elif paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        paragraph_class = "center"
    elif text.startswith("№ "):
        paragraph_class = "doc-number"
    return {"type": "paragraph", "class": paragraph_class, "html": paragraph_html}


def _run_html(run: Any) -> str:
    pieces: list[str] = []
    for child in run._element:
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "t":
            pieces.append(html.escape(child.text or ""))
        elif tag == "br":
            pieces.append("<br>")
    rendered = "".join(pieces)
    if not rendered:
        return ""
    if run.bold:
        rendered = f"<strong>{rendered}</strong>"
    return rendered


def _table_block(table: Table) -> dict[str, Any] | None:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [_normalize_text(cell.text) for cell in row.cells]
        if any(cells):
            rows.append(cells)
    if not rows:
        return None
    table_text = "\n".join("\n".join(row) for row in rows)
    table_class = "price-table" if "Вид работ" in table_text or "Стоимость" in table_text else "doc-table"
    if len(rows) == 1 and len(rows[0]) == 2 and (rows[0][0].startswith("№ ") or "Администра" in rows[0][1]):
        return {"type": "meta", "left": rows[0][0], "right": rows[0][1]}
    return {"type": "table", "class": table_class, "rows": rows}


def _trim_body_blocks(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    trimmed: list[dict[str, Any]] = []
    seen_content = False
    for block in blocks:
        rendered = str(block.get("html") or block.get("left") or block.get("rows") or "")
        if rendered.strip():
            seen_content = True
        if seen_content:
            trimmed.append(block)
    return trimmed


def _table_text(table: Table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def _company_paragraph() -> str:
    return (
        "ООО «Параллельные Решения» специализируется на комплексной разработке документов "
        "территориального планирования и градостроительного зонирования. В основе нашей работы — "
        "научная методология, глубокая градостроительная экспертиза, современные ГИС-технологии "
        "и автоматизация проектных процессов. Это позволяет сформировать качественные, обоснованные "
        "и практически применимые проектные решения, снизить риски замечаний при согласовании и "
        "обеспечить утверждение проекта в кратчайшие сроки."
    )


def _closing_paragraph(context: dict[str, Any]) -> str:
    if str(context.get("WORK_TYPE") or "") == WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES:
        return (
            "Просим Вас ознакомиться с документацией и по возможности направить обратную связь по нашему "
            "предложению. В случае заинтересованности готовы провести рабочую консультацию в формате ВКС, "
            "обсудить состав работ, сроки, порядок взаимодействия и ответить на вопросы по подготовке "
            "нормативов и другой градостроительной документации."
        )
    return (
        "Просим Вас ознакомиться с проектом договора, технического задания и календарного плана и по "
        "возможности направить обратную связь по нашему предложению. В случае заинтересованности готовы "
        "провести рабочую консультацию в формате ВКС, обсудить состав работ, сроки, порядок взаимодействия "
        "и ответить на вопросы по подготовке нормативов и другой градостроительной документации."
    )


def _rels_for_part(archive: ZipFile, part_name: str) -> dict[str, str]:
    basename = posixpath.basename(part_name)
    dirname = posixpath.dirname(part_name)
    rels_name = posixpath.join(dirname, "_rels", f"{basename}.rels")
    if rels_name not in archive.namelist():
        return {}
    rels_root = etree.fromstring(archive.read(rels_name))
    return {rel.get("Id"): rel.get("Target") for rel in rels_root.xpath("//rel:Relationship", namespaces=REL_NS)}


def _first_image_node(root: Any) -> Any | None:
    nodes = root.xpath("//wp:anchor|//wp:inline", namespaces=WORD_NS)
    return nodes[0] if nodes else None


def _data_uri_for_image_node(archive: ZipFile, part_name: str, rels: dict[str, str], node: Any) -> str:
    rid = _preferred_image_rid(node)
    if not rid:
        return ""
    target = rels.get(rid)
    if not target:
        return ""
    media_name = _resolve_part_target(part_name, target)
    if media_name not in archive.namelist():
        return ""
    payload = archive.read(media_name)
    mime_type = _mime_type(media_name, payload)
    encoded = base64.b64encode(payload).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"


def _preferred_image_rid(node: Any) -> str:
    svg_ids = node.xpath(".//asvg:svgBlip/@r:embed", namespaces=WORD_NS)
    if svg_ids:
        return str(svg_ids[0])
    image_ids = node.xpath(".//a:blip/@r:embed", namespaces=WORD_NS)
    return str(image_ids[0]) if image_ids else ""


def _resolve_part_target(part_name: str, target: str) -> str:
    if target.startswith("/"):
        return target.lstrip("/")
    return posixpath.normpath(posixpath.join(posixpath.dirname(part_name), target))


def _mime_type(name: str, payload: bytes) -> str:
    suffix = Path(name).suffix.lower()
    if suffix == ".svg":
        return "image/svg+xml"
    if suffix in {".jpg", ".jpeg"}:
        return "image/jpeg"
    if suffix == ".png":
        return "image/png"
    guessed = mimetypes.guess_type(name)[0]
    return guessed or "application/octet-stream"


def _paragraph_text_for_node(node: Any) -> str:
    parent = node.getparent()
    while parent is not None and etree.QName(parent).localname != "p":
        parent = parent.getparent()
    if parent is None:
        return ""
    return " ".join(parent.xpath(".//w:t/text()", namespaces=WORD_NS))


def _normalize_contact(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\u00a0", " ").replace(" @ ", "@").replace(" . ", ".")).strip()


def _normalize_text(value: str) -> str:
    return re.sub(r"[ \t\u00a0]+", " ", str(value or "")).strip()


def _first_text(values: list[Any]) -> str:
    return str(values[0]) if values else "0"


def _safe_int(value: Any) -> int:
    try:
        return int(str(value or "0"))
    except ValueError:
        return 0


def _css() -> str:
    return """
@page { size: A4; margin: 0; }
* { box-sizing: border-box; }
html, body { margin: 0; padding: 0; width: 210mm; height: 297mm; overflow: hidden; background: #fff; }
body { font-family: Tahoma, Arial, sans-serif; color: #595959; }
.page { width: 210mm; height: 297mm; overflow: hidden; padding: 13.5mm 14.5mm 11mm 14.5mm; position: relative; background: #fff; }
.header { display: grid; grid-template-columns: 1fr 74mm; column-gap: 12mm; align-items: start; min-height: 30mm; color: #152447; }
.logo { width: 92mm; height: auto; display: block; margin-top: 1mm; }
.logo-text { color: #152447; font-size: 18pt; font-weight: 700; line-height: 1; }
.company { font-size: 7.2pt; line-height: 1.26; color: #152447; }
.content { margin-top: 12mm; }
.meta { display: grid; grid-template-columns: 1fr 88mm; column-gap: 12mm; margin: 0 0 8mm; color: #00133f; font-size: 10.2pt; font-weight: 700; line-height: 1.25; }
p { margin: 0 0 4.1mm; font-size: 10.2pt; line-height: 1.19; text-align: justify; }
p.title { margin: 0 0 4.2mm; text-align: center; color: #00133f; font-weight: 700; font-size: 11.4pt; line-height: 1.12; }
p.center { text-align: center; }
p.doc-number { color: #00133f; font-weight: 700; }
strong { font-weight: 700; color: inherit; }
.doc-table, .price-table { width: 100%; border-collapse: collapse; margin: 3.6mm 0 1.6mm; page-break-inside: avoid; font-size: 9.45pt; line-height: 1.12; color: #595959; }
.doc-table td, .price-table td, .price-table th { border: 0.7pt solid #000; padding: 1.4mm 1.7mm; vertical-align: middle; }
.price-table th { background: #d9d9d9; color: #333; font-weight: 700; text-align: center; }
.price-table td:last-child, .price-table th:last-child { width: 31mm; text-align: center; white-space: nowrap; }
.price-table tr:last-child td { font-weight: 700; color: #333; }
.signature { position: relative; height: 45mm; margin-top: 5.4mm; page-break-inside: avoid; color: #00133f; }
.director-title { position: absolute; left: 0; top: 0; width: 78mm; font-size: 12.2pt; line-height: 1.15; font-weight: 700; color: #000; }
.signer { position: absolute; left: 116mm; top: 8.4mm; width: 72mm; font-size: 11.5pt; line-height: 1.1; font-weight: 700; color: #00133f; text-align: center; z-index: 4; }
.stamp { position: absolute; left: 80mm; top: 1.4mm; width: 44mm; height: auto; z-index: 3; }
.contacts { position: absolute; left: 0; top: 13.8mm; width: 78mm; color: #1d2a4d; font-size: 8.8pt; line-height: 1.2; }
.executor { font-weight: 700; margin-bottom: 1.8mm; }
.contact-row { display: flex; align-items: center; gap: 3.7mm; min-height: 4.6mm; margin-top: 0.4mm; color: #595959; }
.contact-icon { width: 3.5mm; max-height: 3.7mm; object-fit: contain; flex: 0 0 3.5mm; }
.contact-icon.mail { max-height: 3.1mm; }
.contact-symbol { width: 3.5mm; flex: 0 0 3.5mm; color: #00133f; font-size: 10pt; line-height: 1; }
.sig-bg { position: absolute; opacity: 1; z-index: 0; pointer-events: none; }
.sig-bg-right { right: -54mm; top: -1mm; width: 119mm; }
.sig-bg-left { left: 56mm; top: 62mm; width: 96mm; }
.signature > *:not(.sig-bg) { z-index: 2; }
"""



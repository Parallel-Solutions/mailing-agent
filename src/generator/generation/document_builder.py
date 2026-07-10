from __future__ import annotations

import math
import re
import struct
import time
import zipfile
import zlib
from pathlib import Path
import shutil
from xml.etree import ElementTree

from lxml import etree as LxmlElementTree

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from docx.table import _Cell, Table

from src.generator.generation.config_generator import BATCH_DOCX_DIR, KP_GENERATION_ENGINE, OUTPUT_DIR, TEMPLATES_DIR
from src.generator.generation.structured_kp import render_structured_kp_docx
from src.generator.generation.transforms import (
    build_district_admin_name,
    build_output_folder_name,
    ensure_official_district_wording,
    sanitize_path_component,
)
from src.generator.generation.work_types import WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES
from src.jobs.json_store import read_json, write_json_atomic


KP_TEMPLATE_FILENAME = "kp_template_source.docx"
KP_TEMPLATE_PDF_FILENAME = "kp_template_source.pdf"
CONTRACT_TEMPLATE_FILENAME = "contract_template_source.docx"
KP_TEMPLATE_PATH = TEMPLATES_DIR / KP_TEMPLATE_FILENAME
KP_TEMPLATE_PDF_PATH = TEMPLATES_DIR / KP_TEMPLATE_PDF_FILENAME
CONTRACT_TEMPLATE_PATH = TEMPLATES_DIR / CONTRACT_TEMPLATE_FILENAME
DOCUMENT_RENDERER_VERSION = "2026-06-23-signature-contact-template-gap-v23"
OUTPUT_FOLDER_MANIFEST_FILENAME = ".mailing_agent_output.json"

SVG_BLIP_PATTERN = re.compile(
    r'<a:blip r:embed="(?P<png>rId\d+)">'
    r'<a:extLst><a:ext uri="\{96DAC541-7B7A-43D3-8B79-37D633B846F1\}">'
    r'<asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" '
    r'r:embed="(?P<svg>rId\d+)"/></a:ext></a:extLst>',
    re.S,
)
SVG_RELATION_PATTERN = re.compile(
    r'<Relationship Id="(?P<id>[^"]+)" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="(?P<target>[^"]+)"\/>'
)
BACKGROUND_ANCHOR_PATTERN = re.compile(
    r'<wp:anchor\b(?=[^>]*\bbehindDoc="1")[\s\S]*?</wp:anchor>',
    re.S,
)
SVG_EXTLST_PATTERN = re.compile(
    r'<a:extLst><a:ext uri="\{96DAC541-7B7A-43D3-8B79-37D633B846F1\}">'
    r'<asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" '
    r'r:embed="rId\d+"/></a:ext></a:extLst>',
    re.S,
)
EMPTY_BLIP_PATTERN = re.compile(r'(<a:blip\b[^>]*)></a:blip>')
MALFORMED_SELF_CLOSING_BLIP_PATTERN = re.compile(r'(<a:blip\b[^>]*?)/+>')
BLIP_EMBED_PATTERN = re.compile(r'<a:blip\b[^>]*\br:embed="(?P<id>rId\d+)"')
SVG_EMBED_PATTERN = re.compile(r'<asvg:svgBlip\b[^>]*\br:embed="(?P<id>rId\d+)"')
ANCHOR_EXTENT_PATTERN = re.compile(r'<wp:extent cx="(?P<cx>\d+)" cy="(?P<cy>\d+)"/>')
ANCHOR_GEOMETRY_PATTERN = re.compile(r'<wp:positionH\b[\s\S]*?</wp:positionH><wp:positionV\b[\s\S]*?</wp:positionV><wp:extent\b[^>]*/>')
SVG_PATH_TOKEN_PATTERN = re.compile(r'[MmLlHhVvCcSsQqTtZz]|[-+]?(?:\d+\.\d*|\.\d+|\d+)(?:[eE][-+]?\d+)?')
EMU_PER_INCH = 914400
BACKGROUND_FALLBACK_PPI = 300
TERRITORIAL_ZONE_SIGNATURE_STAMP_POS_V = "-350000"
SIGNATURE_CONTACT_LEADING_SPACES = ""
SIGNATURE_CONTACT_ICON_POS_V_TARGETS = ("150190", "351165")
SIGNATURE_CONTACT_MNGP_LEFT_ICON_POS_V_TARGETS = SIGNATURE_CONTACT_ICON_POS_V_TARGETS
SIGNATURE_CONTACT_LEFT_ICON_POS_V_TARGETS = ("190000", "380000")
SIGNATURE_CONTACT_MNGP_MARKERS = ("+7 993", "+7\u00a0993", "993 079", "993\u00a0079")
SIGNATURE_CONTACT_ICON_MIN_POS_H_FOR_PDF_SHIFT = 300000
WORD_XML_NAMESPACES = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "a14": "http://schemas.microsoft.com/office/drawing/2010/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "asvg": "http://schemas.microsoft.com/office/drawing/2016/SVG/main",
}
WORD_XML_NAMESPACE_DECLS = " ".join(f'xmlns:{prefix}="{uri}"' for prefix, uri in WORD_XML_NAMESPACES.items())


def resolve_template_paths(templates_dir: Path | None = None) -> tuple[Path, Path]:
    if templates_dir is not None:
        return templates_dir / KP_TEMPLATE_FILENAME, templates_dir / CONTRACT_TEMPLATE_FILENAME
    return KP_TEMPLATE_PATH, CONTRACT_TEMPLATE_PATH


def resolve_kp_template_path(templates_dir: Path | None = None) -> Path:
    root_dir = templates_dir or TEMPLATES_DIR
    pdf_path = root_dir / KP_TEMPLATE_PDF_FILENAME
    if pdf_path.exists():
        return pdf_path
    return root_dir / KP_TEMPLATE_FILENAME



KP_TEMPLATE_AUTO_ENGINES = {"auto", "adaptive"}
KP_TEMPLATE_REQUIRED_RECIPIENT_PLACEHOLDERS = {"ADM", "ADM_NAME", "ADM_NAME_1"}
KP_TEMPLATE_REQUIRED_SCOPE_PLACEHOLDERS = {"MUN_NAME", "MUN_NAME_1", "MUN_NAME_2", "MUN_R_NAME", "MUN_R_NAME_1", "SUB_RF"}


def _docx_main_text_and_table_count(template_path: Path) -> tuple[str, int]:
    with zipfile.ZipFile(template_path, "r") as archive:
        names = archive.namelist()
        part_names = [
            name
            for name in names
            if name == "word/document.xml" or re.match(r"word/(header|footer)\d+\.xml$", name)
        ]
        text_parts: list[str] = []
        table_count = 0
        for part_name in part_names:
            xml_text = archive.read(part_name).decode("utf-8", errors="ignore")
            table_count += len(re.findall(r"<w:tbl\b", xml_text))
            text_parts.extend(re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml_text, flags=re.S))
    text = " ".join(re.sub(r"<[^>]+>", " ", part) for part in text_parts)
    return " ".join(text.split()), table_count


def is_kp_docx_template_compatible(template_path: Path) -> bool:
    if not template_path.exists() or template_path.suffix.lower() != ".docx":
        return False
    try:
        text, table_count = _docx_main_text_and_table_count(template_path)
    except Exception:
        return False
    placeholders = set(re.findall(r"\b[A-ZА-ЯЁ0-9_]{2,}\b", text))
    has_recipient = bool(placeholders & KP_TEMPLATE_REQUIRED_RECIPIENT_PLACEHOLDERS)
    has_scope = bool(placeholders & KP_TEMPLATE_REQUIRED_SCOPE_PLACEHOLDERS)
    has_price_table = table_count >= 2
    return has_recipient and has_scope and has_price_table


def should_use_structured_kp_renderer(template_path: Path, context: dict) -> bool:
    engine = str(KP_GENERATION_ENGINE or "template").strip().lower()
    if engine == "structured":
        return True
    if engine not in KP_TEMPLATE_AUTO_ENGINES:
        return False
    if str(context.get("WORK_TYPE") or "") == WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES:
        # This type has a dedicated structured renderer path elsewhere.
        return False
    return not is_kp_docx_template_compatible(template_path)


def read_output_folder_manifest(folder: Path) -> dict:
    result = read_json(folder / OUTPUT_FOLDER_MANIFEST_FILENAME, default={})
    return result.data if result.ok and isinstance(result.data, dict) else {}


def write_output_folder_manifest(folder: Path, row: dict) -> None:
    payload = {
        "row_id": str(row.get("ID") or "").strip(),
        "folder_name": build_output_folder_name(row),
        "mun_name": str(row.get("MUN_NAME") or row.get("MUN_R_NAME") or "").strip(),
        "renderer_version": DOCUMENT_RENDERER_VERSION,
    }
    write_json_atomic(folder / OUTPUT_FOLDER_MANIFEST_FILENAME, payload)


def ensure_output_folder(row: dict, output_dir: Path | None = None) -> Path:
    root_dir = output_dir or OUTPUT_DIR
    folder = root_dir / build_output_folder_name(row)
    if folder.exists():
        shutil.rmtree(folder, ignore_errors=True)
    folder.mkdir(parents=True, exist_ok=True)
    write_output_folder_manifest(folder, row)
    return folder

def ensure_batch_docx_dir(batch_docx_dir: Path | None = None) -> Path:
    target_dir = batch_docx_dir or BATCH_DOCX_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    return target_dir


def iter_paragraphs(parent: DocumentObject | _Cell | Table):
    if isinstance(parent, Table):
        for row in parent.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)
        return

    for paragraph in parent.paragraphs:
        yield paragraph

    for table in parent.tables:
        yield from iter_paragraphs(table)


def replace_text_in_runs(paragraph, replacements: list[tuple[str, str]]) -> None:
    if not paragraph.runs:
        return

    for target, replacement in replacements:
        if not target:
            continue
        replacement_text = str(replacement)
        while True:
            chunks = [run.text for run in paragraph.runs]
            original_text = "".join(chunks)
            if not original_text:
                return
            start = original_text.find(target)
            if start < 0:
                break
            end = start + len(target)

            positions: list[tuple[int, int]] = []
            cursor = 0
            for chunk in chunks:
                positions.append((cursor, cursor + len(chunk)))
                cursor += len(chunk)

            start_run_index = None
            end_run_index = None
            for index, (run_start, run_end) in enumerate(positions):
                if start_run_index is None and run_start <= start < run_end:
                    start_run_index = index
                if run_start < end <= run_end:
                    end_run_index = index
                    break

            if start_run_index is None:
                start_run_index = 0
            if end_run_index is None:
                end_run_index = len(paragraph.runs) - 1

            start_run_start, _ = positions[start_run_index]
            end_run_start, _ = positions[end_run_index]
            prefix = chunks[start_run_index][: start - start_run_start]
            suffix = chunks[end_run_index][end - end_run_start :]

            if start_run_index == end_run_index:
                paragraph.runs[start_run_index].text = prefix + replacement_text + suffix
                break

            paragraph.runs[start_run_index].text = prefix + replacement_text
            for index in range(start_run_index + 1, end_run_index):
                paragraph.runs[index].text = ""
            paragraph.runs[end_run_index].text = suffix
            break


def format_kp_recipient(value: object) -> str:
    text = str(value or "").strip()
    for index, char in enumerate(text):
        if char.isalpha():
            return f"{text[:index]}{char.upper()}{text[index + 1:]}"
    return text


def build_head_greeting_name(context: dict) -> str:
    head_fio = str(context.get("HEAD_FIO") or "").strip()
    parts = [part for part in head_fio.split() if part]
    if len(parts) >= 3:
        return f"{parts[1]} {parts[2]}"
    if head_fio:
        return head_fio
    return str(context.get("HEAD_FIO_SHORT") or "").strip()


def build_head_greeting(context: dict) -> str:
    name = build_head_greeting_name(context)
    if not name:
        return ""

    full_parts = [part for part in str(context.get("HEAD_FIO") or "").split() if part]
    patronymic = full_parts[2].lower() if len(full_parts) >= 3 else name.split()[-1].lower()
    if patronymic.endswith(("вна", "чна", "шна")):
        prefix = "Уважаемая"
    elif patronymic.endswith(("ич", "оглы")):
        prefix = "Уважаемый"
    else:
        prefix = "Уважаемый(ая)"
    return f"{prefix} {name}!"


def clear_highlights(doc: DocumentObject) -> None:
    for paragraph in iter_paragraphs(doc):
        for run in paragraph.runs:
            run.font.highlight_color = None
            r_pr = run._element.rPr
            if r_pr is not None:
                for child in list(r_pr):
                    if child.tag.endswith("}highlight") or child.tag.endswith("}shd"):
                        r_pr.remove(child)


def rebuild_paragraph(paragraph, segments: list[tuple[str, bool]]) -> None:
    base_style = paragraph.style
    template_run = paragraph.runs[0] if paragraph.runs else None

    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)

    paragraph.style = base_style
    for text, is_bold in segments:
        new_run = paragraph.add_run(text)
        if template_run is not None:
            if template_run.font.name:
                new_run.font.name = template_run.font.name
            if template_run.font.size:
                new_run.font.size = template_run.font.size
        new_run.bold = is_bold


def rebuild_paragraph_with_format(paragraph, segments: list[tuple[str, bool]], font_name: str, color: RGBColor) -> None:
    base_style = paragraph.style
    template_run = paragraph.runs[0] if paragraph.runs else None

    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)

    paragraph.style = base_style
    for text, is_bold in segments:
        new_run = paragraph.add_run(text)
        if template_run is not None:
            if template_run.font.size:
                new_run.font.size = template_run.font.size
        new_run.font.name = font_name
        new_run.font.color.rgb = color
        new_run.bold = is_bold


def apply_paragraph_font(paragraph, font_name: str, font_size_pt: float, color: RGBColor) -> None:
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        run.font.color.rgb = color


def compact_paragraph(paragraph, line_spacing: float = 1.0) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = line_spacing


def set_paragraph_line_spacing(paragraph, line_spacing: float = 1.0) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.line_spacing = line_spacing


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find("w:cantSplit", tr_pr.nsmap) is not None:
        return
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def insert_page_break_before_table(table: Table) -> None:
    tbl = table._tbl
    prev = tbl.getprevious()
    if prev is not None and prev.tag.endswith("}p"):
        texts = [node.text for node in prev.iter() if node.text]
        if "".join(texts).strip() == "":
            for br in prev.iter():
                if br.tag.endswith("}br"):
                    br_type = br.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type")
                    if br_type == "page":
                        return

    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type", "page")
    r.append(br)
    p.append(r)
    tbl.addprevious(p)


def insert_spacer_before_table(table: Table, height_pt: float) -> None:
    tbl = table._tbl
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}before", "0")
    spacing.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}after",
        str(int(height_pt * 20)),
    )
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}line", "240")
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lineRule", "auto")
    p_pr.append(spacing)
    p.append(p_pr)

    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = " "
    r.append(t)
    p.append(r)
    tbl.addprevious(p)


def set_cell_border(cell: _Cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge_name, edge_data in kwargs.items():
        tag = f"w:{edge_name}"
        element = tc_borders.find(tag, tc_borders.nsmap)
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{key}", value)


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)


def write_plain_lines(
    cell: _Cell,
    lines: list[str],
    font_name: str = "Times New Roman",
    font_size_pt: float = 12,
    first_bold: bool = False,
) -> None:
    cell.text = ""
    for index, value in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = index < len(lines) - 1
        run = paragraph.add_run(value)
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        run.bold = first_bold and index == 0


def reset_cell_text(cell, paragraphs: list[str], first_bold: bool = False) -> None:
    cell.text = ""
    target_paragraphs = cell.paragraphs
    for index, value in enumerate(paragraphs):
        paragraph = target_paragraphs[0] if index == 0 else cell.add_paragraph()
        rebuild_paragraph(paragraph, [(value, first_bold if index == 0 else False)])
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = True
        compact_paragraph(paragraph, 1.0)


def normalize_kp_formatting(doc: DocumentObject, context: dict) -> None:
    if str(context.get("WORK_TYPE") or "") == WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES:
        return

    gray = RGBColor(0x59, 0x59, 0x59)
    body_font_size = 10
    compact_body_font_size = 9.5
    body_line_spacing = 1.0
    compact_line_spacing = 1.0
    table_line_spacing = 1.0

    work_scope_fragment = (
        str(context.get("WORK_SCOPE_FRAGMENT", "")).strip()
        or f"{context.get('MUN_NAME_2', '')} {context.get('MUN_R_NAME_1', '')} {context.get('SUB_RF_1', '')}".strip()
    )
    work_scope_fragment = ensure_official_district_wording(work_scope_fragment)
    work_title = str(context.get("WORK_TITLE") or "разработке проекта местных нормативов градостроительного проектирования").strip()

    for section in doc.sections:
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.1)

    for paragraph in doc.paragraphs:
        if "ООО «Параллельные Решения» предлагает выполнить" in paragraph.text:
            rebuild_paragraph_with_format(
                paragraph,
                [
                    ("ООО «Параллельные Решения» предлагает выполнить работы по ", False),
                    (work_title, True),
                    (f" {work_scope_fragment}.", False),
                ],
                "Tahoma",
                gray,
            )
            apply_paragraph_font(paragraph, "Tahoma", compact_body_font_size, gray)
            set_paragraph_line_spacing(paragraph, compact_line_spacing)
            break

    if len(doc.tables) >= 1 and doc.tables[0].rows and len(doc.tables[0].rows[0].cells) > 1:
        rebuild_paragraph_with_format(
            doc.tables[0].rows[0].cells[1].paragraphs[0],
            [(format_kp_recipient(context.get("ADM_NAME_1", "")), False)],
            "Tahoma",
            gray,
        )
        apply_paragraph_font(doc.tables[0].rows[0].cells[1].paragraphs[0], "Tahoma", body_font_size, gray)
        set_paragraph_line_spacing(doc.tables[0].rows[0].cells[1].paragraphs[0], table_line_spacing)

    if len(doc.tables) >= 2 and len(doc.tables[1].rows) > 1 and len(doc.tables[1].rows[1].cells) > 0:
        rebuild_paragraph_with_format(
            doc.tables[1].rows[1].cells[0].paragraphs[0],
                [
                    (
                        "Выполнение работ по ",
                        False,
                    ),
                    (work_title, False),
                    (f" {work_scope_fragment}", False),
                ],
            "Tahoma",
            gray,
        )
        apply_paragraph_font(doc.tables[1].rows[1].cells[0].paragraphs[0], "Tahoma", compact_body_font_size, gray)
        set_paragraph_line_spacing(doc.tables[1].rows[1].cells[0].paragraphs[0], table_line_spacing)

    for paragraph_index in [3, 5, 6, 8, 9, 10]:
        if paragraph_index < len(doc.paragraphs):
            font_size = compact_body_font_size if paragraph_index in {5, 6, 8, 9, 10} else body_font_size
            line_spacing = compact_line_spacing if paragraph_index in {5, 6, 8, 9, 10} else body_line_spacing
            apply_paragraph_font(doc.paragraphs[paragraph_index], "Tahoma", font_size, gray)
            set_paragraph_line_spacing(doc.paragraphs[paragraph_index], line_spacing)

    if len(doc.tables) >= 2:
        for row_index in [0, 1, 2]:
            if row_index < len(doc.tables[1].rows):
                for cell in doc.tables[1].rows[row_index].cells:
                    for paragraph in cell.paragraphs:
                        font_size = compact_body_font_size if row_index == 1 else body_font_size
                        apply_paragraph_font(paragraph, "Tahoma", font_size, gray)
                        set_paragraph_line_spacing(paragraph, table_line_spacing)

    if len(doc.tables) >= 3 and doc.tables[2].rows:
        insert_spacer_before_table(doc.tables[2], 24)


def normalize_contract_formatting(doc: DocumentObject, context: dict) -> None:
    head_mo_fragment = str(context.get("HEAD_MO_FRAGMENT", "")).strip() or str(context.get("MUN_NAME_1", "")).strip()
    work_scope_fragment = (
        str(context.get("WORK_SCOPE_FRAGMENT", "")).strip()
        or f"{context.get('MUN_NAME_2', '')} {context.get('MUN_R_NAME_1', '')} {context.get('SUB_RF_1', '')}".strip()
    )
    work_scope_fragment = ensure_official_district_wording(work_scope_fragment)
    population_with_unit = str(context.get("POPULATION_WITH_UNIT", "")).strip()
    work_title = str(context.get("WORK_TITLE") or "разработке проекта местных нормативов градостроительного проектирования").strip()

    for section in doc.sections:
        section.bottom_margin = Cm(1.2)

    marker = "именуемая в дальнейшем «Заказчик»"
    for paragraph in doc.paragraphs:
        if marker in paragraph.text:
            rebuild_paragraph(
                paragraph,
                [
                    (f"{context.get('ADM_NAME', '')}, ", True),
                    ("именуемая в дальнейшем «Заказчик», в лице ", False),
                    ("главы ", True),
                    (f"{head_mo_fragment} ", False),
                    (f"{context.get('HEAD_FIO_1', '')},", True),
                    (" действующего на основании Устава, с одной стороны и ", False),
                    ("Общество с ограниченной ответственностью «ПАРАЛЛЕЛЬНЫЕ РЕШЕНИЯ»", True),
                    (", именуемое в дальнейшем «Исполнитель», в лице ", False),
                    ("Исполнительного директора Крашенинникова Константина Ивановича,", True),
                    (
                        " действующего на основании Устава, с другой стороны, в дальнейшем именуемые "
                        "«Стороны», с соблюдением требований Гражданского кодекса Российской Федерации, "
                        "Федерального закона от 05.04.2013 г. № 44-ФЗ «О контрактной системе в сфере "
                        "закупок товаров, работ, услуг для обеспечения государственных и муниципальных "
                        "нужд» (далее – Закон о контрактной системе) заключили настоящий договор "
                        "(далее - Договор) о нижеследующем:",
                        False,
                    ),
                ],
            )
            break

    work_marker = "1.1. Исполнитель обязуется своевременно оказать услуги"
    for paragraph in doc.paragraphs:
        if work_marker in paragraph.text:
            rebuild_paragraph(
                paragraph,
                [
                    ("1.1. Исполнитель обязуется своевременно оказать услуги ", False),
                    ("по ", False),
                    (work_title, True),
                    (" ", False),
                    (work_scope_fragment, False),
                    (
                        " (далее – Работы) и сдать результат Работ Заказчику, а Заказчик обязуется "
                        "принять результат работ и оплатить его.",
                        False,
                    ),
                ],
            )
            break

    execution_marker = "1.2. Выполнение"
    for paragraph in doc.paragraphs:
        if execution_marker in paragraph.text and "неотъемлемой частью настоящего" in paragraph.text:
            rebuild_paragraph(
                paragraph,
                [
                    (
                        "1.2. Выполнение работ осуществляется по месту нахождения исполнителя "
                        "на условиях и в сроки, установленные настоящим договором, "
                        "техническим заданием (Приложение № 1 к договору), "
                        "календарным планом выполнения работ (Приложение № 2 к договору), "
                        "которые являются неотъемлемой частью настоящего договора.",
                        False,
                    )
                ],
            )
            break

    population_marker = "Численность населения проектируемой территории составляет"
    if population_with_unit:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if population_marker in paragraph.text:
                            rebuild_paragraph(
                                paragraph,
                                [
                                    (
                                        f"Численность населения проектируемой территории составляет {population_with_unit}.",
                                        False,
                                    )
                                ],
                            )
                            break

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text.startswith("Глава "):
                        for run in paragraph.runs:
                            run.bold = False
                    if text.startswith("________________") or text.startswith("____________________"):
                        for run in paragraph.runs:
                            run.bold = False

    # Appendix 1 tends to spill onto the next page because the table and signature block
    # keep Word defaults for paragraph spacing. Compact these tables without changing
    # the overall template structure.
    for table_index in [4, 5]:
        if len(doc.tables) <= table_index:
            continue
        for row in doc.tables[table_index].rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.keep_together = True
                    for run in paragraph.runs:
                        if run.font.size is None or run.font.size.pt > 11:
                            run.font.size = Pt(11)

    if len(doc.tables) > 3:
        signature_table = doc.tables[3]
        keep_table_row_together(signature_table.rows[0])
        keep_table_row_together(signature_table.rows[1])

    table_index = 3
    if len(doc.tables) > table_index and len(doc.tables[table_index].rows) > 1:
        table = doc.tables[table_index]
        header_row = table.rows[0]
        text_row = table.rows[1]

        header_cells = header_row.cells
        text_cells = text_row.cells
        if len(header_cells) < 2 or len(text_cells) < 2:
            return

        left_header = header_cells[0]
        right_header = header_cells[-1]
        left_text = text_cells[0]
        right_text = text_cells[-1]
        left_header_spacer = header_cells[1] if len(header_cells) > 2 else left_header
        left_text_spacer = text_cells[1] if len(text_cells) > 2 else left_text

        write_plain_lines(left_header, ["Заказчик:"], font_size_pt=12, first_bold=True)
        if left_header._tc is not left_header_spacer._tc:
            write_plain_lines(left_header_spacer, [""], font_size_pt=12)
        write_plain_lines(right_header, ["Исполнитель:"], font_size_pt=12, first_bold=True)

        write_plain_lines(
            left_text,
            [f"Глава {context.get('ADM_NAME_1', '')}"],
            font_size_pt=12,
        )
        if left_text._tc is not left_text_spacer._tc:
            write_plain_lines(left_text_spacer, [""], font_size_pt=12)
        write_plain_lines(
            right_text,
            ['Исполнительный директор', 'ООО "Параллельные решения"'],
            font_size_pt=12,
        )

        # Hide the inner border between the two left cells so it looks like one wide cell in PDF.
        for left_cell, spacer_cell in [
            (left_header, left_header_spacer),
            (left_text, left_text_spacer),
        ]:
            set_cell_border(left_cell, right={"val": "nil"})
            if left_cell._tc is not spacer_cell._tc:
                set_cell_border(spacer_cell, left={"val": "nil"})

        signature_row = table.add_row()
        keep_table_row_together(signature_row)
        signature_cells = signature_row.cells
        left_signature = signature_cells[0]
        right_signature = signature_cells[-1]
        left_signature_spacer = signature_cells[1] if len(signature_cells) > 2 else left_signature

        write_plain_lines(
            left_signature,
            [
                f"________________ {context.get('HEAD_FIO_SHORT', '')}",
                "М.П.            «__» ____ 2026 г.",
            ],
        )
        if left_signature._tc is not left_signature_spacer._tc:
            write_plain_lines(left_signature_spacer, [""], font_size_pt=12)
        write_plain_lines(
            right_signature,
            [
                "________________ К.И. Крашенинников",
                "М.П.            «__» ____ 2026 г.",
            ],
        )

        set_cell_border(left_signature, right={"val": "nil"}, top={"val": "nil"})
        if left_signature._tc is not left_signature_spacer._tc:
            set_cell_border(left_signature_spacer, left={"val": "nil"}, top={"val": "nil"})
        set_cell_border(right_signature, top={"val": "nil"})

        for cell in signature_row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM


def background_anchor_pixel_size(anchor_fragment: str, *, ppi: int = BACKGROUND_FALLBACK_PPI) -> tuple[int, int] | None:
    match = ANCHOR_EXTENT_PATTERN.search(anchor_fragment)
    if not match:
        return None
    width_px = max(1, round(int(match.group("cx")) / EMU_PER_INCH * ppi))
    height_px = max(1, round(int(match.group("cy")) / EMU_PER_INCH * ppi))
    return width_px, height_px


def build_pdf_background_png_payload(
    png_payload: bytes,
    *,
    svg_payload: bytes | None = None,
    width_px: int | None = None,
    height_px: int | None = None,
) -> bytes:
    if svg_payload and width_px and height_px:
        rendered_payload = render_svg_paths_to_png(svg_payload, width_px=width_px, height_px=height_px)
        if rendered_payload:
            return rendered_payload
    return png_payload


def render_svg_paths_to_png(svg_payload: bytes, *, width_px: int, height_px: int, supersample: int = 2) -> bytes | None:
    try:
        import numpy as np
    except ImportError:
        return None

    if width_px <= 0 or height_px <= 0:
        return None

    try:
        root = ElementTree.fromstring(svg_payload)
        view_box = parse_svg_view_box(root.attrib.get("viewBox", ""))
    except Exception:
        return None
    if view_box is None:
        return None

    view_x, view_y, view_width, view_height = view_box
    if view_width <= 0 or view_height <= 0:
        return None

    render_width = max(1, width_px * supersample)
    render_height = max(1, height_px * supersample)
    scale_x = render_width / view_width
    scale_y = render_height / view_height
    canvas = np.zeros((render_height, render_width, 4), dtype=np.uint8)
    class_fills = svg_class_fills(root)
    rendered_any = False

    for element in root.iter():
        if _local_name(element) != "path":
            continue
        path_data = element.attrib.get("d", "")
        fill = resolve_svg_path_fill(element, class_fills)
        if not path_data or fill is None:
            continue
        subpaths = [
            [((x - view_x) * scale_x, (y - view_y) * scale_y) for x, y in subpath]
            for subpath in flatten_svg_path(path_data)
        ]
        if rasterize_compound_path(canvas, subpaths, fill, fill_rule=resolve_svg_fill_rule(element)):
            rendered_any = True

    if not rendered_any:
        return None

    if supersample > 1:
        canvas = canvas.reshape(height_px, supersample, width_px, supersample, 4).mean(axis=(1, 3)).astype(np.uint8)
    return encode_png_rgba(canvas)


def parse_svg_view_box(value: str) -> tuple[float, float, float, float] | None:
    parts = [part for part in re.split(r"[\s,]+", value.strip()) if part]
    if len(parts) != 4:
        return None
    try:
        return tuple(float(part) for part in parts)  # type: ignore[return-value]
    except ValueError:
        return None


def svg_class_fills(root: ElementTree.Element) -> dict[str, tuple[int, int, int, int]]:
    fills: dict[str, tuple[int, int, int, int]] = {}
    for element in root.iter():
        if _local_name(element) != "style" or not element.text:
            continue
        for match in re.finditer(r"\.([A-Za-z0-9_-]+)\s*\{[^}]*fill\s*:\s*([^;\s}]+)", element.text):
            color = parse_svg_color(match.group(2))
            if color is not None:
                fills[match.group(1)] = color
    return fills


def resolve_svg_path_fill(element: ElementTree.Element, class_fills: dict[str, tuple[int, int, int, int]]) -> tuple[int, int, int, int] | None:
    style = element.attrib.get("style", "")
    if style:
        match = re.search(r"(?:^|;)\s*fill\s*:\s*([^;]+)", style)
        if match:
            return parse_svg_color(match.group(1))
    for class_name in element.attrib.get("class", "").split():
        if class_name in class_fills:
            return class_fills[class_name]
    return parse_svg_color(element.attrib.get("fill"))


def resolve_svg_fill_rule(element: ElementTree.Element) -> str:
    style = element.attrib.get("style", "")
    if style:
        match = re.search(r"(?:^|;)\s*fill-rule\s*:\s*([^;]+)", style)
        if match and match.group(1).strip().lower() == "evenodd":
            return "evenodd"
    return "evenodd" if element.attrib.get("fill-rule", "").strip().lower() == "evenodd" else "nonzero"


def parse_svg_color(value: str | None) -> tuple[int, int, int, int] | None:
    if not value:
        return None
    value = value.strip()
    if value.lower() == "none":
        return None
    if value.startswith("#"):
        hex_value = value[1:]
        if len(hex_value) == 3:
            hex_value = "".join(char * 2 for char in hex_value)
        if len(hex_value) == 6:
            try:
                return int(hex_value[0:2], 16), int(hex_value[2:4], 16), int(hex_value[4:6], 16), 255
            except ValueError:
                return None
    match = re.match(r"rgb\((\d+),\s*(\d+),\s*(\d+)\)", value)
    if match:
        return tuple(max(0, min(255, int(part))) for part in match.groups()) + (255,)  # type: ignore[return-value]
    return None


def flatten_svg_path(path_data: str, *, curve_steps: int = 18) -> list[list[tuple[float, float]]]:
    tokens = SVG_PATH_TOKEN_PATTERN.findall(path_data.replace(",", " "))
    index = 0
    command = ""
    current = (0.0, 0.0)
    start = (0.0, 0.0)
    subpath: list[tuple[float, float]] = []
    subpaths: list[list[tuple[float, float]]] = []

    def is_command(token: str) -> bool:
        return len(token) == 1 and token.isalpha()

    def has_number() -> bool:
        return index < len(tokens) and not is_command(tokens[index])

    def read_float() -> float:
        nonlocal index
        value = float(tokens[index])
        index += 1
        return value

    def append_finished_subpath() -> None:
        if len(subpath) >= 3:
            subpaths.append(subpath.copy())

    while index < len(tokens):
        if is_command(tokens[index]):
            command = tokens[index]
            index += 1
        if not command:
            break

        relative = command.islower()
        command_upper = command.upper()
        if command_upper == "M":
            first_pair = True
            while has_number():
                x = read_float()
                y = read_float()
                if relative:
                    x += current[0]
                    y += current[1]
                current = (x, y)
                if first_pair:
                    append_finished_subpath()
                    subpath = [current]
                    start = current
                    first_pair = False
                else:
                    subpath.append(current)
        elif command_upper == "L":
            while has_number():
                x = read_float()
                y = read_float()
                if relative:
                    x += current[0]
                    y += current[1]
                current = (x, y)
                subpath.append(current)
        elif command_upper == "H":
            while has_number():
                x = read_float()
                if relative:
                    x += current[0]
                current = (x, current[1])
                subpath.append(current)
        elif command_upper == "V":
            while has_number():
                y = read_float()
                if relative:
                    y += current[1]
                current = (current[0], y)
                subpath.append(current)
        elif command_upper == "C":
            while has_number():
                x1 = read_float()
                y1 = read_float()
                x2 = read_float()
                y2 = read_float()
                x = read_float()
                y = read_float()
                if relative:
                    x1 += current[0]
                    y1 += current[1]
                    x2 += current[0]
                    y2 += current[1]
                    x += current[0]
                    y += current[1]
                for step in range(1, curve_steps + 1):
                    subpath.append(cubic_bezier_point(current, (x1, y1), (x2, y2), (x, y), step / curve_steps))
                current = (x, y)
        elif command_upper == "Z":
            if subpath and subpath[-1] != start:
                subpath.append(start)
            append_finished_subpath()
            subpath = []
            current = start
            command = ""
        else:
            break

    append_finished_subpath()
    return subpaths


def cubic_bezier_point(
    p0: tuple[float, float],
    p1: tuple[float, float],
    p2: tuple[float, float],
    p3: tuple[float, float],
    t: float,
) -> tuple[float, float]:
    inverse = 1 - t
    x = inverse**3 * p0[0] + 3 * inverse**2 * t * p1[0] + 3 * inverse * t**2 * p2[0] + t**3 * p3[0]
    y = inverse**3 * p0[1] + 3 * inverse**2 * t * p1[1] + 3 * inverse * t**2 * p2[1] + t**3 * p3[1]
    return x, y


def rasterize_compound_path(
    canvas,
    subpaths: list[list[tuple[float, float]]],
    fill: tuple[int, int, int, int],
    *,
    fill_rule: str = "nonzero",
) -> bool:
    polygons = [points for points in subpaths if len(points) >= 3]
    if not polygons:
        return False
    height, width, _ = canvas.shape
    min_y = max(0, math.floor(min(y for points in polygons for _, y in points)))
    max_y = min(height - 1, math.ceil(max(y for points in polygons for _, y in points)))
    if max_y < min_y:
        return False

    edges: list[tuple[float, float, float, float]] = []
    for points in polygons:
        edges.extend((x1, y1, x2, y2) for (x1, y1), (x2, y2) in zip(points, points[1:] + points[:1]) if y1 != y2)
    if not edges:
        return False

    rendered = False
    for y in range(min_y, max_y + 1):
        scan_y = y + 0.5
        events: list[tuple[float, int]] = []
        for x1, y1, x2, y2 in edges:
            if (y1 <= scan_y < y2) or (y2 <= scan_y < y1):
                x = x1 + (scan_y - y1) * (x2 - x1) / (y2 - y1)
                events.append((x, 1 if y2 > y1 else -1))
        if len(events) < 2:
            continue
        events.sort(key=lambda item: item[0])
        segments = evenodd_fill_segments(events) if fill_rule == "evenodd" else nonzero_fill_segments(events)
        for start, end in segments:
            start_x = max(0, math.floor(min(start, end)))
            end_x = min(width, math.ceil(max(start, end)))
            if end_x > start_x:
                canvas[y, start_x:end_x] = fill
                rendered = True
    return rendered


def evenodd_fill_segments(events: list[tuple[float, int]]) -> list[tuple[float, float]]:
    intersections = [x for x, _ in events]
    return [
        (intersections[index], intersections[index + 1])
        for index in range(0, len(intersections) - 1, 2)
        if intersections[index + 1] > intersections[index]
    ]


def nonzero_fill_segments(events: list[tuple[float, int]]) -> list[tuple[float, float]]:
    segments: list[tuple[float, float]] = []
    winding = 0
    segment_start: float | None = None
    for x, delta in events:
        previous = winding
        winding += delta
        if previous == 0 and winding != 0:
            segment_start = x
        elif previous != 0 and winding == 0 and segment_start is not None:
            if x > segment_start:
                segments.append((segment_start, x))
            segment_start = None
    return segments


def encode_png_rgba(array) -> bytes:
    height, width, _ = array.shape
    raw_rows = b"".join(b"\x00" + array[row_index].tobytes() for row_index in range(height))

    def chunk(tag: bytes, data: bytes) -> bytes:
        checksum = zlib.crc32(tag + data) & 0xFFFFFFFF
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", checksum)

    header = struct.pack(">IIBBBBB", width, height, 8, 6, 0, 0, 0)
    pixels_per_meter = round(BACKGROUND_FALLBACK_PPI / 0.0254)
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", header)
        + chunk(b"sRGB", b"\x00")
        + chunk(b"gAMA", struct.pack(">I", 45455))
        + chunk(b"pHYs", struct.pack(">IIB", pixels_per_meter, pixels_per_meter, 1))
        + chunk(b"IDAT", zlib.compress(raw_rows, 9))
        + chunk(b"IEND", b"")
    )


def restore_missing_background_runs_from_template(
    *,
    part_name: str,
    template_text: str,
    output_text: str,
    template_rels_text: str,
    output_rels_text: str,
    template_names: set[str],
    payloads: dict[str, bytes],
    template_zip: zipfile.ZipFile,
) -> tuple[str, str, bool, bool]:
    if 'behindDoc="1"' not in template_text:
        return output_text, output_rels_text, False, False

    background_runs: list[str] = []
    needed_relation_ids: set[str] = set()
    for anchor_match in BACKGROUND_ANCHOR_PATTERN.finditer(template_text):
        anchor_fragment = anchor_match.group(0)
        blip_matches = list(SVG_BLIP_PATTERN.finditer(anchor_fragment))
        if not blip_matches:
            continue
        background_runs.append(f"<w:r><w:drawing>{anchor_fragment}</w:drawing></w:r>")
        needed_relation_ids.update(match.group("png") for match in blip_matches)
        needed_relation_ids.update(match.group("svg") for match in blip_matches)

    if not background_runs:
        return output_text, output_rels_text, False, False

    original_output_text = output_text
    output_text, removed_existing = remove_background_runs(output_text)
    output_text, inserted = insert_background_runs_into_signature_table_cell(output_text, background_runs)
    if not inserted:
        output_text = original_output_text
        removed_existing = False
    part_changed = removed_existing or inserted
    rels_changed = False
    template_relations = {
        match.group("id"): match.group(0)
        for match in SVG_RELATION_PATTERN.finditer(template_rels_text)
    }
    template_targets = {
        match.group("id"): match.group("target")
        for match in SVG_RELATION_PATTERN.finditer(template_rels_text)
    }
    for relation_id in sorted(needed_relation_ids):
        if relation_id not in output_rels_text and relation_id in template_relations:
            output_rels_text = output_rels_text.replace("</Relationships>", template_relations[relation_id] + "</Relationships>")
            rels_changed = True
        target = template_targets.get(relation_id, "")
        if target:
            media_part = str((Path(part_name).parent / target).as_posix())
            if media_part in template_names and media_part not in payloads:
                payloads[media_part] = template_zip.read(media_part)

    return output_text, output_rels_text, part_changed, rels_changed


def restore_signature_foreground_runs_from_template(
    *,
    part_name: str,
    template_text: str,
    output_text: str,
    template_rels_text: str,
    output_rels_text: str,
    template_names: set[str],
    payloads: dict[str, bytes],
    template_zip: zipfile.ZipFile,
) -> tuple[str, str, bool, bool]:
    if "<wp:anchor" not in template_text:
        return output_text, output_rels_text, False, False
    try:
        template_root = LxmlElementTree.fromstring(ensure_word_xml_namespaces(template_text).encode("utf-8"))
        output_root = LxmlElementTree.fromstring(ensure_word_xml_namespaces(output_text).encode("utf-8"))
    except LxmlElementTree.XMLSyntaxError:
        return output_text, output_rels_text, False, False

    template_table = find_signature_table(template_root)
    output_table = find_signature_table(output_root)
    if template_table is None or output_table is None:
        return output_text, output_rels_text, False, False

    runs_by_cell: dict[tuple[int, int], list[tuple[int, str]]] = {}
    needed_relation_ids: set[str] = set()
    template_rows = template_table.xpath("./w:tr", namespaces=WORD_XML_NAMESPACES)
    for row_index, row in enumerate(template_rows):
        cells = row.xpath("./w:tc", namespaces=WORD_XML_NAMESPACES)
        for cell_index, cell in enumerate(cells):
            paragraphs = cell.xpath("./w:p", namespaces=WORD_XML_NAMESPACES)
            for paragraph_index, paragraph in enumerate(paragraphs):
                runs = paragraph.xpath(
                    "./w:r[w:drawing/wp:anchor[not(@behindDoc='1')]]",
                    namespaces=WORD_XML_NAMESPACES,
                )
                if not runs:
                    continue
                key = (row_index, cell_index)
                for run in runs:
                    runs_by_cell.setdefault(key, []).append(
                        (paragraph_index, LxmlElementTree.tostring(run, encoding="unicode"))
                    )
                    needed_relation_ids.update(run.xpath(".//a:blip/@r:embed", namespaces=WORD_XML_NAMESPACES))
                    needed_relation_ids.update(run.xpath(".//asvg:svgBlip/@r:embed", namespaces=WORD_XML_NAMESPACES))
    if not runs_by_cell:
        return output_text, output_rels_text, False, False

    output_rows = output_table.xpath("./w:tr", namespaces=WORD_XML_NAMESPACES)
    part_changed = False
    for (row_index, cell_index), run_items in runs_by_cell.items():
        if row_index >= len(output_rows):
            continue
        output_cells = output_rows[row_index].xpath("./w:tc", namespaces=WORD_XML_NAMESPACES)
        if cell_index >= len(output_cells):
            continue
        target_cell = output_cells[cell_index]
        part_changed = remove_foreground_drawing_runs(target_cell) or part_changed
        target_paragraphs = target_cell.xpath("./w:p", namespaces=WORD_XML_NAMESPACES)
        if not target_paragraphs:
            target_paragraph = LxmlElementTree.Element(f"{{{WORD_XML_NAMESPACES['w']}}}p")
            target_cell.append(target_paragraph)
            target_paragraphs = [target_paragraph]
        for paragraph_index, run_fragment in run_items:
            target_index = min(paragraph_index, len(target_paragraphs) - 1)
            target_run = LxmlElementTree.fromstring(run_fragment.encode("utf-8"))
            target_paragraphs[target_index].append(target_run)
            part_changed = True

    if not part_changed:
        return output_text, output_rels_text, False, False

    rels_changed = False
    template_relations = {
        match.group("id"): match.group(0)
        for match in SVG_RELATION_PATTERN.finditer(template_rels_text)
    }
    template_targets = {
        match.group("id"): match.group("target")
        for match in SVG_RELATION_PATTERN.finditer(template_rels_text)
    }
    for relation_id in sorted(needed_relation_ids):
        if relation_id not in output_rels_text and relation_id in template_relations:
            output_rels_text = output_rels_text.replace("</Relationships>", template_relations[relation_id] + "</Relationships>")
            rels_changed = True
        target = template_targets.get(relation_id, "")
        if target:
            media_part = str((Path(part_name).parent / target).as_posix())
            if media_part in template_names and media_part not in payloads:
                payloads[media_part] = template_zip.read(media_part)

    return LxmlElementTree.tostring(output_root, encoding="unicode"), output_rels_text, True, rels_changed


def remove_foreground_drawing_runs(cell) -> bool:
    changed = False
    runs = cell.xpath(
        ".//w:r[w:drawing/wp:anchor[not(@behindDoc='1')]]",
        namespaces=WORD_XML_NAMESPACES,
    )
    for run in runs:
        parent = run.getparent()
        if parent is None:
            continue
        parent.remove(run)
        changed = True
    return changed


def normalize_malformed_self_closing_blips(output_text: str) -> tuple[str, bool]:
    updated_text = MALFORMED_SELF_CLOSING_BLIP_PATTERN.sub(r"\1/>", output_text)
    return updated_text, updated_text != output_text


def restore_background_anchor_geometry_from_template(template_text: str, output_text: str) -> tuple[str, bool]:
    template_geometry_by_relation: dict[str, str] = {}
    for anchor_match in BACKGROUND_ANCHOR_PATTERN.finditer(template_text):
        anchor_fragment = anchor_match.group(0)
        relation_match = BLIP_EMBED_PATTERN.search(anchor_fragment)
        geometry_match = ANCHOR_GEOMETRY_PATTERN.search(anchor_fragment)
        if relation_match and geometry_match:
            template_geometry_by_relation[relation_match.group("id")] = geometry_match.group(0)
    if not template_geometry_by_relation:
        return output_text, False

    changed = False

    def replace_anchor(match: re.Match[str]) -> str:
        nonlocal changed
        anchor_fragment = match.group(0)
        relation_match = BLIP_EMBED_PATTERN.search(anchor_fragment)
        if not relation_match:
            return anchor_fragment
        template_geometry = template_geometry_by_relation.get(relation_match.group("id"))
        if not template_geometry:
            return anchor_fragment
        updated_anchor = ANCHOR_GEOMETRY_PATTERN.sub(template_geometry, anchor_fragment, count=1)
        if updated_anchor != anchor_fragment:
            changed = True
        return updated_anchor

    return BACKGROUND_ANCHOR_PATTERN.sub(replace_anchor, output_text), changed

def normalize_signature_contact_spacing_for_pdf(output_text: str) -> tuple[str, bool]:
    xml_text = ensure_word_xml_namespaces(output_text)
    try:
        root = LxmlElementTree.fromstring(xml_text.encode("utf-8"))
    except LxmlElementTree.XMLSyntaxError:
        return output_text, False
    signature_table = find_signature_table(root)
    if signature_table is None:
        return output_text, False

    changed = False
    contact_markers = (
        "\u0438\u0441\u043f.",
        "\u0442\u0435\u043b.",
        "@",
    )
    rows = signature_table.xpath("./w:tr", namespaces=WORD_XML_NAMESPACES)
    for row in rows:
        row_text = word_node_text(row).casefold()
        if not any(marker in row_text for marker in contact_markers):
            continue
        is_mngp_contact_row = any(marker in row_text for marker in SIGNATURE_CONTACT_MNGP_MARKERS)
        cells = row.xpath("./w:tc", namespaces=WORD_XML_NAMESPACES)
        if not cells:
            continue
        for text_node in cells[0].xpath(".//w:t", namespaces=WORD_XML_NAMESPACES):
            if not text_node.text or text_node.text.strip():
                continue
            if not is_mngp_contact_row and len(text_node.text) < 7:
                continue
            if text_node.text != SIGNATURE_CONTACT_LEADING_SPACES:
                text_node.text = SIGNATURE_CONTACT_LEADING_SPACES
                text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                changed = True
    if not changed:
        return output_text, False
    return LxmlElementTree.tostring(root, encoding="unicode"), True


def adjust_signature_contact_icon_positions_for_pdf(output_text: str) -> tuple[str, bool]:
    xml_text = ensure_word_xml_namespaces(output_text)
    try:
        root = LxmlElementTree.fromstring(xml_text.encode("utf-8"))
    except LxmlElementTree.XMLSyntaxError:
        return output_text, False
    signature_table = find_signature_table(root)
    if signature_table is None:
        return output_text, False

    changed = False
    contact_markers = (
        "\u0438\u0441\u043f.",
        "\u0442\u0435\u043b.",
        "@",
    )
    rows = signature_table.xpath("./w:tr", namespaces=WORD_XML_NAMESPACES)
    for row in rows:
        row_text = word_node_text(row).casefold()
        if not any(marker in row_text for marker in contact_markers):
            continue
        is_mngp_contact_row = any(marker in row_text for marker in SIGNATURE_CONTACT_MNGP_MARKERS)
        cells = row.xpath("./w:tc", namespaces=WORD_XML_NAMESPACES)
        if not cells:
            continue
        anchors = cells[0].xpath(
            ".//wp:anchor[not(@behindDoc='1')]",
            namespaces=WORD_XML_NAMESPACES,
        )
        for anchor_index, anchor in enumerate(anchors[:2]):
            extent = anchor.xpath("./wp:extent", namespaces=WORD_XML_NAMESPACES)
            if not extent:
                continue
            try:
                cx = int(extent[0].get("cx", "0"))
                cy = int(extent[0].get("cy", "0"))
            except ValueError:
                continue
            if cx > 250000 or cy > 250000:
                continue
            pos_h = anchor.xpath("./wp:positionH/wp:posOffset", namespaces=WORD_XML_NAMESPACES)
            try:
                horizontal_offset = int(pos_h[0].text or "0") if pos_h else 0
            except ValueError:
                continue
            if horizontal_offset < SIGNATURE_CONTACT_ICON_MIN_POS_H_FOR_PDF_SHIFT:
                left_targets = (
                    SIGNATURE_CONTACT_MNGP_LEFT_ICON_POS_V_TARGETS
                    if is_mngp_contact_row
                    else SIGNATURE_CONTACT_LEFT_ICON_POS_V_TARGETS
                )
                target_pos_v = left_targets[anchor_index]
            else:
                target_pos_v = SIGNATURE_CONTACT_ICON_POS_V_TARGETS[anchor_index]
            pos_v = anchor.xpath("./wp:positionV/wp:posOffset", namespaces=WORD_XML_NAMESPACES)
            if not pos_v:
                continue
            if pos_v[0].text != target_pos_v:
                pos_v[0].text = target_pos_v
                changed = True
    if not changed:
        return output_text, False
    return LxmlElementTree.tostring(root, encoding="unicode"), True


def adjust_territorial_signature_stamp_position(output_text: str) -> tuple[str, bool]:
    if "описания местоположения границ территориальных зон" not in output_text.casefold():
        return output_text, False
    xml_text = ensure_word_xml_namespaces(output_text)
    try:
        root = LxmlElementTree.fromstring(xml_text.encode("utf-8"))
    except LxmlElementTree.XMLSyntaxError:
        return output_text, False
    signature_table = find_signature_table(root)
    if signature_table is None:
        signature_table = find_territorial_signature_stamp_table(root)
    if signature_table is None:
        return output_text, False

    changed = False
    for anchor in signature_table.xpath(
        ".//wp:anchor[not(@behindDoc='1')]",
        namespaces=WORD_XML_NAMESPACES,
    ):
        extent = anchor.xpath("./wp:extent", namespaces=WORD_XML_NAMESPACES)
        if not extent:
            continue
        try:
            width = int(extent[0].get("cx") or "0")
            height = int(extent[0].get("cy") or "0")
        except ValueError:
            continue
        if width < 1_000_000 or height < 1_000_000:
            continue
        offsets = anchor.xpath("./wp:positionV/wp:posOffset", namespaces=WORD_XML_NAMESPACES)
        if not offsets:
            continue
        if offsets[0].text != TERRITORIAL_ZONE_SIGNATURE_STAMP_POS_V:
            offsets[0].text = TERRITORIAL_ZONE_SIGNATURE_STAMP_POS_V
            changed = True
    if not changed:
        return output_text, False
    return LxmlElementTree.tostring(root, encoding="unicode"), True

def enhance_background_png_payloads(
    *,
    part_name: str,
    output_text: str,
    output_rels_text: str,
    payloads: dict[str, bytes],
) -> None:
    if 'behindDoc="1"' not in output_text:
        return

    output_targets = {
        match.group("id"): match.group("target")
        for match in SVG_RELATION_PATTERN.finditer(output_rels_text)
    }
    for anchor_match in BACKGROUND_ANCHOR_PATTERN.finditer(output_text):
        anchor_fragment = anchor_match.group(0)
        pixel_size = background_anchor_pixel_size(anchor_fragment)
        png_relation_ids = [match.group("id") for match in BLIP_EMBED_PATTERN.finditer(anchor_fragment)]
        svg_relation_ids = [match.group("id") for match in SVG_EMBED_PATTERN.finditer(anchor_fragment)]
        for index, relation_id in enumerate(png_relation_ids):
            target = output_targets.get(relation_id, "")
            if not target.lower().endswith(".png"):
                continue
            media_part = str((Path(part_name).parent / target).as_posix())
            if media_part not in payloads:
                continue
            svg_relation_id = svg_relation_ids[index] if index < len(svg_relation_ids) else (svg_relation_ids[0] if svg_relation_ids else "")
            svg_target = output_targets.get(svg_relation_id, "")
            svg_part = str((Path(part_name).parent / svg_target).as_posix()) if svg_target else ""
            width_px, height_px = pixel_size or (None, None)
            payloads[media_part] = build_pdf_background_png_payload(
                payloads[media_part],
                svg_payload=payloads.get(svg_part),
                width_px=width_px,
                height_px=height_px,
            )

def strip_background_svg_extensions_for_pdf(output_text: str, output_rels_text: str) -> tuple[str, str, bool, bool]:
    text_changed = False

    def replace_anchor(match: re.Match[str]) -> str:
        nonlocal text_changed
        anchor_fragment = match.group(0)
        png_anchor = SVG_EXTLST_PATTERN.sub("", anchor_fragment)
        png_anchor = EMPTY_BLIP_PATTERN.sub(r"\1/>", png_anchor)
        png_anchor = re.sub(r"(<a:blip\b[^>]*?)/+>(?=<a:stretch>)", r"\1/>", png_anchor)
        png_anchor = re.sub(r"(<a:blip\b[^>]*?)(?<!/)>(?=<a:stretch>)", r"\1/>", png_anchor)
        if png_anchor != anchor_fragment:
            text_changed = True
        return png_anchor

    updated_text = BACKGROUND_ANCHOR_PATTERN.sub(replace_anchor, output_text)
    return updated_text, output_rels_text, text_changed, False

def remove_background_runs(output_text: str) -> tuple[str, bool]:
    changed = False
    pieces: list[str] = []
    cursor = 0
    for match in re.finditer(r"<w:r\b[\s\S]*?</w:r>", output_text):
        run_fragment = match.group(0)
        if BACKGROUND_ANCHOR_PATTERN.search(run_fragment):
            pieces.append(output_text[cursor : match.start()])
            cursor = match.end()
            changed = True
    if not changed:
        return output_text, False
    pieces.append(output_text[cursor:])
    return "".join(pieces), True


def insert_background_runs_into_signature_table_cell(output_text: str, run_fragments: list[str]) -> tuple[str, bool]:
    updated_text = insert_background_runs_with_lxml(output_text, run_fragments)
    if updated_text is not None:
        return updated_text, updated_text != output_text

    payload = "".join(run_fragments)
    for table_match in re.finditer(r"<w:tbl\b[\s\S]*?</w:tbl>", output_text):
        table_fragment = table_match.group(0)
        table_text = xml_fragment_text(table_fragment)
        if not is_signature_table_text(table_text):
            continue
        updated_table = insert_runs_into_signature_table_fragment(table_fragment, payload)
        if updated_table != table_fragment:
            return output_text[: table_match.start()] + updated_table + output_text[table_match.end() :], True
    fallback_text = insert_runs_into_last_table_cell(output_text, run_fragments)
    return fallback_text, fallback_text != output_text


def insert_background_runs_with_lxml(output_text: str, run_fragments: list[str]) -> str | None:
    xml_text = ensure_word_xml_namespaces(output_text)
    try:
        root = LxmlElementTree.fromstring(xml_text.encode("utf-8"))
    except LxmlElementTree.XMLSyntaxError:
        return None

    signature_table = find_signature_table(root)
    if signature_table is None:
        return None
    target_cell = find_signature_background_cell(signature_table)
    if target_cell is None:
        return None

    paragraphs = target_cell.xpath("./w:p", namespaces=WORD_XML_NAMESPACES)
    if paragraphs:
        target_paragraph = paragraphs[-1]
    else:
        target_paragraph = LxmlElementTree.Element(f"{{{WORD_XML_NAMESPACES['w']}}}p")
        target_cell.append(target_paragraph)

    for run_fragment in run_fragments:
        for run_element in parse_word_run_fragment(run_fragment):
            target_paragraph.append(run_element)
    return LxmlElementTree.tostring(root, encoding="unicode")


def ensure_word_xml_namespaces(output_text: str) -> str:
    document_match = re.search(r"<w:document\b[^>]*>", output_text)
    if not document_match:
        return output_text

    document_tag = document_match.group(0)
    missing_declarations = [
        f'xmlns:{prefix}="{uri}"'
        for prefix, uri in WORD_XML_NAMESPACES.items()
        if f"xmlns:{prefix}=" not in document_tag
    ]
    if not missing_declarations:
        return output_text

    updated_tag = document_tag[:-1] + " " + " ".join(missing_declarations) + document_tag[-1]
    return output_text[: document_match.start()] + updated_tag + output_text[document_match.end() :]

def parse_word_run_fragment(run_fragment: str) -> list:
    wrapper = LxmlElementTree.fromstring(f"<root {WORD_XML_NAMESPACE_DECLS}>{run_fragment}</root>".encode("utf-8"))
    return list(wrapper)


def find_signature_table(root) -> object | None:
    for table in root.xpath(".//w:tbl", namespaces=WORD_XML_NAMESPACES):
        if is_signature_table_text(word_node_text(table)):
            return table
    return None



def find_territorial_signature_stamp_table(root) -> object | None:
    for table in root.xpath(".//w:tbl", namespaces=WORD_XML_NAMESPACES):
        if "с уважением" not in word_node_text(table).casefold():
            continue
        for extent in table.xpath(".//wp:anchor[not(@behindDoc='1')]/wp:extent", namespaces=WORD_XML_NAMESPACES):
            try:
                width = int(extent.get("cx") or "0")
                height = int(extent.get("cy") or "0")
            except ValueError:
                continue
            if width >= 1_000_000 and height >= 1_000_000:
                return table
    return None

def find_signature_background_cell(table) -> object | None:
    rows = table.xpath("./w:tr", namespaces=WORD_XML_NAMESPACES)
    contact_rows = [row for row in rows if is_signature_contact_text(word_node_text(row))]
    for row in reversed(contact_rows or rows):
        cells = row.xpath("./w:tc", namespaces=WORD_XML_NAMESPACES)
        for cell in reversed(cells):
            if not word_node_text(cell).strip():
                return cell
    return None


def is_signature_table_text(text: str) -> bool:
    normalized = text.casefold()
    if "с уважением" not in normalized:
        return False
    return any(marker in normalized for marker in ("исполнитель", "директор", "исп.")) or is_signature_contact_text(text)


def is_signature_contact_text(text: str) -> bool:
    normalized = text.casefold()
    if any(marker in normalized for marker in ("исп.", "тел", "@")):
        return True
    return bool(re.search(r"(?:\+7|8)\s*[\d\s()\-]{7,}", text))


def word_node_text(node) -> str:
    return " ".join("".join(node.xpath(".//w:t/text()", namespaces=WORD_XML_NAMESPACES)).split())

def insert_runs_into_signature_table_fragment(table_fragment: str, payload: str) -> str:
    rows = list(re.finditer(r"<w:tr\b[\s\S]*?</w:tr>", table_fragment))
    for row_match in reversed(rows):
        row_fragment = row_match.group(0)
        cells = list(re.finditer(r"<w:tc\b[\s\S]*?</w:tc>", row_fragment))
        for cell_match in reversed(cells):
            cell_fragment = cell_match.group(0)
            if xml_fragment_text(cell_fragment).strip():
                continue
            updated_cell = insert_runs_into_cell_fragment(cell_fragment, payload)
            updated_row = row_fragment[: cell_match.start()] + updated_cell + row_fragment[cell_match.end() :]
            return table_fragment[: row_match.start()] + updated_row + table_fragment[row_match.end() :]
    return table_fragment


def insert_runs_into_cell_fragment(cell_fragment: str, payload: str) -> str:
    paragraph_end = cell_fragment.rfind("</w:p>")
    if paragraph_end >= 0:
        return cell_fragment[:paragraph_end] + payload + cell_fragment[paragraph_end:]
    cell_end = cell_fragment.rfind("</w:tc>")
    if cell_end < 0:
        return cell_fragment
    return cell_fragment[:cell_end] + f"<w:p>{payload}</w:p>" + cell_fragment[cell_end:]


def xml_fragment_text(fragment: str) -> str:
    return " ".join(re.sub(r"<[^>]+>", " ", fragment).split())


def insert_runs_into_last_table_cell(output_text: str, run_fragments: list[str]) -> str:
    table_start = output_text.rfind("<w:tbl")
    if table_start < 0:
        return output_text
    table_end = output_text.find("</w:tbl>", table_start)
    if table_end < 0:
        return output_text
    cell_end = output_text.rfind("</w:tc>", table_start, table_end)
    if cell_end < 0:
        return output_text
    paragraph_end = output_text.rfind("</w:p>", table_start, cell_end)
    insert_at = paragraph_end if paragraph_end >= 0 else cell_end
    return output_text[:insert_at] + "".join(run_fragments) + output_text[insert_at:]


def restore_svg_assets_from_template(template_path: Path, output_path: Path) -> None:
    with zipfile.ZipFile(template_path, "r") as template_zip, zipfile.ZipFile(output_path, "r") as output_zip:
        output_items = output_zip.infolist()
        payloads: dict[str, bytes] = {item.filename: output_zip.read(item.filename) for item in output_items}
        template_names = set(template_zip.namelist())

        candidate_parts = ["word/document.xml"] if "word/document.xml" in template_names and "word/document.xml" in payloads else []

        for part_name in candidate_parts:
            template_text = template_zip.read(part_name).decode("utf-8", errors="ignore")
            output_text = payloads[part_name].decode("utf-8", errors="ignore")
            if "svgBlip" not in template_text:
                continue

            rels_name = str(Path(part_name).parent / "_rels" / f"{Path(part_name).name}.rels").replace("\\", "/")
            template_rels_text = (
                template_zip.read(rels_name).decode("utf-8", errors="ignore")
                if rels_name in template_names
                else ""
            )
            output_rels_text = (
                payloads[rels_name].decode("utf-8", errors="ignore")
                if rels_name in payloads
                else ""
            )
            template_rel_map = {
                match.group("id"): match.group("target")
                for match in SVG_RELATION_PATTERN.finditer(template_rels_text)
            }

            part_changed = False
            rels_changed = False
            output_text, normalized_blips = normalize_malformed_self_closing_blips(output_text)
            part_changed = part_changed or normalized_blips
            for match in SVG_BLIP_PATTERN.finditer(template_text):
                full_fragment = match.group(0)
                png_rid = match.group("png")
                svg_rid = match.group("svg")
                open_tag = f'<a:blip r:embed="{png_rid}">'
                self_closing_tag = f'<a:blip r:embed="{png_rid}"/>'

                fragment_restored = False
                if open_tag in output_text and full_fragment not in output_text:
                    output_text = output_text.replace(open_tag, full_fragment, 1)
                    part_changed = True
                    fragment_restored = True
                elif self_closing_tag in output_text and full_fragment not in output_text:
                    output_text = output_text.replace(self_closing_tag, full_fragment, 1)
                    part_changed = True
                    fragment_restored = True

                if fragment_restored and svg_rid and svg_rid not in output_rels_text and svg_rid in template_rel_map:
                    relation_line = next(
                        (
                            rel_match.group(0)
                            for rel_match in SVG_RELATION_PATTERN.finditer(template_rels_text)
                            if rel_match.group("id") == svg_rid
                        ),
                        "",
                    )
                    if relation_line:
                        output_rels_text = output_rels_text.replace("</Relationships>", relation_line + "</Relationships>")
                        rels_changed = True

                    target = template_rel_map.get(svg_rid, "")
                    if target:
                        media_part = str((Path(part_name).parent / target).as_posix())
                        if media_part in template_names and media_part not in payloads:
                            payloads[media_part] = template_zip.read(media_part)

            output_text, output_rels_text, foreground_part_changed, foreground_rels_changed = restore_signature_foreground_runs_from_template(
                part_name=part_name,
                template_text=template_text,
                output_text=output_text,
                template_rels_text=template_rels_text,
                output_rels_text=output_rels_text,
                template_names=template_names,
                payloads=payloads,
                template_zip=template_zip,
            )
            part_changed = part_changed or foreground_part_changed
            rels_changed = rels_changed or foreground_rels_changed
            output_text, output_rels_text, background_part_changed, background_rels_changed = restore_missing_background_runs_from_template(
                part_name=part_name,
                template_text=template_text,
                output_text=output_text,
                template_rels_text=template_rels_text,
                output_rels_text=output_rels_text,
                template_names=template_names,
                payloads=payloads,
                template_zip=template_zip,
            )
            part_changed = part_changed or background_part_changed
            rels_changed = rels_changed or background_rels_changed
            output_text, geometry_changed = restore_background_anchor_geometry_from_template(template_text, output_text)
            part_changed = part_changed or geometry_changed
            output_text, contact_spacing_changed = normalize_signature_contact_spacing_for_pdf(output_text)
            part_changed = part_changed or contact_spacing_changed
            output_text, contact_icon_position_changed = adjust_signature_contact_icon_positions_for_pdf(output_text)
            part_changed = part_changed or contact_icon_position_changed
            output_text, stamp_position_changed = adjust_territorial_signature_stamp_position(output_text)
            part_changed = part_changed or stamp_position_changed
            enhance_background_png_payloads(
                part_name=part_name,
                output_text=output_text,
                output_rels_text=output_rels_text,
                payloads=payloads,
            )
            output_text, output_rels_text, strip_part_changed, strip_rels_changed = strip_background_svg_extensions_for_pdf(
                output_text,
                output_rels_text,
            )
            part_changed = part_changed or strip_part_changed
            rels_changed = rels_changed or strip_rels_changed
            if part_changed:
                payloads[part_name] = output_text.encode("utf-8")
            if rels_changed:
                payloads[rels_name] = output_rels_text.encode("utf-8")

    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as target_zip:
        known_names = {item.filename for item in output_items}
        for item in output_items:
            target_zip.writestr(item, payloads[item.filename])
        for extra_name, extra_payload in payloads.items():
            if extra_name not in known_names:
                target_zip.writestr(extra_name, extra_payload)


def render_docx(template_path: Path, replacements: list[tuple[str, str]], output_path: Path, context: dict) -> Path:
    doc = None
    last_exc: Exception | None = None
    for _attempt in range(4):
        try:
            doc = Document(template_path)
            break
        except (zipfile.BadZipFile, EOFError, OSError) as exc:
            last_exc = exc
            time.sleep(0.3)
    if doc is None:
        raise last_exc  # type: ignore[misc]

    for paragraph in iter_paragraphs(doc):
        replace_text_in_runs(paragraph, replacements)

    if template_path.name.startswith("kp_"):
        stabilize_kp_pdf_layout(doc, context)

    doc.save(output_path)
    if template_path.name.startswith("kp_"):
        restore_svg_assets_from_template(template_path, output_path)
    return output_path


def stabilize_kp_pdf_layout(doc: DocumentObject, context: dict | None = None) -> None:
    preserve_template_spacing = should_preserve_kp_template_spacing(doc, context)
    if not preserve_template_spacing:
        compact_kp_price_note_spacer(doc)
    remove_blank_paragraphs_before_kp_signature_table(doc)
    compact_kp_signature_table(doc, compact_spacing=True)
    remove_trailing_blank_body_paragraphs(doc)


def should_preserve_kp_template_spacing(doc: DocumentObject, context: dict | None = None) -> bool:
    if find_kp_signature_table_object(doc) is None:
        return False
    if str((context or {}).get("WORK_TYPE") or "") == WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES:
        return True
    document_text = "\n".join(paragraph.text for paragraph in doc.paragraphs).casefold()
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells).casefold()
    return "описания местоположения границ территориальных зон" in f"{document_text}\n{table_text}"


def compact_kp_price_note_spacer(doc: DocumentObject) -> None:
    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs[:-1]):
        if "Стоимость выполнения работ" not in paragraph.text:
            continue
        paragraph.paragraph_format.space_after = Pt(0)
        next_paragraph = paragraphs[index + 1]
        if next_paragraph.text.strip():
            continue
        compact_blank_paragraph(next_paragraph, line_spacing_pt=4)


def compact_kp_signature_table(doc: DocumentObject, *, compact_spacing: bool = True) -> None:
    table = find_kp_signature_table_object(doc)
    if table is None:
        return
    for row_index, row in enumerate(table.rows):
        keep_table_row_together(row)
        if row_index < len(table.rows) - 1:
            keep_table_row_with_next(row)
        if compact_spacing:
            for cell in row.cells:
                compact_signature_cell_content(cell)


def find_kp_signature_table_object(doc: DocumentObject) -> Table | None:
    for table in reversed(doc.tables):
        signature_text = "\n".join(cell.text for row in table.rows for cell in row.cells).casefold()
        if is_signature_table_text(signature_text):
            return table
    return None


def remove_blank_paragraphs_before_kp_signature_table(doc: DocumentObject) -> None:
    table = find_kp_signature_table_object(doc)
    if table is None:
        return
    previous = table._tbl.getprevious()
    kept_spacer = False
    while previous is not None and _local_name(previous) == "p" and _is_removable_blank_paragraph(previous):
        current = previous
        previous = current.getprevious()
        if not kept_spacer:
            compact_blank_paragraph_element(current, line_spacing_pt=3)
            kept_spacer = True
        else:
            doc._body._element.remove(current)


def keep_table_row_with_next(row) -> None:
    for cell in row.cells:
        for paragraph in iter_cell_paragraphs_deep(cell):
            paragraph.paragraph_format.keep_with_next = True


def compact_signature_cell_content(cell: _Cell) -> None:
    for paragraph in cell.paragraphs:
        compact_signature_paragraph(paragraph)
    for nested_table in cell.tables:
        for row in nested_table.rows:
            keep_table_row_together(row)
            for nested_cell in row.cells:
                compact_signature_cell_content(nested_cell)


def iter_cell_paragraphs_deep(cell: _Cell):
    for paragraph in cell.paragraphs:
        yield paragraph
    for nested_table in cell.tables:
        for row in nested_table.rows:
            for nested_cell in row.cells:
                yield from iter_cell_paragraphs_deep(nested_cell)


def compact_signature_paragraph(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if paragraph.text.strip():
        paragraph.paragraph_format.line_spacing = 1.0
    else:
        compact_blank_paragraph(paragraph, line_spacing_pt=1)


def compact_blank_paragraph(paragraph, *, line_spacing_pt: float) -> None:
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(" ")
    run.text = " "
    run.font.size = Pt(1)
    for extra_run in paragraph.runs[1:]:
        extra_run.text = ""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(line_spacing_pt)


def remove_trailing_blank_body_paragraphs(doc: DocumentObject) -> None:
    body = doc._body._element
    for child in reversed(list(body)):
        if _local_name(child) == "sectPr":
            continue
        if _local_name(child) != "p":
            break
        if not _is_removable_blank_paragraph(child):
            break
        body.remove(child)


def compact_blank_paragraph_element(paragraph_element, *, line_spacing_pt: float) -> None:
    p_pr = paragraph_element.find("w:pPr", paragraph_element.nsmap)
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        paragraph_element.insert(0, p_pr)
    spacing = p_pr.find("w:spacing", paragraph_element.nsmap)
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}before", "0")
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}after", "0")
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}line", str(int(line_spacing_pt * 20)))
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lineRule", "exact")
    for run in list(paragraph_element.findall("w:r", paragraph_element.nsmap)):
        paragraph_element.remove(run)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    size = OxmlElement("w:sz")
    size.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "2")
    run_pr.append(size)
    text = OxmlElement("w:t")
    text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text.text = " "
    run.append(run_pr)
    run.append(text)
    paragraph_element.append(run)

def _is_removable_blank_paragraph(paragraph_element) -> bool:
    if any(_local_name(node) == "sectPr" for node in paragraph_element.iter()):
        return False
    if any(_local_name(node) in {"drawing", "pict", "object", "br", "fldChar", "instrText"} for node in paragraph_element.iter()):
        return False
    text = "".join(node.text or "" for node in paragraph_element.iter() if _local_name(node) == "t")
    return not text.strip()


def _local_name(element) -> str:
    return str(element.tag).rsplit("}", 1)[-1]


def ensure_render_context(context: dict) -> dict:
    if context.get("DOCUMENT_ENTITY_TYPE") != "district":
        return context
    district_name = str(context.get("MUN_R_NAME") or context.get("MUN_NAME") or "").strip()
    if not district_name:
        return context
    if not str(context.get("MUN_NAME") or "").strip():
        context["MUN_NAME"] = district_name
    if not str(context.get("ADM_NAME") or "").strip():
        context["ADM_NAME"] = build_district_admin_name(district_name)
    if not str(context.get("ADM_NAME_1") or "").strip():
        context["ADM_NAME_1"] = f"Администрации муниципального образования {district_name}"
    if not str(context.get("HEAD_MO_FRAGMENT") or "").strip():
        context["HEAD_MO_FRAGMENT"] = str(context.get("MUN_R_NAME_1") or district_name).strip()
    if not str(context.get("WORK_SCOPE_FRAGMENT") or "").strip():
        context["WORK_SCOPE_FRAGMENT"] = ensure_official_district_wording(
            f"{context.get('MUN_R_NAME_1') or district_name} {context.get('SUB_RF_1') or context.get('SUB_RF') or ''}".strip()
        )
    return context


def build_kp_replacements(context: dict) -> list[tuple[str, str]]:
    context = ensure_render_context(context)
    head_greeting = build_head_greeting(context)
    head_fio = str(context.get("HEAD_FIO") or "")
    head_fio_short = str(context.get("HEAD_FIO_SHORT") or head_fio)
    district_scope_fragment = ensure_official_district_wording(
        f"{context.get('MUN_R_NAME_1', '')} {context.get('SUB_RF_1', '')}".strip()
    )
    placeholder_scope_fragment = district_scope_fragment
    if str(context.get("WORK_TYPE") or "") == WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES:
        placeholder_scope_fragment = ensure_official_district_wording(
            str(context.get("WORK_SCOPE_FRAGMENT") or "").strip()
            or f"{context.get('MUN_NAME_2', '')} {context.get('MUN_R_NAME_1', '')} {context.get('SUB_RF_1', '')}".strip()
        )
    return [
        ("№ 101-КП от 12.05.2026", f"№ {context['OUTGOING_NUMBER']}-КП от {context['DATE']}"),
        ("WORK_TYPE_LABEL", str(context.get("WORK_TYPE_LABEL", ""))),
        ("WORK_SHORT_NAME", str(context.get("WORK_SHORT_NAME", ""))),
        ("WORK_TITLE_NOMINATIVE", str(context.get("WORK_TITLE_NOMINATIVE", ""))),
        ("WORK_TITLE_1", str(context.get("WORK_TITLE_1", context.get("WORK_TITLE", "")))),
        ("WORK_TITLE", str(context.get("WORK_TITLE", ""))),
        ("WORK_RESULT_NAME", str(context.get("WORK_RESULT_NAME", ""))),
        ("HEAD_GREETING", head_greeting),
        ("Уважаемый (ая) HEAD_FIO  !", head_greeting),
        ("Уважаемый (ая) HEAD_FIO !", head_greeting),
        ("HEAD_FIO_SHORT", head_fio_short),
        ("HEAD_FIO", head_fio_short),
        ("ADM_NAME_1", format_kp_recipient(context.get("ADM_NAME_1", ""))),
        ("ADM_NAME", format_kp_recipient(context.get("ADM_NAME_1") or context.get("ADM_NAME", ""))),
        ("MUN_NAME_2", str(context.get("MUN_NAME_2", ""))),
        ("MUN_NAME_1", str(context.get("MUN_NAME_1", ""))),
        ("MUN_NAME", str(context.get("MUN_NAME", ""))),
        ("MUN_R_NAME_1", str(context.get("MUN_R_NAME_1", ""))),
        ("MUN_R_NAME  SUB_RF_1", placeholder_scope_fragment),
        ("MUN_R_NAME  SUB_RF", placeholder_scope_fragment),
        ("MUN_R_NAME SUB_RF_1", placeholder_scope_fragment),
        ("MUN_R_NAME SUB_RF", placeholder_scope_fragment),
        ("MUN_R_NAME", str(context.get("MUN_R_NAME", ""))),
        ("SUB_RF_1", str(context.get("SUB_RF_1", ""))),
        ("SUB_RF", str(context.get("SUB_RF", ""))),
    ]


def build_contract_replacements(context: dict) -> list[tuple[str, str]]:
    context = ensure_render_context(context)
    contract_number = str(context.get("CONTRACT_NUMBER", ""))
    date = str(context.get("DATE", ""))
    work_scope_fragment = ensure_official_district_wording(
        str(context.get("WORK_SCOPE_FRAGMENT", "")).strip()
        or f"{context.get('MUN_NAME_2', '')} {context.get('MUN_R_NAME_1', '')} {context.get('SUB_RF_1', '')}".strip()
    )
    district_scope_fragment = ensure_official_district_wording(
        f"{context.get('MUN_R_NAME_1', '')} {context.get('SUB_RF_1', '')}".strip()
    )
    return [
        ("№ 101", f"№ {contract_number}"),
        ("« » мая 2026 г.", date),
        ("от «  » мая 2026 г.", f"от {date}"),
        ("WORK_TYPE_LABEL", str(context.get("WORK_TYPE_LABEL", ""))),
        ("WORK_SHORT_NAME", str(context.get("WORK_SHORT_NAME", ""))),
        ("WORK_TITLE_NOMINATIVE", str(context.get("WORK_TITLE_NOMINATIVE", ""))),
        ("WORK_TITLE_1", str(context.get("WORK_TITLE_1", context.get("WORK_TITLE", "")))),
        ("WORK_TITLE", str(context.get("WORK_TITLE", ""))),
        ("WORK_RESULT_NAME", str(context.get("WORK_RESULT_NAME", ""))),
        ("ADM_NAME_1", str(context.get("ADM_NAME_1", ""))),
        ("Глава ADM_NAME", f"Глава {context.get('ADM_NAME_1', '')}"),
        ("ADM_NAME", str(context.get("ADM_NAME", ""))),
        ("HEAD_FIO_1", str(context.get("HEAD_FIO_1", ""))),
        ("HEAD_FIO", str(context.get("HEAD_FIO_SHORT", context.get("HEAD_FIO", "")))),
        ("MUN_NAME_2", str(context.get("MUN_NAME_2", ""))),
        ("MUN_NAME_1", str(context.get("MUN_NAME_1", ""))),
        ("Глава MUN_NAME", f"Глава {context.get('HEAD_MO_FRAGMENT') or context.get('MUN_NAME_1', '')}"),
        (
            "MUN_NAME_2 MUN_R_NAME SUB_RF_1",
            work_scope_fragment,
        ),
        (
            "MUN_NAME_2 MUN_R_NAME SUB_RF",
            work_scope_fragment,
        ),
        (
            "MUN_NAME MUN_R_NAME SUB_RF",
            work_scope_fragment,
        ),
        ("MUN_R_NAME  SUB_RF_1", district_scope_fragment),
        ("MUN_R_NAME  SUB_RF", district_scope_fragment),
        ("MUN_R_NAME SUB_RF_1", district_scope_fragment),
        ("MUN_R_NAME SUB_RF", district_scope_fragment),
        ("MUN_NAME", str(context.get("MUN_NAME", ""))),
        ("MUN_R_NAME_1", str(context.get("MUN_R_NAME_1", ""))),
        ("MUN_R_NAME", str(context.get("MUN_R_NAME", ""))),
        ("SUB_RF_1", str(context.get("SUB_RF_1", ""))),
        ("SUB_RF", str(context.get("SUB_RF", ""))),
        ("ADRES", str(context.get("ADRES", ""))),
        ("REQUISITES_INN", f"ИНН {context.get('REQUISITES_INN', '')}"),
        ("REQUISITES_KPP", f"КПП {context.get('REQUISITES_KPP', '')}"),
        ("REQUISITES_OGRN", f"ОГРН {context.get('REQUISITES_OGRN', '')}"),
        ("REQUISITES_OKPO", f"ОКПО {context.get('REQUISITES_OKPO', '')}"),
        ("REQUISITES_OKTNO", f"ОКТМО {context.get('REQUISITES_OKTNO', '')}"),
        ("TEL_OSN", str(context.get("TEL_OSN", ""))),
        ("EMAIL_OSN", str(context.get("EMAIL_OSN", ""))),
        ("POPULATION человек", str(context.get("POPULATION_WITH_UNIT", ""))),
        ("POPULATION", str(context.get("POPULATION", ""))),
    ]


def build_kp_filename(row: dict, context: dict | None = None) -> str:
    mun_name = sanitize_path_component(row.get("MUN_NAME") or row.get("MUN_R_NAME") or "unknown")
    filename_label = sanitize_path_component((context or {}).get("WORK_FILENAME_LABEL") or "МНГП", preserve_case=True)
    return f"КП_{filename_label}_{mun_name}.docx"


def build_kp_pdf_filename(row: dict, context: dict | None = None) -> str:
    return build_kp_filename(row, context).removesuffix(".docx") + ".pdf"


def build_contract_filename(row: dict, context: dict | None = None) -> str:
    mun_name = sanitize_path_component(row.get("MUN_NAME") or row.get("MUN_R_NAME") or "unknown")
    filename_label = sanitize_path_component((context or {}).get("WORK_FILENAME_LABEL") or "МНГП", preserve_case=True)
    return f"Договор_{filename_label}_{mun_name}.docx"


def build_staged_filename(row: dict, kind: str, extension: str = ".docx") -> str:
    row_id = sanitize_path_component(row.get("ID", "unknown"))
    safe_kind = sanitize_path_component(kind)
    mun_name = sanitize_path_component(row.get("MUN_NAME") or row.get("MUN_R_NAME") or "unknown")
    safe_extension = extension if extension.startswith(".") else f".{extension}"
    return f"{row_id}_{safe_kind}_{mun_name}{safe_extension}"


DOCUMENT_MODE_KP = "kp"
DOCUMENT_MODE_CONTRACT = "contract"
DOCUMENT_MODE_BOTH = "both"
DOCUMENT_MODE_VALUES = {DOCUMENT_MODE_KP, DOCUMENT_MODE_CONTRACT, DOCUMENT_MODE_BOTH}


def normalize_document_mode(value: str | None) -> str:
    mode = str(value or "").strip().lower()
    return mode if mode in DOCUMENT_MODE_VALUES else DOCUMENT_MODE_BOTH


def document_mode_kinds(value: str | None) -> tuple[str, ...]:
    mode = normalize_document_mode(value)
    if mode == DOCUMENT_MODE_KP:
        return ("kp",)
    if mode == DOCUMENT_MODE_CONTRACT:
        return ("contract",)
    return ("kp", "contract")


def cleanup_batch_docx_dir(batch_docx_dir: Path | None = None) -> None:
    target_dir = batch_docx_dir or BATCH_DOCX_DIR
    if target_dir.exists():
        shutil.rmtree(target_dir)


def generate_documents_for_row(
    row: dict,
    context: dict,
    *,
    output_dir: Path | None = None,
    batch_docx_dir: Path | None = None,
    templates_dir: Path | None = None,
    document_mode: str | None = None,
) -> dict[str, Path]:
    output_folder = ensure_output_folder(row, output_dir=output_dir)
    batch_docx_dir = ensure_batch_docx_dir(batch_docx_dir=batch_docx_dir)
    kp_docx_template_path, contract_template_path = resolve_template_paths(templates_dir)
    kp_template_path = resolve_kp_template_path(templates_dir)
    kp_path = batch_docx_dir / build_staged_filename(row, "kp")
    kp_pdf_path = batch_docx_dir / build_staged_filename(row, "kp", extension=".pdf")
    contract_path = batch_docx_dir / build_staged_filename(row, "contract")
    requested_kinds = set(document_mode_kinds(document_mode))
    use_structured_kp = should_use_structured_kp_renderer(kp_docx_template_path, context)

    generated_files: dict[str, Path] = {}

    if "kp" in requested_kinds and (kp_template_path.exists() or use_structured_kp):
        if kp_template_path.suffix.lower() == ".pdf":
            from src.generator.generation.pdf_template_renderer import can_render_kp_pdf_template, render_kp_pdf_template

            if not can_render_kp_pdf_template(work_type=context.get("WORK_TYPE"), template_path=kp_template_path):
                raise ValueError("PDF-шаблон КП сейчас поддержан только для вида работ «Случайный лес».")
            generated_files["kp_pdf"] = render_kp_pdf_template(kp_template_path, context, kp_pdf_path)
            generated_files["kp_final_pdf"] = output_folder / build_kp_pdf_filename(row, context)
        else:
            kp_filename = build_kp_filename(row, context)
            if use_structured_kp:
                generated_files["kp"] = render_structured_kp_docx(
                    context,
                    kp_path,
                    style_template_path=kp_docx_template_path if kp_docx_template_path.exists() else None,
                )
            else:
                generated_files["kp"] = render_docx(kp_docx_template_path, build_kp_replacements(context), kp_path, context)
            generated_files["kp_final_docx"] = output_folder / kp_filename
            generated_files["kp_final_pdf"] = output_folder / kp_filename.replace(".docx", ".pdf")
    if "contract" in requested_kinds and contract_template_path.exists():
        contract_filename = build_contract_filename(row, context)
        generated_files["contract"] = render_docx(
            contract_template_path,
            build_contract_replacements(context),
            contract_path,
            context,
        )
        generated_files["contract_final_docx"] = output_folder / contract_filename
        generated_files["contract_final_pdf"] = output_folder / contract_filename.replace(".docx", ".pdf")

    return generated_files

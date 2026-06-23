from __future__ import annotations

from dataclasses import dataclass
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path

from src.generator.generation.work_types import WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from docx.table import _Cell

@dataclass(frozen=True)
class KPPrice:
    amount_rubles: int
    amount_words: str
    vat_rate_percent: int = 5

    @property
    def vat_amount(self) -> Decimal:
        return (Decimal(self.amount_rubles) * Decimal(self.vat_rate_percent) / Decimal(105)).quantize(
            Decimal("0.01"),
            rounding=ROUND_HALF_UP,
        )


@dataclass(frozen=True)
class KPDocumentModel:
    outgoing_number: str
    date: str
    recipient: str
    work_title: str
    work_scope: str
    work_result: str
    price: KPPrice
    validity_date: str = "31.07.2026"

    @property
    def title(self) -> str:
        return "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ"

    @property
    def work_table_title(self) -> str:
        return f"Выполнение работ по {self.work_title} {self.work_scope}."

    @property
    def intro(self) -> str:
        return (
            "ООО «Параллельные Решения» предлагает выполнить работы по "
            f"{self.work_title} {self.work_scope}."
        )

    @property
    def included_services(self) -> str:
        if "местных нормативов" in self.work_title:
            return (
                "В стоимость работ включено консультационное сопровождение Заказчика на всех этапах, "
                "сбор и анализ исходных данных, подготовка проектных материалов и сопровождение "
                "согласования проекта до его утверждения."
            )
        return (
            "В стоимость работ включена подготовка электронных XML-документов с описанием "
            "местоположения границ территориальных зон в строгом соответствии с актуальной схемой "
            "Росреестра и законодательными требованиями к точности координат, а также поддержка "
            "при внесении данных в ЕГРН."
        )

    @property
    def price_note(self) -> str:
        return (
            "Стоимость выполнения работ составляет "
            f"{format_money_rubles(self.price.amount_rubles)} ({self.price.amount_words}) рублей 00 копеек, "
            f"в том числе НДС {self.price.vat_rate_percent}% — {format_decimal_rubles(self.price.vat_amount)} руб."
        )


def build_structured_kp_model(context: dict) -> KPDocumentModel:
    work_title = str(context.get("WORK_TITLE") or "").strip()
    if not work_title:
        work_title = "разработке проекта местных нормативов градостроительного проектирования"

    work_scope = str(context.get("WORK_SCOPE_FRAGMENT") or "").strip()
    if not work_scope:
        work_scope = " ".join(
            part
            for part in [
                str(context.get("MUN_NAME_2") or "").strip(),
                str(context.get("MUN_R_NAME_1") or "").strip(),
                str(context.get("SUB_RF_1") or "").strip(),
            ]
            if part
        )

    price = _price_from_context(context)
    return KPDocumentModel(
        outgoing_number=str(context.get("OUTGOING_NUMBER") or ""),
        date=str(context.get("DATE") or ""),
        recipient=str(context.get("ADM_NAME_1") or context.get("ADM_NAME") or "").strip(),
        work_title=work_title,
        work_scope=work_scope,
        work_result=str(context.get("WORK_RESULT_NAME") or "").strip(),
        price=price,
    )


def render_structured_kp_docx(
    context: dict,
    output_path: Path,
    *,
    style_template_path: Path | None = None,
) -> Path:
    model = build_structured_kp_model(context)
    doc = _new_document(style_template_path)
    _configure_document(doc)
    _render_header_block(doc, model)
    _render_title_block(doc, model)
    _render_price_table(doc, model)
    _render_company_text(doc, model)
    _render_signature(doc)
    doc.save(output_path)
    return output_path


def _new_document(style_template_path: Path | None) -> DocumentObject:
    if style_template_path and style_template_path.exists():
        doc = Document(style_template_path)
        _clear_body(doc)
        return doc
    return Document()


def _clear_body(doc: DocumentObject) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _configure_document(doc: DocumentObject) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(0.9)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Tahoma"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def _render_header_block(doc: DocumentObject, model: KPDocumentModel) -> None:
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Cm(9.0)
    header_table.columns[1].width = Cm(7.0)
    _clear_borders(header_table)

    left = header_table.rows[0].cells[0]
    right = header_table.rows[0].cells[1]
    _write_cell_lines(
        left,
        ["ПАРАЛЛЕЛЬНЫЕ РЕШЕНИЯ", "AI-технологии в градостроительстве"],
        font_size=11,
        bold_first=True,
        color=RGBColor(0x23, 0x2E, 0x50),
    )
    _write_cell_lines(
        right,
        [
            "Общество с ограниченной ответственностью",
            "«Параллельные Решения» (ООО «ПР»)",
            "",
            "195220, г. Санкт-Петербург,",
            "Кушелевская дор, д. 12 литера А, помещ. 2-н",
            "ИНН 5038110107, КПП 780401001,",
            "ОГРН 1145038110458,",
            "т. +7 (812) 242-93-12, parresh.ru",
        ],
        font_size=7.5,
        color=RGBColor(0x23, 0x2E, 0x50),
    )

    doc.add_paragraph("")
    meta_table = doc.add_table(rows=1, cols=2)
    _clear_borders(meta_table)
    _write_cell_lines(
        meta_table.rows[0].cells[0],
        [f"№ {model.outgoing_number}-КП от {model.date}"],
        font_size=11,
        bold_first=True,
        color=RGBColor(0x23, 0x2E, 0x50),
    )
    _write_cell_lines(meta_table.rows[0].cells[1], [model.recipient], font_size=10.5, bold_first=True)


def _render_title_block(doc: DocumentObject, model: KPDocumentModel) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(model.title)
    run.bold = True
    run.font.name = "Tahoma"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x23, 0x2E, 0x50)

    _add_body_paragraph(doc, model.intro, bold_fragments=[model.work_title])
    _add_body_paragraph(doc, model.included_services)


def _render_price_table(doc: DocumentObject, model: KPDocumentModel) -> None:
    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(13.8)
    table.columns[1].width = Cm(2.6)

    _set_cell_text(table.rows[0].cells[0], "Вид работ", bold=True, center=True, shade="D8D8D8")
    _set_cell_text(table.rows[0].cells[1], "Стоимость,\nруб.", bold=True, center=True, shade="D8D8D8")
    _set_cell_text(table.rows[1].cells[0], model.work_table_title)
    _set_cell_text(table.rows[1].cells[1], format_money_decimal(model.price.amount_rubles), center=True)
    _set_cell_text(table.rows[2].cells[0], "ИТОГО:", bold=True)
    _set_cell_text(table.rows[2].cells[1], format_money_decimal(model.price.amount_rubles), bold=True, center=True)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    note = _add_body_paragraph(doc, model.price_note, bold_fragments=[format_money_rubles(model.price.amount_rubles), format_decimal_rubles(model.price.vat_amount)])
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(10)


def _render_company_text(doc: DocumentObject, model: KPDocumentModel) -> None:
    _add_body_paragraph(
        doc,
        (
            "ООО «Параллельные Решения» специализируется на комплексной разработке документов "
            "территориального планирования и градостроительного зонирования. В основе нашей работы — "
            "научная методология, глубокая градостроительная экспертиза, современные ГИС-технологии "
            "и автоматизация проектных процессов. Это позволяет сформировать качественные, "
            "обоснованные и практически применимые проектные решения, снизить риски замечаний "
            "при согласовании и обеспечить утверждение проекта в кратчайшие сроки."
        ),
    )
    _add_body_paragraph(
        doc,
        (
            "Просим Вас ознакомиться с проектом договора, технического задания и календарного плана "
            "и по возможности направить обратную связь по нашему предложению. В случае заинтересованности "
            "готовы провести рабочую консультацию в формате ВКС, обсудить состав работ, сроки, порядок "
            "взаимодействия и ответить на вопросы по подготовке нормативов и другой градостроительной документации."
        ),
    )
    _add_body_paragraph(doc, f"Срок действия коммерческого предложения: до {model.validity_date}.")


def _render_signature(doc: DocumentObject) -> None:
    table = doc.add_table(rows=2, cols=3)
    _clear_borders(table)
    _write_cell_lines(table.rows[0].cells[0], ["С уважением,", "исполнительный директор"], font_size=10.5, bold_first=True)
    _write_cell_lines(table.rows[0].cells[2], ["К.И. Крашенинников"], font_size=10.5, bold_first=True)
    _write_cell_lines(
        table.rows[1].cells[0],
        ["Исп. Черкашина Наталья", "тел. +7 963 912-74-25", "ks@parresh.ru"],
        font_size=8.5,
        bold_first=True,
        color=RGBColor(0x23, 0x2E, 0x50),
    )


def _add_body_paragraph(doc: DocumentObject, text: str, *, bold_fragments: list[str] | None = None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.right_indent = Cm(0)
    _add_runs_with_bold_fragments(paragraph, text, bold_fragments or [])
    return paragraph


def _add_runs_with_bold_fragments(paragraph, text: str, bold_fragments: list[str]) -> None:
    remaining = text
    fragments = [fragment for fragment in bold_fragments if fragment]
    while remaining:
        matches = [(remaining.find(fragment), fragment) for fragment in fragments if remaining.find(fragment) >= 0]
        if not matches:
            _add_run(paragraph, remaining)
            return
        index, fragment = min(matches, key=lambda item: item[0])
        if index > 0:
            _add_run(paragraph, remaining[:index])
        _add_run(paragraph, fragment, bold=True)
        remaining = remaining[index + len(fragment) :]


def _add_run(paragraph, text: str, *, bold: bool = False):
    run = paragraph.add_run(text)
    run.font.name = "Tahoma"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    run.bold = bold
    return run


def _set_cell_text(cell: _Cell, text: str, *, bold: bool = False, center: bool = False, shade: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(text.split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Tahoma"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        run.bold = bold
    if shade:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill", shade)
        tc_pr.append(shd)


def _write_cell_lines(
    cell: _Cell,
    lines: list[str],
    *,
    font_size: float,
    bold_first: bool = False,
    color: RGBColor | None = None,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    for index, line in enumerate(lines):
        if index:
            paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = "Tahoma"
        run.font.size = Pt(font_size)
        if color is not None:
            run.font.color.rgb = color
        run.bold = bold_first and index == 0


def _clear_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.first_child_found_in("w:tcBorders")
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                edge = tc_borders.find(f"w:{edge_name}", tc_borders.nsmap)
                if edge is None:
                    edge = OxmlElement(f"w:{edge_name}")
                    tc_borders.append(edge)
                edge.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "nil")


def _price_from_context(context: dict) -> KPPrice:
    raw_amount = str(context.get("KP_PRICE_RUBLES") or context.get("PRICE_RUBLES") or "").strip()
    default_amount = "10000" if str(context.get("WORK_TYPE") or "") == WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES else "99000"
    amount = int("".join(ch for ch in raw_amount if ch.isdigit()) or default_amount)
    words = str(context.get("KP_PRICE_WORDS") or context.get("PRICE_WORDS") or "").strip()
    if not words:
        words = number_to_russian_words(amount)
    return KPPrice(amount_rubles=amount, amount_words=words)


def format_money_rubles(amount: int) -> str:
    return f"{amount:,}".replace(",", " ")


def format_money_decimal(amount: int) -> str:
    return f"{format_money_rubles(amount)},00"


def format_decimal_rubles(amount: Decimal) -> str:
    return f"{amount:,.2f}".replace(",", " ").replace(".", ",")


def number_to_russian_words(value: int) -> str:
    if value == 99000:
        return "девяносто девять тысяч"
    if value == 10000:
        return "десять тысяч"
    if value == 0:
        return "ноль"
    if value % 1000 == 0 and 1 <= value // 1000 <= 99:
        return f"{_number_under_100(value // 1000)} тысяч"
    return format_money_rubles(value)


def _number_under_100(value: int) -> str:
    ones = {
        1: "одна",
        2: "две",
        3: "три",
        4: "четыре",
        5: "пять",
        6: "шесть",
        7: "семь",
        8: "восемь",
        9: "девять",
        10: "десять",
        11: "одиннадцать",
        12: "двенадцать",
        13: "тринадцать",
        14: "четырнадцать",
        15: "пятнадцать",
        16: "шестнадцать",
        17: "семнадцать",
        18: "восемнадцать",
        19: "девятнадцать",
    }
    tens = {
        20: "двадцать",
        30: "тридцать",
        40: "сорок",
        50: "пятьдесят",
        60: "шестьдесят",
        70: "семьдесят",
        80: "восемьдесят",
        90: "девяносто",
    }
    if value in ones:
        return ones[value]
    ten = value // 10 * 10
    one = value % 10
    return f"{tens.get(ten, str(ten))} {ones.get(one, str(one))}".strip()


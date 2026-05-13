from __future__ import annotations

import re
import zipfile
from pathlib import Path
import shutil

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import _Cell, Table

from src.generator.config_generator import BATCH_DOCX_DIR, OUTPUT_DIR, TEMPLATES_DIR
from src.generator.transforms import build_output_folder_name, ensure_official_district_wording


KP_TEMPLATE_FILENAME = "kp_template_source.docx"
CONTRACT_TEMPLATE_FILENAME = "contract_template_source.docx"
KP_TEMPLATE_PATH = TEMPLATES_DIR / KP_TEMPLATE_FILENAME
CONTRACT_TEMPLATE_PATH = TEMPLATES_DIR / CONTRACT_TEMPLATE_FILENAME

SVG_BLIP_PATTERN = re.compile(
    r'<a:blip r:embed="(?P<png>rId\d+)">'
    r'<a:extLst><a:ext uri="\{96DAC541-7B7A-43D3-8B79-37D633B846F1\}">'
    r'<asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" '
    r'r:embed="(?P<svg>rId\d+)"/></a:ext></a:extLst>',
    re.S,
)
SVG_RELATION_PATTERN = re.compile(
    r'<Relationship Id="(?P<id>[^"]+)" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="(?P<target>[^"]+)"\/>'
)


def resolve_template_paths(templates_dir: Path | None = None) -> tuple[Path, Path]:
    if templates_dir is not None:
        return templates_dir / KP_TEMPLATE_FILENAME, templates_dir / CONTRACT_TEMPLATE_FILENAME
    return KP_TEMPLATE_PATH, CONTRACT_TEMPLATE_PATH


def ensure_output_folder(row: dict, output_dir: Path | None = None) -> Path:
    root_dir = output_dir or OUTPUT_DIR
    folder = root_dir / build_output_folder_name(row)
    row_id = str(row.get("ID", "")).strip()
    if row_id and root_dir.exists():
        prefix = f"{row_id}_"
        for existing_path in root_dir.iterdir():
            if not existing_path.is_dir():
                continue
            if existing_path == folder:
                continue
            if existing_path.name.startswith(prefix):
                shutil.rmtree(existing_path, ignore_errors=True)
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def ensure_batch_docx_dir(batch_docx_dir: Path | None = None) -> Path:
    target_dir = batch_docx_dir or BATCH_DOCX_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    return target_dir


def iter_paragraphs(parent: DocumentObject | _Cell | Table):
    if isinstance(parent, Table):
        for row in parent.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)
        return

    for paragraph in parent.paragraphs:
        yield paragraph

    for table in parent.tables:
        yield from iter_paragraphs(table)


def iter_cells(table: Table):
    for row in table.rows:
        for cell in row.cells:
            yield cell
            for nested_table in cell.tables:
                yield from iter_cells(nested_table)


def replace_text_in_runs(paragraph, replacements: list[tuple[str, str]]) -> None:
    if not paragraph.runs:
        return

    original_chunks = [run.text for run in paragraph.runs]
    original_text = "".join(original_chunks)
    if not original_text:
        return

    new_text = original_text
    for target, replacement in replacements:
        if target:
            new_text = new_text.replace(target, replacement)

    if new_text == original_text:
        return

    remaining = new_text
    for index, run in enumerate(paragraph.runs):
        if index == len(paragraph.runs) - 1:
            run.text = remaining
            break

        original_length = len(original_chunks[index])
        run.text = remaining[:original_length]
        remaining = remaining[original_length:]


def clear_highlights(doc: DocumentObject) -> None:
    for paragraph in iter_paragraphs(doc):
        for run in paragraph.runs:
            run.font.highlight_color = None
            r_pr = run._element.rPr
            if r_pr is not None:
                for child in list(r_pr):
                    if child.tag.endswith("}highlight") or child.tag.endswith("}shd"):
                        r_pr.remove(child)


def rebuild_paragraph(paragraph, segments: list[tuple[str, bool]]) -> None:
    base_style = paragraph.style
    template_run = paragraph.runs[0] if paragraph.runs else None

    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)

    paragraph.style = base_style
    for text, is_bold in segments:
        new_run = paragraph.add_run(text)
        if template_run is not None:
            if template_run.font.name:
                new_run.font.name = template_run.font.name
            if template_run.font.size:
                new_run.font.size = template_run.font.size
        new_run.bold = is_bold


def rebuild_paragraph_with_format(paragraph, segments: list[tuple[str, bool]], font_name: str, color: RGBColor) -> None:
    base_style = paragraph.style
    template_run = paragraph.runs[0] if paragraph.runs else None

    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)

    paragraph.style = base_style
    for text, is_bold in segments:
        new_run = paragraph.add_run(text)
        if template_run is not None:
            if template_run.font.size:
                new_run.font.size = template_run.font.size
        new_run.font.name = font_name
        new_run.font.color.rgb = color
        new_run.bold = is_bold


def apply_paragraph_font(paragraph, font_name: str, font_size_pt: float, color: RGBColor) -> None:
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        run.font.color.rgb = color


def compact_paragraph(paragraph, line_spacing: float = 1.0) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = line_spacing


def set_paragraph_line_spacing(paragraph, line_spacing: float = 1.0) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.line_spacing = line_spacing


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def insert_page_break_before_table(table: Table) -> None:
    tbl = table._tbl
    prev = tbl.getprevious()
    if prev is not None and prev.tag.endswith("}p"):
        texts = [node.text for node in prev.iter() if node.text]
        if "".join(texts).strip() == "":
            for br in prev.iter():
                if br.tag.endswith("}br"):
                    br_type = br.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type")
                    if br_type == "page":
                        return

    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type", "page")
    r.append(br)
    p.append(r)
    tbl.addprevious(p)


def insert_spacer_before_table(table: Table, height_pt: float) -> None:
    tbl = table._tbl
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}before", "0")
    spacing.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}after",
        str(int(height_pt * 20)),
    )
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}line", "240")
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lineRule", "auto")
    p_pr.append(spacing)
    p.append(p_pr)

    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = " "
    r.append(t)
    p.append(r)
    tbl.addprevious(p)


def set_cell_border(cell: _Cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge_name, edge_data in kwargs.items():
        tag = f"w:{edge_name}"
        element = tc_borders.find(tag, tc_borders.nsmap)
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{key}", value)


def set_cell_margins(cell: _Cell, *, start: int | None = None, end: int | None = None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for edge_name, value in {"start": start, "end": end}.items():
        if value is None:
            continue
        edge = tc_mar.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_mar.append(edge)
        edge.set(qn("w:w"), str(value))
        edge.set(qn("w:type"), "dxa")


def shift_small_anchored_drawings(paragraph, shift_emu: int) -> None:
    for anchor in paragraph._p.iter(qn("wp:anchor")):
        extent = anchor.find(qn("wp:extent"))
        if extent is None:
            continue
        try:
            cx = int(extent.get("cx", "0"))
            cy = int(extent.get("cy", "0"))
        except ValueError:
            continue
        # Only move the small phone/mail icons in the executor contact block.
        if cx > 250000 or cy > 250000:
            continue
        position_h = anchor.find(qn("wp:positionH"))
        if position_h is None:
            continue
        pos_offset = position_h.find(qn("wp:posOffset"))
        if pos_offset is None or not pos_offset.text:
            continue
        try:
            current_offset = int(pos_offset.text)
        except ValueError:
            continue
        if current_offset < 70000:
            pos_offset.text = str(current_offset + shift_emu)


def normalize_signature_contact_block(signature_table: Table) -> None:
    left_padding_dxa = 260
    icon_shift_emu = 90000
    contact_markers = ("Исп.", "тел.", "parresh")

    for cell in iter_cells(signature_table):
        cell_text = "\n".join(paragraph.text for paragraph in cell.paragraphs)
        if not any(marker in cell_text for marker in contact_markers):
            continue

        set_cell_margins(cell, start=left_padding_dxa, end=108)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            if any(marker in paragraph.text for marker in contact_markers):
                shift_small_anchored_drawings(paragraph, icon_shift_emu)


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)


def write_plain_lines(
    cell: _Cell,
    lines: list[str],
    font_name: str = "Times New Roman",
    font_size_pt: float = 12,
    first_bold: bool = False,
) -> None:
    cell.text = ""
    for index, value in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = index < len(lines) - 1
        run = paragraph.add_run(value)
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        run.bold = first_bold and index == 0


def reset_cell_text(cell, paragraphs: list[str], first_bold: bool = False) -> None:
    cell.text = ""
    target_paragraphs = cell.paragraphs
    for index, value in enumerate(paragraphs):
        paragraph = target_paragraphs[0] if index == 0 else cell.add_paragraph()
        rebuild_paragraph(paragraph, [(value, first_bold if index == 0 else False)])
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = True
        compact_paragraph(paragraph, 1.0)


def normalize_kp_formatting(doc: DocumentObject, context: dict) -> None:
    gray = RGBColor(0x59, 0x59, 0x59)
    body_font_size = 10
    compact_body_font_size = 9.5
    body_line_spacing = 1.0
    compact_line_spacing = 1.0
    table_line_spacing = 1.0

    work_scope_fragment = (
        str(context.get("WORK_SCOPE_FRAGMENT", "")).strip()
        or f"{context.get('MUN_NAME_2', '')} {context.get('MUN_R_NAME_1', '')} {context.get('SUB_RF_1', '')}".strip()
    )
    work_scope_fragment = ensure_official_district_wording(work_scope_fragment)

    for section in doc.sections:
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.1)

    for paragraph in doc.paragraphs:
        if "ООО «Параллельные Решения» предлагает выполнить" in paragraph.text:
            rebuild_paragraph_with_format(
                paragraph,
                [
                    ("ООО «Параллельные Решения» предлагает выполнить работы по разработке ", False),
                    ("проекта местных нормативов градостроительного проектирования", True),
                    (f" {work_scope_fragment}.", False),
                ],
                "Tahoma",
                gray,
            )
            apply_paragraph_font(paragraph, "Tahoma", compact_body_font_size, gray)
            set_paragraph_line_spacing(paragraph, compact_line_spacing)
            break

    if len(doc.tables) >= 1 and doc.tables[0].rows and len(doc.tables[0].rows[0].cells) > 1:
        rebuild_paragraph_with_format(
            doc.tables[0].rows[0].cells[1].paragraphs[0],
            [(str(context.get("ADM_NAME_1", "")), False)],
            "Tahoma",
            gray,
        )
        apply_paragraph_font(doc.tables[0].rows[0].cells[1].paragraphs[0], "Tahoma", body_font_size, gray)
        set_paragraph_line_spacing(doc.tables[0].rows[0].cells[1].paragraphs[0], table_line_spacing)

    if len(doc.tables) >= 2 and len(doc.tables[1].rows) > 1 and len(doc.tables[1].rows[1].cells) > 0:
        rebuild_paragraph_with_format(
            doc.tables[1].rows[1].cells[0].paragraphs[0],
                [
                    (
                        "Выполнение работ по разработке проекта местных нормативов градостроительного "
                        "проектирования",
                        False,
                    ),
                    (f" {work_scope_fragment}", False),
                ],
            "Tahoma",
            gray,
        )
        apply_paragraph_font(doc.tables[1].rows[1].cells[0].paragraphs[0], "Tahoma", compact_body_font_size, gray)
        set_paragraph_line_spacing(doc.tables[1].rows[1].cells[0].paragraphs[0], table_line_spacing)

    for paragraph_index in [3, 5, 6, 8, 9, 10]:
        if paragraph_index < len(doc.paragraphs):
            font_size = compact_body_font_size if paragraph_index in {5, 6, 8, 9, 10} else body_font_size
            line_spacing = compact_line_spacing if paragraph_index in {5, 6, 8, 9, 10} else body_line_spacing
            apply_paragraph_font(doc.paragraphs[paragraph_index], "Tahoma", font_size, gray)
            set_paragraph_line_spacing(doc.paragraphs[paragraph_index], line_spacing)

    if len(doc.tables) >= 2:
        for row_index in [0, 1, 2]:
            if row_index < len(doc.tables[1].rows):
                for cell in doc.tables[1].rows[row_index].cells:
                    for paragraph in cell.paragraphs:
                        font_size = compact_body_font_size if row_index == 1 else body_font_size
                        apply_paragraph_font(paragraph, "Tahoma", font_size, gray)
                        set_paragraph_line_spacing(paragraph, table_line_spacing)

    if len(doc.tables) >= 3 and doc.tables[2].rows:
        insert_spacer_before_table(doc.tables[2], 24)
        signature_table = doc.tables[2]
        left_padding_dxa = 260
        for row in signature_table.rows:
            if not row.cells:
                continue
            left_cell = row.cells[0]
            set_cell_margins(left_cell, start=left_padding_dxa, end=108)
            for paragraph in left_cell.paragraphs:
                paragraph.paragraph_format.left_indent = Pt(0)
                paragraph.paragraph_format.first_line_indent = Pt(0)
        normalize_signature_contact_block(signature_table)


def normalize_contract_formatting(doc: DocumentObject, context: dict) -> None:
    head_mo_fragment = str(context.get("HEAD_MO_FRAGMENT", "")).strip() or str(context.get("MUN_NAME_1", "")).strip()
    work_scope_fragment = (
        str(context.get("WORK_SCOPE_FRAGMENT", "")).strip()
        or f"{context.get('MUN_NAME_2', '')} {context.get('MUN_R_NAME_1', '')} {context.get('SUB_RF_1', '')}".strip()
    )
    work_scope_fragment = ensure_official_district_wording(work_scope_fragment)
    population_with_unit = str(context.get("POPULATION_WITH_UNIT", "")).strip()

    for section in doc.sections:
        section.bottom_margin = Cm(1.2)

    marker = "именуемая в дальнейшем «Заказчик»"
    for paragraph in doc.paragraphs:
        if marker in paragraph.text:
            rebuild_paragraph(
                paragraph,
                [
                    (f"{context.get('ADM_NAME', '')}, ", True),
                    ("именуемая в дальнейшем «Заказчик», в лице ", False),
                    ("главы ", True),
                    (f"{head_mo_fragment} ", False),
                    (f"{context.get('HEAD_FIO_1', '')},", True),
                    (" действующего на основании Устава, с одной стороны и ", False),
                    ("Общество с ограниченной ответственностью «ПАРАЛЛЕЛЬНЫЕ РЕШЕНИЯ»", True),
                    (", именуемое в дальнейшем «Исполнитель», в лице ", False),
                    ("Исполнительного директора Крашенинникова Константина Ивановича,", True),
                    (
                        " действующего на основании Устава, с другой стороны, в дальнейшем именуемые "
                        "«Стороны», с соблюдением требований Гражданского кодекса Российской Федерации, "
                        "Федерального закона от 05.04.2013 г. № 44-ФЗ «О контрактной системе в сфере "
                        "закупок товаров, работ, услуг для обеспечения государственных и муниципальных "
                        "нужд» (далее – Закон о контрактной системе) заключили настоящий договор "
                        "(далее - Договор) о нижеследующем:",
                        False,
                    ),
                ],
            )
            break

    work_marker = "1.1. Исполнитель обязуется своевременно оказать услуги"
    for paragraph in doc.paragraphs:
        if work_marker in paragraph.text:
            rebuild_paragraph(
                paragraph,
                [
                    ("1.1. Исполнитель обязуется своевременно оказать услуги ", False),
                    ("по разработке проекта местных нормативов градостроительного проектирования", True),
                    (" ", False),
                    (work_scope_fragment, False),
                    (
                        " (далее – Работы) и сдать результат Работ Заказчику, а Заказчик обязуется "
                        "принять результат работ и оплатить его.",
                        False,
                    ),
                ],
            )
            break

    execution_marker = "1.2. Выполнение"
    for paragraph in doc.paragraphs:
        if execution_marker in paragraph.text and "неотъемлемой частью настоящего" in paragraph.text:
            rebuild_paragraph(
                paragraph,
                [
                    (
                        "1.2. Выполнение работ осуществляется по месту нахождения исполнителя "
                        "на условиях и в сроки, установленные настоящим договором, "
                        "техническим заданием (Приложение № 1 к договору), "
                        "календарным планом выполнения работ (Приложение № 2 к договору), "
                        "которые являются неотъемлемой частью настоящего договора.",
                        False,
                    )
                ],
            )
            break

    population_marker = "Численность населения проектируемой территории составляет"
    if population_with_unit:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if population_marker in paragraph.text:
                            rebuild_paragraph(
                                paragraph,
                                [
                                    (
                                        f"Численность населения проектируемой территории составляет {population_with_unit}.",
                                        False,
                                    )
                                ],
                            )
                            break

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text.startswith("Глава "):
                        for run in paragraph.runs:
                            run.bold = False
                    if text.startswith("________________") or text.startswith("____________________"):
                        for run in paragraph.runs:
                            run.bold = False

    # Appendix 1 tends to spill onto the next page because the table and signature block
    # keep Word defaults for paragraph spacing. Compact these tables without changing
    # the overall template structure.
    for table_index in [4, 5]:
        if len(doc.tables) <= table_index:
            continue
        for row in doc.tables[table_index].rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.keep_together = True
                    for run in paragraph.runs:
                        if run.font.size is None or run.font.size.pt > 11:
                            run.font.size = Pt(11)

    if len(doc.tables) > 3:
        signature_table = doc.tables[3]
        keep_table_row_together(signature_table.rows[0])
        keep_table_row_together(signature_table.rows[1])

    table_index = 3
    if len(doc.tables) > table_index and len(doc.tables[table_index].rows) > 1:
        table = doc.tables[table_index]
        header_row = table.rows[0]
        text_row = table.rows[1]

        header_cells = header_row.cells
        text_cells = text_row.cells
        if len(header_cells) < 2 or len(text_cells) < 2:
            return

        left_header = header_cells[0]
        right_header = header_cells[-1]
        left_text = text_cells[0]
        right_text = text_cells[-1]
        left_header_spacer = header_cells[1] if len(header_cells) > 2 else left_header
        left_text_spacer = text_cells[1] if len(text_cells) > 2 else left_text

        write_plain_lines(left_header, ["Заказчик:"], font_size_pt=12, first_bold=True)
        if left_header._tc is not left_header_spacer._tc:
            write_plain_lines(left_header_spacer, [""], font_size_pt=12)
        write_plain_lines(right_header, ["Исполнитель:"], font_size_pt=12, first_bold=True)

        write_plain_lines(
            left_text,
            [f"Глава {context.get('ADM_NAME_1', '')}"],
            font_size_pt=12,
        )
        if left_text._tc is not left_text_spacer._tc:
            write_plain_lines(left_text_spacer, [""], font_size_pt=12)
        write_plain_lines(
            right_text,
            ['Исполнительный директор', 'ООО "Параллельные решения"'],
            font_size_pt=12,
        )

        # Hide the inner border between the two left cells so it looks like one wide cell in PDF.
        for left_cell, spacer_cell in [
            (left_header, left_header_spacer),
            (left_text, left_text_spacer),
        ]:
            set_cell_border(left_cell, right={"val": "nil"})
            if left_cell._tc is not spacer_cell._tc:
                set_cell_border(spacer_cell, left={"val": "nil"})

        signature_row = table.add_row()
        keep_table_row_together(signature_row)
        signature_cells = signature_row.cells
        left_signature = signature_cells[0]
        right_signature = signature_cells[-1]
        left_signature_spacer = signature_cells[1] if len(signature_cells) > 2 else left_signature

        write_plain_lines(
            left_signature,
            [
                f"________________ {context.get('HEAD_FIO_SHORT', '')}",
                "М.П.            «__» ____ 2026 г.",
            ],
        )
        if left_signature._tc is not left_signature_spacer._tc:
            write_plain_lines(left_signature_spacer, [""], font_size_pt=12)
        write_plain_lines(
            right_signature,
            [
                "________________ К.И. Крашенинников",
                "М.П.            «__» ____ 2026 г.",
            ],
        )

        set_cell_border(left_signature, right={"val": "nil"}, top={"val": "nil"})
        if left_signature._tc is not left_signature_spacer._tc:
            set_cell_border(left_signature_spacer, left={"val": "nil"}, top={"val": "nil"})
        set_cell_border(right_signature, top={"val": "nil"})

        for cell in signature_row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM


def force_svg_blips_in_docx(docx_path: Path) -> None:
    removed_rel_ids_by_part: dict[str, list[str]] = {}

    def iter_svg_candidate_parts(names: list[str]) -> list[str]:
        return sorted(
            name
            for name in names
            if name.startswith("word/")
            and name.endswith(".xml")
            and "/_rels/" not in name
        )

    def rels_path_for_part(part_name: str) -> str:
        part_path = Path(part_name)
        return str(part_path.parent / "_rels" / f"{part_path.name}.rels").replace("\\", "/")

    with zipfile.ZipFile(docx_path, "r") as source_zip:
        items = source_zip.infolist()
        payloads: dict[str, bytes] = {}
        candidate_parts = iter_svg_candidate_parts([item.filename for item in items])

        for item in items:
            data = source_zip.read(item.filename)

            if item.filename in candidate_parts:
                text = data.decode("utf-8", errors="ignore")
                removed_rel_ids: list[str] = []

                def replace_svg_fallback(match: re.Match[str]) -> str:
                    removed_rel_ids.append(match.group("png"))
                    return f'<a:blip r:embed="{match.group("svg")}">'

                updated_text = SVG_BLIP_PATTERN.sub(replace_svg_fallback, text)
                if removed_rel_ids:
                    removed_rel_ids_by_part[item.filename] = removed_rel_ids
                data = updated_text.encode("utf-8")

            elif item.filename.endswith(".rels"):
                text = data.decode("utf-8", errors="ignore")
                for parent_part, relation_ids in removed_rel_ids_by_part.items():
                    if rels_path_for_part(parent_part) != item.filename:
                        continue
                    for relation_id in relation_ids:
                        text = re.sub(
                            rf'<Relationship Id="{relation_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="[^"]+"\/>',
                            "",
                            text,
                        )
                data = text.encode("utf-8")

            payloads[item.filename] = data

    with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as target_zip:
        for item in items:
            target_zip.writestr(item, payloads[item.filename])


def restore_svg_assets_from_template(template_path: Path, output_path: Path) -> None:
    with zipfile.ZipFile(template_path, "r") as template_zip, zipfile.ZipFile(output_path, "r") as output_zip:
        output_items = output_zip.infolist()
        payloads: dict[str, bytes] = {item.filename: output_zip.read(item.filename) for item in output_items}
        template_names = set(template_zip.namelist())

        candidate_parts = sorted(
            name
            for name in template_names
            if name.startswith("word/")
            and name.endswith(".xml")
            and "/_rels/" not in name
            and name in payloads
        )

        for part_name in candidate_parts:
            template_text = template_zip.read(part_name).decode("utf-8", errors="ignore")
            output_text = payloads[part_name].decode("utf-8", errors="ignore")
            if "svgBlip" not in template_text:
                continue

            rels_name = str(Path(part_name).parent / "_rels" / f"{Path(part_name).name}.rels").replace("\\", "/")
            template_rels_text = (
                template_zip.read(rels_name).decode("utf-8", errors="ignore")
                if rels_name in template_names
                else ""
            )
            output_rels_text = (
                payloads[rels_name].decode("utf-8", errors="ignore")
                if rels_name in payloads
                else ""
            )
            template_rel_map = {
                match.group("id"): match.group("target")
                for match in SVG_RELATION_PATTERN.finditer(template_rels_text)
            }

            part_changed = False
            rels_changed = False
            for match in SVG_BLIP_PATTERN.finditer(template_text):
                full_fragment = match.group(0)
                png_rid = match.group("png")
                svg_rid = match.group("svg")
                open_tag = f'<a:blip r:embed="{png_rid}">'
                self_closing_tag = f'<a:blip r:embed="{png_rid}"/>'

                if open_tag in output_text and full_fragment not in output_text:
                    output_text = output_text.replace(open_tag, full_fragment, 1)
                    part_changed = True
                elif self_closing_tag in output_text and full_fragment not in output_text:
                    output_text = output_text.replace(self_closing_tag, full_fragment, 1)
                    part_changed = True

                if svg_rid and svg_rid not in output_rels_text and svg_rid in template_rel_map:
                    relation_line = next(
                        (
                            rel_match.group(0)
                            for rel_match in SVG_RELATION_PATTERN.finditer(template_rels_text)
                            if rel_match.group("id") == svg_rid
                        ),
                        "",
                    )
                    if relation_line:
                        output_rels_text = output_rels_text.replace("</Relationships>", relation_line + "</Relationships>")
                        rels_changed = True

                    target = template_rel_map.get(svg_rid, "")
                    if target:
                        media_part = str((Path(part_name).parent / target).as_posix())
                        if media_part in template_names and media_part not in payloads:
                            payloads[media_part] = template_zip.read(media_part)

            if part_changed:
                payloads[part_name] = output_text.encode("utf-8")
            if rels_changed:
                payloads[rels_name] = output_rels_text.encode("utf-8")

    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as target_zip:
        known_names = {item.filename for item in output_items}
        for item in output_items:
            target_zip.writestr(item, payloads[item.filename])
        for extra_name, extra_payload in payloads.items():
            if extra_name not in known_names:
                target_zip.writestr(extra_name, extra_payload)


def render_docx(template_path: Path, replacements: list[tuple[str, str]], output_path: Path, context: dict) -> Path:
    doc = Document(template_path)

    for paragraph in iter_paragraphs(doc):
        replace_text_in_runs(paragraph, replacements)

    clear_highlights(doc)
    if template_path.name.startswith("kp_"):
        normalize_kp_formatting(doc, context)
    if template_path.name.startswith("contract_"):
        normalize_contract_formatting(doc, context)
    doc.save(output_path)
    if template_path.name.startswith("kp_"):
        restore_svg_assets_from_template(template_path, output_path)
        force_svg_blips_in_docx(output_path)
    return output_path


def build_kp_replacements(context: dict) -> list[tuple[str, str]]:
    return [
        ("№ 101-КП от 12.05.2026", f"№ {context['OUTGOING_NUMBER']}-КП от {context['DATE']}"),
        ("ADM_NAME_1", str(context.get("ADM_NAME_1", ""))),
        ("MUN_NAME_2", str(context.get("MUN_NAME_2", ""))),
        ("MUN_R_NAME_1", str(context.get("MUN_R_NAME_1", ""))),
        ("SUB_RF_1", str(context.get("SUB_RF_1", ""))),
    ]


def build_contract_replacements(context: dict) -> list[tuple[str, str]]:
    contract_number = str(context.get("CONTRACT_NUMBER", ""))
    date = str(context.get("DATE", ""))
    work_scope_fragment = ensure_official_district_wording(
        str(context.get("WORK_SCOPE_FRAGMENT", "")).strip()
        or f"{context.get('MUN_NAME_2', '')} {context.get('MUN_R_NAME_1', '')} {context.get('SUB_RF_1', '')}".strip()
    )
    return [
        ("№ 101", f"№ {contract_number}"),
        ("« » мая 2026 г.", date),
        ("от «  » мая 2026 г.", f"от {date}"),
        ("ADM_NAME_1", str(context.get("ADM_NAME_1", ""))),
        ("ADM_NAME", str(context.get("ADM_NAME", ""))),
        ("HEAD_FIO_1", str(context.get("HEAD_FIO_1", ""))),
        ("HEAD_FIO", str(context.get("HEAD_FIO_SHORT", context.get("HEAD_FIO", "")))),
        ("MUN_NAME_2", str(context.get("MUN_NAME_2", ""))),
        ("MUN_NAME_1", str(context.get("MUN_NAME_1", ""))),
        (
            "MUN_NAME_2 MUN_R_NAME SUB_RF_1",
            work_scope_fragment,
        ),
        (
            "MUN_NAME_2 MUN_R_NAME SUB_RF",
            work_scope_fragment,
        ),
        (
            "MUN_NAME MUN_R_NAME SUB_RF",
            work_scope_fragment,
        ),
        ("MUN_NAME", str(context.get("MUN_NAME", ""))),
        ("MUN_R_NAME_1", str(context.get("MUN_R_NAME_1", ""))),
        ("MUN_R_NAME", str(context.get("MUN_R_NAME_1", context.get("MUN_R_NAME", "")))),
        ("SUB_RF_1", str(context.get("SUB_RF_1", ""))),
        ("SUB_RF", str(context.get("SUB_RF", ""))),
        ("ADRES", str(context.get("ADRES", ""))),
        ("REQUISITES_INN", f"ИНН {context.get('REQUISITES_INN', '')}"),
        ("REQUISITES_KPP", f"КПП {context.get('REQUISITES_KPP', '')}"),
        ("REQUISITES_OGRN", f"ОГРН {context.get('REQUISITES_OGRN', '')}"),
        ("REQUISITES_OKPO", f"ОКПО {context.get('REQUISITES_OKPO', '')}"),
        ("REQUISITES_OKTNO", f"ОКТМО {context.get('REQUISITES_OKTNO', '')}"),
        ("TEL_OSN", str(context.get("TEL_OSN", ""))),
        ("EMAIL_OSN", str(context.get("EMAIL_OSN", ""))),
        ("POPULATION человек", str(context.get("POPULATION_WITH_UNIT", ""))),
        ("POPULATION", str(context.get("POPULATION", ""))),
    ]


def build_kp_filename(row: dict) -> str:
    mun_name = str(row.get("MUN_NAME", "unknown")).strip()
    return f"КП_МНГП_{mun_name}.docx"


def build_contract_filename(row: dict) -> str:
    mun_name = str(row.get("MUN_NAME", "unknown")).strip()
    return f"Договор_МНГП_{mun_name}.docx"


def build_staged_filename(row: dict, kind: str) -> str:
    row_id = str(row.get("ID", "unknown")).strip()
    mun_name = str(row.get("MUN_NAME", "unknown")).strip().replace("/", "-").replace("\\", "-")
    return f"{row_id}_{kind}_{mun_name}.docx"


def cleanup_batch_docx_dir(batch_docx_dir: Path | None = None) -> None:
    target_dir = batch_docx_dir or BATCH_DOCX_DIR
    if target_dir.exists():
        shutil.rmtree(target_dir)


def generate_documents_for_row(
    row: dict,
    context: dict,
    *,
    output_dir: Path | None = None,
    batch_docx_dir: Path | None = None,
    templates_dir: Path | None = None,
) -> dict[str, Path]:
    output_folder = ensure_output_folder(row, output_dir=output_dir)
    batch_docx_dir = ensure_batch_docx_dir(batch_docx_dir=batch_docx_dir)
    kp_template_path, contract_template_path = resolve_template_paths(templates_dir)
    kp_path = batch_docx_dir / build_staged_filename(row, "kp")
    contract_path = batch_docx_dir / build_staged_filename(row, "contract")

    generated_files: dict[str, Path] = {}

    if kp_template_path.exists():
        generated_files["kp"] = render_docx(kp_template_path, build_kp_replacements(context), kp_path, context)
        generated_files["kp_final_docx"] = output_folder / build_kp_filename(row)
        generated_files["kp_final_pdf"] = output_folder / build_kp_filename(row).replace(".docx", ".pdf")

    if contract_template_path.exists():
        generated_files["contract"] = render_docx(
            contract_template_path,
            build_contract_replacements(context),
            contract_path,
            context,
        )
        generated_files["contract_final_docx"] = output_folder / build_contract_filename(row)
        generated_files["contract_final_pdf"] = output_folder / build_contract_filename(row).replace(".docx", ".pdf")

    return generated_files

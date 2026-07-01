from __future__ import annotations

import imaplib
import smtplib
import json
import re
import base64
import hashlib
import mimetypes
import secrets
import threading
from concurrent.futures import FIRST_COMPLETED, Future, ThreadPoolExecutor, wait
from datetime import datetime
from email.message import EmailMessage
from html import escape
from pathlib import Path
from time import perf_counter, sleep
from typing import Any, Callable
from urllib.error import HTTPError, URLError
from urllib.parse import urlencode
from urllib.request import Request, urlopen

from docx import Document

from src.generator.inflection.ai_case_agent import (
    OpenAI,
    _resolve_openai_api_key,
    _resolve_openai_base_url,
)
from src.generator.orchestration.agent_handoff import (
    count_tasks_for_agent,
    create_task,
    get_recent_events,
    get_tasks_for_agent,
)
from src.generator.generation.config_generator import DATA_DIR, DATA_XLSX_PATH, OUTPUT_DIR, START_OUTGOING_NUMBER, TEMPLATES_DIR
from src.generator.generation.excel_io import load_rows, save_workbook, update_status
from src.generator.generation.generator_agent import run_generator_agent
from src.generator.generation.work_types import DEFAULT_WORK_TYPE, get_work_type_profile, normalize_work_type
from src.generator.delivery.consent_store import (
    CONSENT_TEXT,
    has_confirmed_consent,
    mark_consent_request_sent,
    prepare_consent_request,
)
from src.generator.case_engine import build_inflected_fields_with_trace
from src.generator.philologist.philologist_agent import run_philologist
from src.generator.orchestration.responsibility_matrix import diagnose_responsibility
from src.generator.knowledge.service_knowledge import find_relevant_service_docs, format_service_rag_context
from src.jobs import load_agent_state, resolve_job_paths, save_agent_state
from src.utils.config import settings


EMAIL_PATTERN = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
MAIL_TEMPLATE_PATH = TEMPLATES_DIR / "mail_template.txt"
MAIL_TEMPLATE_DOCX_PATH = TEMPLATES_DIR / "mail_template.docx"
SENT_MAIL_LOG_PATH = DATA_DIR / "sent_mail_log.jsonl"
DEFAULT_MAIL_SUBJECT = "Коммерческое предложение на разработку МНГП."
CONSENT_REQUEST_SUBJECT = "МНГП для {MUN_R_NAME}: согласие на получение КП и проекта договора"
DEFAULT_MAIL_BODY = (
    "Добрый день!\n"
    "Направляем для рассмотрения коммерческое предложение на выполнение работ по разработке проекта "
    "местных нормативов градостроительного проектирования.\n"
    "Во вложении:\n"
    "— коммерческое предложение;\n"
    "— проект договора;\n"
    "— проект технического задания;\n"
    "— календарный план выполнения работ.\n"
    "\n"
    "ООО «Параллельные Решения» специализируется на разработке документов территориального планирования "
    "и градостроительного зонирования. В состав работ входят сбор и анализ исходных данных, подготовка "
    "проектных материалов и сопровождение согласования проекта до его утверждения.\n"
    "Просим передать материалы должностному лицу, курирующему вопросы архитектуры и градостроительства. "
    "Готовы провести рабочую консультацию по составу работ, срокам, порядку взаимодействия и ответить "
    "на вопросы в формате ВКС."
)
DEFAULT_MAIL_FOOTER_TEXT = (
    "С уважением,\n"
    "Крашенинников Константин Иванович\n"
    "+7 (921) 409-4561\n"
    "ООО «Параллельные Решения»\n"
    "https://www.parresh.ru/"
)
MAIL_FOOTER_LOGO_PATH = Path(__file__).resolve().parents[1] / "assets" / "parresh-signature-logo-KI.png"
MAIL_FOOTER_LOGO_CID = "parresh-signature-logo"
MAIL_FOOTER_HTML_TEMPLATE = """
<table role="presentation" cellpadding="0" cellspacing="0" style="margin-top:16px;border-collapse:collapse">
  <tr><td style="padding:0"><img src="{image_src}" alt="Параллельные Решения" width="340" style="display:block;width:340px;max-width:340px;height:auto;border:0;outline:none;text-decoration:none"></td></tr>
  <tr><td style="padding:6px 0 0 0;font-family:Arial,'Helvetica Neue',Helvetica,sans-serif;font-size:13px;line-height:1.3"><a href="https://www.parresh.ru/" style="color:#1f5da8;text-decoration:underline">https://www.parresh.ru/</a></td></tr>
</table>
""".strip()
STATUS_SENT_VALUE = "Отправлено"
STATUS_CONSENT_REQUEST_SENT_VALUE = "Запрос согласия отправлен"
STATUS_ERROR_VALUE = "Ошибка"
STATUS_OK_VALUES = {"ОК", "OK", "SENT", "ОТПРАВЛЕНО", "ОТПРАВЛЕНО (ОК)"}
ATTACHMENT_MODE_KP = "kp"
ATTACHMENT_MODE_CONTRACT = "contract"
ATTACHMENT_MODE_BOTH = "both"
ATTACHMENT_MODE_VALUES = {ATTACHMENT_MODE_KP, ATTACHMENT_MODE_CONTRACT, ATTACHMENT_MODE_BOTH}
MAX_STATUS_ERROR_LENGTH = 240
UNISENDER_GO_SEND_PATH = "email/send.json"
UNISENDER_CLASSIC_SEND_PATH = "sendEmail"
UNISENDER_CLASSIC_CHECK_PATH = "checkEmail"
RUSENDER_SEND_PATH = "external-mails/send"
MAILOPOST_SEND_PATH = "email/messages"

RECIPIENT_STRATEGY_ALL = "all"
RECIPIENT_STRATEGY_PRIMARY_THEN_FALLBACK = "primary_then_fallback"
DEFAULT_RECIPIENT_STRATEGY = RECIPIENT_STRATEGY_ALL
DELIVERY_FALLBACK_FAILURE_STATUSES = {
    "hard_bounced",
    "err_delivery_failed",
    "err_user_unknown",
    "err_user_inactive",
    "err_mailbox_full",
    "err_spam_rejected",
    "err_lost",
    "spam",
    "skipped",
    "failed",
    "failure",
    "error",
    "delivery_failed",
    "not_delivered",
    "undelivered",
    "rejected",
    "bounced",
    "bounce",
}
_DELIVERY_FALLBACK_LOCK = threading.Lock()
_DELIVERY_FALLBACK_RUNNING: set[str] = set()

SENDER_STATE: dict[str, Any] = {
    "status": "idle",
    "mode": "dry_run",
    "started_at": None,
    "completed_at": None,
    "processed_rows": 0,
    "ready_rows": 0,
    "sent_rows": 0,
    "error_rows": 0,
    "skipped_rows": 0,
    "handoff_rows": 0,
    "total_rows": 0,
    "summary_text": "Проверка перед отправкой ещё не запускалась.",
    "rows": [],
    "stats": {},
    "warning_rows": 0,
    "task_stats": {"total": 0, "pending": 0, "in_progress": 0, "done": 0, "blocked": 0},
    "tasks": [],
    "recent_events": [],
    "generator_handoff_rows": 0,
    "philology_blocked_rows": 0,
    "autonomous_recovery_rows": 0,
    "effective_limit": None,
    "remaining_rows": 0,
    "stop_requested": False,
    "stop_requested_at": None,
    "transport": "smtp",
    "attachment_mode": ATTACHMENT_MODE_KP,
    "recipient_strategy": DEFAULT_RECIPIENT_STRATEGY,
}

SENDER_STATE_ROWS_LIMIT = 200
SENDER_WORKBOOK_SAVE_EVERY = 25
UNISENDER_RETRY_ATTEMPTS = 3
UNISENDER_RETRY_BASE_SECONDS = 2.0
UNISENDER_RETRYABLE_HTTP_CODES = {408, 409, 425, 429, 500, 502, 503, 504}
UNISENDER_REQUESTS_PER_MINUTE = 60
UNISENDER_MIN_REQUEST_INTERVAL_SECONDS = 60.0 / UNISENDER_REQUESTS_PER_MINUTE
UNISENDER_RATE_LIMIT_RETRY_SECONDS = 65.0
_UNISENDER_RATE_LIMIT_LOCK = threading.Lock()
_last_unisender_request_at = 0.0
RUSENDER_RETRY_ATTEMPTS = 3
RUSENDER_RETRY_BASE_SECONDS = 2.0
RUSENDER_RETRYABLE_HTTP_CODES = {408, 409, 425, 429, 500, 502, 503, 504}
RUSENDER_REQUESTS_PER_SECOND = 10
RUSENDER_MIN_REQUEST_INTERVAL_SECONDS = 1.0 / RUSENDER_REQUESTS_PER_SECOND
_RUSENDER_RATE_LIMIT_LOCK = threading.Lock()
_last_rusender_request_at = 0.0
MAILOPOST_RETRY_ATTEMPTS = 3
MAILOPOST_RETRY_BASE_SECONDS = 2.0
MAILOPOST_RETRYABLE_HTTP_CODES = {408, 409, 425, 429, 500, 502, 503, 504}
MAILOPOST_REQUESTS_PER_SECOND = 2
MAILOPOST_MIN_REQUEST_INTERVAL_SECONDS = 1.0 / MAILOPOST_REQUESTS_PER_SECOND
MAILOPOST_RATE_LIMIT_EXTRA_SECONDS = 5.0
MAILOPOST_RATE_LIMIT_FALLBACK_SECONDS = 60.0
MAILOPOST_MAX_ATTACHMENT_BYTES = 5 * 1024 * 1024
_MAILOPOST_RATE_LIMIT_LOCK = threading.Lock()
_last_mailopost_request_at = 0.0
_EXCEL_STATS_CACHE_LOCK = threading.Lock()
_EXCEL_STATS_CACHE: dict[str, dict[str, Any]] = {}


class MailoPostRateLimitError(RuntimeError):
    def __init__(self, message: str, retry_after_seconds: float):
        super().__init__(message)
        self.retry_after_seconds = max(0.0, float(retry_after_seconds))


def _load_sender_state(job_id: str | None = None) -> dict[str, Any]:
    return load_agent_state("sender", SENDER_STATE, job_id)


def _save_sender_state(state: dict[str, Any], job_id: str | None = None) -> dict[str, Any]:
    return save_agent_state("sender", state, job_id)


def _state_rows_snapshot(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if len(rows) <= SENDER_STATE_ROWS_LIMIT:
        return list(rows)
    return list(rows[-SENDER_STATE_ROWS_LIMIT:])


def _flush_sender_workbook(workbook: Any, data_xlsx_path: Path) -> str:
    try:
        save_workbook(workbook, data_xlsx_path)
        return ""
    except Exception as exc:
        return f"Не удалось сохранить изменения в data.xlsx: {_safe_text(exc) or exc}"


def _should_flush_sender_workbook(*, dirty: bool, processed_rows: int, total_rows: int) -> bool:
    if not dirty:
        return False
    if processed_rows <= 0:
        return False
    if processed_rows >= total_rows:
        return True
    return processed_rows % SENDER_WORKBOOK_SAVE_EVERY == 0


def _file_signature(path: Path) -> tuple[int, int] | None:
    try:
        stat = path.stat()
    except OSError:
        return None
    return (int(stat.st_mtime_ns), int(stat.st_size))


def _sleep_sender_retry(delay_seconds: float) -> None:
    if delay_seconds <= 0:
        return
    sleep(delay_seconds)


def _is_retryable_unisender_exception(exc: Exception) -> bool:
    if isinstance(exc, HTTPError):
        return int(exc.code) in UNISENDER_RETRYABLE_HTTP_CODES
    return isinstance(exc, (TimeoutError, URLError, OSError))


def _is_unisender_rate_limit_text(value: Any) -> bool:
    text = _safe_text(value).lower()
    return "api call limit exceeded" in text or "limit 60 calls" in text


def _wait_unisender_api_slot() -> None:
    global _last_unisender_request_at
    with _UNISENDER_RATE_LIMIT_LOCK:
        now = perf_counter()
        wait_seconds = UNISENDER_MIN_REQUEST_INTERVAL_SECONDS - (now - _last_unisender_request_at)
        if wait_seconds > 0:
            sleep(wait_seconds)
        _last_unisender_request_at = perf_counter()


def _run_unisender_request(request: Request, *, timeout: float, request_label: str) -> str:
    last_error: Exception | None = None
    for attempt in range(1, UNISENDER_RETRY_ATTEMPTS + 1):
        try:
            _wait_unisender_api_slot()
            with urlopen(request, timeout=timeout) as response:
                return response.read().decode("utf-8", errors="replace")
        except HTTPError as exc:
            raw = exc.read().decode("utf-8", errors="replace")
            if attempt < UNISENDER_RETRY_ATTEMPTS and _is_unisender_rate_limit_text(raw):
                _sleep_sender_retry(UNISENDER_RATE_LIMIT_RETRY_SECONDS)
                last_error = RuntimeError(f"{request_label} временно ограничил API-запросы, повтор {attempt}.")
                continue
            if attempt < UNISENDER_RETRY_ATTEMPTS and _is_retryable_unisender_exception(exc):
                _sleep_sender_retry(UNISENDER_RETRY_BASE_SECONDS * attempt)
                last_error = RuntimeError(
                    f"{request_label} временно недоступен (HTTP {exc.code}), повтор {attempt}."
                )
                continue
            exc.raw_body = raw  # type: ignore[attr-defined]
            raise
        except Exception as exc:
            if attempt < UNISENDER_RETRY_ATTEMPTS and _is_unisender_rate_limit_text(exc):
                _sleep_sender_retry(UNISENDER_RATE_LIMIT_RETRY_SECONDS)
                last_error = exc
                continue
            if attempt < UNISENDER_RETRY_ATTEMPTS and _is_retryable_unisender_exception(exc):
                _sleep_sender_retry(UNISENDER_RETRY_BASE_SECONDS * attempt)
                last_error = exc
                continue
            raise
    if last_error is not None:
        raise last_error
    raise RuntimeError(f"{request_label} не ответил после повторных попыток.")


def _resolve_sender_data_xlsx_path(job_id: str | None = None) -> Path:
    job_paths = resolve_job_paths(job_id)
    if job_id:
        return job_paths.data_xlsx
    return job_paths.data_xlsx if job_paths.data_xlsx.exists() else DATA_XLSX_PATH


def _refresh_sender_stop_flag(state: dict[str, Any], job_id: str | None = None) -> dict[str, Any]:
    persisted = _load_sender_state(job_id)
    state["stop_requested"] = bool(persisted.get("stop_requested", False))
    state["stop_requested_at"] = persisted.get("stop_requested_at")
    return state


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


class _SafeMailTemplateValues(dict[str, str]):
    def __missing__(self, key: str) -> str:
        return ""


def _read_docx_mail_template(template_path: Path) -> str:
    document = Document(template_path)
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
        elif blocks and blocks[-1] != "":
            blocks.append("")

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append("\t".join(cells))

    return "\n".join(blocks).strip()


def _read_mail_template(mail_template_path: Path | None = None) -> str:
    if mail_template_path is None:
        return DEFAULT_MAIL_BODY
    template_paths = [mail_template_path]
    for template_path in [path for path in template_paths if path is not None]:
        if not template_path.exists():
            continue
        try:
            if template_path.suffix.lower() == ".docx":
                text = _read_docx_mail_template(template_path)
            else:
                text = template_path.read_text(encoding="utf-8-sig").strip()
            if text:
                return text
        except OSError:
            continue
    return DEFAULT_MAIL_BODY


def _mail_outgoing_number(row: dict[str, Any]) -> str:
    raw_id = _safe_text(row.get("ID"))
    try:
        row_number = int(float(raw_id))
    except (TypeError, ValueError):
        return raw_id
    return str(START_OUTGOING_NUMBER + row_number - 1)


def _mail_mun_name_genitive(mun_name: str, fallback: str = "") -> str:
    normalized = re.sub(r"\s+", " ", _safe_text(mun_name)).strip()
    fallback = _safe_text(fallback)
    if fallback:
        return fallback
    if normalized.startswith("Городское поселение "):
        return f"Городского поселения {normalized[len('Городское поселение ') :].strip()}".strip()
    if normalized.startswith("Сельское поселение "):
        return f"Сельского поселения {normalized[len('Сельское поселение ') :].strip()}".strip()
    return normalized


def _mail_template_values(row: dict[str, Any]) -> dict[str, str]:
    outgoing_number = _mail_outgoing_number(row)
    try:
        inflected, _ = build_inflected_fields_with_trace(row)
    except Exception:
        inflected = {}
    mun_name = _safe_text(row.get("MUN_NAME"))
    mun_name_genitive = (
        _safe_text(inflected.get("MUN_R_NAME_1"))
        or _safe_text(row.get("MUN_R_NAME_1"))
        or _safe_text(row.get("MUN_R_NAME"))
        or _mail_mun_name_genitive(mun_name, _safe_text(inflected.get("MUN_NAME_1")))
    )
    subject_name_genitive = (
        _safe_text(inflected.get("SUB_RF_1"))
        or _safe_text(row.get("SUB_RF_1"))
        or _safe_text(row.get("SUB_RF"))
    )
    return {
        "HEAD_FIO": _safe_text(row.get("HEAD_FIO")),
        "ADM_NAME": _safe_text(row.get("ADM_NAME")),
        "MUN_NAME": mun_name,
        "MUN_NAME_GENITIVE": mun_name_genitive,
        "MUN_R_NAME": mun_name_genitive,
        "SUB_RF": _safe_text(row.get("SUB_RF")),
        "SUB_RF_1": subject_name_genitive,
        "DATE": datetime.now().strftime("%d.%m.%Y"),
        "OUTGOING_NUMBER": outgoing_number,
        "OUTGOING_NUMBER_KP": f"{outgoing_number}-КП" if outgoing_number else "",
    }


def _render_mail_template(template: str, row: dict[str, Any]) -> str:
    values = _mail_template_values(row)
    bracket_replacements = {
        "наименование муниципального образования": values["MUN_NAME"],
        "наименование муниципального образования в родительном падеже": values["MUN_NAME_GENITIVE"],
        "муниципальное образование в родительном падеже": values["MUN_NAME_GENITIVE"],
        "мун р": values["MUN_NAME_GENITIVE"],
        "номер": values["OUTGOING_NUMBER"],
        "номер кп": values["OUTGOING_NUMBER_KP"],
        "дата": values["DATE"],
    }

    rendered = template
    # Existing mail templates used the nominative placeholder after "проектирования",
    # where Russian grammar requires the genitive form.
    rendered = re.sub(
        r"(градостроительного\s+проектирования\s*)\[\s*наименование\s+муниципального\s+образования\s*\]",
        rf"\1{values['MUN_R_NAME']}",
        rendered,
        flags=re.IGNORECASE,
    )
    for placeholder, value in bracket_replacements.items():
        rendered = re.sub(
            rf"\[\s*{re.escape(placeholder)}\s*\]",
            value,
            rendered,
            flags=re.IGNORECASE,
        )

    return rendered.format_map(_SafeMailTemplateValues(values))


def _build_mail_footer_html(*, inline_image: bool = False) -> str:
    return MAIL_FOOTER_HTML_TEMPLATE.format(image_src=_mail_footer_image_src(inline=inline_image))


def _mail_footer_html_markers() -> list[str]:
    lines = [line.strip() for line in DEFAULT_MAIL_FOOTER_TEXT.splitlines() if line.strip()]
    markers: list[str] = []
    if len(lines) >= 2:
        markers.append(f"{escape(lines[0], quote=False)}<br>{escape(lines[1], quote=False)}")
    markers.append("С уважением,<br>Черкашина Наталья Александровна")
    return markers

def _mail_footer_image_src(*, inline: bool = False) -> str:
    if inline:
        return f"cid:{MAIL_FOOTER_LOGO_CID}"
    signature_image_url = _safe_text(settings.mail_signature_image_url).rstrip("/")
    if signature_image_url:
        return _append_mail_footer_image_version(signature_image_url)
    public_base_url = _safe_text(settings.public_base_url).rstrip("/")
    if public_base_url:
        return _append_mail_footer_image_version(f"{public_base_url}/public/mail-signature.png")
    return ""


def _append_mail_footer_image_version(url: str) -> str:
    signature = _file_signature(MAIL_FOOTER_LOGO_PATH)
    if not signature:
        return url
    separator = "&" if "?" in url else "?"
    return f"{url}{separator}sig={signature[0]}-{signature[1]}"


def _append_mail_footer_text(body: str) -> str:
    body_text = body.rstrip()
    footer_text = DEFAULT_MAIL_FOOTER_TEXT.strip()
    body_text = re.sub(
        r"(?:\n{2,}|^)С уважением,\s*\nЧеркашина\s+Наталья\s+Александровна[\s\S]*?(?:https?://www\.parresh\.ru/?\s*)?$",
        "",
        body_text,
        flags=re.IGNORECASE,
    ).rstrip()
    if footer_text in body_text:
        return body_text
    body_text = re.sub(
        r"\n{2,}С уважением,\s*\nООО\s+«ПР»\s*$",
        "",
        body_text,
        flags=re.IGNORECASE,
    ).rstrip()
    return f"{body_text}\n\n{footer_text}" if body_text else footer_text


def _parse_emails(raw_value: Any) -> list[str]:
    raw_text = _safe_text(raw_value)
    if not raw_text:
        return []
    parts = re.split(r"[;,]", raw_text)
    result: list[str] = []
    for part in parts:
        email = part.strip()
        if email and email not in result:
            result.append(email)
    return result


def _is_valid_email(value: str) -> bool:
    return bool(EMAIL_PATTERN.match(value))


def _choose_recipient(row: dict[str, Any]) -> dict[str, Any]:
    primary_candidates = _parse_emails(row.get("EMAIL_OSN"))
    extra_candidates = _parse_emails(row.get("EMAIL_DOP"))
    all_candidates = primary_candidates + [item for item in extra_candidates if item not in primary_candidates]
    valid_emails = [item for item in all_candidates if _is_valid_email(item)]
    invalid_emails = [item for item in all_candidates if item and item not in valid_emails]

    strategy = "no_valid_email"
    recipient = None
    fallback_candidates: list[str] = []
    decision_reason = "В строке не найдено ни одного валидного email."

    valid_primary = [item for item in primary_candidates if _is_valid_email(item)]
    valid_extra = [item for item in extra_candidates if _is_valid_email(item)]

    if valid_primary:
        recipient = valid_primary[0]
        fallback_candidates = [item for item in valid_emails if item != recipient]
        strategy = "primary"
        if fallback_candidates:
            decision_reason = "Использую основной и дополнительные email из EMAIL_OSN/EMAIL_DOP."
        else:
            decision_reason = "Использую основной email из EMAIL_OSN."
    elif valid_extra:
        recipient = valid_extra[0]
        fallback_candidates = [item for item in valid_emails if item != recipient]
        strategy = "fallback_extra"
        if primary_candidates:
            decision_reason = "Основной email отсутствует или невалиден, использую EMAIL_DOP."
        else:
            decision_reason = "Основной email пустой, использую EMAIL_DOP."

    return {
        "recipient": recipient,
        "valid_emails": valid_emails,
        "invalid_emails": invalid_emails,
        "strategy": strategy,
        "fallback_candidates": fallback_candidates,
        "decision_reason": decision_reason,
        "needs_parser": recipient is None,
    }


def _normalize_recipient_strategy(value: str | None) -> str:
    cleaned = _safe_text(value).lower().replace("-", "_")
    if cleaned in {"primary", "primary_then_fallback", "main_then_backup", "main_then_fallback"}:
        return RECIPIENT_STRATEGY_PRIMARY_THEN_FALLBACK
    if cleaned in {"all", "all_recipients", "all_emails"}:
        return RECIPIENT_STRATEGY_ALL
    return DEFAULT_RECIPIENT_STRATEGY


def _allowed_send_recipients(
    email_decision: dict[str, Any],
    *,
    recipient_strategy: str | None = None,
) -> list[str]:
    strategy = _normalize_recipient_strategy(recipient_strategy)
    if strategy == RECIPIENT_STRATEGY_PRIMARY_THEN_FALLBACK:
        recipient = _safe_text(email_decision.get("recipient"))
        return [recipient] if recipient else []
    return list(email_decision.get("valid_emails") or [])


def _fallback_send_recipients(
    email_decision: dict[str, Any],
    *,
    recipient_strategy: str | None = None,
) -> list[str]:
    if _normalize_recipient_strategy(recipient_strategy) != RECIPIENT_STRATEGY_PRIMARY_THEN_FALLBACK:
        return []
    primary_key = _mail_key(email_decision.get("recipient"))
    return [
        recipient
        for recipient in list(email_decision.get("fallback_candidates") or [])
        if _mail_key(recipient) and _mail_key(recipient) != primary_key
    ]


def _unique_send_recipients(recipients: list[str]) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for recipient in recipients:
        key = _mail_key(recipient)
        if not key or key in seen:
            continue
        result.append(_safe_text(recipient))
        seen.add(key)
    return result


def _consent_candidate_recipients(
    email_decision: dict[str, Any],
    *,
    recipient_strategy: str | None = None,
) -> list[str]:
    return _unique_send_recipients(
        [
            *_allowed_send_recipients(email_decision, recipient_strategy=recipient_strategy),
            *_fallback_send_recipients(email_decision, recipient_strategy=recipient_strategy),
        ]
    )


def _combine_send_results(results: list[dict[str, Any]]) -> dict[str, Any]:
    attempts: list[dict[str, Any]] = []
    sent_recipients: list[str] = []
    warnings: list[str] = []
    errors: list[str] = []
    consent_urls: dict[str, str] = {}
    first_success = ""

    for result in results:
        attempts.extend(result.get("attempts") or [])
        for recipient in result.get("recipients") or []:
            if recipient not in sent_recipients:
                sent_recipients.append(recipient)
        if result.get("recipient") and not first_success:
            first_success = _safe_text(result.get("recipient"))
        if _safe_text(result.get("warning")):
            warnings.append(_safe_text(result.get("warning")))
        if _safe_text(result.get("error")):
            errors.append(_safe_text(result.get("error")))
        urls = result.get("consent_urls")
        if isinstance(urls, dict):
            consent_urls.update({str(key): str(value) for key, value in urls.items()})

    return {
        "recipient": first_success or None,
        "recipients": sent_recipients,
        "attempts": attempts,
        "error": "" if first_success else ("; ".join(errors) or "Не удалось отправить письмо."),
        "warning": " ".join(warnings).strip(),
        "consent_urls": consent_urls,
    }


def _send_recipient_sequence_until_success(
    recipients: list[str],
    send_one: Callable[[str], dict[str, Any]],
    *,
    wait_between_recipients: Callable[[], bool] | None = None,
) -> dict[str, Any]:
    results: list[dict[str, Any]] = []
    attempted_recipients = 0
    for recipient in recipients:
        if attempted_recipients > 0 and wait_between_recipients is not None:
            if not wait_between_recipients():
                results.append(
                    {
                        "recipient": None,
                        "recipients": [],
                        "attempts": [
                            {
                                "recipient": recipient,
                                "status": "error",
                                "error": "Отправка остановлена во время паузы между письмами.",
                            }
                        ],
                        "error": "Отправка остановлена во время паузы между письмами.",
                        "warning": "",
                    }
                )
                break
        attempted_recipients += 1
        result = send_one(recipient)
        results.append(result)
        if result.get("recipient"):
            break
    return _combine_send_results(results)


def _send_consent_requests_with_transport(
    row: dict[str, Any],
    recipients: list[str],
    subject: str,
    *,
    transport: str,
    mail_template_path: Path | None = None,
    job_id: str | None = None,
    send_run_id: str = "",
    attachment_mode: str = ATTACHMENT_MODE_KP,
    subject_template: str | None = None,
    work_type: str | None = None,
    recipient_strategy: str | None = None,
    wait_between_recipients: Callable[[], bool] | None = None,
    wait_after_rate_limit: Callable[[float, str], bool] | None = None,
) -> dict[str, Any]:
    attempts: list[dict[str, Any]] = []
    sent_recipients: list[str] = []
    warnings: list[str] = []
    consent_urls: dict[str, str] = {}

    attempted_recipients = 0
    for recipient in recipients:
        if attempted_recipients > 0 and wait_between_recipients is not None:
            if not wait_between_recipients():
                attempts.append(
                    {
                        "recipient": recipient,
                        "status": "error",
                        "error": "Отправка остановлена во время паузы между письмами.",
                    }
                )
                break
        attempted_recipients += 1
        try:
            consent_record = prepare_consent_request(
                job_id=job_id,
                row=row,
                recipient=recipient,
                transport=transport,
                attachment_mode=attachment_mode,
                subject_template=subject_template,
                work_type=work_type,
                recipient_strategy=recipient_strategy,
            )
            consent_url = _safe_text(consent_record.get("consent_url"))
            consent_urls[recipient] = consent_url
            body_override = _build_consent_request_body(
                row,
                consent_url=consent_url,
                attachment_mode=attachment_mode,
                work_type=work_type,
            )
            result = _send_with_transport(
                row,
                [recipient],
                [],
                subject,
                transport=transport,
                mail_template_path=mail_template_path,
                body_override=body_override,
                job_id=job_id,
                send_run_id=send_run_id,
                send_mode="consent_request",
                attachment_mode=attachment_mode,
                wait_after_rate_limit=wait_after_rate_limit,
            )
        except Exception as exc:
            attempts.append(
                {
                    "recipient": recipient,
                    "status": "error",
                    "error": _safe_text(exc) or f"{transport} error",
                }
            )
            continue

        attempts.extend(result.get("attempts") or [])
        if _safe_text(result.get("warning")):
            warnings.append(_safe_text(result.get("warning")))
        if result.get("recipient"):
            sent_recipients.extend(result.get("recipients") or [result["recipient"]])

    if sent_recipients and len(sent_recipients) == len(recipients):
        return {
            "recipient": sent_recipients[0],
            "recipients": sent_recipients,
            "attempts": attempts,
            "error": "",
            "warning": " ".join(warnings).strip(),
            "consent_urls": consent_urls,
        }

    failed_errors = [attempt["error"] for attempt in attempts if attempt.get("status") == "error" and attempt.get("error")]
    return {
        "recipient": None,
        "recipients": sent_recipients,
        "attempts": attempts,
        "error": "; ".join(failed_errors) or "Не удалось отправить запрос согласия.",
        "warning": " ".join(warnings).strip(),
        "consent_urls": consent_urls,
    }


def _build_output_folder_index(
    row_ids: set[str],
    *,
    output_dir: Path | None = None,
) -> tuple[dict[str, Path], dict[str, str]]:
    root_dir = output_dir or OUTPUT_DIR
    if not root_dir.exists():
        return {}, {row_id: f"Не найдена папка output для ID={row_id}." for row_id in row_ids}

    matches: dict[str, list[Path]] = {row_id: [] for row_id in row_ids}
    for path in root_dir.iterdir():
        if not path.is_dir() or "_" not in path.name:
            continue
        row_id = path.name.split("_", 1)[0]
        if row_id in matches:
            matches[row_id].append(path)

    folder_index: dict[str, Path] = {}
    folder_errors: dict[str, str] = {}
    for row_id in row_ids:
        row_matches = matches.get(row_id) or []
        if not row_matches:
            folder_errors[row_id] = f"Не найдена папка output для ID={row_id}."
        elif len(row_matches) > 1:
            folder_errors[row_id] = f"Найдено несколько папок output для ID={row_id}."
        else:
            folder_index[row_id] = row_matches[0]
    return folder_index, folder_errors


def _resolve_output_folder(
    row_id: Any,
    *,
    output_dir: Path | None = None,
    folder_index: dict[str, Path] | None = None,
    folder_errors: dict[str, str] | None = None,
) -> tuple[Path | None, str | None]:
    if row_id in (None, ""):
        return None, "Не указан ID строки."

    row_id_text = _safe_text(row_id)
    if folder_errors and row_id_text in folder_errors:
        return None, folder_errors[row_id_text]
    if folder_index is not None:
        folder = folder_index.get(row_id_text)
        if folder:
            return folder, None
        return None, f"Не найдена папка output для ID={row_id_text}."

    root_dir = output_dir or OUTPUT_DIR
    prefix = f"{row_id_text}_"
    matches = [path for path in root_dir.iterdir() if path.is_dir() and path.name.startswith(prefix)] if root_dir.exists() else []
    if not matches:
        return None, f"Не найдена папка output для ID={row_id_text}."
    if len(matches) > 1:
        return None, f"Найдено несколько папок output для ID={row_id_text}."
    return matches[0], None


def _normalize_attachment_mode(attachment_mode: str | None) -> str:
    value = _safe_text(attachment_mode).lower()
    return value if value in ATTACHMENT_MODE_VALUES else ATTACHMENT_MODE_KP


def _resolve_pdf_attachments(
    folder: Path | None,
    *,
    attachment_mode: str = ATTACHMENT_MODE_KP,
) -> tuple[list[str], str | None]:
    if folder is None:
        return [], "Папка документов не определена."

    mode = _normalize_attachment_mode(attachment_mode)
    pdf_files = sorted(folder.glob("*.pdf"))
    docx_files = sorted(folder.glob("*.docx"))
    kp_pdf = next(
        (
            path
            for path in pdf_files
            if "кп" in path.name.lower() or "коммер" in path.name.lower()
        ),
        None,
    )
    contract_docx = next(
        (path for path in docx_files if "договор" in path.name.lower()),
        None,
    )

    requested_paths: list[tuple[Path | None, str]] = []
    if mode in {ATTACHMENT_MODE_KP, ATTACHMENT_MODE_BOTH}:
        requested_paths.append((kp_pdf, "КП в PDF"))
    if mode in {ATTACHMENT_MODE_CONTRACT, ATTACHMENT_MODE_BOTH}:
        requested_paths.append((contract_docx, "договор в Word"))

    attachments: list[str] = []
    missing: list[str] = []
    for path, label in requested_paths:
        if not path:
            missing.append(label)
            continue
        try:
            if not path.exists():
                missing.append(f"{label} (файл исчез из папки)")
                continue
            if path.stat().st_size <= 0:
                missing.append(f"{label} (файл пустой)")
                continue
        except OSError as exc:
            missing.append(f"{label} (не удалось проверить файл: {_safe_text(exc)})")
            continue
        attachments.append(str(path))
    if not missing:
        return attachments, None
    return attachments, f"Перед отправкой не прошла проверка вложений. В папке {folder.name} проблема: {', '.join(missing)}."


def _status_class(raw_status: Any) -> str:
    status_text = _safe_text(raw_status).upper()
    if not status_text:
        return "pending"
    if "ЗАПРОС СОГЛАСИЯ" in status_text:
        return "consent_requested"
    if status_text in STATUS_OK_VALUES:
        return "sent"
    if "ОШИБ" in status_text or "ERROR" in status_text:
        return "error"
    return "other"


def _format_error_status(error: Any) -> str:
    error_text = re.sub(r"\s+", " ", _safe_text(error)).strip()
    if not error_text:
        return STATUS_ERROR_VALUE
    if len(error_text) > MAX_STATUS_ERROR_LENGTH:
        error_text = error_text[: MAX_STATUS_ERROR_LENGTH - 3].rstrip() + "..."
    return f"{STATUS_ERROR_VALUE}: {error_text}"


def _persist_row_status(
    workbook: Any,
    worksheet: Any,
    data_xlsx_path: Path,
    row: dict[str, Any],
    status_value: str,
    *,
    flush: bool = True,
) -> str:
    try:
        update_status(worksheet, row["_row_index"], status_value)
        row["STATUS"] = status_value
        if flush:
            save_workbook(workbook, data_xlsx_path)
        return ""
    except Exception as exc:
        return f"Не удалось сохранить статус строки в data.xlsx: {_safe_text(exc) or exc}"


def _wait_sender_delay(delay_seconds: float, state: dict[str, Any], job_id: str | None = None) -> bool:
    deadline = perf_counter() + max(0.0, delay_seconds)
    while perf_counter() < deadline:
        _refresh_sender_stop_flag(state, job_id)
        if state.get("stop_requested"):
            return False
        sleep(min(1.0, max(0.0, deadline - perf_counter())))
    return True


def _collect_excel_stats(data_xlsx_path: Path | None = None) -> dict[str, int]:
    source_path = data_xlsx_path or DATA_XLSX_PATH
    if not source_path.exists():
        return {"total": 0, "sent": 0, "error": 0, "pending": 0}

    signature = _file_signature(source_path)
    cache_key = str(source_path.resolve())
    if signature is not None:
        with _EXCEL_STATS_CACHE_LOCK:
            cached = _EXCEL_STATS_CACHE.get(cache_key)
            if cached and cached.get("signature") == signature:
                return dict(cached.get("stats") or {"total": 0, "sent": 0, "error": 0, "pending": 0})

    workbook, _, rows = load_rows(source_path)
    stats = {"total": len(rows), "sent": 0, "error": 0, "pending": 0}
    for row in rows:
        status_class = _status_class(row.get("STATUS"))
        if status_class == "sent":
            stats["sent"] += 1
        elif status_class == "error":
            stats["error"] += 1
        else:
            stats["pending"] += 1
    close = getattr(workbook, "close", None)
    if callable(close):
        close()
    if signature is not None:
        with _EXCEL_STATS_CACHE_LOCK:
            _EXCEL_STATS_CACHE[cache_key] = {"signature": signature, "stats": dict(stats)}
    return stats


def _format_sender_summary(state: dict[str, Any]) -> str:
    def _plural_rows(value: Any) -> str:
        count = int(value or 0)
        return f"{count} строк"

    def _error_type_summary() -> str:
        rows = state.get("rows") or []
        if not isinstance(rows, list):
            return ""
        counters = {
            "email": 0,
            "documents": 0,
            "attachments": 0,
            "text_review": 0,
            "provider": 0,
            "other": 0,
        }
        samples: dict[str, str] = {}
        details: dict[str, str] = {}
        for row in rows:
            if not isinstance(row, dict):
                continue
            result = _safe_text(row.get("result"))
            error = _safe_text(row.get("error"))
            if not error and result not in {
                "needs_enrichment",
                "error_missing_output",
                "error_missing_attachments",
                "blocked_by_philologist",
                "error",
            }:
                continue
            if result == "needs_enrichment" or "email" in error.lower():
                key = "email"
                human = "у части строк нет корректного email получателя"
            elif result == "error_missing_output":
                key = "documents"
                human = "для части строк не найдена папка с готовыми документами"
            elif result == "error_missing_attachments":
                key = "attachments"
                human = "для части строк не найдены нужные вложения"
            elif result == "blocked_by_philologist":
                key = "text_review"
                human = "часть документов требует ручной проверки текста"
            elif "unisender" in error.lower() or "smtp" in error.lower() or row.get("provider_status"):
                key = "provider"
                human = "провайдер отправки вернул ошибку"
            else:
                key = "other"
                human = "часть строк требует ручной проверки"
            counters[key] += 1
            samples.setdefault(key, human)
            if error and key not in details:
                row_id = _safe_text(row.get("id")) or "без ID"
                details[key] = f"пример: ID {row_id}: {error}"

        parts = []
        for key in ("email", "documents", "attachments", "text_review", "provider", "other"):
            if counters[key] > 0:
                detail = f" ({details[key]})." if key in details else "."
                parts.append(f"{_plural_rows(counters[key])}: {samples[key]}{detail}")
        return " ".join(parts)

    if state.get("status") == "stopped":
        summary = (
            f"Отправка остановлена. Уже отправлено {state.get('sent_rows', 0)}, "
            f"готово к отправке {state.get('ready_rows', 0)}, ошибок {state.get('error_rows', 0)}."
        )
        if state.get("warning_rows", 0) > 0:
            summary += (
                f" У {state.get('warning_rows', 0)} писем отправка прошла, "
                "но копию не удалось сохранить в папку «Отправленные»."
            )
        if state.get("remaining_rows", 0) > 0:
            summary += f" Осталось отправить: {state.get('remaining_rows', 0)}."
        return summary
    if state.get("total_rows", 0) == 0:
        return "В таблице пока нет получателей для отправки."
    is_consent_mode = _safe_text(state.get("send_mode")) == "consent_request"
    if state.get("mode") == "dry_run":
        if int(state.get("error_rows") or 0) <= 0:
            if is_consent_mode:
                summary = f"Проверка завершена: запросы согласия готовы. Готово писем: {state.get('ready_rows', 0)}."
            else:
                summary = (
                    "Проверка завершена: всё в порядке, письма готовы к отправке. "
                    f"Готово писем: {state.get('ready_rows', 0)}."
                )
        else:
            summary = (
                "Проверка завершена, но не все письма готовы. "
                f"Готово к отправке: {state.get('ready_rows', 0)}. "
                f"Проблемных строк: {state.get('error_rows', 0)}."
            )
    else:
        if int(state.get("error_rows") or 0) <= 0 and int(state.get("remaining_rows") or 0) <= 0:
            if is_consent_mode:
                summary = f"Запросы согласия отправлены. Отправлено: {state.get('sent_rows', 0)}."
            else:
                summary = f"Отправка завершена: все письма ушли. Отправлено: {state.get('sent_rows', 0)}."
        else:
            summary = (
                "Отправка завершена не полностью. "
                f"Отправлено: {state.get('sent_rows', 0)}. "
                f"Не отправлено: {state.get('error_rows', 0)}. "
                f"Осталось в очереди: {state.get('remaining_rows', 0)}."
            )
    error_summary = _error_type_summary()
    if error_summary:
        summary = f"{summary} Что нужно проверить: {error_summary}"
    if state.get("generator_handoff_rows", 0) > 0:
        summary += f" Нужно заново подготовить документы: {state.get('generator_handoff_rows', 0)}."
    if state.get("philology_blocked_rows", 0) > 0:
        summary += f" Нужно проверить текст перед отправкой: {state.get('philology_blocked_rows', 0)}."
    if state.get("autonomous_recovery_rows", 0) > 0:
        summary += f" Автоматически восстановлено комплектов: {state.get('autonomous_recovery_rows', 0)}."
    if state.get("warning_rows", 0) > 0:
        summary += (
            f" У {state.get('warning_rows', 0)} писем отправка прошла, "
            "но копию не удалось сохранить в папку «Отправленные»."
        )
    if state.get("remaining_rows", 0) > 0:
        summary += f" Можно отправить оставшиеся письма: {state.get('remaining_rows', 0)}."
    return summary


def _normalize_limit(limit: int | None, *, dry_run: bool) -> int | None:
    if limit in (None, 0):
        return None

    normalized = max(1, int(limit))
    return min(normalized, max(1, int(settings.sender_max_batch_size)))


def _normalize_transport(transport: str | None) -> str:
    value = _safe_text(transport).lower()
    if value in {"unisender", "smtp", "rusender", "mailopost"}:
        return value
    configured = _safe_text(settings.sender_transport).lower()
    return configured if configured in {"unisender", "smtp", "rusender", "mailopost"} else "smtp"


def _normalize_send_mode(send_mode: str | None) -> str:
    value = _safe_text(send_mode).lower()
    if value in {"consent_request", "materials"}:
        return value
    return "materials"


def request_sender_stop(job_id: str | None = None) -> dict[str, Any]:
    state = _load_sender_state(job_id)
    state["stop_requested"] = True
    state["stop_requested_at"] = datetime.now().isoformat(timespec="seconds")
    if state.get("status") == "running":
        state["summary_text"] = (
            "Получен запрос на остановку. Отправщик завершит текущую строку и больше не будет брать новые."
        )
    _save_sender_state(state, job_id)
    return get_sender_status(job_id)


def clear_sender_stop_request(job_id: str | None = None) -> None:
    state = _load_sender_state(job_id)
    state["stop_requested"] = False
    state["stop_requested_at"] = None
    _save_sender_state(state, job_id)


def _active_sender_review_task(row_id: Any, *, job_id: str | None = None) -> dict[str, Any] | None:
    if not settings.inter_agent_handoffs_enabled:
        return None
    row_id_text = _safe_text(row_id)
    for item in get_tasks_for_agent("sender", job_id):
        if _safe_text(item.get("task_type")) != "review_before_send":
            continue
        if _safe_text(item.get("row_id")) != row_id_text:
            continue
        if _safe_text(item.get("status")) not in {"pending", "in_progress"}:
            continue
        return item
    return None


def _retry_row_resources(
    row_id: Any,
    *,
    output_dir: Path | None = None,
    attachment_mode: str = ATTACHMENT_MODE_KP,
) -> tuple[Path | None, str | None, list[str], str | None]:
    folder, folder_error = _resolve_output_folder(row_id, output_dir=output_dir)
    attachments, attachment_error = _resolve_pdf_attachments(folder, attachment_mode=attachment_mode)
    return folder, folder_error, attachments, attachment_error


def _delegate_sender_problem(
    *,
    symptom: str,
    row_id: Any,
    mun_name: str,
    details: dict[str, Any],
    job_id: str | None = None,
) -> dict[str, Any]:
    if not settings.inter_agent_handoffs_enabled:
        return {}
    diagnosis = diagnose_responsibility(symptom=symptom, context=details)
    task_type = {
        "missing_output": "generate_missing_documents",
        "missing_attachments": "restore_missing_attachments",
        "recipient_data_missing": "enrich_email",
        "review_before_send": "review_before_send",
    }.get(diagnosis["problem_type"], diagnosis["problem_type"])
    return create_task(
        source_agent="sender",
        target_agent=diagnosis["owner_agent"],
        owner_agent=diagnosis["owner_agent"],
        task_type=task_type,
        problem_type=diagnosis["problem_type"],
        symptom=symptom,
        root_cause=diagnosis["root_cause"],
        priority=diagnosis["priority"],
        blocking=diagnosis["blocking"],
        can_retry_after=diagnosis["can_retry_after"],
        row_id=row_id,
        mun_name=mun_name,
        details=details,
        job_id=job_id,
    )


def _run_autonomous_recovery_for_generator(
    *,
    row_id: Any,
    work_type: str | None = None,
    job_id: str | None = None,
) -> dict[str, Any]:
    row_id_text = _safe_text(row_id)
    generator_result = run_generator_agent(
        row_ids=[row_id_text] if row_id_text else None,
        work_type=work_type,
        job_id=job_id,
    )
    result = {
        "generator_result": generator_result,
    }
    if settings.philologist_auto_run_enabled:
        result["philologist_result"] = run_philologist(
            ai_enabled=True,
            row_ids=[row_id_text] if row_id_text else None,
            job_id=job_id,
        )
    return result


def _run_autonomous_recovery_for_philologist(*, row_id: Any, job_id: str | None = None) -> dict[str, Any]:
    if not settings.philologist_auto_run_enabled:
        return {"philologist_result": {"status": "skipped", "summary_text": "Автозапуск филолога отключён."}}
    row_id_text = _safe_text(row_id)
    philologist_result = run_philologist(ai_enabled=True, row_ids=[row_id_text] if row_id_text else None, job_id=job_id)
    return {"philologist_result": philologist_result}


def preview_recipients(*, limit: int | None = None, job_id: str | None = None) -> dict[str, Any]:
    data_xlsx_path = _resolve_sender_data_xlsx_path(job_id)
    if not data_xlsx_path.exists():
        return {
            "status": "error",
            "summary_text": "Файл data.xlsx не найден.",
            "rows": [],
            "total_rows": 0,
        }

    workbook, _, rows = load_rows(data_xlsx_path)
    close = getattr(workbook, "close", None)
    if callable(close):
        close()
    effective_limit = _normalize_limit(limit, dry_run=True)
    candidates = rows[:effective_limit] if effective_limit else rows
    preview_rows: list[dict[str, Any]] = []
    missing_count = 0
    fallback_count = 0
    invalid_count = 0
    extra_rows_count = 0
    extra_addresses_count = 0

    for row in candidates:
        email_decision = _choose_recipient(row)
        status_class = _status_class(row.get("STATUS"))
        primary_emails = _parse_emails(row.get("EMAIL_OSN"))
        extra_emails = _parse_emails(row.get("EMAIL_DOP"))
        entry = {
            "id": row.get("ID"),
            "mun_name": _safe_text(row.get("MUN_NAME")),
            "status_before": _safe_text(row.get("STATUS")),
            "status_class": status_class,
            "recipient": email_decision["recipient"],
            "email_strategy": email_decision["strategy"],
            "decision_reason": email_decision["decision_reason"],
            "primary_emails": primary_emails,
            "extra_emails": extra_emails,
            "invalid_emails": email_decision["invalid_emails"],
            "fallback_candidates": email_decision["fallback_candidates"],
        }
        if not entry["recipient"]:
            missing_count += 1
        if entry["email_strategy"] == "fallback_extra":
            fallback_count += 1
        if extra_emails:
            extra_rows_count += 1
            extra_addresses_count += len(extra_emails)
        if entry["invalid_emails"]:
            invalid_count += 1
        preview_rows.append(entry)

    summary_text = (
        f"Предпросмотр адресов: строк {len(preview_rows)}, "
        f"без адреса {missing_count}, "
        f"с дополнительными адресами {extra_rows_count}, "
        f"дополнительных адресов всего {extra_addresses_count}, "
        f"где EMAIL_DOP стал основным {fallback_count}, "
        f"с невалидными значениями {invalid_count}."
    )
    return {
        "status": "ok",
        "summary_text": summary_text,
        "rows": preview_rows,
        "total_rows": len(preview_rows),
        "effective_limit": effective_limit,
        "missing_count": missing_count,
        "fallback_count": fallback_count,
        "extra_rows_count": extra_rows_count,
        "extra_addresses_count": extra_addresses_count,
        "invalid_count": invalid_count,
    }


def _format_preview_rows(rows: list[dict[str, Any]], *, limit: int = 5) -> str:
    if not rows:
        return "Подходящих строк для показа не нашлось."
    lines: list[str] = []
    for item in rows[:limit]:
        recipient = item.get("recipient") or "нет валидного адреса"
        strategy = item.get("email_strategy") or "unknown"
        lines.append(
            f"{item.get('id')} {item.get('mun_name')}: {recipient} "
            f"({item.get('decision_reason')}; стратегия: {strategy})"
        )
    return "\n".join(lines)


def _build_message(
    row: dict[str, Any],
    recipient: str,
    attachments: list[str],
    subject: str,
    *,
    mail_template_path: Path | None = None,
    body_override: str | None = None,
) -> EmailMessage:
    body = body_override if body_override is not None else _build_mail_body(row, mail_template_path=mail_template_path)
    message = EmailMessage()
    message["From"] = settings.smtp_sender_email
    message["To"] = recipient
    message["Subject"] = subject
    message.set_content(body)
    message.add_alternative(
        _htmlify_mail_body(body, inline_footer_image=True, include_unsubscribe=False),
        subtype="html",
    )
    if MAIL_FOOTER_LOGO_PATH.exists():
        html_part = message.get_payload()[-1]
        html_part.add_related(
            MAIL_FOOTER_LOGO_PATH.read_bytes(),
            maintype="image",
            subtype="png",
            cid=f"<{MAIL_FOOTER_LOGO_CID}>",
            filename=MAIL_FOOTER_LOGO_PATH.name,
        )

    for attachment_path in attachments:
        path = Path(attachment_path)
        content_type, _ = mimetypes.guess_type(str(path))
        maintype, subtype = (content_type or "application/octet-stream").split("/", 1)
        message.add_attachment(
            path.read_bytes(),
            maintype=maintype,
            subtype=subtype,
            filename=path.name,
        )
    return message


def _build_mail_body(row: dict[str, Any], *, mail_template_path: Path | None = None) -> str:
    body = _render_mail_template(_read_mail_template(mail_template_path), row)
    return _append_mail_footer_text(body.strip())


def _build_mail_subject(subject_template: str, row: dict[str, Any]) -> str:
    return _render_mail_template(_safe_text(subject_template) or DEFAULT_MAIL_SUBJECT, row).strip()


def _materials_subject(attachment_mode: str, work_type: str | None = None) -> str:
    profile = get_work_type_profile(work_type)
    mode = _normalize_attachment_mode(attachment_mode)
    if mode == ATTACHMENT_MODE_CONTRACT:
        return f"Проект договора: {profile.service_title_nominative}."
    return profile.mail_subject


def _consent_request_subject(attachment_mode: str, work_type: str | None = None) -> str:
    profile = get_work_type_profile(work_type)
    mode = _normalize_attachment_mode(attachment_mode)
    if mode == ATTACHMENT_MODE_CONTRACT:
        return profile.consent_subject_contract
    if mode == ATTACHMENT_MODE_BOTH:
        return profile.consent_subject_both
    return profile.consent_subject_kp


def _consent_request_material_text(attachment_mode: str, work_type: str | None = None) -> tuple[str, str, str]:
    profile = get_work_type_profile(work_type)
    mode = _normalize_attachment_mode(attachment_mode)
    if mode == ATTACHMENT_MODE_CONTRACT:
        return (
            "проект договора",
            "проект договора",
            profile.consent_button_contract,
        )
    if mode == ATTACHMENT_MODE_BOTH:
        return (
            "коммерческое предложение и проект договора",
            "полный пакет документов: описание, условия, проект договора, техническое задание, календарный план",
            profile.consent_button_both,
        )
    return (
        "коммерческое предложение",
        "коммерческое предложение с описанием состава работ и условий",
        profile.consent_button_kp,
    )


def _consent_request_action_text(attachment_mode: str) -> tuple[str, str]:
    mode = _normalize_attachment_mode(attachment_mode)
    if mode == ATTACHMENT_MODE_CONTRACT:
        return (
            "Чтобы получить проект договора, просто кликните:",
            "После этого мы отправим вам файл отдельным письмом.",
        )
    if mode == ATTACHMENT_MODE_BOTH:
        return (
            "Чтобы получить документы, просто кликните:",
            "После этого мы отправим вам файлы отдельным письмом.",
        )
    return (
        "Чтобы получить КП, просто кликните:",
        "После этого мы отправим вам файл отдельным письмом.",
    )


def _build_consent_request_body(
    row: dict[str, Any],
    *,
    consent_url: str,
    attachment_mode: str = ATTACHMENT_MODE_KP,
    work_type: str | None = None,
) -> str:
    profile = get_work_type_profile(work_type)
    values = _mail_template_values(row)
    mun_name = _safe_text(values.get("MUN_R_NAME")) or _safe_text(values.get("MUN_NAME"))
    subject_name = _safe_text(values.get("SUB_RF_1")) or _safe_text(values.get("SUB_RF"))
    prepared_materials, package_text, button_text = _consent_request_material_text(attachment_mode, work_type)
    action_text, dispatch_text = _consent_request_action_text(attachment_mode)
    object_text = (
        f"для {mun_name} ({subject_name})"
        if mun_name and subject_name
        else f"для {mun_name}"
        if mun_name
        else f"для муниципального образования ({subject_name})"
        if subject_name
        else "для муниципального образования"
    )
    return _append_mail_footer_text(
        "\n".join(
            [
                "Здравствуйте!",
                "",
                f"ООО «Параллельные Решения» уже подготовило {object_text} {prepared_materials} {profile.consent_prepared_phrase}.",
                "",
                f"Если это направление вам актуально, мы можем направить вам {package_text}.",
                action_text,
                "",
                f"{button_text} {consent_url}",
                "",
                dispatch_text,
                "Если тема неактуальна — просто удалите или проигнорируйте это сообщение.",
                "",
                "Вы получили это письмо, так как ваш контакт указан в открытых источниках информации о муниципальных образованиях.",
            ]
        )
    )


def _append_sent_mail_log(
    *,
    row: dict[str, Any],
    recipient: str,
    attachments: list[str],
    subject: str,
    transport: str,
    warning: str = "",
    provider: dict[str, Any] | None = None,
    sent_mail_log_path: Path | None = None,
    send_run_id: str = "",
    send_run_started_at: str = "",
    send_mode: str = "",
    attachment_mode: str = "",
    subject_template: str = "",
    work_type: str = "",
    recipient_strategy: str = "",
) -> str | None:
    safe_provider = _safe_provider_payload(provider)
    record = {
        "sent_at": datetime.now().isoformat(timespec="seconds"),
        "transport": _safe_text(transport) or "smtp",
        "row_id": _safe_text(row.get("ID")),
        "mun_name": _safe_text(row.get("MUN_NAME")),
        "recipient": _safe_text(recipient),
        "subject": _safe_text(subject),
        "attachments": [Path(path).name for path in attachments if _safe_text(path)],
        "attachment_paths": [str(Path(path)) for path in attachments if _safe_text(path)],
        "warning": _safe_text(warning),
        "send_mode": _safe_text(send_mode),
        "attachment_mode": _safe_text(attachment_mode),
        "subject_template": _safe_text(subject_template),
        "work_type": _safe_text(work_type),
        "recipient_strategy": _safe_text(recipient_strategy) or DEFAULT_RECIPIENT_STRATEGY,
    }
    if _safe_text(send_run_id):
        record["send_run_id"] = _safe_text(send_run_id)
    if _safe_text(send_run_started_at):
        record["send_run_started_at"] = _safe_text(send_run_started_at)
    if safe_provider:
        record["provider"] = safe_provider
        if safe_provider.get("message_id"):
            record["provider_message_id"] = safe_provider["message_id"]
        if safe_provider.get("job_id"):
            record["provider_job_id"] = safe_provider["job_id"]
        provider_idempotency_key = _safe_text(
            safe_provider.get("idempotency_key") or safe_provider.get("idempotence_key")
        )
        if provider_idempotency_key:
            record["provider_idempotency_key"] = provider_idempotency_key
    try:
        log_path = sent_mail_log_path or SENT_MAIL_LOG_PATH
        log_path.parent.mkdir(parents=True, exist_ok=True)
        with log_path.open("a", encoding="utf-8") as handle:
            handle.write(json.dumps(record, ensure_ascii=False) + "\n")
    except OSError as exc:
        return (
            "Письмо отправлено, но не удалось записать его в локальный журнал: "
            f"{_safe_text(exc) or 'ошибка записи файла'}."
        )
    return None


def _mail_key(value: Any) -> str:
    return _safe_text(value).lower()


def _build_provider_idempotency_key(
    *,
    provider: str,
    job_id: str | None,
    send_run_id: str = "",
    row_id: Any = "",
    recipient: Any = "",
    send_mode: str = "",
    attachment_mode: str = "",
) -> str:
    payload = {
        "provider": _safe_text(provider).lower(),
        "job_id": _safe_text(job_id) or "__legacy__",
        "send_run_id": _safe_text(send_run_id) or "__no_send_run__",
        "row_id": _safe_text(row_id) or "__no_row__",
        "recipient": _mail_key(recipient) or "__no_recipient__",
        "send_mode": _safe_text(send_mode).lower() or "materials",
        "attachment_mode": _safe_text(attachment_mode).lower() or ATTACHMENT_MODE_KP,
    }
    raw = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    digest = hashlib.sha256(raw.encode("utf-8")).hexdigest()[:40]
    return f"mailing-agent:{payload['provider']}:{digest}"


def _load_sent_mail_recipients(
    sent_mail_log_path: Path | None = None,
    *,
    send_run_id: str = "",
    send_run_started_at: str = "",
) -> dict[str, set[str]]:
    log_path = sent_mail_log_path or SENT_MAIL_LOG_PATH
    sent_by_row: dict[str, set[str]] = {}
    if not log_path.exists():
        return sent_by_row

    try:
        lines = log_path.read_text(encoding="utf-8").splitlines()
    except OSError:
        return sent_by_row

    for line in lines:
        try:
            item = json.loads(line)
        except json.JSONDecodeError:
            continue
        if not _sent_log_item_in_send_scope(
            item,
            send_run_id=send_run_id,
            send_run_started_at=send_run_started_at,
        ):
            continue
        row_id = _safe_text(item.get("row_id"))
        recipient = _mail_key(item.get("recipient"))
        if not row_id or not recipient:
            continue
        sent_by_row.setdefault(row_id, set()).add(recipient)
    return sent_by_row


def _sent_log_item_in_send_scope(
    item: dict[str, Any],
    *,
    send_run_id: str = "",
    send_run_started_at: str = "",
) -> bool:
    expected_run_id = _safe_text(send_run_id)
    if expected_run_id:
        item_run_id = _safe_text(item.get("send_run_id"))
        if item_run_id:
            return item_run_id == expected_run_id
    started_at = _safe_text(send_run_started_at)
    sent_at = _safe_text(item.get("sent_at"))
    if started_at and sent_at:
        return sent_at >= started_at
    return True


def _load_sent_mail_log_items(sent_mail_log_path: Path | None = None) -> list[dict[str, Any]]:
    log_path = sent_mail_log_path or SENT_MAIL_LOG_PATH
    if not log_path.exists():
        return []

    items: list[dict[str, Any]] = []
    try:
        lines = log_path.read_text(encoding="utf-8").splitlines()
    except OSError:
        return []

    for line in lines:
        try:
            item = json.loads(line)
        except json.JSONDecodeError:
            continue
        if isinstance(item, dict):
            items.append(item)
    return items


def schedule_delivery_fallback_check(job_ids: Any, *, provider: str = "") -> None:
    if isinstance(job_ids, str):
        normalized_job_ids = [job_ids]
    else:
        try:
            normalized_job_ids = list(job_ids or [])
        except TypeError:
            normalized_job_ids = []
    provider_key = _safe_text(provider).lower() or "provider"
    for raw_job_id in normalized_job_ids:
        job_id = _safe_text(raw_job_id)
        if not job_id:
            continue
        key = f"{provider_key}:{job_id}"
        with _DELIVERY_FALLBACK_LOCK:
            if key in _DELIVERY_FALLBACK_RUNNING:
                continue
            _DELIVERY_FALLBACK_RUNNING.add(key)
        thread = threading.Thread(
            target=_run_scheduled_delivery_fallback_check,
            args=(job_id, provider_key, key),
            name=f"delivery-fallback-{job_id[:12]}",
            daemon=True,
        )
        thread.start()


def _run_scheduled_delivery_fallback_check(job_id: str, provider: str, running_key: str) -> None:
    try:
        for _ in range(30):
            state = _load_sender_state(job_id)
            if _safe_text(state.get("status")) != "running":
                break
            sleep(10)
        process_delivery_fallbacks(job_id=job_id, provider=provider)
    finally:
        with _DELIVERY_FALLBACK_LOCK:
            _DELIVERY_FALLBACK_RUNNING.discard(running_key)


def process_delivery_fallbacks(*, job_id: str, provider: str = "") -> dict[str, Any]:
    state = _load_sender_state(job_id)
    if _safe_text(state.get("status")) == "running":
        return {"status": "skipped_running", "job_id": job_id, "dispatched_rows": []}

    job_paths = resolve_job_paths(job_id)
    sent_mail_log_path = None if job_paths.uses_legacy_layout else job_paths.sent_mail_log_path
    sent_items = _load_sent_mail_log_items(sent_mail_log_path)
    if not sent_items:
        return {"status": "no_sent_log", "job_id": job_id, "dispatched_rows": []}

    rows_by_id = _load_sender_rows_by_id(job_id)
    latest_events = _latest_delivery_events_by_row_recipient(job_id, sent_items, provider=provider)
    sent_by_row = _sent_items_by_row(sent_items)
    dispatch_groups: dict[tuple[str, str, str, str, str], set[str]] = {}
    dispatch_rows: list[dict[str, str]] = []

    for row_id, row_items in sent_by_row.items():
        strategy_items = [
            item
            for item in row_items
            if _normalize_recipient_strategy(item.get("recipient_strategy"))
            == RECIPIENT_STRATEGY_PRIMARY_THEN_FALLBACK
        ]
        if not strategy_items:
            continue
        last_item = max(strategy_items, key=_sent_item_order_key)
        failed_recipient_key = _mail_key(last_item.get("recipient"))
        if not failed_recipient_key:
            continue
        event = latest_events.get((row_id, failed_recipient_key))
        if not event or not _is_delivery_failure_status(event.get("provider_status") or event.get("event_type")):
            continue
        row = rows_by_id.get(row_id)
        if row is None:
            continue
        email_decision = _choose_recipient(row)
        fallback_recipients = _fallback_send_recipients(
            email_decision,
            recipient_strategy=RECIPIENT_STRATEGY_PRIMARY_THEN_FALLBACK,
        )
        logged_recipients = {_mail_key(item.get("recipient")) for item in row_items if _mail_key(item.get("recipient"))}
        next_fallbacks = [recipient for recipient in fallback_recipients if _mail_key(recipient) not in logged_recipients]
        if not next_fallbacks:
            continue
        transport = _safe_text(last_item.get("transport")) or _safe_text(provider) or _safe_text(state.get("transport"))
        send_mode = _safe_text(last_item.get("send_mode")) or _safe_text(state.get("send_mode")) or "consent_request"
        attachment_mode = _safe_text(last_item.get("attachment_mode")) or _safe_text(state.get("attachment_mode")) or ATTACHMENT_MODE_KP
        subject_template = _safe_text(last_item.get("subject_template")) or _safe_text(state.get("subject_template"))
        work_type = _safe_text(last_item.get("work_type")) or _safe_text(state.get("work_type")) or DEFAULT_WORK_TYPE
        group_key = (transport, send_mode, attachment_mode, subject_template, work_type)
        dispatch_groups.setdefault(group_key, set()).add(row_id)
        dispatch_rows.append(
            {
                "row_id": row_id,
                "failed_recipient": _safe_text(last_item.get("recipient")),
                "next_recipient": next_fallbacks[0],
                "provider_status": _safe_text(event.get("provider_status") or event.get("event_type")),
            }
        )

    if not dispatch_groups:
        return {"status": "no_fallback_needed", "job_id": job_id, "dispatched_rows": []}

    results: list[dict[str, Any]] = []
    for (transport, send_mode, attachment_mode, subject_template, work_type), row_ids in dispatch_groups.items():
        result = run_sender(
            dry_run=False,
            row_ids=sorted(row_ids, key=_sort_row_id_text),
            transport=transport,
            send_mode=send_mode,
            attachment_mode=attachment_mode,
            subject_template=subject_template,
            require_confirmed_consent=send_mode == "materials",
            work_type=work_type,
            recipient_strategy=RECIPIENT_STRATEGY_PRIMARY_THEN_FALLBACK,
            job_id=job_id,
        )
        results.append(
            {
                "transport": transport,
                "send_mode": send_mode,
                "attachment_mode": attachment_mode,
                "row_ids": sorted(row_ids, key=_sort_row_id_text),
                "status": result.get("status"),
                "summary_text": result.get("summary_text"),
            }
        )
    return {"status": "dispatched", "job_id": job_id, "dispatched_rows": dispatch_rows, "results": results}


def _load_sender_rows_by_id(job_id: str) -> dict[str, dict[str, Any]]:
    data_xlsx_path = _resolve_sender_data_xlsx_path(job_id)
    if not data_xlsx_path.exists():
        return {}
    workbook = None
    try:
        workbook, _, rows = load_rows(data_xlsx_path)
    except Exception:
        return {}
    finally:
        close = getattr(workbook, "close", None)
        if callable(close):
            close()
    return {_safe_text(row.get("ID")): row for row in rows if _safe_text(row.get("ID"))}


def _sent_items_by_row(sent_items: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    grouped: dict[str, list[dict[str, Any]]] = {}
    for item in sent_items:
        row_id = _safe_text(item.get("row_id"))
        if row_id:
            grouped.setdefault(row_id, []).append(item)
    return grouped


def _sent_item_order_key(item: dict[str, Any]) -> tuple[str, str]:
    return (_safe_text(item.get("sent_at")), _safe_text(item.get("recipient")).lower())


def _sort_row_id_text(value: Any) -> tuple[int, str]:
    text = _safe_text(value)
    try:
        return (int(float(text)), text)
    except ValueError:
        return (10**9, text)


def _is_delivery_failure_status(status: Any) -> bool:
    normalized = _safe_text(status).lower()
    if not normalized:
        return False
    if normalized in DELIVERY_FALLBACK_FAILURE_STATUSES:
        return True
    return normalized.startswith("err_") and normalized != "err_will_retry"


def _latest_delivery_events_by_row_recipient(
    job_id: str,
    sent_items: list[dict[str, Any]],
    *,
    provider: str = "",
) -> dict[tuple[str, str], dict[str, Any]]:
    provider_filter = _safe_text(provider).lower()
    events = _load_delivery_events(job_id, provider_filter)
    provider_id_index, row_recipient_index, recipient_index = _sent_item_indexes(sent_items)
    latest: dict[tuple[str, str], dict[str, Any]] = {}
    for event in events:
        matched_item = _match_delivery_event_to_sent_item(
            event,
            provider_id_index=provider_id_index,
            row_recipient_index=row_recipient_index,
            recipient_index=recipient_index,
        )
        row_id = _safe_text(event.get("row_id")) or _safe_text(matched_item.get("row_id") if matched_item else "")
        recipient = _mail_key(event.get("recipient") or (matched_item.get("recipient") if matched_item else ""))
        if not row_id or not recipient:
            continue
        normalized = dict(event)
        normalized["row_id"] = row_id
        normalized["recipient"] = recipient
        normalized.setdefault("provider_status", _safe_text(event.get("event_type")))
        key = (row_id, recipient)
        current = latest.get(key)
        if current is None or _delivery_event_order_key(normalized) >= _delivery_event_order_key(current):
            latest[key] = normalized
    return latest


def _load_delivery_events(job_id: str, provider: str) -> list[dict[str, Any]]:
    events: list[dict[str, Any]] = []
    providers = {provider} if provider and provider != "provider" else {"mailopost", "rusender", "unisender"}
    if "mailopost" in providers:
        try:
            from src.generator.delivery.mailopost_events import load_mailopost_events

            events.extend(load_mailopost_events(job_id))
        except Exception:
            pass
    if "rusender" in providers:
        try:
            from src.generator.delivery.rusender_events import load_rusender_events

            events.extend(load_rusender_events(job_id))
        except Exception:
            pass
    if "unisender" in providers:
        try:
            from src.generator.delivery.unisender_go_events import load_unisender_go_events

            for event in load_unisender_go_events(job_id):
                normalized = dict(event)
                normalized.setdefault("provider_status", _safe_text(event.get("event_type")))
                events.append(normalized)
        except Exception:
            pass
    return events


def _sent_item_indexes(
    sent_items: list[dict[str, Any]],
) -> tuple[dict[str, dict[str, Any]], dict[tuple[str, str], dict[str, Any]], dict[str, list[dict[str, Any]]]]:
    provider_id_index: dict[str, dict[str, Any]] = {}
    row_recipient_index: dict[tuple[str, str], dict[str, Any]] = {}
    recipient_index: dict[str, list[dict[str, Any]]] = {}
    for item in sent_items:
        for provider_id in _provider_ids_from_sent_item(item):
            provider_id_index[provider_id] = item
        row_id = _safe_text(item.get("row_id"))
        recipient = _mail_key(item.get("recipient"))
        if row_id and recipient:
            row_recipient_index[(row_id, recipient)] = item
        if recipient:
            recipient_index.setdefault(recipient, []).append(item)
    return provider_id_index, row_recipient_index, recipient_index


def _match_delivery_event_to_sent_item(
    event: dict[str, Any],
    *,
    provider_id_index: dict[str, dict[str, Any]],
    row_recipient_index: dict[tuple[str, str], dict[str, Any]],
    recipient_index: dict[str, list[dict[str, Any]]],
) -> dict[str, Any] | None:
    for provider_id in _provider_ids_from_delivery_event(event):
        item = provider_id_index.get(provider_id)
        if item:
            return item
    row_id = _safe_text(event.get("row_id"))
    recipient = _mail_key(event.get("recipient"))
    if row_id and recipient:
        item = row_recipient_index.get((row_id, recipient))
        if item:
            return item
    if recipient:
        candidates = recipient_index.get(recipient) or []
        if candidates:
            return max(candidates, key=_sent_item_order_key)
    return None


def _provider_ids_from_sent_item(item: dict[str, Any]) -> set[str]:
    provider = item.get("provider") if isinstance(item.get("provider"), dict) else {}
    values = (
        item.get("provider_message_id"),
        item.get("message_id"),
        item.get("provider_job_id"),
        item.get("provider_idempotency_key"),
        item.get("idempotency_key"),
        item.get("idempotencyKey"),
        provider.get("message_id"),
        provider.get("id"),
        provider.get("uuid"),
        provider.get("job_id"),
        provider.get("idempotency_key"),
        provider.get("idempotencyKey"),
    )
    return {_safe_text(value) for value in values if _safe_text(value)}


def _provider_ids_from_delivery_event(event: dict[str, Any]) -> set[str]:
    values = (
        event.get("message_id"),
        event.get("task_id"),
        event.get("provider_job_id"),
        event.get("email_id"),
        event.get("id"),
        event.get("uuid"),
    )
    raw_event = event.get("event") if isinstance(event.get("event"), dict) else {}
    nested_values = [
        raw_event.get("message_id"),
        raw_event.get("messageId"),
        raw_event.get("taskId"),
        raw_event.get("task_id"),
        raw_event.get("email_id"),
        raw_event.get("id"),
        raw_event.get("uuid"),
    ]
    payload = raw_event.get("payload")
    if isinstance(payload, dict):
        nested_values.extend(
            [
                payload.get("message_id"),
                payload.get("messageId"),
                payload.get("taskId"),
                payload.get("task_id"),
                payload.get("uuid"),
                payload.get("id"),
            ]
        )
    return {_safe_text(value) for value in (*values, *nested_values) if _safe_text(value)}


def _delivery_event_order_key(event: dict[str, Any]) -> tuple[str, str, str]:
    return (
        _safe_text(event.get("occurred_at")),
        _safe_text(event.get("received_at")),
        _safe_text(event.get("event_key")),
    )
def _check_unisender_classic_messages(email_ids: list[str]) -> dict[str, dict[str, Any]]:
    api_key = _safe_text(settings.unisender_api_key)
    unique_ids = [_safe_text(item) for item in dict.fromkeys(email_ids) if _safe_text(item)]
    if not api_key or not unique_ids:
        return {}

    params = {
        "format": "json",
        "api_key": api_key,
        "email_id": ",".join(unique_ids[:500]),
    }
    request = Request(
        _build_unisender_classic_url(UNISENDER_CLASSIC_CHECK_PATH),
        data=urlencode(params).encode("utf-8", errors="ignore"),
        method="POST",
        headers={"Content-Type": "application/x-www-form-urlencoded; charset=utf-8"},
    )
    with urlopen(request, timeout=30) as response:
        raw = response.read().decode("utf-8", errors="replace")
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"UniSender checkEmail вернул непонятный ответ: {raw[:300]}") from exc
    if data.get("error"):
        raise RuntimeError(_safe_text(data.get("error")) or "UniSender checkEmail вернул ошибку.")

    statuses = (data.get("result") or {}).get("statuses") if isinstance(data.get("result"), dict) else data.get("statuses")
    result: dict[str, dict[str, Any]] = {}
    if isinstance(statuses, list):
        for item in statuses:
            if not isinstance(item, dict):
                continue
            message_id = _safe_text(item.get("id"))
            if message_id:
                result[message_id] = {
                    "provider_status": _safe_text(item.get("status")),
                    "checked_at": datetime.now().isoformat(timespec="seconds"),
                }
    return result


def _unisender_status_label(status: str) -> str:
    normalized = _safe_text(status)
    labels = {
        "accepted": "принято UniSender в обработку",
        "success": "принято UniSender Go в обработку",
        "not_sent": "ещё не обработано",
        "ok_sent": "отправлено, ждём подтверждение доставки",
        "ok_delivered": "доставлено",
        "ok_read": "прочитано",
        "ok_link_visited": "перешли по ссылке",
        "ok_unsubscribed": "получатель отписался",
        "ok_spam_folder": "попало в спам",
        "err_user_unknown": "ошибка: неизвестный пользователь",
        "err_user_inactive": "ошибка: ящик неактивен",
        "err_mailbox_full": "ошибка: ящик переполнен",
        "err_spam_rejected": "ошибка: отклонено как спам",
        "err_delivery_failed": "ошибка доставки",
        "err_will_retry": "временная ошибка, UniSender повторит",
        "err_lost": "статус потерян, нужна проверка вручную",
    }
    return labels.get(normalized, normalized or "статус неизвестен")


def get_unisender_history(
    *,
    job_id: str | None = None,
    limit: int = 50,
    refresh: bool = False,
) -> dict[str, Any]:
    job_paths = resolve_job_paths(job_id)
    sent_mail_log_path = None if job_paths.uses_legacy_layout else job_paths.sent_mail_log_path
    items = [
        item
        for item in _load_sent_mail_log_items(sent_mail_log_path)
        if _safe_text(item.get("transport")) == "unisender"
    ]
    items = list(reversed(items))[: max(1, min(int(limit or 50), 200))]

    classic_ids: list[str] = []
    for item in items:
        provider = item.get("provider") if isinstance(item.get("provider"), dict) else {}
        if _safe_text(provider.get("provider")) != "unisender_classic":
            continue
        message_id = _safe_text(item.get("provider_message_id") or provider.get("message_id"))
        if message_id:
            classic_ids.append(message_id)
    provider_statuses: dict[str, dict[str, Any]] = {}
    refresh_error = ""
    if refresh and classic_ids:
        try:
            provider_statuses = _check_unisender_classic_messages(classic_ids)
        except Exception as exc:
            refresh_error = _safe_text(exc) or "не удалось обновить статусы UniSender"

    result_items: list[dict[str, Any]] = []
    for item in items:
        provider = item.get("provider") if isinstance(item.get("provider"), dict) else {}
        provider_name = _safe_text(provider.get("provider")) or "unisender"
        message_id = _safe_text(item.get("provider_message_id") or provider.get("message_id"))
        job_provider_id = _safe_text(item.get("provider_job_id") or provider.get("job_id"))
        provider_status = _safe_text(provider.get("status")) or "accepted"
        checked_at = ""
        if message_id and message_id in provider_statuses:
            provider_status = provider_statuses[message_id].get("provider_status") or provider_status
            checked_at = provider_statuses[message_id].get("checked_at") or ""
        result_items.append(
            {
                "sent_at": item.get("sent_at"),
                "row_id": item.get("row_id"),
                "mun_name": item.get("mun_name"),
                "recipient": item.get("recipient"),
                "subject": item.get("subject"),
                "attachments": item.get("attachments") or [],
                "provider": provider_name,
                "provider_message_id": message_id,
                "provider_job_id": job_provider_id,
                "provider_status": provider_status,
                "provider_status_label": _unisender_status_label(provider_status),
                "checked_at": checked_at,
                "warning": item.get("warning") or "",
            }
        )

    summary = f"История UniSender: найдено {len(result_items)} писем."
    if refresh_error:
        summary += f" Статусы доставки не удалось обновить: {refresh_error}."
    elif refresh and classic_ids:
        summary += " Статусы доставки обновлены через UniSender checkEmail."
    elif refresh:
        summary += " Для этих писем нет classic email_id для проверки через checkEmail."
    return {
        "status": "ok",
        "summary_text": summary,
        "items": result_items,
        "refresh_error": refresh_error,
    }


def _save_sent_copy(message: EmailMessage) -> str | None:
    if not settings.smtp_save_sent_copy:
        return None
    if not settings.smtp_sender_email or not settings.smtp_sender_password:
        return "Не удалось сохранить копию письма: не заданы почтовые учётные данные."
    if not settings.imap_host:
        return "Не удалось сохранить копию письма: не указан IMAP host."

    def _imap_utf7_decode(value: str) -> str:
        def _decode_match(match: re.Match[str]) -> str:
            chunk = match.group(1)
            if chunk == "-":
                return "&"
            payload = chunk[:-1].replace(",", "/")
            padding = "=" * ((4 - len(payload) % 4) % 4)
            raw = base64.b64decode(payload + padding)
            return raw.decode("utf-16-be", errors="ignore")

        return re.sub(r"&([A-Za-z0-9+,]+-)", _decode_match, value)

    def _imap_utf7_encode(value: str) -> str:
        result: list[str] = []
        buffer = ""

        def _flush_buffer() -> None:
            nonlocal buffer
            if not buffer:
                return
            encoded = base64.b64encode(buffer.encode("utf-16-be")).decode("ascii").rstrip("=")
            result.append("&" + encoded.replace("/", ",") + "-")
            buffer = ""

        for char in value:
            code = ord(char)
            if 0x20 <= code <= 0x7E and char != "&":
                _flush_buffer()
                result.append(char)
            elif char == "&":
                _flush_buffer()
                result.append("&-")
            else:
                buffer += char
        _flush_buffer()
        return "".join(result)

    def _extract_imap_mailbox_name(line: bytes | str) -> tuple[str, str]:
        text = line.decode("ascii", errors="ignore") if isinstance(line, bytes) else str(line)
        match = re.match(r'^\((?P<flags>.*?)\)\s+"[^"]*"\s+(?P<name>.+)$', text.strip())
        if not match:
            return "", ""
        flags = (match.group("flags") or "").strip()
        name = (match.group("name") or "").strip()
        if name.startswith('"') and name.endswith('"') and len(name) >= 2:
            name = name[1:-1]
        return flags, _imap_utf7_decode(name)

    def _discover_imap_sent_folders(client: imaplib.IMAP4) -> list[str]:
        try:
            status, mailboxes = client.list()
        except Exception:
            return []
        if status != "OK" or not mailboxes:
            return []

        discovered: list[str] = []
        for item in mailboxes:
            flags, name = _extract_imap_mailbox_name(item)
            if not name:
                continue
            haystack = f"{flags} {name}".lower()
            if "\\sent" in haystack or "sent" in haystack or "отправ" in haystack:
                if name not in discovered:
                    discovered.append(name)
        return discovered

    tried_folders: list[str] = []
    try:
        if settings.imap_use_ssl:
            client = imaplib.IMAP4_SSL(settings.imap_host, settings.imap_port)
        else:
            client = imaplib.IMAP4(settings.imap_host, settings.imap_port)
        try:
            client.login(settings.smtp_sender_email, settings.smtp_sender_password)
            payload = message.as_bytes()
            internal_date = imaplib.Time2Internaldate(datetime.now().timestamp())
            folder_candidates = [
                _safe_text(settings.imap_sent_folder) or "Отправленные",
                "Sent",
                "Sent Messages",
                *_discover_imap_sent_folders(client),
            ]
            for folder in folder_candidates:
                if folder in tried_folders:
                    continue
                tried_folders.append(folder)
                try:
                    status, _ = client.append(_imap_utf7_encode(folder), "\\Seen", internal_date, payload)
                except imaplib.IMAP4.error:
                    status = "NO"
                if status == "OK":
                    return None
        finally:
            try:
                client.logout()
            except Exception:
                pass
    except Exception as exc:
        return f"Не удалось сохранить копию письма в 'Отправленные': {_safe_text(exc) or 'ошибка IMAP'}."

    return (
        "Не удалось сохранить копию письма в папку 'Отправленные'. "
        f"Пробовал папки: {', '.join(tried_folders)}."
    )


def _send_via_smtp(
    row: dict[str, Any],
    recipient: str,
    attachments: list[str],
    subject: str,
    *,
    mail_template_path: Path | None = None,
    body_override: str | None = None,
) -> str | None:
    if not settings.smtp_allow_real_send:
        raise RuntimeError(
            "Реальная SMTP-отправка запрещена настройкой smtp_allow_real_send. "
            "Сейчас доступен только dry-run режим."
        )
    if not settings.smtp_sender_email or not settings.smtp_sender_password:
        raise RuntimeError("Не настроены SMTP-учётные данные отправителя.")
    if not settings.smtp_host:
        raise RuntimeError("Не указан SMTP host.")

    message = _build_message(
        row,
        recipient,
        attachments,
        subject,
        mail_template_path=mail_template_path,
        body_override=body_override,
    )

    if settings.smtp_use_ssl:
        with smtplib.SMTP_SSL(settings.smtp_host, settings.smtp_port, timeout=30) as server:
            server.login(settings.smtp_sender_email, settings.smtp_sender_password)
            server.send_message(message)
        return _save_sent_copy(message)

    with smtplib.SMTP(settings.smtp_host, settings.smtp_port, timeout=30) as server:
        server.ehlo()
        if settings.smtp_use_starttls:
            server.starttls()
            server.ehlo()
        server.login(settings.smtp_sender_email, settings.smtp_sender_password)
        server.send_message(message)
    return _save_sent_copy(message)


def _htmlify_mail_body(
    body: str,
    *,
    inline_footer_image: bool = False,
    include_unsubscribe: bool = True,
) -> str:
    parts: list[str] = []
    for line in body.splitlines():
        stripped = line.strip()
        consent_match = re.match(r"^(Получить .+?\.?)\s*:?\s*(https?://\S+)\s*$", stripped)
        if consent_match:
            button_text = escape(consent_match.group(1), quote=False)
            consent_url = escape(consent_match.group(2), quote=True)
            parts.append(
                "<div style=\"margin:18px 0;\">"
                "<a "
                f"href=\"{consent_url}\" "
                "style=\"display:inline-block;padding:12px 18px;background:#2d720d;"
                "color:#ffffff;text-decoration:none;border-radius:8px;font-weight:700;\""
                f">{button_text}</a>"
                "<div style=\"margin-top:8px;font-size:12px;line-height:1.45;color:#60705a;\">"
                "Нажимая на кнопку, вы просто даёте нам знать, что документы нужны. Мы вышлем их сразу."
                "</div>"
                "</div>"
            )
            continue
        parts.append(escape(stripped))
    non_empty = [line for line in parts if line]
    html = "<br>".join(non_empty)
    for marker in _mail_footer_html_markers():
        marker_index = html.find(marker)
        if marker_index >= 0:
            html = html[:marker_index] + _build_mail_footer_html(inline_image=inline_footer_image)
            break
    if include_unsubscribe and "{{UnsubscribeUrl}}" not in html:
        html += "<br><br><a href='{{UnsubscribeUrl}}'>Отписаться от писем</a>"
    return html


def _build_unisender_go_url(path: str) -> str:
    base_url = _safe_text(settings.unisender_api_base_url).rstrip("/")
    if not base_url:
        base_url = "https://goapi.unisender.ru/ru/transactional/api/v1"
    return f"{base_url}/{path.lstrip('/')}"


def _uses_unisender_go_api() -> bool:
    base_url = _safe_text(settings.unisender_api_base_url).lower()
    return "goapi.unisender.ru" in base_url


def _build_unisender_classic_url(path: str) -> str:
    base_url = _safe_text(settings.unisender_api_base_url).rstrip("/")
    if not base_url:
        base_url = "https://api.unisender.com/ru/api"
    return f"{base_url}/{path.lstrip('/')}"


def _build_rusender_url(path: str) -> str:
    base_url = _safe_text(settings.rusender_api_base_url).rstrip("/")
    if not base_url:
        base_url = "https://api.rusender.ru/api/v1"
    return f"{base_url}/{path.lstrip('/')}"


def _build_mailopost_url(path: str) -> str:
    base_url = _safe_text(settings.mailopost_api_base_url).rstrip("/")
    if not base_url:
        base_url = "https://api.mailopost.ru/v1"
    return f"{base_url}/{path.lstrip('/')}"

def _is_retryable_rusender_exception(exc: Exception) -> bool:
    if isinstance(exc, HTTPError):
        return int(exc.code) in RUSENDER_RETRYABLE_HTTP_CODES
    return isinstance(exc, (TimeoutError, URLError, OSError))


def _wait_rusender_api_slot() -> None:
    global _last_rusender_request_at
    with _RUSENDER_RATE_LIMIT_LOCK:
        now = perf_counter()
        wait_seconds = RUSENDER_MIN_REQUEST_INTERVAL_SECONDS - (now - _last_rusender_request_at)
        if wait_seconds > 0:
            sleep(wait_seconds)
        _last_rusender_request_at = perf_counter()


def _run_rusender_request(request: Request, *, timeout: float) -> str:
    last_error: Exception | None = None
    for attempt in range(1, RUSENDER_RETRY_ATTEMPTS + 1):
        try:
            _wait_rusender_api_slot()
            with urlopen(request, timeout=timeout) as response:
                return response.read().decode("utf-8", errors="replace")
        except HTTPError as exc:
            raw = exc.read().decode("utf-8", errors="replace")
            exc.raw_body = raw  # type: ignore[attr-defined]
            if attempt < RUSENDER_RETRY_ATTEMPTS and _is_retryable_rusender_exception(exc):
                _sleep_sender_retry(RUSENDER_RETRY_BASE_SECONDS * attempt)
                last_error = exc
                continue
            raise
        except Exception as exc:
            if attempt < RUSENDER_RETRY_ATTEMPTS and _is_retryable_rusender_exception(exc):
                _sleep_sender_retry(RUSENDER_RETRY_BASE_SECONDS * attempt)
                last_error = exc
                continue
            raise
    if last_error is not None:
        raise last_error
    raise RuntimeError("RuSender не ответил после повторных попыток.")


def _parse_retry_after_seconds(value: Any) -> float | None:
    text = _safe_text(value)
    if not text:
        return None
    try:
        seconds = float(text)
    except ValueError:
        match = re.search(r"try\s+again\s+in\s+(\d+(?:\.\d+)?)\s+seconds?", text, flags=re.IGNORECASE)
        if not match:
            match = re.search(r"(?:retry|повтор)\D{0,20}(\d+(?:\.\d+)?)", text, flags=re.IGNORECASE)
        if not match:
            return None
        seconds = float(match.group(1))
    if seconds < 0:
        return None
    return seconds


def _mailopost_retry_after_seconds(exc: HTTPError, raw: str, message: str) -> float:
    header_value = exc.headers.get("Retry-After") if exc.headers else None
    for candidate in (header_value, message, raw):
        seconds = _parse_retry_after_seconds(candidate)
        if seconds is not None:
            return seconds + MAILOPOST_RATE_LIMIT_EXTRA_SECONDS
    return MAILOPOST_RATE_LIMIT_FALLBACK_SECONDS

def _is_retryable_mailopost_exception(exc: Exception) -> bool:
    if isinstance(exc, HTTPError):
        return int(exc.code) in MAILOPOST_RETRYABLE_HTTP_CODES
    return isinstance(exc, (TimeoutError, URLError, OSError))


def _wait_mailopost_api_slot() -> None:
    global _last_mailopost_request_at
    with _MAILOPOST_RATE_LIMIT_LOCK:
        now = perf_counter()
        wait_seconds = MAILOPOST_MIN_REQUEST_INTERVAL_SECONDS - (now - _last_mailopost_request_at)
        if wait_seconds > 0:
            sleep(wait_seconds)
        _last_mailopost_request_at = perf_counter()


def _run_mailopost_request(request: Request, *, timeout: float) -> str:
    last_error: Exception | None = None
    for attempt in range(1, MAILOPOST_RETRY_ATTEMPTS + 1):
        try:
            _wait_mailopost_api_slot()
            with urlopen(request, timeout=timeout) as response:
                return response.read().decode("utf-8", errors="replace")
        except HTTPError as exc:
            raw = exc.read().decode("utf-8", errors="replace")
            exc.raw_body = raw  # type: ignore[attr-defined]
            if int(exc.code) == 429:
                raise
            if attempt < MAILOPOST_RETRY_ATTEMPTS and _is_retryable_mailopost_exception(exc):
                _sleep_sender_retry(MAILOPOST_RETRY_BASE_SECONDS * attempt)
                last_error = exc
                continue
            raise
        except Exception as exc:
            if attempt < MAILOPOST_RETRY_ATTEMPTS and _is_retryable_mailopost_exception(exc):
                _sleep_sender_retry(MAILOPOST_RETRY_BASE_SECONDS * attempt)
                last_error = exc
                continue
            raise
    if last_error is not None:
        raise last_error
    raise RuntimeError("MailoPost не ответил после повторных попыток.")

def _safe_provider_payload(provider: dict[str, Any] | None) -> dict[str, Any]:
    if not isinstance(provider, dict):
        return {}
    allowed_keys = {
        "provider",
        "status",
        "message_id",
        "email_id",
        "job_id",
        "uuid",
        "recipient",
        "accepted_emails",
        "failed_emails",
        "idempotence_key",
        "idempotency_key",
    }
    safe: dict[str, Any] = {}
    for key in allowed_keys:
        if key not in provider:
            continue
        value = provider.get(key)
        if value in (None, "", [], {}):
            continue
        safe[key] = value
    return safe


def _rusender_auth_headers(api_key: str) -> dict[str, str]:
    if api_key.startswith("rs_ck_"):
        return {"Authorization": f"Bearer {api_key}"}
    return {"X-Api-Key": api_key}


def _send_via_rusender(
    row: dict[str, Any],
    recipient: str,
    attachments: list[str],
    subject: str,
    *,
    mail_template_path: Path | None = None,
    body_override: str | None = None,
    job_id: str | None = None,
    send_run_id: str = "",
    send_mode: str = "",
    attachment_mode: str = "",
) -> dict[str, Any]:
    api_key = _safe_text(settings.rusender_api_key)
    sender_email = _safe_text(settings.rusender_sender_email or settings.smtp_sender_email)
    sender_name = _safe_text(settings.rusender_sender_name) or "ООО «ПР»"
    if not api_key:
        raise RuntimeError("Не указан API-ключ RuSender.")
    if not sender_email:
        raise RuntimeError("Не указан подтверждённый email отправителя RuSender.")

    plaintext_body = body_override if body_override is not None else _build_mail_body(row, mail_template_path=mail_template_path)
    html_body = _htmlify_mail_body(plaintext_body, include_unsubscribe=False)
    idempotency_key = _build_provider_idempotency_key(
        provider="rusender",
        job_id=job_id,
        send_run_id=send_run_id,
        row_id=row.get("ID"),
        recipient=recipient,
        send_mode=send_mode,
        attachment_mode=attachment_mode,
    )
    encoded_attachments: list[dict[str, str]] = []
    for attachment_path in attachments:
        path = Path(attachment_path)
        encoded_attachments.append({path.name: base64.b64encode(path.read_bytes()).decode("ascii")})

    payload: dict[str, Any] = {
        "idempotencyKey": idempotency_key,
        "mail": {
            "to": {
                "email": recipient,
                "name": _safe_text(row.get("MUN_NAME")),
            },
            "from": {
                "email": sender_email,
                "name": sender_name,
            },
            "subject": subject,
            "html": html_body,
            "text": plaintext_body,
        },
    }
    if encoded_attachments:
        payload["mail"]["attachments"] = encoded_attachments

    request = Request(
        _build_rusender_url(RUSENDER_SEND_PATH),
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        method="POST",
        headers={
            "Content-Type": "application/json",
            "Accept": "application/json",
            **_rusender_auth_headers(api_key),
        },
    )
    try:
        raw = _run_rusender_request(request, timeout=60)
    except HTTPError as exc:
        raw = getattr(exc, "raw_body", "")
        try:
            data = json.loads(raw)
        except json.JSONDecodeError:
            raise RuntimeError(f"RuSender вернул HTTP {exc.code}: {raw[:300]}") from exc
        message = _safe_text(data.get("message") or data.get("error"))
        status_code = data.get("statusCode") or exc.code
        if message:
            raise RuntimeError(f"{message} (code {status_code})") from exc
        raise RuntimeError(f"RuSender вернул HTTP {exc.code}: {raw[:300]}") from exc

    try:
        data = json.loads(raw) if raw else {}
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"RuSender вернул непонятный ответ: {raw[:300]}") from exc

    uuid = _safe_text(data.get("uuid"))
    if not uuid:
        message = _safe_text(data.get("message") or data.get("error"))
        raise RuntimeError(message or "RuSender не подтвердил отправку письма.")

    failed_recipients: list[str] = []
    for item in data.get("additionalRecipients") or []:
        if isinstance(item, dict) and item.get("error"):
            failed_recipients.append(_safe_text(item.get("error")))
    if failed_recipients:
        raise RuntimeError("; ".join(failed_recipients))

    return {
        "provider": "rusender",
        "status": "accepted",
        "message_id": uuid,
        "uuid": uuid,
        "recipient": recipient,
        "idempotency_key": idempotency_key,
    }


def _mailopost_error_message(raw: str, *, fallback: str = "MailoPost отклонил письмо.") -> str:
    try:
        data = json.loads(raw) if raw else {}
    except json.JSONDecodeError:
        return raw[:300] or fallback
    errors = data.get("errors")
    if isinstance(errors, list) and errors:
        messages: list[str] = []
        for item in errors:
            if not isinstance(item, dict):
                continue
            detail = _safe_text(item.get("detail") or item.get("message") or item.get("title"))
            code = item.get("code") or item.get("status")
            if detail:
                messages.append(f"{detail} (code {code})" if code else detail)
        if messages:
            return "; ".join(messages)
    return _safe_text(data.get("message") or data.get("error") or data.get("detail")) or fallback


def _mailopost_smtp_headers(
    row: dict[str, Any],
    *,
    idempotency_key: str,
    recipient: str,
    job_id: str | None,
) -> dict[str, str]:
    return {
        "Client-Id": idempotency_key,
        "X-Mailing-Agent-Job": _safe_text(job_id),
        "X-Mailing-Agent-Row": _safe_text(row.get("ID")),
        "X-Mailing-Agent-Recipient": _safe_text(recipient),
    }


def _build_mailopost_json_request(*, api_token: str, payload: dict[str, Any]) -> Request:
    return Request(
        _build_mailopost_url(MAILOPOST_SEND_PATH),
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        method="POST",
        headers={
            "Content-Type": "application/json",
            "Accept": "application/json",
            "Authorization": f"Bearer {api_token}",
        },
    )


def _build_mailopost_multipart_request(
    *,
    api_token: str,
    fields: dict[str, Any],
    attachments: list[str],
) -> Request:
    boundary = f"----mailing-agent-{secrets.token_hex(16)}"
    parts: list[bytes] = []

    def add_field(name: str, value: Any) -> None:
        if value in (None, ""):
            return
        parts.append(
            (
                f"--{boundary}\r\n"
                f"Content-Disposition: form-data; name=\"{name}\"\r\n\r\n"
                f"{_safe_text(value)}\r\n"
            ).encode("utf-8")
        )

    for name, value in fields.items():
        if isinstance(value, dict):
            for nested_name, nested_value in value.items():
                add_field(f"{name}[{nested_name}]", nested_value)
        else:
            add_field(name, value)

    total_attachment_bytes = 0
    for attachment_path in attachments:
        path = Path(attachment_path)
        file_bytes = path.read_bytes()
        total_attachment_bytes += len(file_bytes)
        if total_attachment_bytes > MAILOPOST_MAX_ATTACHMENT_BYTES:
            raise RuntimeError("MailoPost принимает вложения суммарно до 5 МБ.")
        mime_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
        parts.append(
            (
                f"--{boundary}\r\n"
                f"Content-Disposition: form-data; name=\"attachments[]\"; filename=\"{path.name}\"\r\n"
                f"Content-Type: {mime_type}\r\n\r\n"
            ).encode("utf-8")
        )
        parts.append(file_bytes)
        parts.append(b"\r\n")

    parts.append(f"--{boundary}--\r\n".encode("utf-8"))
    return Request(
        _build_mailopost_url(MAILOPOST_SEND_PATH),
        data=b"".join(parts),
        method="POST",
        headers={
            "Content-Type": f"multipart/form-data; boundary={boundary}",
            "Accept": "application/json",
            "Authorization": f"Bearer {api_token}",
        },
    )


def _send_via_mailopost(
    row: dict[str, Any],
    recipient: str,
    attachments: list[str],
    subject: str,
    *,
    mail_template_path: Path | None = None,
    body_override: str | None = None,
    job_id: str | None = None,
    send_run_id: str = "",
    send_mode: str = "",
    attachment_mode: str = "",
) -> dict[str, Any]:
    api_token = _safe_text(settings.mailopost_api_token)
    sender_email = _safe_text(settings.mailopost_sender_email or settings.smtp_sender_email)
    sender_name = _safe_text(settings.mailopost_sender_name) or "ООО «ПР»"
    if not api_token:
        raise RuntimeError("Не указан API-токен MailoPost.")
    if not sender_email:
        raise RuntimeError("Не указан подтверждённый email отправителя MailoPost.")

    plaintext_body = body_override if body_override is not None else _build_mail_body(row, mail_template_path=mail_template_path)
    html_body = _htmlify_mail_body(plaintext_body, include_unsubscribe=False)
    idempotency_key = _build_provider_idempotency_key(
        provider="mailopost",
        job_id=job_id,
        send_run_id=send_run_id,
        row_id=row.get("ID"),
        recipient=recipient,
        send_mode=send_mode,
        attachment_mode=attachment_mode,
    )
    common_payload: dict[str, Any] = {
        "from_email": sender_email,
        "from_name": sender_name,
        "to": recipient,
        "subject": subject,
        "text": plaintext_body,
        "html": html_body,
        "payment": "credit",
        "smtp_headers": _mailopost_smtp_headers(row, idempotency_key=idempotency_key, recipient=recipient, job_id=job_id),
    }
    request = (
        _build_mailopost_multipart_request(api_token=api_token, fields=common_payload, attachments=attachments)
        if attachments
        else _build_mailopost_json_request(api_token=api_token, payload=common_payload)
    )
    try:
        raw = _run_mailopost_request(request, timeout=60)
    except HTTPError as exc:
        raw = getattr(exc, "raw_body", "")
        message = _mailopost_error_message(raw, fallback=f"MailoPost вернул HTTP {exc.code}")
        if int(exc.code) == 429:
            retry_after_seconds = _mailopost_retry_after_seconds(exc, raw, message)
            raise MailoPostRateLimitError(
                f"{message} (HTTP {exc.code}). Повтор через {int(retry_after_seconds)} сек.",
                retry_after_seconds,
            ) from exc
        raise RuntimeError(f"{message} (HTTP {exc.code})") from exc

    try:
        data = json.loads(raw) if raw else {}
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"MailoPost вернул непонятный ответ: {raw[:300]}") from exc

    message_id = _safe_text(data.get("id") or data.get("message_id") or data.get("uuid"))
    provider_status = _safe_text(data.get("status")) or "queued"
    if data.get("errors"):
        raise RuntimeError(_mailopost_error_message(raw))
    if not message_id and provider_status.lower() not in {"queued", "sent", "accepted"}:
        raise RuntimeError(_mailopost_error_message(raw, fallback="MailoPost не подтвердил отправку письма."))
    return {
        "provider": "mailopost",
        "status": provider_status,
        "message_id": message_id or idempotency_key,
        "recipient": recipient,
        "idempotency_key": idempotency_key,
    }

def _send_via_unisender_classic(
    row: dict[str, Any],
    recipient: str,
    attachments: list[str],
    subject: str,
    *,
    mail_template_path: Path | None = None,
    body_override: str | None = None,
) -> dict[str, Any]:
    api_key = _safe_text(settings.unisender_api_key)
    sender_email = _safe_text(settings.unisender_sender_email or settings.smtp_sender_email)
    sender_name = _safe_text(settings.unisender_sender_name) or "ООО «ПР»"
    list_id = int(settings.unisender_list_id or 1)
    if not api_key:
        raise RuntimeError("Не указан API-ключ UniSender.")
    if not sender_email:
        raise RuntimeError("Не указан подтверждённый email отправителя UniSender.")

    body = _htmlify_mail_body(
        body_override if body_override is not None else _build_mail_body(row, mail_template_path=mail_template_path),
        include_unsubscribe=False,
    )
    payload: dict[str, Any] = {
        "format": "json",
        "api_key": api_key,
        "email": recipient,
        "sender_name": sender_name,
        "sender_email": sender_email,
        "subject": subject,
        "body": body,
        "list_id": list_id,
        "lang": "ru",
        "error_checking": 1,
    }

    for attachment_path in attachments:
        path = Path(attachment_path)
        payload[f"attachments[{path.name}]"] = path.read_bytes()

    request = Request(
        _build_unisender_classic_url(UNISENDER_CLASSIC_SEND_PATH),
        data=urlencode(payload).encode("utf-8", errors="ignore"),
        method="POST",
        headers={"Content-Type": "application/x-www-form-urlencoded; charset=utf-8"},
    )
    try:
        raw = _run_unisender_request(request, timeout=60, request_label="UniSender")
    except HTTPError as exc:
        raw = getattr(exc, "raw_body", "")
        try:
            data = json.loads(raw)
        except json.JSONDecodeError:
            raise RuntimeError(f"UniSender вернул HTTP {exc.code}: {raw[:300]}") from exc
        error_message = _safe_text(data.get("error"))
        if error_message:
            raise RuntimeError(error_message) from exc
        raise RuntimeError(f"UniSender вернул HTTP {exc.code}: {raw[:300]}") from exc
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"UniSender вернул непонятный ответ: {raw[:300]}") from exc

    result_items = data.get("result")
    if isinstance(result_items, list) and result_items:
        first = result_items[0]
        errors = first.get("errors") or []
        if errors:
            message = "; ".join(_safe_text(err.get("message")) for err in errors if isinstance(err, dict))
            raise RuntimeError(message or "UniSender отклонил письмо.")
        if first.get("id"):
            message_id = _safe_text(first.get("id"))
            return {
                "provider": "unisender_classic",
                "status": "accepted",
                "message_id": message_id,
                "email_id": message_id,
                "recipient": _safe_text(first.get("email")) or recipient,
            }

    legacy_result = data.get("result")
    if isinstance(legacy_result, dict) and legacy_result.get("email_id"):
        message_id = _safe_text(legacy_result.get("email_id"))
        return {
            "provider": "unisender_classic",
            "status": "accepted",
            "message_id": message_id,
            "email_id": message_id,
            "recipient": recipient,
        }

    if data.get("error"):
        raise RuntimeError(_safe_text(data.get("error")) or "UniSender вернул ошибку.")
    raise RuntimeError("UniSender не подтвердил отправку письма.")


def _send_via_unisender(
    row: dict[str, Any],
    recipient: str,
    attachments: list[str],
    subject: str,
    *,
    mail_template_path: Path | None = None,
    body_override: str | None = None,
    job_id: str | None = None,
    send_run_id: str = "",
    send_mode: str = "",
    attachment_mode: str = "",
) -> dict[str, Any]:
    if not _uses_unisender_go_api():
        return _send_via_unisender_classic(
            row,
            recipient,
            attachments,
            subject,
            mail_template_path=mail_template_path,
            body_override=body_override,
        )

    api_key = _safe_text(settings.unisender_api_key)
    sender_email = _safe_text(settings.unisender_sender_email or settings.smtp_sender_email)
    sender_name = _safe_text(settings.unisender_sender_name) or "ООО «ПР»"
    if not api_key:
        raise RuntimeError("Не указан API-ключ UniSender Go.")
    if not sender_email:
        raise RuntimeError("Не указан email отправителя UniSender Go.")

    plaintext_body = body_override if body_override is not None else _build_mail_body(row, mail_template_path=mail_template_path)
    html_body = _htmlify_mail_body(plaintext_body, include_unsubscribe=False)
    idempotence_key = _build_provider_idempotency_key(
        provider="unisender_go",
        job_id=job_id,
        send_run_id=send_run_id,
        row_id=row.get("ID"),
        recipient=recipient,
        send_mode=send_mode,
        attachment_mode=attachment_mode,
    )
    payload: dict[str, Any] = {
        "message": {
            "recipients": [{"email": recipient}],
            "body": {
                "html": html_body,
                "plaintext": plaintext_body,
            },
            "subject": subject,
            "from_email": sender_email,
            "from_name": sender_name,
            "reply_to": sender_email,
            "reply_to_name": sender_name,
            "global_language": "ru",
            "template_engine": "simple",
            "track_links": 1,
            "track_read": 1,
            "idempotence_key": idempotence_key,
            "global_metadata": _build_unisender_go_metadata(row, recipient=recipient, job_id=job_id),
        }
    }

    encoded_attachments: list[dict[str, str]] = []
    for attachment_path in attachments:
        path = Path(attachment_path)
        mime_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
        encoded_attachments.append(
            {
                "type": mime_type,
                "name": path.name,
                "content": base64.b64encode(path.read_bytes()).decode("ascii"),
            }
        )
    if encoded_attachments:
        payload["message"]["attachments"] = encoded_attachments

    request = Request(
        _build_unisender_go_url(UNISENDER_GO_SEND_PATH),
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        method="POST",
        headers={
            "Content-Type": "application/json",
            "Accept": "application/json",
            "X-API-KEY": api_key,
        },
    )
    try:
        raw = _run_unisender_request(request, timeout=60, request_label="UniSender Go")
    except HTTPError as exc:
        raw = getattr(exc, "raw_body", "")
        try:
            data = json.loads(raw)
            message = _safe_text(data.get("message"))
            code = data.get("code")
            if message:
                suffix = f" (code {code})" if code is not None else ""
                raise RuntimeError(message + suffix) from exc
        except json.JSONDecodeError:
            pass
        raise RuntimeError(f"UniSender Go вернул HTTP {exc.code}: {raw[:300]}") from exc
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"UniSender Go вернул непонятный ответ: {raw[:300]}") from exc

    if _safe_text(data.get("status")).lower() != "success":
        raise RuntimeError(_safe_text(data.get("message")) or "UniSender Go не подтвердил отправку письма.")

    failed_emails = data.get("failed_emails") or {}
    if isinstance(failed_emails, dict) and recipient in failed_emails:
        raise RuntimeError(f"UniSender Go не принял адрес {recipient}: {_safe_text(failed_emails.get(recipient))}")

    accepted_emails = data.get("emails") or []
    if accepted_emails and recipient not in accepted_emails:
        raise RuntimeError("UniSender Go не подтвердил адрес получателя в ответе.")
    return {
        "provider": "unisender_go",
        "status": _safe_text(data.get("status")) or "success",
        "job_id": _safe_text(data.get("job_id") or data.get("id")),
        "recipient": recipient,
        "accepted_emails": accepted_emails,
        "failed_emails": failed_emails if isinstance(failed_emails, dict) else {},
        "idempotence_key": idempotence_key,
    }


def _send_via_unisender_go_bulk(
    row: dict[str, Any],
    recipients: list[str],
    attachments: list[str],
    subject: str,
    *,
    mail_template_path: Path | None = None,
    body_override: str | None = None,
    job_id: str | None = None,
    send_run_id: str = "",
    send_mode: str = "",
    attachment_mode: str = "",
) -> dict[str, Any]:
    api_key = _safe_text(settings.unisender_api_key)
    sender_email = _safe_text(settings.unisender_sender_email or settings.smtp_sender_email)
    sender_name = _safe_text(settings.unisender_sender_name) or "ООО «ПР»"
    if not api_key:
        raise RuntimeError("Не указан API-ключ UniSender Go.")
    if not sender_email:
        raise RuntimeError("Не указан email отправителя UniSender Go.")

    cleaned_recipients = [recipient for recipient in recipients if recipient]
    if not cleaned_recipients:
        raise RuntimeError("Не найден получатель для отправки.")

    plaintext_body = body_override if body_override is not None else _build_mail_body(row, mail_template_path=mail_template_path)
    html_body = _htmlify_mail_body(plaintext_body, include_unsubscribe=False)
    idempotence_key = _build_provider_idempotency_key(
        provider="unisender_go",
        job_id=job_id,
        send_run_id=send_run_id,
        row_id=row.get("ID"),
        recipient=",".join(sorted(_mail_key(recipient) for recipient in cleaned_recipients)),
        send_mode=send_mode,
        attachment_mode=attachment_mode,
    )
    payload: dict[str, Any] = {
        "message": {
            "recipients": [{"email": recipient} for recipient in cleaned_recipients],
            "body": {
                "html": html_body,
                "plaintext": plaintext_body,
            },
            "subject": subject,
            "from_email": sender_email,
            "from_name": sender_name,
            "reply_to": sender_email,
            "reply_to_name": sender_name,
            "global_language": "ru",
            "template_engine": "simple",
            "track_links": 1,
            "track_read": 1,
            "idempotence_key": idempotence_key,
            "global_metadata": _build_unisender_go_metadata(row, recipient=", ".join(cleaned_recipients), job_id=job_id),
        }
    }

    encoded_attachments: list[dict[str, str]] = []
    for attachment_path in attachments:
        path = Path(attachment_path)
        mime_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
        encoded_attachments.append(
            {
                "type": mime_type,
                "name": path.name,
                "content": base64.b64encode(path.read_bytes()).decode("ascii"),
            }
        )
    if encoded_attachments:
        payload["message"]["attachments"] = encoded_attachments

    request = Request(
        _build_unisender_go_url(UNISENDER_GO_SEND_PATH),
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        method="POST",
        headers={
            "Content-Type": "application/json",
            "Accept": "application/json",
            "X-API-KEY": api_key,
        },
    )
    data: dict[str, Any] = {}
    for attempt in range(1, UNISENDER_RETRY_ATTEMPTS + 1):
        try:
            raw = _run_unisender_request(request, timeout=60, request_label="UniSender Go")
        except HTTPError as exc:
            raw = getattr(exc, "raw_body", "")
            try:
                error_data = json.loads(raw)
                message = _safe_text(error_data.get("message"))
                code = error_data.get("code")
                if message:
                    suffix = f" (code {code})" if code is not None else ""
                    raise RuntimeError(message + suffix) from exc
            except json.JSONDecodeError:
                pass
            raise RuntimeError(f"UniSender Go вернул HTTP {exc.code}: {raw[:300]}") from exc
        try:
            data = json.loads(raw)
        except json.JSONDecodeError:
            raise RuntimeError(f"UniSender Go вернул непонятный ответ: {raw[:300]}")
        if _safe_text(data.get("status")).lower() == "success":
            break
        message = _safe_text(data.get("message")) or "UniSender Go не подтвердил отправку письма."
        if attempt < UNISENDER_RETRY_ATTEMPTS and _is_unisender_rate_limit_text(message):
            _sleep_sender_retry(UNISENDER_RATE_LIMIT_RETRY_SECONDS)
            continue
        raise RuntimeError(message)

    failed_emails = data.get("failed_emails") or {}
    accepted_emails = data.get("emails") or []
    if not accepted_emails:
        accepted_emails = [
            recipient
            for recipient in cleaned_recipients
            if not (isinstance(failed_emails, dict) and recipient in failed_emails)
        ]
    return {
        "provider": "unisender_go",
        "status": _safe_text(data.get("status")) or "success",
        "job_id": _safe_text(data.get("job_id") or data.get("id")),
        "recipients": accepted_emails,
        "accepted_emails": accepted_emails,
        "failed_emails": failed_emails if isinstance(failed_emails, dict) else {},
        "idempotence_key": idempotence_key,
    }


def _build_unisender_go_metadata(row: dict[str, Any], *, recipient: str, job_id: str | None) -> dict[str, str]:
    return {
        "app": "mailing-agent",
        "app_job_id": _safe_text(job_id),
        "app_row_id": _safe_text(row.get("ID")),
        "app_mun_name": _safe_text(row.get("MUN_NAME")),
        "app_recipient": _safe_text(recipient),
    }


def _send_with_transport(
    row: dict[str, Any],
    recipients: list[str],
    attachments: list[str],
    subject: str,
    *,
    transport: str,
    mail_template_path: Path | None = None,
    body_override: str | None = None,
    job_id: str | None = None,
    send_run_id: str = "",
    send_mode: str = "",
    attachment_mode: str = "",
    wait_between_recipients: Callable[[], bool] | None = None,
    wait_after_rate_limit: Callable[[float, str], bool] | None = None,
) -> dict[str, Any]:
    attempts: list[dict[str, Any]] = []
    sent_recipients: list[str] = []
    warnings: list[str] = []
    if transport == "unisender" and _uses_unisender_go_api():
        try:
            provider = _send_via_unisender_go_bulk(
                row,
                recipients,
                attachments,
                subject,
                mail_template_path=mail_template_path,
                body_override=body_override,
                job_id=job_id,
                send_run_id=send_run_id,
                send_mode=send_mode,
                attachment_mode=attachment_mode,
            )
        except Exception as exc:
            return {
                "recipient": None,
                "recipients": [],
                "attempts": [
                    {
                        "recipient": recipient,
                        "status": "error",
                        "error": _safe_text(exc) or "UniSender error",
                    }
                    for recipient in recipients
                ],
                "error": _safe_text(exc) or "UniSender error",
                "warning": "",
            }
        failed_emails = provider.get("failed_emails") if isinstance(provider.get("failed_emails"), dict) else {}
        accepted_emails = provider.get("accepted_emails") or provider.get("recipients") or []
        accepted_keys = {_mail_key(recipient) for recipient in accepted_emails}
        for recipient in recipients:
            if _mail_key(recipient) in accepted_keys:
                attempt = {"recipient": recipient, "status": "sent", "error": "", "provider": provider}
                if provider.get("job_id"):
                    attempt["provider_job_id"] = provider["job_id"]
                attempts.append(attempt)
                sent_recipients.append(recipient)
            else:
                attempts.append(
                    {
                        "recipient": recipient,
                        "status": "error",
                        "error": _safe_text(failed_emails.get(recipient)) or "UniSender Go не подтвердил адрес получателя.",
                    }
                )
        if sent_recipients and len(sent_recipients) == len(recipients):
            return {
                "recipient": sent_recipients[0],
                "recipients": sent_recipients,
                "attempts": attempts,
                "error": "",
                "warning": "",
            }
        failed_errors = [attempt["error"] for attempt in attempts if attempt.get("status") == "error" and attempt.get("error")]
        return {
            "recipient": None,
            "recipients": sent_recipients,
            "attempts": attempts,
            "error": "; ".join(failed_errors) or "UniSender Go не подтвердил отправку.",
            "warning": "",
        }

    attempted_recipients = 0
    for recipient in recipients:
        if attempted_recipients > 0 and wait_between_recipients is not None:
            if not wait_between_recipients():
                attempts.append(
                    {
                        "recipient": recipient,
                        "status": "error",
                        "error": "Отправка остановлена во время паузы между письмами.",
                    }
                )
                break
        attempted_recipients += 1
        try:
            warning = None
            provider: dict[str, Any] = {}
            if transport == "unisender":
                provider = _send_via_unisender(
                    row,
                    recipient,
                    attachments,
                    subject,
                    mail_template_path=mail_template_path,
                    body_override=body_override,
                    job_id=job_id,
                    send_run_id=send_run_id,
                    send_mode=send_mode,
                    attachment_mode=attachment_mode,
                )
            elif transport == "rusender":
                provider = _send_via_rusender(
                    row,
                    recipient,
                    attachments,
                    subject,
                    mail_template_path=mail_template_path,
                    body_override=body_override,
                    job_id=job_id,
                    send_run_id=send_run_id,
                    send_mode=send_mode,
                    attachment_mode=attachment_mode,
                )
            elif transport == "mailopost":
                while True:
                    try:
                        provider = _send_via_mailopost(
                            row,
                            recipient,
                            attachments,
                            subject,
                            mail_template_path=mail_template_path,
                            body_override=body_override,
                            job_id=job_id,
                            send_run_id=send_run_id,
                            send_mode=send_mode,
                            attachment_mode=attachment_mode,
                        )
                        break
                    except MailoPostRateLimitError as exc:
                        retry_after_seconds = max(1.0, float(exc.retry_after_seconds))
                        if wait_after_rate_limit is None:
                            raise
                        warnings.append(f"MailoPost ограничил скорость отправки. Ждали {int(retry_after_seconds)} сек.")
                        if wait_after_rate_limit(retry_after_seconds, _safe_text(exc)):
                            continue
                        attempts.append(
                            {
                                "recipient": recipient,
                                "status": "error",
                                "error": "Отправка остановлена во время ожидания лимита MailoPost.",
                            }
                        )
                        return {
                            "recipient": None,
                            "recipients": sent_recipients,
                            "attempts": attempts,
                            "error": "Отправка остановлена во время ожидания лимита MailoPost.",
                            "warning": " ".join(warnings).strip(),
                        }
            else:
                warning = _send_via_smtp(
                    row,
                    recipient,
                    attachments,
                    subject,
                    mail_template_path=mail_template_path,
                    body_override=body_override,
                )
        except Exception as exc:
            attempts.append({"recipient": recipient, "status": "error", "error": _safe_text(exc) or f"{transport} error"})
            continue
        attempt: dict[str, Any] = {"recipient": recipient, "status": "sent", "error": ""}
        safe_provider = _safe_provider_payload(provider)
        if safe_provider:
            attempt["provider"] = safe_provider
            if safe_provider.get("message_id"):
                attempt["provider_message_id"] = safe_provider["message_id"]
            if safe_provider.get("job_id"):
                attempt["provider_job_id"] = safe_provider["job_id"]
        attempts.append(attempt)
        sent_recipients.append(recipient)
        if warning:
            warnings.append(warning)
    if sent_recipients and len(sent_recipients) == len(recipients):
        return {
            "recipient": sent_recipients[0],
            "recipients": sent_recipients,
            "attempts": attempts,
            "error": "",
            "warning": " ".join(warnings).strip(),
        }
    failed_errors = [attempt["error"] for attempt in attempts if attempt.get("status") == "error" and attempt.get("error")]
    last_error = "; ".join(failed_errors) or "Не найден получатель для отправки."
    return {
        "recipient": None,
        "recipients": sent_recipients,
        "attempts": attempts,
        "error": last_error,
        "warning": " ".join(warnings).strip(),
    }


def _provider_for_recipient(attempts: list[dict[str, Any]], recipient: str) -> dict[str, Any]:
    recipient_key = _mail_key(recipient)
    for attempt in attempts:
        if _mail_key(attempt.get("recipient")) != recipient_key:
            continue
        provider = attempt.get("provider")
        return _safe_provider_payload(provider if isinstance(provider, dict) else None)
    return {}


def _unisender_parallel_workers(*, dry_run: bool, transport: str) -> int:
    if dry_run or transport != "unisender":
        return 1
    try:
        return max(1, int(settings.sender_unisender_concurrency or 1))
    except (TypeError, ValueError):
        return 1


def _build_parallel_send_job(
    *,
    row: dict[str, Any],
    entry: dict[str, Any],
    email_decision: dict[str, Any],
    recipients_to_send: list[str],
    attachments: list[str],
    subject: str,
    transport: str,
    mail_template_path: Path | None,
    job_id: str | None,
    send_run_id: str = "",
    send_mode: str = "",
    attachment_mode: str = "",
    subject_template: str = "",
    work_type: str = "",
    recipient_strategy: str = "",
    body_override: str | None = None,
    success_status_value: str = STATUS_SENT_VALUE,
) -> dict[str, Any]:
    return {
        "row": row,
        "entry": entry,
        "email_decision": email_decision,
        "recipients_to_send": recipients_to_send,
        "attachments": attachments,
        "subject": subject,
        "transport": transport,
        "mail_template_path": mail_template_path,
        "body_override": body_override,
        "success_status_value": success_status_value,
        "job_id": job_id,
        "send_run_id": send_run_id,
        "send_mode": send_mode,
        "attachment_mode": attachment_mode,
        "subject_template": subject_template,
        "work_type": work_type,
        "recipient_strategy": recipient_strategy,
    }


def _run_parallel_send_job(job: dict[str, Any]) -> dict[str, Any]:
    send_result = _send_with_transport(
        job["row"],
        job["recipients_to_send"],
        job["attachments"],
        job["subject"],
        transport=job["transport"],
        mail_template_path=job["mail_template_path"],
        body_override=job.get("body_override"),
        job_id=job.get("job_id"),
        send_run_id=job.get("send_run_id") or "",
        send_mode=job.get("send_mode") or "",
        attachment_mode=job.get("attachment_mode") or "",
    )
    return {"job": job, "send_result": send_result}


def _restore_sent_from_local_log(
    *,
    entry: dict[str, Any],
    intended_recipients: list[str],
    row: dict[str, Any],
    workbook: Any,
    worksheet: Any,
    data_xlsx_path: Path,
    success_status_value: str = STATUS_SENT_VALUE,
) -> bool:
    entry["result"] = "skipped_logged_sent"
    entry["attempts"] = [
        {
            "recipient": recipient,
            "status": "already_sent",
            "error": "",
        }
        for recipient in intended_recipients
    ]
    entry["sent_recipients"] = intended_recipients
    entry["recipient"] = intended_recipients[0] if intended_recipients else entry["recipient"]
    entry["decision_reason"] = "Статус восстановлен по локальному журналу: письмо уже было отправлено ранее."
    status_warning = _persist_row_status(
        workbook,
        worksheet,
        data_xlsx_path,
        row,
        success_status_value,
        flush=False,
    )
    if status_warning:
        entry["warning"] = f"{entry['warning']} {status_warning}".strip()
        entry["next_action"] = f"{entry['next_action']} {status_warning}".strip()
        return False
    return True


def _apply_send_result_to_entry(
    *,
    entry: dict[str, Any],
    send_result: dict[str, Any],
    row: dict[str, Any],
    email_decision: dict[str, Any],
    attachments: list[str],
    row_subject: str,
    effective_transport: str,
    effective_send_mode: str,
    effective_attachment_mode: str,
    effective_subject_template: str,
    effective_work_type: str,
    effective_recipient_strategy: str,
    sent_mail_log_path: Path | None,
    sent_mail_recipients: dict[str, set[str]],
    send_run_id: str,
    send_run_started_at: str,
    workbook: Any,
    worksheet: Any,
    data_xlsx_path: Path,
    success_status_value: str = STATUS_SENT_VALUE,
) -> bool:
    entry["attempts"] = send_result["attempts"]
    entry["warning"] = _safe_text(send_result.get("warning"))
    partial_recipients = send_result.get("recipients") or []
    row_id_text = _safe_text(row.get("ID"))

    if partial_recipients and not send_result["recipient"]:
        for sent_recipient in partial_recipients:
            log_warning = _append_sent_mail_log(
                row=row,
                recipient=sent_recipient,
                attachments=attachments,
                subject=row_subject,
                transport=effective_transport,
                warning=entry["warning"],
                provider=_provider_for_recipient(entry["attempts"], sent_recipient),
                sent_mail_log_path=sent_mail_log_path,
                send_run_id=send_run_id,
                send_run_started_at=send_run_started_at,
                send_mode=effective_send_mode,
                attachment_mode=effective_attachment_mode,
                subject_template=effective_subject_template,
                work_type=effective_work_type,
                recipient_strategy=effective_recipient_strategy,
            )
            sent_mail_recipients.setdefault(row_id_text, set()).add(_mail_key(sent_recipient))
            if log_warning:
                entry["warning"] = f"{entry['warning']} {log_warning}".strip() if entry["warning"] else log_warning

    if not send_result["recipient"]:
        raise RuntimeError(send_result["error"])

    entry["recipient"] = send_result["recipient"]
    entry["sent_recipients"] = send_result.get("recipients") or [entry["recipient"]]
    if len(entry["sent_recipients"]) > 1:
        entry["decision_reason"] = "Письмо отправлено на основной и дополнительный email."
    elif entry["email_strategy"] == "fallback_extra" or entry["recipient"] != email_decision["recipient"]:
        entry["decision_reason"] = "Письмо отправлено по резервному email после выбора лучшего доступного адреса."

    for sent_recipient in entry["sent_recipients"]:
        log_warning = _append_sent_mail_log(
            row=row,
            recipient=sent_recipient,
            attachments=attachments,
            subject=row_subject,
            transport=effective_transport,
            warning=entry["warning"],
            provider=_provider_for_recipient(entry["attempts"], sent_recipient),
            sent_mail_log_path=sent_mail_log_path,
            send_run_id=send_run_id,
            send_run_started_at=send_run_started_at,
            send_mode=effective_send_mode,
            attachment_mode=effective_attachment_mode,
            subject_template=effective_subject_template,
            work_type=effective_work_type,
            recipient_strategy=effective_recipient_strategy,
        )
        if log_warning:
            entry["warning"] = f"{entry['warning']} {log_warning}".strip() if entry["warning"] else log_warning
        sent_mail_recipients.setdefault(row_id_text, set()).add(_mail_key(sent_recipient))

    status_warning = _persist_row_status(
        workbook,
        worksheet,
        data_xlsx_path,
        row,
        success_status_value,
        flush=False,
    )
    if status_warning:
        entry["warning"] = f"{entry['warning']} {status_warning}".strip()
        entry["next_action"] = f"{entry['next_action']} {status_warning}".strip()
        return False
    return True


def run_sender(
    *,
    dry_run: bool = True,
    limit: int | None = None,
    auto_recover: bool = False,
    row_ids: list[str] | None = None,
    transport: str | None = None,
    send_mode: str | None = None,
    attachment_mode: str | None = None,
    subject_template: str | None = None,
    require_confirmed_consent: bool = False,
    work_type: str | None = None,
    recipient_strategy: str | None = None,
    job_id: str | None = None,
) -> dict[str, Any]:
    job_paths = resolve_job_paths(job_id)
    data_xlsx_path = _resolve_sender_data_xlsx_path(job_id)
    output_dir = None if job_paths.uses_legacy_layout else job_paths.output_dir
    job_template_docx_path = job_paths.templates_dir / "mail_template.docx"
    job_template_txt_path = job_paths.templates_dir / "mail_template.txt"
    if job_paths.uses_legacy_layout:
        mail_template_path = None
    elif job_template_docx_path.exists():
        mail_template_path = job_template_docx_path
    elif job_template_txt_path.exists():
        mail_template_path = job_template_txt_path
    else:
        mail_template_path = None
    sent_mail_log_path = None if job_paths.uses_legacy_layout else job_paths.sent_mail_log_path
    state = _load_sender_state(job_id)
    clear_sender_stop_request(job_id)
    state = _load_sender_state(job_id)
    effective_limit = _normalize_limit(limit, dry_run=dry_run)
    effective_transport = _normalize_transport(transport)
    effective_send_mode = _normalize_send_mode(send_mode)
    if effective_send_mode == "materials":
        require_confirmed_consent = True
    effective_attachment_mode = _normalize_attachment_mode(attachment_mode)
    effective_recipient_strategy = _normalize_recipient_strategy(recipient_strategy or state.get("recipient_strategy"))
    effective_work_type = normalize_work_type(work_type or state.get("work_type") or DEFAULT_WORK_TYPE)
    effective_subject_template = _safe_text(subject_template)
    if effective_work_type != DEFAULT_WORK_TYPE and effective_subject_template == DEFAULT_MAIL_SUBJECT:
        effective_subject_template = ""
    requested_row_ids = {str(item).strip() for item in (row_ids or []) if str(item).strip()}
    stats = _collect_excel_stats(data_xlsx_path)
    started_at = datetime.now().isoformat(timespec="seconds")
    send_run_id = ""
    send_run_started_at = ""
    if not dry_run:
        send_run_id = _safe_text(state.get("send_run_id"))
        send_run_started_at = _safe_text(state.get("send_run_started_at"))
        if _safe_text(state.get("send_mode")) != effective_send_mode:
            send_run_id = ""
            send_run_started_at = ""
        if not send_run_id:
            send_run_id = f"send-{started_at.replace(':', '').replace('-', '')}-{secrets.token_hex(4)}"
        if not send_run_started_at:
            send_run_started_at = started_at
    state.update(
        {
            "status": "running",
            "mode": "dry_run" if dry_run else "send",
            "send_mode": effective_send_mode,
            "started_at": started_at,
            "completed_at": None,
            "processed_rows": 0,
            "ready_rows": 0,
            "sent_rows": (
                0
                if effective_send_mode == "consent_request" or requested_row_ids or effective_limit
                else stats["sent"]
            ),
            "error_rows": 0,
            "skipped_rows": 0,
            "total_rows": 0,
            "summary_text": (
                "Начинаю проверку запроса согласия."
                if effective_send_mode == "consent_request"
                else "Начинаю проверку перед отправкой."
            ),
            "rows": [],
            "stats": stats,
            "warning_rows": 0,
            "handoff_rows": 0,
            "task_stats": count_tasks_for_agent("sender", job_id),
            "tasks": get_tasks_for_agent("sender", job_id)[:20],
            "recent_events": get_recent_events(agent_name="sender", limit=20, job_id=job_id),
            "generator_handoff_rows": 0,
            "philology_blocked_rows": 0,
            "autonomous_recovery_rows": 0,
            "effective_limit": effective_limit,
            "selection_scoped": bool(requested_row_ids or effective_limit),
            "requested_row_ids": sorted(requested_row_ids),
            "remaining_rows": 0,
            "stop_requested": False,
            "stop_requested_at": None,
            "transport": effective_transport,
            "attachment_mode": effective_attachment_mode,
            "recipient_strategy": effective_recipient_strategy,
            "work_type": effective_work_type,
            "subject_template": effective_subject_template,
        }
    )
    if not dry_run:
        state["send_run_id"] = send_run_id
        state["send_run_started_at"] = send_run_started_at
    _save_sender_state(state, job_id)

    if not data_xlsx_path.exists():
        state["status"] = "error"
        state["summary_text"] = "Файл data.xlsx не найден."
        _save_sender_state(state, job_id)
        return dict(state)

    workbook, worksheet, rows = load_rows(data_xlsx_path)
    if requested_row_ids:
        rows = [row for row in rows if str(row.get("ID")).strip() in requested_row_ids]
    candidates = rows[:effective_limit] if effective_limit else rows
    state["total_rows"] = len(candidates)
    _save_sender_state(state, job_id)
    candidate_row_ids = {
        _safe_text(row.get("ID"))
        for row in candidates
        if _safe_text(row.get("ID"))
    }
    output_folder_index, output_folder_errors = _build_output_folder_index(
        candidate_row_ids,
        output_dir=output_dir,
    )

    processed_entries: list[dict[str, Any]] = []
    runtime_warnings: list[str] = []
    workbook_dirty = False
    started_at = perf_counter()
    subject = (
        _consent_request_subject(effective_attachment_mode, effective_work_type)
        if effective_send_mode == "consent_request"
        else (effective_subject_template or _materials_subject(effective_attachment_mode, effective_work_type))
    )
    sent_mail_recipients = (
        _load_sent_mail_recipients(
            sent_mail_log_path,
            send_run_id=send_run_id,
            send_run_started_at=send_run_started_at,
        )
        if not dry_run
        else {}
    )
    parallel_workers = _unisender_parallel_workers(dry_run=dry_run, transport=effective_transport)
    parallel_send_jobs: list[dict[str, Any]] = []

    for row in candidates:
        _refresh_sender_stop_flag(state, job_id)
        if state.get("stop_requested"):
            break
        row_id = row.get("ID")
        row_status = _safe_text(row.get("STATUS"))
        entry: dict[str, Any] = {
            "id": row_id,
            "row_index": row.get("_row_index"),
            "mun_name": _safe_text(row.get("MUN_NAME")),
            "status_before": row_status,
            "result": "",
            "recipient": None,
            "emails": [],
            "invalid_emails": [],
            "email_strategy": "",
            "decision_reason": "",
            "fallback_candidates": [],
            "recipient_strategy": effective_recipient_strategy,
            "folder": None,
            "attachments": [],
            "error": "",
            "warning": "",
            "next_action": "",
            "attempts": [],
            "consent_urls": {},
        }

        status_class = _status_class(row_status)
        if status_class == "sent" and effective_send_mode == "materials":
            entry["result"] = "skipped_duplicate"
            entry["decision_reason"] = "Строка уже помечена как успешно отправленная."
            processed_entries.append(entry)
            state["skipped_rows"] += 1
            state["processed_rows"] += 1
            state["rows"] = _state_rows_snapshot(processed_entries)
            _save_sender_state(state, job_id)
            continue
        email_decision = _choose_recipient(row)
        entry["recipient"] = email_decision["recipient"]
        entry["emails"] = email_decision["valid_emails"]
        entry["invalid_emails"] = email_decision["invalid_emails"]
        entry["email_strategy"] = email_decision["strategy"]
        entry["decision_reason"] = email_decision["decision_reason"]
        entry["fallback_candidates"] = email_decision["fallback_candidates"]
        allowed_recipients = _allowed_send_recipients(
            email_decision,
            recipient_strategy=effective_recipient_strategy,
        )
        fallback_recipients = _fallback_send_recipients(
            email_decision,
            recipient_strategy=effective_recipient_strategy,
        )
        confirmed_material_recipients: list[str] = []
        if effective_send_mode == "materials" and require_confirmed_consent:
            consent_candidates = _consent_candidate_recipients(
                email_decision,
                recipient_strategy=effective_recipient_strategy,
            )
            confirmed_material_recipients = [
                recipient
                for recipient in consent_candidates
                if has_confirmed_consent(
                    job_id=job_id,
                    row_id=row_id,
                    recipient=recipient,
                    attachment_mode=effective_attachment_mode,
                )
            ]
        missing_confirmed_consent = (
            effective_send_mode == "materials"
            and require_confirmed_consent
            and bool(allowed_recipients)
            and not confirmed_material_recipients
        )
        folder: Path | None = None
        folder_error: str | None = None
        attachments: list[str] = []
        attachment_error: str | None = None
        review_task: dict[str, Any] | None = None
        if effective_send_mode == "materials" and not missing_confirmed_consent:
            folder, folder_error = _resolve_output_folder(
                row_id,
                output_dir=output_dir,
                folder_index=output_folder_index,
                folder_errors=output_folder_errors,
            )
            entry["folder"] = str(folder) if folder else None
            attachments, attachment_error = _resolve_pdf_attachments(
                folder,
                attachment_mode=effective_attachment_mode,
            )
            entry["attachments"] = attachments
            review_task = _active_sender_review_task(row_id, job_id=job_id)
        else:
            entry["folder"] = None
            entry["attachments"] = []
        recovery_info: dict[str, Any] | None = None

        if not entry["recipient"]:
            entry["result"] = "needs_enrichment"
            entry["error"] = "Не найден валидный email получателя."
            entry["next_action"] = (
                "Нужно вручную проверить и заполнить email получателя."
                if not settings.inter_agent_handoffs_enabled
                else "Запросить у агента-парсера поиск или уточнение email."
            )
            task = _delegate_sender_problem(
                symptom="missing_recipient_data",
                row_id=row_id,
                mun_name=entry["mun_name"],
                details={
                    "email_osn": _safe_text(row.get("EMAIL_OSN")),
                    "email_dop": _safe_text(row.get("EMAIL_DOP")),
                    "reason": entry["error"],
                    "has_primary_email": bool(_parse_emails(row.get("EMAIL_OSN"))),
                    "has_any_valid_email": bool(email_decision["valid_emails"]),
                },
                job_id=job_id,
            )
            if task:
                entry["handoff_task_id"] = task.get("id")
                state["handoff_rows"] += 1
        elif folder_error:
            entry["result"] = "error_missing_output"
            entry["error"] = folder_error
            entry["next_action"] = (
                "Нужно вручную перезапустить генератор и пересобрать комплект документов."
                if not settings.inter_agent_handoffs_enabled
                else "Передать генератору задачу на пересборку комплекта документов."
            )
            task = _delegate_sender_problem(
                symptom="missing_output",
                row_id=row_id,
                mun_name=entry["mun_name"],
                details={
                    "reason": folder_error,
                    "folder_exists": False,
                    "attachment_count": 0,
                },
                job_id=job_id,
            )
            if task:
                entry["handoff_task_id"] = task.get("id")
                state["generator_handoff_rows"] += 1
            if auto_recover:
                recovery_info = _run_autonomous_recovery_for_generator(
                    row_id=row_id,
                    work_type=effective_work_type,
                    job_id=job_id,
                )
                folder, folder_error, attachments, attachment_error = _retry_row_resources(
                    row_id,
                    output_dir=output_dir,
                    attachment_mode=effective_attachment_mode,
                )
                entry["folder"] = str(folder) if folder else None
                entry["attachments"] = attachments
                if not folder_error and not attachment_error and entry["recipient"]:
                    entry["result"] = "ready_after_recovery" if dry_run else "sent"
                    entry["decision_reason"] += " Генератор автоматически пересобрал комплект документов."
                    state["autonomous_recovery_rows"] += 1
        elif missing_confirmed_consent:
            entry["result"] = "blocked_no_consent"
            entry["error"] = "Нет подтверждённого согласия на отправку КП и договора."
            entry["next_action"] = "Сначала отправьте запрос согласия и дождитесь подтверждения."
        elif attachment_error:
            entry["result"] = "error_missing_attachments"
            entry["error"] = attachment_error
            entry["next_action"] = (
                "Нужно вручную перезапустить генератор и восстановить вложения."
                if not settings.inter_agent_handoffs_enabled
                else "Передать генератору задачу на восстановление вложений."
            )
            task = _delegate_sender_problem(
                symptom="missing_attachments",
                row_id=row_id,
                mun_name=entry["mun_name"],
                details={
                    "reason": attachment_error,
                    "folder": entry["folder"],
                    "folder_exists": bool(folder),
                    "attachment_count": len(attachments),
                },
                job_id=job_id,
            )
            if task:
                entry["handoff_task_id"] = task.get("id")
                state["generator_handoff_rows"] += 1
            if auto_recover:
                recovery_info = _run_autonomous_recovery_for_generator(
                    row_id=row_id,
                    work_type=effective_work_type,
                    job_id=job_id,
                )
                folder, folder_error, attachments, attachment_error = _retry_row_resources(
                    row_id,
                    output_dir=output_dir,
                    attachment_mode=effective_attachment_mode,
                )
                entry["folder"] = str(folder) if folder else None
                entry["attachments"] = attachments
                if not folder_error and not attachment_error and entry["recipient"]:
                    entry["result"] = "ready_after_recovery" if dry_run else "sent"
                    entry["decision_reason"] += " Генератор автоматически восстановил вложения."
                    state["autonomous_recovery_rows"] += 1
        elif review_task:
            entry["result"] = "blocked_by_philologist"
            entry["error"] = _safe_text(review_task.get("details", {}).get("note")) or (
                "Перед отправкой нужно учесть замечания филолога."
            )
            diagnosis = diagnose_responsibility(
                symptom="philology_review_block",
                context={
                    "unresolved_issue_count": review_task.get("details", {}).get("unresolved_issue_count", 1),
                },
            )
            entry["next_action"] = (
                "Сначала завершить языковую проверку и согласовать замечания филолога. "
                + diagnosis["root_cause"]
            )
            state["philology_blocked_rows"] += 1
            if auto_recover and settings.philologist_auto_run_enabled:
                recovery_info = _run_autonomous_recovery_for_philologist(row_id=row_id, job_id=job_id)
                review_task = _active_sender_review_task(row_id, job_id=job_id)
                if not review_task:
                    entry["result"] = "ready_after_recovery" if dry_run else "sent"
                    entry["error"] = ""
                    entry["next_action"] = ""
                    entry["decision_reason"] += " Филолог автоматически перепроверил документы и снял блокер."
                    state["autonomous_recovery_rows"] += 1
            elif auto_recover:
                entry["next_action"] = (
                    "Нужна ручная проверка филолога. Автоматический запуск филолога сейчас отключён."
                )
        else:
            entry["result"] = "ready" if dry_run else "sent"

        if recovery_info:
            entry["recovery"] = {
                key: {
                    "summary_text": value.get("summary_text"),
                    "status": value.get("status"),
                }
                for key, value in recovery_info.items()
                if isinstance(value, dict)
            }

        if entry["result"] in {
            "needs_enrichment",
            "error_missing_output",
            "error_missing_attachments",
            "blocked_by_philologist",
            "blocked_no_consent",
            "error",
        }:
            if not dry_run:
                status_warning = _persist_row_status(
                    workbook,
                    worksheet,
                    data_xlsx_path,
                    row,
                    _format_error_status(entry["error"] or entry["result"]),
                    flush=False,
                )
                workbook_dirty = workbook_dirty or not bool(status_warning)
                if status_warning:
                    entry["warning"] = f"{entry['warning']} {status_warning}".strip()
            state["error_rows"] += 1
            processed_entries.append(entry)
            state["processed_rows"] += 1
            if _should_flush_sender_workbook(
                dirty=workbook_dirty,
                processed_rows=state["processed_rows"],
                total_rows=state["total_rows"],
            ):
                flush_warning = _flush_sender_workbook(workbook, data_xlsx_path)
                if flush_warning:
                    entry["warning"] = f"{entry['warning']} {flush_warning}".strip()
                    entry["next_action"] = f"{entry['next_action']} {flush_warning}".strip()
                    runtime_warnings.append(flush_warning)
                else:
                    workbook_dirty = False
            state["rows"] = _state_rows_snapshot(processed_entries)
            _save_sender_state(state, job_id)
            continue

        state["ready_rows"] += 1
        row_subject = _build_mail_subject(subject, row)
        row_body_override: str | None = None
        smtp_recipient_delay: Callable[[], bool] | None = None
        if not dry_run and effective_transport == "smtp":
            delay_seconds = max(0.0, float(settings.sender_delay_seconds or 0))
            if delay_seconds > 0:
                smtp_recipient_delay = lambda: _wait_sender_delay(delay_seconds, state, job_id)
        mailopost_rate_limit_delay: Callable[[float, str], bool] | None = None
        if not dry_run and effective_transport == "mailopost":
            def mailopost_rate_limit_delay(wait_seconds: float, message: str) -> bool:
                delay_seconds = max(1.0, float(wait_seconds))
                state["summary_text"] = (
                    "MailoPost ограничил скорость отправки. "
                    f"Жду {int(delay_seconds)} сек. Потом продолжу с того же письма."
                )
                state["mailopost_rate_limited_at"] = datetime.now().isoformat(timespec="seconds")
                state["mailopost_retry_after_seconds"] = int(delay_seconds)
                state["mailopost_rate_limit_message"] = _safe_text(message)
                _save_sender_state(state, job_id)
                return _wait_sender_delay(delay_seconds, state, job_id)
        success_status_value = (
            STATUS_CONSENT_REQUEST_SENT_VALUE
            if effective_send_mode == "consent_request"
            else STATUS_SENT_VALUE
        )

        if not dry_run:
            _refresh_sender_stop_flag(state, job_id)
            if state.get("stop_requested"):
                entry["result"] = "stopped_before_send"
                entry["next_action"] = "Отправка этой и следующих строк остановлена по запросу пользователя."
                state["ready_rows"] -= 1
                processed_entries.append(entry)
                state["processed_rows"] += 1
                state["rows"] = _state_rows_snapshot(processed_entries)
                _save_sender_state(state, job_id)
                break
            try:
                row_id_text = _safe_text(row_id)
                already_logged = sent_mail_recipients.get(row_id_text, set())
                intended_recipients = list(allowed_recipients)
                if effective_send_mode == "materials" and require_confirmed_consent:
                    intended_recipients = list(confirmed_material_recipients)
                recipients_to_send = [
                    recipient
                    for recipient in intended_recipients
                    if _mail_key(recipient) not in already_logged
                ]
                fallback_recipients_to_send = [
                    recipient
                    for recipient in fallback_recipients
                    if _mail_key(recipient) not in already_logged
                ]
                attachments_to_send = [] if effective_send_mode == "consent_request" else attachments
                has_deferred_fallback_recipients = (
                    effective_recipient_strategy == RECIPIENT_STRATEGY_PRIMARY_THEN_FALLBACK
                    and bool(fallback_recipients_to_send)
                )
                if not recipients_to_send and not has_deferred_fallback_recipients:
                    workbook_dirty = (
                        _restore_sent_from_local_log(
                            entry=entry,
                            intended_recipients=intended_recipients,
                            row=row,
                            workbook=workbook,
                            worksheet=worksheet,
                            data_xlsx_path=data_xlsx_path,
                            success_status_value=success_status_value,
                        )
                        or workbook_dirty
                    )
                    state["sent_rows"] += 1
                elif (
                    parallel_workers > 1
                    and effective_send_mode != "consent_request"
                    and effective_recipient_strategy != RECIPIENT_STRATEGY_PRIMARY_THEN_FALLBACK
                ):
                    parallel_send_jobs.append(
                        _build_parallel_send_job(
                            row=row,
                            entry=entry,
                            email_decision=email_decision,
                            recipients_to_send=recipients_to_send,
                            attachments=attachments_to_send,
                            subject=row_subject,
                            transport=effective_transport,
                            mail_template_path=mail_template_path,
                            body_override=row_body_override,
                            success_status_value=success_status_value,
                            job_id=job_id,
                            send_run_id=send_run_id,
                            send_mode=effective_send_mode,
                            attachment_mode=effective_attachment_mode,
                            subject_template=effective_subject_template,
                            work_type=effective_work_type,
                            recipient_strategy=effective_recipient_strategy,
                        )
                    )
                    entry["result"] = "queued_parallel_send"
                    entry["next_action"] = "Письмо поставлено в очередь на параллельную отправку через UniSender."
                else:
                    def send_one_recipient(recipient: str) -> dict[str, Any]:
                        if effective_send_mode == "consent_request":
                            return _send_consent_requests_with_transport(
                                row,
                                [recipient],
                                row_subject,
                                transport=effective_transport,
                                mail_template_path=mail_template_path,
                                job_id=job_id,
                                send_run_id=send_run_id,
                                attachment_mode=effective_attachment_mode,
                                subject_template=effective_subject_template,
                                work_type=effective_work_type,
                                recipient_strategy=effective_recipient_strategy,
                                wait_after_rate_limit=mailopost_rate_limit_delay,
                            )
                        return _send_with_transport(
                            row,
                            [recipient],
                            attachments_to_send,
                            row_subject,
                            transport=effective_transport,
                            mail_template_path=mail_template_path,
                            body_override=row_body_override,
                            job_id=job_id,
                            send_run_id=send_run_id,
                            send_mode=effective_send_mode,
                            attachment_mode=effective_attachment_mode,
                            wait_after_rate_limit=mailopost_rate_limit_delay,
                        )

                    if (
                        effective_recipient_strategy == RECIPIENT_STRATEGY_PRIMARY_THEN_FALLBACK
                        and fallback_recipients_to_send
                    ):
                        send_result = _send_recipient_sequence_until_success(
                            _unique_send_recipients([*recipients_to_send, *fallback_recipients_to_send]),
                            send_one_recipient,
                            wait_between_recipients=smtp_recipient_delay,
                        )
                    elif effective_send_mode == "consent_request":
                        send_result = _send_consent_requests_with_transport(
                            row,
                            recipients_to_send,
                            row_subject,
                            transport=effective_transport,
                            mail_template_path=mail_template_path,
                            job_id=job_id,
                            send_run_id=send_run_id,
                            attachment_mode=effective_attachment_mode,
                            subject_template=effective_subject_template,
                            work_type=effective_work_type,
                            recipient_strategy=effective_recipient_strategy,
                            wait_after_rate_limit=mailopost_rate_limit_delay,
                            wait_between_recipients=smtp_recipient_delay,
                        )
                    else:
                        send_result = _send_with_transport(
                            row,
                            recipients_to_send,
                            attachments_to_send,
                            row_subject,
                            transport=effective_transport,
                            mail_template_path=mail_template_path,
                            body_override=row_body_override,
                            job_id=job_id,
                            send_run_id=send_run_id,
                            send_mode=effective_send_mode,
                            attachment_mode=effective_attachment_mode,
                            wait_between_recipients=smtp_recipient_delay,
                            wait_after_rate_limit=mailopost_rate_limit_delay,
                        )
                    if effective_send_mode == "consent_request":
                        entry["consent_urls"] = send_result.get("consent_urls") or {}
                        if len(entry["consent_urls"]) == 1:
                            entry["consent_url"] = next(iter(entry["consent_urls"].values()))
                    workbook_dirty = (
                        _apply_send_result_to_entry(
                            entry=entry,
                            send_result=send_result,
                            row=row,
                            email_decision=email_decision,
                            attachments=attachments_to_send,
                            row_subject=row_subject,
                            effective_transport=effective_transport,
                            effective_send_mode=effective_send_mode,
                            effective_attachment_mode=effective_attachment_mode,
                            effective_subject_template=effective_subject_template,
                            effective_work_type=effective_work_type,
                            effective_recipient_strategy=effective_recipient_strategy,
                            sent_mail_log_path=sent_mail_log_path,
                            sent_mail_recipients=sent_mail_recipients,
                            send_run_id=send_run_id,
                            send_run_started_at=send_run_started_at,
                            workbook=workbook,
                            worksheet=worksheet,
                            data_xlsx_path=data_xlsx_path,
                            success_status_value=success_status_value,
                        )
                        or workbook_dirty
                    )
                    if effective_send_mode == "consent_request" and entry.get("sent_recipients"):
                        for sent_recipient in entry["sent_recipients"]:
                            mark_consent_request_sent(
                                job_id=job_id,
                                row_id=row_id,
                                recipient=sent_recipient,
                                provider=_provider_for_recipient(entry.get("attempts") or [], sent_recipient),
                                attachment_mode=effective_attachment_mode,
                            )
                    state["sent_rows"] += 1
                if entry["warning"]:
                    state["warning_rows"] += 1
                    entry["next_action"] = entry["warning"]
            except Exception as exc:
                entry["result"] = "error_send"
                entry["error"] = _safe_text(exc) or "Ошибка SMTP-отправки."
                entry["next_action"] = "Повторить отправку позже или передать строку на ручную проверку."
                state["ready_rows"] -= 1
                state["error_rows"] += 1
                status_warning = _persist_row_status(
                    workbook,
                    worksheet,
                    data_xlsx_path,
                    row,
                    _format_error_status(entry["error"]),
                    flush=False,
                )
                workbook_dirty = workbook_dirty or not bool(status_warning)
                if status_warning:
                    entry["warning"] = f"{entry['warning']} {status_warning}".strip()
                    entry["next_action"] = f"{entry['next_action']} {status_warning}".strip()
        else:
            preview_recipients = list(allowed_recipients)
            if effective_send_mode == "materials" and require_confirmed_consent:
                preview_recipients = list(confirmed_material_recipients)
            entry["attempts"] = [
                {
                    "recipient": recipient,
                    "status": "consent_request_ready" if effective_send_mode == "consent_request" else "ready",
                    "error": "",
                }
                for recipient in preview_recipients
            ]
        if entry.get("result") == "queued_parallel_send":
            state["rows"] = _state_rows_snapshot(processed_entries)
            _save_sender_state(state, job_id)
            continue

        processed_entries.append(entry)
        state["processed_rows"] += 1
        if _should_flush_sender_workbook(
            dirty=workbook_dirty,
            processed_rows=state["processed_rows"],
            total_rows=state["total_rows"],
        ):
            flush_warning = _flush_sender_workbook(workbook, data_xlsx_path)
            if flush_warning:
                entry["warning"] = f"{entry['warning']} {flush_warning}".strip()
                entry["next_action"] = f"{entry['next_action']} {flush_warning}".strip()
                runtime_warnings.append(flush_warning)
            else:
                workbook_dirty = False
        state["rows"] = _state_rows_snapshot(processed_entries)
        _save_sender_state(state, job_id)

        if not dry_run and entry.get("result") == "sent" and effective_transport == "smtp":
            delay_seconds = max(0.0, float(settings.sender_delay_seconds or 0))
            if delay_seconds > 0 and state["processed_rows"] < state["total_rows"]:
                if not _wait_sender_delay(delay_seconds, state, job_id):
                    state["summary_text"] = (
                        "Отправка остановлена пользователем во время паузы между письмами."
                    )
                    _save_sender_state(state, job_id)
                    break

    if parallel_send_jobs and not dry_run:
        max_workers = parallel_workers
        with ThreadPoolExecutor(max_workers=max_workers, thread_name_prefix="unisender-send") as executor:
            pending_futures: dict[Future, dict[str, Any]] = {}
            next_job_index = 0
            stop_submitting = False

            while next_job_index < len(parallel_send_jobs) and len(pending_futures) < max_workers:
                job = parallel_send_jobs[next_job_index]
                pending_futures[executor.submit(_run_parallel_send_job, job)] = job
                next_job_index += 1

            while pending_futures:
                _refresh_sender_stop_flag(state, job_id)
                if state.get("stop_requested"):
                    stop_submitting = True
                done_futures, _ = wait(list(pending_futures.keys()), return_when=FIRST_COMPLETED)
                for future in done_futures:
                    job = pending_futures.pop(future)
                    entry = job["entry"]
                    row = job["row"]
                    try:
                        outcome = future.result()
                        send_result = outcome["send_result"]
                        workbook_dirty = (
                            _apply_send_result_to_entry(
                                entry=entry,
                                send_result=send_result,
                                row=row,
                                email_decision=job["email_decision"],
                                attachments=job["attachments"],
                                row_subject=job["subject"],
                                effective_transport=effective_transport,
                                effective_send_mode=job.get("send_mode") or effective_send_mode,
                                effective_attachment_mode=job.get("attachment_mode") or effective_attachment_mode,
                                effective_subject_template=job.get("subject_template") or effective_subject_template,
                                effective_work_type=job.get("work_type") or effective_work_type,
                                effective_recipient_strategy=job.get("recipient_strategy") or effective_recipient_strategy,
                                sent_mail_log_path=sent_mail_log_path,
                                sent_mail_recipients=sent_mail_recipients,
                                send_run_id=send_run_id,
                                send_run_started_at=send_run_started_at,
                                workbook=workbook,
                                worksheet=worksheet,
                                data_xlsx_path=data_xlsx_path,
                                success_status_value=job.get("success_status_value", STATUS_SENT_VALUE),
                            )
                            or workbook_dirty
                        )
                        state["sent_rows"] += 1
                        if job.get("success_status_value") == STATUS_CONSENT_REQUEST_SENT_VALUE and entry.get("sent_recipients"):
                            for sent_recipient in entry["sent_recipients"]:
                                mark_consent_request_sent(
                                    job_id=job_id,
                                    row_id=row.get("ID"),
                                    recipient=sent_recipient,
                                    provider=_provider_for_recipient(entry.get("attempts") or [], sent_recipient),
                                    attachment_mode=effective_attachment_mode,
                                )
                        if entry["warning"]:
                            state["warning_rows"] += 1
                            entry["next_action"] = entry["warning"]
                        entry["result"] = "sent"
                    except Exception as exc:
                        entry["result"] = "error_send"
                        entry["error"] = _safe_text(exc) or "Ошибка UniSender-отправки."
                        entry["next_action"] = "Повторить отправку позже или передать строку на ручную проверку."
                        state["ready_rows"] -= 1
                        state["error_rows"] += 1
                        status_warning = _persist_row_status(
                            workbook,
                            worksheet,
                            data_xlsx_path,
                            row,
                            _format_error_status(entry["error"]),
                            flush=False,
                        )
                        workbook_dirty = workbook_dirty or not bool(status_warning)
                        if status_warning:
                            entry["warning"] = f"{entry['warning']} {status_warning}".strip()
                            entry["next_action"] = f"{entry['next_action']} {status_warning}".strip()

                    processed_entries.append(entry)
                    state["processed_rows"] += 1
                    if _should_flush_sender_workbook(
                        dirty=workbook_dirty,
                        processed_rows=state["processed_rows"],
                        total_rows=state["total_rows"],
                    ):
                        flush_warning = _flush_sender_workbook(workbook, data_xlsx_path)
                        if flush_warning:
                            entry["warning"] = f"{entry['warning']} {flush_warning}".strip()
                            entry["next_action"] = f"{entry['next_action']} {flush_warning}".strip()
                            runtime_warnings.append(flush_warning)
                        else:
                            workbook_dirty = False
                    state["rows"] = _state_rows_snapshot(processed_entries)
                    _save_sender_state(state, job_id)

                while not stop_submitting and next_job_index < len(parallel_send_jobs) and len(pending_futures) < max_workers:
                    job = parallel_send_jobs[next_job_index]
                    pending_futures[executor.submit(_run_parallel_send_job, job)] = job
                    next_job_index += 1

                if stop_submitting and next_job_index < len(parallel_send_jobs):
                    remaining_jobs = parallel_send_jobs[next_job_index:]
                    for job in remaining_jobs:
                        entry = job["entry"]
                        entry["result"] = "stopped_before_send"
                        entry["next_action"] = "Отправка этой и следующих строк остановлена по запросу пользователя."
                        state["ready_rows"] -= 1
                        processed_entries.append(entry)
                        state["processed_rows"] += 1
                    state["rows"] = _state_rows_snapshot(processed_entries)
                    _save_sender_state(state, job_id)
                    next_job_index = len(parallel_send_jobs)

    if not dry_run:
        if workbook_dirty:
            final_flush_warning = _flush_sender_workbook(workbook, data_xlsx_path)
            if final_flush_warning:
                runtime_warnings.append(final_flush_warning)
        close = getattr(workbook, "close", None)
        if callable(close):
            close()
        state["stats"] = _collect_excel_stats(data_xlsx_path)
        if state.get("selection_scoped") or effective_send_mode == "consent_request":
            state["remaining_rows"] = max(0, int(state.get("total_rows") or 0) - int(state.get("processed_rows") or 0))
        else:
            state["remaining_rows"] = int(state["stats"].get("pending", 0)) + int(state["stats"].get("error", 0))
    else:
        state["remaining_rows"] = max(0, len(rows) - len(candidates))

    state["rows"] = _state_rows_snapshot(processed_entries)
    state["completed_at"] = datetime.now().isoformat(timespec="seconds")
    state["elapsed_seconds"] = round(perf_counter() - started_at, 2)
    state["status"] = "stopped" if state.get("stop_requested") else "completed"
    state["task_stats"] = count_tasks_for_agent("sender", job_id)
    state["tasks"] = get_tasks_for_agent("sender", job_id)[:20]
    state["recent_events"] = get_recent_events(agent_name="sender", limit=20, job_id=job_id)
    state["summary_text"] = _format_sender_summary(state)
    if runtime_warnings:
        unique_warnings = list(dict.fromkeys(item for item in runtime_warnings if item))
        if unique_warnings:
            state["summary_text"] = f"{state['summary_text']} {' '.join(unique_warnings)}".strip()
    _save_sender_state(state, job_id)
    close = getattr(workbook, "close", None)
    if callable(close):
        close()
    return dict(state)


def get_sender_status(job_id: str | None = None) -> dict[str, Any]:
    state = _load_sender_state(job_id)
    data_xlsx_path = _resolve_sender_data_xlsx_path(job_id)
    if state.get("status") == "running" and isinstance(state.get("stats"), dict):
        state["stats"] = dict(state.get("stats") or {})
    else:
        state["stats"] = _collect_excel_stats(data_xlsx_path)
    state["task_stats"] = count_tasks_for_agent("sender", job_id)
    state["tasks"] = get_tasks_for_agent("sender", job_id)[:20]
    state["recent_events"] = get_recent_events(agent_name="sender", limit=20, job_id=job_id)
    return state


def _fallback_sender_chat(message: str, state: dict[str, Any], *, job_id: str | None = None) -> str:
    rows = state.get("rows") or []
    preview = preview_recipients(limit=10, job_id=job_id)
    tasks = state.get("tasks") or []
    recent_events = state.get("recent_events") or []
    rag_docs = find_relevant_service_docs(message, limit=1)
    rag_hint = f"\nСправка: {rag_docs[0].get('answer')}" if rag_docs else ""
    if not rows:
        data_xlsx_path = _resolve_sender_data_xlsx_path(job_id)
        stats = state.get("stats") or _collect_excel_stats(data_xlsx_path)
        extra = ""
        if tasks:
            extra = (
                "\nВнутренние задачи между агентами:\n"
                + "\n".join(
                    f"- {item.get('source_agent')} -> sender: {item.get('task_type')} "
                    f"для строки {item.get('row_id')} ({item.get('status')})"
                    for item in tasks[:5]
                )
            )
        if recent_events:
            extra += (
                "\nПоследние внутренние сообщения:\n"
                + "\n".join(
                    f"- {item.get('source_agent')} -> {item.get('target_agent') or 'system'}: {item.get('message')}"
                    for item in recent_events[-3:]
                )
            )
        return (
            f"Рассылка ещё не запускалась. Сейчас в Excel: всего {stats['total']}, "
            f"отправлено {stats['sent']}, ошибок {stats['error']}, ожидают {stats['pending']}.\n"
            f"{preview.get('summary_text')}\n"
            f"{_format_preview_rows(preview.get('rows') or [], limit=5)}"
            f"{extra}"
            f"{rag_hint}"
        )
    return (
        (state.get("summary_text") or "Статус отправщика пока недоступен.")
        + "\nПоследние обработанные строки:\n"
        + _format_preview_rows(
            [
                {
                    "id": item.get("id"),
                    "mun_name": item.get("mun_name"),
                    "recipient": item.get("recipient"),
                    "decision_reason": item.get("decision_reason"),
                    "email_strategy": item.get("email_strategy"),
                }
                for item in rows[:5]
            ],
            limit=5,
        )
        + (
            "\nПоследние внутренние сообщения:\n"
            + "\n".join(
                f"- {item.get('source_agent')} -> {item.get('target_agent') or 'system'}: {item.get('message')}"
                for item in recent_events[-3:]
            )
            if recent_events else ""
        )
        + rag_hint
    )


def _build_openai_client() -> OpenAI | None:
    api_key = _resolve_openai_api_key()
    if not api_key:
        return None
    base_url = _resolve_openai_base_url()
    kwargs = {"api_key": api_key}
    if base_url:
        kwargs["base_url"] = base_url
    try:
        return OpenAI(**kwargs)
    except Exception:
        return None


def chat_with_sender(message: str, *, job_id: str | None = None) -> dict[str, Any]:
    state = get_sender_status(job_id)
    client = _build_openai_client()
    preview = preview_recipients(limit=30, job_id=job_id)
    if not client:
        return {"reply": _fallback_sender_chat(message, state, job_id=job_id), "state": state}

    compact_rows = []
    for item in (state.get("rows") or [])[:20]:
        compact_rows.append(
            {
                "id": item.get("id"),
                "mun_name": item.get("mun_name"),
                "result": item.get("result"),
                "recipient": item.get("recipient"),
                "error": item.get("error"),
            }
        )
    rag_context = format_service_rag_context(find_relevant_service_docs(message, limit=3))

    prompt = (
        "Ты агент-отправщик писем с выбранными вложениями: КП, договором или КП и договором. "
        "Отвечай кратко, по-русски, только на основе текущего состояния запуска и предпросмотра адресов из data.xlsx. "
        "Если пользователь спрашивает про адреса или почты до рассылки, опирайся на предпросмотр, а не проси запускать отправку. "
        "Не выдумывай информацию, которой нет в данных. "
        "Если справка RAG противоречит состоянию запуска, главным источником правды является состояние запуска.\n\n"
        f"Справка RAG по сервису:\n{rag_context}\n\n"
        f"Состояние последнего запуска:\n{json.dumps({'summary_text': state.get('summary_text'), 'stats': state.get('stats'), 'rows': compact_rows, 'task_stats': state.get('task_stats'), 'tasks': (state.get('tasks') or [])[:10], 'recent_events': (state.get('recent_events') or [])[:10]}, ensure_ascii=False, indent=2)}\n\n"
        f"Предпросмотр адресов из data.xlsx:\n{json.dumps(preview, ensure_ascii=False, indent=2)}\n\n"
        f"Вопрос пользователя:\n{message}"
    )

    request_kwargs = {
        "model": settings.case_agent_model,
        "messages": [{"role": "user", "content": prompt}],
    }
    if not _resolve_openai_base_url():
        request_kwargs["response_format"] = {"type": "text"}

    try:
        response = client.chat.completions.create(**request_kwargs)
        reply = _safe_text(response.choices[0].message.content)
        if not reply:
            reply = _fallback_sender_chat(message, state, job_id=job_id)
    except Exception:
        reply = _fallback_sender_chat(message, state, job_id=job_id)

    return {"reply": reply, "state": state}

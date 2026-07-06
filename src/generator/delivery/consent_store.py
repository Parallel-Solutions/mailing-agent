from __future__ import annotations

from contextlib import contextmanager
import hashlib
import os
import re
import secrets
import threading
import time
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

from docx import Document

from src.jobs import resolve_job_paths
from src.jobs.json_store import read_json, write_json_atomic
from src.jobs.access import read_job_owner
from src.utils.config import settings
from src.generator.generation.work_types import DEFAULT_WORK_TYPE, normalize_work_type


CONSENT_FILENAME = "consents.json"
CONSENT_TEXT = "Согласен получить коммерческое предложение и проект договора от ООО «Параллельные Решения»."
CONSENT_OPERATOR_NAME = "ООО «Параллельные Решения»"
CONSENT_OPERATOR_INN = "5038110107"
ATTACHMENT_MODE_KP = "kp"
ATTACHMENT_MODE_CONTRACT = "contract"
ATTACHMENT_MODE_BOTH = "both"
CONSENT_FILE_LOCK_TIMEOUT_SECONDS = 10.0
CONSENT_FILE_LOCK_STALE_SECONDS = 300.0
DEFAULT_CONSENT_TOKEN_TTL_HOURS = 720
MATERIALS_DISPATCH_STALE_RETRY_SECONDS = 10 * 60

_CONSENT_THREAD_LOCKS: dict[str, threading.RLock] = {}
_CONSENT_THREAD_LOCKS_GUARD = threading.Lock()


def _safe_text(value: Any) -> str:
    return str(value or "").strip()


def _normalize_attachment_mode(value: Any) -> str:
    mode = _safe_text(value).lower()
    if mode in {ATTACHMENT_MODE_KP, ATTACHMENT_MODE_CONTRACT, ATTACHMENT_MODE_BOTH}:
        return mode
    return ATTACHMENT_MODE_KP


def _consent_text_for_attachment_mode(attachment_mode: Any) -> str:
    mode = _normalize_attachment_mode(attachment_mode)
    if mode == ATTACHMENT_MODE_CONTRACT:
        return "Согласен получить проект договора от ООО «Параллельные Решения»."
    if mode == ATTACHMENT_MODE_BOTH:
        return CONSENT_TEXT
    return "Согласен получить коммерческое предложение от ООО «Параллельные Решения»."


def _materials_list_for_attachment_mode(attachment_mode: Any) -> str:
    mode = _normalize_attachment_mode(attachment_mode)
    if mode == ATTACHMENT_MODE_CONTRACT:
        return "проекта договора и сопутствующих материалов"
    if mode == ATTACHMENT_MODE_BOTH:
        return "коммерческого предложения, проекта договора, технического задания, календарного плана и иных материалов"
    return "коммерческого предложения и сопутствующих материалов"


def _consent_path(job_id: str | None) -> Path:
    return resolve_job_paths(job_id).root_dir / "state" / CONSENT_FILENAME


def _consent_documents_dir(job_id: str | None) -> Path:
    return resolve_job_paths(job_id).consents_dir


def _records_thread_lock(path: Path) -> threading.RLock:
    key = str(path.resolve(strict=False))
    with _CONSENT_THREAD_LOCKS_GUARD:
        lock = _CONSENT_THREAD_LOCKS.get(key)
        if lock is None:
            lock = threading.RLock()
            _CONSENT_THREAD_LOCKS[key] = lock
        return lock


def _consent_file_lock_is_stale(lock_path: Path) -> bool:
    try:
        return (time.time() - lock_path.stat().st_mtime) > CONSENT_FILE_LOCK_STALE_SECONDS
    except OSError:
        return False


def _acquire_consent_file_lock(lock_path: Path) -> int:
    deadline = time.monotonic() + CONSENT_FILE_LOCK_TIMEOUT_SECONDS
    while True:
        try:
            fd = os.open(str(lock_path), os.O_CREAT | os.O_EXCL | os.O_RDWR)
            os.write(fd, f"pid={os.getpid()} acquired_at={time.time()}\n".encode("ascii", errors="ignore"))
            return fd
        except FileExistsError:
            if _consent_file_lock_is_stale(lock_path):
                try:
                    lock_path.unlink()
                    continue
                except OSError:
                    pass
            if time.monotonic() >= deadline:
                raise TimeoutError(f"Не удалось получить lock файла consent: {lock_path}")
            time.sleep(0.05)


@contextmanager
def _locked_records(job_id: str | None):
    path = _consent_path(job_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    thread_lock = _records_thread_lock(path)
    lock_path = path.with_name(f".{path.name}.lock")
    with thread_lock:
        fd = _acquire_consent_file_lock(lock_path)
        try:
            yield
        finally:
            try:
                os.close(fd)
            finally:
                try:
                    lock_path.unlink(missing_ok=True)
                except OSError:
                    pass


def _safe_filename_part(value: Any, fallback: str = "item") -> str:
    text = _safe_text(value).lower()
    text = re.sub(r"[^a-zа-яё0-9_-]+", "-", text, flags=re.IGNORECASE)
    text = text.strip("-_")
    return (text or fallback)[:80]


def _recipient_hash(value: Any) -> str:
    return hashlib.sha256(_safe_text(value).lower().encode("utf-8")).hexdigest()[:12]


def _relative_to_job_root(path: Path, *, job_id: str | None) -> str:
    root_dir = resolve_job_paths(job_id).root_dir
    try:
        return path.relative_to(root_dir).as_posix()
    except ValueError:
        return str(path)


def _consent_document_path(record: dict[str, Any], *, job_id: str | None) -> Path:
    confirmed_at = _safe_text(record.get("confirmed_at")) or datetime.now().isoformat(timespec="seconds")
    date_part = confirmed_at[:10] if len(confirmed_at) >= 10 else datetime.now().strftime("%Y-%m-%d")
    row_part = _safe_filename_part(record.get("row_id"), fallback="row")
    recipient_part = _recipient_hash(record.get("recipient"))
    timestamp_part = re.sub(r"[^0-9]+", "", confirmed_at)[:14] or datetime.now().strftime("%Y%m%d%H%M%S")
    return (
        _consent_documents_dir(job_id)
        / date_part
        / f"row-{row_part}_{recipient_part}_{timestamp_part}_consent.docx"
    )


def _add_consent_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(text)


def _save_consent_document(record: dict[str, Any], *, job_id: str | None) -> Path:
    existing_path = _safe_text(record.get("consent_document_path"))
    if existing_path:
        candidate = resolve_job_paths(job_id).root_dir / existing_path
        if candidate.exists():
            return candidate

    path = _consent_document_path(record, job_id=job_id)
    path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_heading("СОГЛАСИЕ", level=1)
    document.add_heading("НА ПОЛУЧЕНИЕ МАТЕРИАЛОВ И ОБРАБОТКУ ПЕРСОНАЛЬНЫХ ДАННЫХ", level=2)

    _add_consent_paragraph(
        document,
        (
            f"Настоящим я, перейдя по ссылке из письма, направленного {CONSENT_OPERATOR_NAME}, "
            f"ИНН {CONSENT_OPERATOR_INN} (далее — Оператор), как субъект персональных данных, "
            "во исполнение требований Федерального закона от 27.07.2006 № 152-ФЗ "
            "«О персональных данных», добровольно, своей волей и в своем интересе предоставляю "
            "своё согласие на:"
        ),
    )
    _add_consent_paragraph(
        document,
        f"1. Получение от Оператора по электронной почте {_materials_list_for_attachment_mode(record.get('attachment_mode'))}.",
    )
    _add_consent_paragraph(
        document,
        "2. Обработку моих персональных данных: адрес электронной почты, IP-адрес, время и факт "
        "перехода по ссылке, данные об устройстве.",
    )
    _add_consent_paragraph(
        document,
        "Цель обработки: направление запрошенных мной материалов по электронной почте.",
    )
    _add_consent_paragraph(
        document,
        "Перечень действий с персональными данными: сбор, запись, систематизация, накопление, "
        "хранение, уточнение (обновление, изменение), извлечение, использование, передача "
        "(распространение, предоставление, доступ) в объёме, необходимом для достижения указанной цели, "
        "а также блокирование, удаление, уничтожение.",
    )
    _add_consent_paragraph(
        document,
        "Согласие действует до момента направления запрошенных материалов либо до момента моего отзыва "
        "по запросу на email: personal.offer@parresh.ru.",
    )
    _add_consent_paragraph(
        document,
        "Я подтверждаю, что переход по ссылке является аналогом собственноручной подписи и полностью "
        "заменяет её для целей фиксации согласия в информационной системе Оператора.",
    )

    document.add_heading("Фиксация согласия произведена:", level=2)
    _add_consent_paragraph(document, f"Дата и время: {_safe_text(record.get('confirmed_at'))}")
    _add_consent_paragraph(document, f"Уникальный ID получателя: {_safe_text(record.get('token'))}")
    _add_consent_paragraph(document, f"Email получателя: {_safe_text(record.get('recipient'))}")
    _add_consent_paragraph(document, f"Муниципальное образование: {_safe_text(record.get('mun_name'))}")
    _add_consent_paragraph(document, f"ID строки: {_safe_text(record.get('row_id'))}")
    _add_consent_paragraph(document, f"IP-адрес: {_safe_text(record.get('confirmed_ip'))}")
    _add_consent_paragraph(document, f"User-Agent: {_safe_text(record.get('confirmed_user_agent'))}")

    document.save(path)
    return path


def _load_records(job_id: str | None) -> list[dict[str, Any]]:
    result = read_json(_consent_path(job_id), default={"records": []})
    if not result.ok:
        return []
    data = result.data
    if isinstance(data, list):
        return [item for item in data if isinstance(item, dict)]
    records = data.get("records") if isinstance(data, dict) else None
    return [item for item in records if isinstance(item, dict)] if isinstance(records, list) else []


def load_consent_records(job_id: str | None) -> list[dict[str, Any]]:
    return [dict(record) for record in _load_records(job_id)]


def _save_records(job_id: str | None, records: list[dict[str, Any]]) -> None:
    write_json_atomic(_consent_path(job_id), {"records": records})


def _record_matches(record: dict[str, Any], *, row_id: Any, recipient: str) -> bool:
    return (
        _safe_text(record.get("row_id")) == _safe_text(row_id)
        and _safe_text(record.get("recipient")).lower() == _safe_text(recipient).lower()
    )


def _consent_covers_attachment_mode(record_mode: Any, requested_mode: Any) -> bool:
    confirmed_mode = _normalize_attachment_mode(record_mode)
    requested = _normalize_attachment_mode(requested_mode)
    return confirmed_mode == ATTACHMENT_MODE_BOTH or confirmed_mode == requested


def _record_matches_attachment_mode(record: dict[str, Any], attachment_mode: Any) -> bool:
    return _normalize_attachment_mode(record.get("attachment_mode")) == _normalize_attachment_mode(attachment_mode)


def _recipient_key(value: Any) -> str:
    return _safe_text(value).lower()


def _consent_token_ttl_hours() -> int:
    try:
        configured = int(getattr(settings, "consent_token_ttl_hours", DEFAULT_CONSENT_TOKEN_TTL_HOURS) or 0)
    except (TypeError, ValueError):
        configured = DEFAULT_CONSENT_TOKEN_TTL_HOURS
    return configured if configured > 0 else DEFAULT_CONSENT_TOKEN_TTL_HOURS


def _consent_expires_at(now_text: str) -> str:
    try:
        now = datetime.fromisoformat(now_text)
    except ValueError:
        now = datetime.now()
    return (now + timedelta(hours=_consent_token_ttl_hours())).isoformat(timespec="seconds")


def _record_is_expired(record: dict[str, Any], *, now: datetime | None = None) -> bool:
    if _safe_text(record.get("status")) == "confirmed":
        return False
    expires_at = _safe_text(record.get("expires_at"))
    if not expires_at:
        return False
    try:
        expires_at_dt = datetime.fromisoformat(expires_at)
    except ValueError:
        return False
    return (now or datetime.now()) > expires_at_dt


def _parse_datetime(value: Any) -> datetime | None:
    text = _safe_text(value)
    if not text:
        return None
    try:
        return datetime.fromisoformat(text)
    except ValueError:
        return None


def _materials_dispatch_sent(record: dict[str, Any]) -> bool:
    return _safe_text(record.get("materials_status")) == "sent" or bool(_safe_text(record.get("materials_sent_at")))


def _materials_dispatch_should_start(record: dict[str, Any], *, now: datetime) -> bool:
    if _materials_dispatch_sent(record):
        return False

    status = _safe_text(record.get("materials_status")).lower()
    if status == "error" or _safe_text(record.get("materials_error")):
        return True

    if not _safe_text(record.get("materials_dispatch_requested_at")):
        return True

    if _safe_text(record.get("materials_dispatch_completed_at")):
        return True

    requested_at = _parse_datetime(record.get("materials_dispatch_requested_at"))
    if requested_at is None:
        return True
    return (now - requested_at).total_seconds() >= MATERIALS_DISPATCH_STALE_RETRY_SECONDS


def _job_owner_metadata(job_id: str | None) -> dict[str, str]:
    owner = read_job_owner(job_id)
    return {
        "owner_username": _safe_text(owner.get("owner_username")),
        "tenant_id": _safe_text(owner.get("tenant_id")),
    }


def _record_matches_job_scope(record: dict[str, Any], *, job_id: str | None) -> bool:
    if _safe_text(record.get("job_id")) != _safe_text(job_id):
        return False
    owner = _job_owner_metadata(job_id)
    owner_tenant = _safe_text(owner.get("tenant_id"))
    record_tenant = _safe_text(record.get("tenant_id"))
    if owner_tenant and record_tenant and owner_tenant != record_tenant:
        return False
    recipient_key = _safe_text(record.get("recipient_key"))
    if recipient_key and recipient_key != _recipient_key(record.get("recipient")):
        return False
    return True

def public_consent_url(token: str) -> str:
    base_url = _safe_text(settings.public_base_url).rstrip("/")
    if not base_url:
        raise ValueError("PUBLIC_BASE_URL must be set to build consent links.")
    return f"{base_url}/consent/confirm/{token}"


def prepare_consent_request(
    *,
    job_id: str | None,
    row: dict[str, Any],
    recipient: str,
    transport: str,
    attachment_mode: str = "kp",
    subject_template: str | None = None,
    work_type: str | None = None,
    recipient_strategy: str | None = None,
    sender_email: str | None = None,
    campaign_name: str | None = None,
) -> dict[str, Any]:
    with _locked_records(job_id):
        records = _load_records(job_id)
        row_id = row.get("ID")
        now = datetime.now().isoformat(timespec="seconds")
        effective_attachment_mode = _normalize_attachment_mode(attachment_mode)
        effective_work_type = normalize_work_type(work_type or DEFAULT_WORK_TYPE)
        effective_recipient_strategy = _safe_text(recipient_strategy)
        effective_sender_email = _safe_text(sender_email)
        effective_campaign_name = _safe_text(campaign_name)
        owner_metadata = _job_owner_metadata(job_id)
        expires_at = _consent_expires_at(now)
        for record in records:
            if not _record_matches(record, row_id=row_id, recipient=recipient):
                continue
            record_status = _safe_text(record.get("status"))
            if record_status == "confirmed" and not _consent_covers_attachment_mode(
                record.get("attachment_mode"),
                effective_attachment_mode,
            ):
                continue

            record.setdefault("token", secrets.token_urlsafe(24))
            record["job_id"] = _safe_text(job_id)
            record["recipient_key"] = _recipient_key(recipient)
            if owner_metadata.get("tenant_id"):
                record["tenant_id"] = owner_metadata["tenant_id"]
            if owner_metadata.get("owner_username"):
                record["owner_username"] = owner_metadata["owner_username"]
            record["last_request_prepared_at"] = now
            if record_status != "confirmed":
                record["status"] = "pending" if record_status == "expired" else (record_status or "pending")
                record["transport"] = _safe_text(transport)
                record["attachment_mode"] = effective_attachment_mode
                record["work_type"] = effective_work_type
                record["recipient_strategy"] = effective_recipient_strategy
                record["consent_text"] = _consent_text_for_attachment_mode(effective_attachment_mode)
                record["subject_template"] = _safe_text(subject_template)
                record["sender_email"] = effective_sender_email
                record["campaign_name"] = effective_campaign_name
                record["expires_at"] = expires_at
            _save_records(job_id, records)
            return dict(record, consent_url=public_consent_url(record["token"]))

        token = secrets.token_urlsafe(24)
        record = {
            "token": token,
            "status": "pending",
            "job_id": _safe_text(job_id),
            "row_id": _safe_text(row_id),
            "mun_name": _safe_text(row.get("MUN_NAME")),
            "recipient": _safe_text(recipient),
            "recipient_key": _recipient_key(recipient),
            "tenant_id": _safe_text(owner_metadata.get("tenant_id")),
            "owner_username": _safe_text(owner_metadata.get("owner_username")),
            "consent_text": _consent_text_for_attachment_mode(effective_attachment_mode),
            "created_at": now,
            "last_request_prepared_at": now,
            "expires_at": expires_at,
            "transport": _safe_text(transport),
            "attachment_mode": effective_attachment_mode,
            "work_type": effective_work_type,
            "recipient_strategy": effective_recipient_strategy,
            "subject_template": _safe_text(subject_template),
            "sender_email": effective_sender_email,
            "campaign_name": effective_campaign_name,
        }
        records.append(record)
        _save_records(job_id, records)
        return dict(record, consent_url=public_consent_url(token))


def mark_consent_request_sent(
    *,
    job_id: str | None,
    row_id: Any,
    recipient: str,
    provider: dict[str, Any] | None = None,
    attachment_mode: str | None = None,
) -> None:
    with _locked_records(job_id):
        records = _load_records(job_id)
        now = datetime.now().isoformat(timespec="seconds")
        for record in records:
            if not _record_matches(record, row_id=row_id, recipient=recipient):
                continue
            if attachment_mode is not None and not _record_matches_attachment_mode(record, attachment_mode):
                continue
            if _safe_text(record.get("status")) != "confirmed":
                record["status"] = "request_sent"
            record["request_sent_at"] = now
            if provider:
                record["provider"] = provider
            _save_records(job_id, records)
            return


def mark_materials_dispatch_result(
    *,
    job_id: str | None,
    row_id: Any,
    recipient: str,
    sent: bool,
    error: str = "",
    summary: str = "",
    attachment_mode: str | None = None,
) -> None:
    with _locked_records(job_id):
        records = _load_records(job_id)
        now = datetime.now().isoformat(timespec="seconds")
        for record in records:
            if not _record_matches(record, row_id=row_id, recipient=recipient):
                continue
            if attachment_mode is not None and not _record_matches_attachment_mode(record, attachment_mode):
                continue
            record["materials_dispatch_completed_at"] = now
            record["materials_dispatch_summary"] = _safe_text(summary)
            if sent:
                if not _safe_text(record.get("materials_sent_at")):
                    record["materials_sent_at"] = now
                record["materials_error"] = ""
                record["materials_status"] = "sent"
            else:
                if _safe_text(record.get("materials_status")) == "sent" or _safe_text(record.get("materials_sent_at")):
                    record["materials_status"] = "sent"
                    record["materials_error"] = ""
                    if not _safe_text(record.get("materials_dispatch_summary")):
                        record["materials_dispatch_summary"] = "Материалы уже были отправлены ранее."
                    _save_records(job_id, records)
                    return
                record["materials_error"] = _safe_text(error) or "Материалы не отправлены."
                record["materials_status"] = "error"
            _save_records(job_id, records)
            return


def has_confirmed_consent(
    *,
    job_id: str | None,
    row_id: Any,
    recipient: str,
    attachment_mode: str | None = None,
) -> bool:
    for record in _load_records(job_id):
        if not _record_matches(record, row_id=row_id, recipient=recipient):
            continue
        if _safe_text(record.get("status")) != "confirmed":
            continue
        if attachment_mode is not None and not _consent_covers_attachment_mode(record.get("attachment_mode"), attachment_mode):
            continue
        return True
    return False


def get_consent_by_token(token: str) -> dict[str, Any] | None:
    clean_token = _safe_text(token)
    for job_dir in _iter_job_dirs():
        job_id = job_dir.name if job_dir.name.startswith("job-") else None
        for record in _load_records(job_id):
            if _safe_text(record.get("token")) != clean_token:
                continue
            if not _record_matches_job_scope(record, job_id=job_id):
                continue
            result = dict(record)
            if _record_is_expired(record):
                result["_expired"] = True
            return result
    return None


def confirm_consent(token: str, *, ip: str = "", user_agent: str = "") -> dict[str, Any] | None:
    clean_token = _safe_text(token)
    for job_dir in _iter_job_dirs():
        job_id = job_dir.name if job_dir.name.startswith("job-") else None
        with _locked_records(job_id):
            records = _load_records(job_id)
            for record in records:
                if _safe_text(record.get("token")) != clean_token:
                    continue
                if not _record_matches_job_scope(record, job_id=job_id):
                    continue
                now = datetime.now().isoformat(timespec="seconds")
                changed = False
                already_confirmed = _safe_text(record.get("status")) == "confirmed"
                if _record_is_expired(record) and not already_confirmed:
                    record["status"] = "expired"
                    record["expired_at"] = now
                    _save_records(job_id, records)
                    return dict(record, _expired=True, _already_confirmed=False, _dispatch_materials=False)
                if not already_confirmed:
                    record["status"] = "confirmed"
                    changed = True
                if not _safe_text(record.get("confirmed_at")):
                    record["confirmed_at"] = now
                    changed = True
                if not _safe_text(record.get("confirmed_ip")):
                    record["confirmed_ip"] = _safe_text(ip)
                    changed = True
                if not _safe_text(record.get("confirmed_user_agent")):
                    record["confirmed_user_agent"] = _safe_text(user_agent)
                    changed = True

                dispatch_materials = _materials_dispatch_should_start(record, now=datetime.fromisoformat(now))
                if dispatch_materials:
                    record["materials_dispatch_requested_at"] = now
                    try:
                        record["materials_dispatch_attempts"] = int(record.get("materials_dispatch_attempts") or 0) + 1
                    except (TypeError, ValueError):
                        record["materials_dispatch_attempts"] = 1
                    if _safe_text(record.get("materials_status")) != "sent":
                        record["materials_status"] = "queued"
                    record["materials_error"] = ""
                    record["materials_dispatch_completed_at"] = ""
                    record["materials_dispatch_summary"] = ""
                    changed = True

                document_path = _save_consent_document(record, job_id=job_id)
                relative_document_path = _relative_to_job_root(document_path, job_id=job_id)
                if _safe_text(record.get("consent_document_path")) != relative_document_path:
                    record["consent_document_path"] = relative_document_path
                    changed = True
                if changed:
                    _save_records(job_id, records)
                return dict(
                    record,
                    _already_confirmed=already_confirmed,
                    _dispatch_materials=dispatch_materials,
                )
    return None


def _iter_job_dirs() -> list[Path]:
    jobs_root = resolve_job_paths("job-placeholder").root_dir.parent
    candidates: list[Path] = []
    if jobs_root.exists():
        candidates.extend(path for path in jobs_root.iterdir() if path.is_dir())
    legacy_root = resolve_job_paths(None).root_dir
    if legacy_root.exists():
        candidates.append(legacy_root)
    return candidates

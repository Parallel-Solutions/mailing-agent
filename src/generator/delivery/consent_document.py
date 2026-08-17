"""Shared DOCX renderer for consent-to-receive-materials evidence."""

from __future__ import annotations

import hashlib
import os
import uuid
from pathlib import Path
from typing import Any

from docx import Document


CONSENT_DOCUMENT_TEXT_VERSION = "materials-consent-v1"
CONSENT_OPERATOR_NAME = "ООО «Параллельные Решения»"
CONSENT_OPERATOR_INN = "5038110107"


def _safe_text(value: Any) -> str:
    return str(value or "").strip()


def _add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(text)


def _materials_description(record: dict[str, Any]) -> str:
    material_names = record.get("material_names")
    if isinstance(material_names, list):
        names = [_safe_text(item) for item in material_names]
        names = [item for item in names if item]
        if names:
            return "следующих материалов: " + "; ".join(names)
    return _safe_text(record.get("materials_description")) or "запрошенных материалов"


def write_consent_document(path: Path, record: dict[str, Any]) -> str:
    """Write a consent DOCX atomically and return its SHA-256 digest."""

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("СОГЛАСИЕ", level=1)
    document.add_heading(
        "НА ПОЛУЧЕНИЕ МАТЕРИАЛОВ И ОБРАБОТКУ ПЕРСОНАЛЬНЫХ ДАННЫХ",
        level=2,
    )

    _add_paragraph(
        document,
        (
            f"Настоящим я, перейдя по ссылке из письма, направленного {CONSENT_OPERATOR_NAME}, "
            f"ИНН {CONSENT_OPERATOR_INN} (далее — Оператор), как субъект персональных данных, "
            "во исполнение требований Федерального закона от 27.07.2006 № 152-ФЗ "
            "«О персональных данных», добровольно, своей волей и в своем интересе предоставляю "
            "своё согласие на:"
        ),
    )
    _add_paragraph(
        document,
        f"1. Получение от Оператора по электронной почте {_materials_description(record)}.",
    )
    _add_paragraph(
        document,
        "2. Обработку моих персональных данных: адрес электронной почты, IP-адрес, время и факт "
        "перехода по ссылке, данные об устройстве.",
    )
    _add_paragraph(
        document,
        "Цель обработки: направление запрошенных мной материалов по электронной почте.",
    )
    _add_paragraph(
        document,
        "Перечень действий с персональными данными: сбор, запись, систематизация, накопление, "
        "хранение, уточнение (обновление, изменение), извлечение, использование, передача "
        "(распространение, предоставление, доступ) в объёме, необходимом для достижения указанной цели, "
        "а также блокирование, удаление, уничтожение.",
    )
    _add_paragraph(
        document,
        "Согласие действует до момента направления запрошенных материалов либо до момента моего отзыва "
        "по запросу на email: personal.offer@parresh.ru.",
    )
    _add_paragraph(
        document,
        "Я подтверждаю, что переход по ссылке является аналогом собственноручной подписи и полностью "
        "заменяет её для целей фиксации согласия в информационной системе Оператора.",
    )

    document.add_heading("Фиксация согласия произведена:", level=2)
    _add_paragraph(document, f"Дата и время: {_safe_text(record.get('confirmed_at'))}")
    _add_paragraph(document, f"Уникальный ID получателя: {_safe_text(record.get('token'))}")
    _add_paragraph(document, f"Email получателя: {_safe_text(record.get('recipient'))}")
    _add_paragraph(document, f"Муниципальное образование: {_safe_text(record.get('mun_name'))}")
    _add_paragraph(document, f"ID строки: {_safe_text(record.get('row_id'))}")
    _add_paragraph(document, f"IP-адрес: {_safe_text(record.get('confirmed_ip'))}")
    _add_paragraph(document, f"User-Agent: {_safe_text(record.get('confirmed_user_agent'))}")

    optional_fields = (
        ("Кампания", "campaign_name"),
        ("Блок цепочки", "target_node_name"),
        ("Версия текста согласия", "consent_text_version"),
    )
    for label, key in optional_fields:
        value = _safe_text(record.get(key))
        if value:
            _add_paragraph(document, f"{label}: {value}")

    temporary_path = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
    try:
        document.save(temporary_path)
        os.replace(temporary_path, path)
    finally:
        temporary_path.unlink(missing_ok=True)
    return hashlib.sha256(path.read_bytes()).hexdigest()

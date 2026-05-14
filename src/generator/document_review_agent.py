from __future__ import annotations

import json
import re
import sys
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable, List

from docx import Document

try:
    from src.generator.config_generator import DOCUMENT_REVIEW_MODEL, ENABLE_DOCUMENT_REVIEW_AI
    from src.generator.ai_case_agent import (
        _extract_json_payload,
        _resolve_openai_api_key,
        _resolve_openai_base_url,
    )
    from src.generator.philology_knowledge import find_relevant_rules, format_rules_context
except ImportError:  # pragma: no cover
    from generator.config_generator import DOCUMENT_REVIEW_MODEL, ENABLE_DOCUMENT_REVIEW_AI
    from generator.ai_case_agent import (
        _extract_json_payload,
        _resolve_openai_api_key,
        _resolve_openai_base_url,
    )
    from generator.philology_knowledge import find_relevant_rules, format_rules_context

try:
    from openai import OpenAI  # type: ignore
except ImportError:  # pragma: no cover
    OpenAI = None

try:
    import httpx  # type: ignore
except ImportError:  # pragma: no cover
    httpx = None

PLACEHOLDER_PATTERN = re.compile(
    r"\b(?:ADM_NAME|MUN_NAME(?:_[12])?|MUN_R_NAME(?:_1)?|SUB_RF(?:_1)?|HEAD_FIO(?:_1)?|POPULATION|ADRES|EMAIL_OSN|TEL_OSN|REQUISITES_[A-Z]+)\b"
)
POPULATION_PATTERN = re.compile(
    r"(Численность населения проектируемой территории составляет )(\d+)\s+(человек(?:а)?)",
    re.IGNORECASE,
)
LOWERCASE_LOCALITY_PATTERN = re.compile(
    r"\b(?P<prefix>города|город|села|село|поселка|посёлка|поселок|посёлок)\s+"
    r"(?P<name>[а-яё]+(?:-[а-яё]+)*)",
)
REPEATED_GEO_FRAGMENT_PATTERN = re.compile(
    r"(?P<fragment>"
    r"[А-ЯЁа-яё-]+(?:\s+[А-ЯЁа-яё-]+){0,4}\s+муниципального\s+района\s+"
    r"[А-ЯЁа-яё-]+(?:\s+[А-ЯЁа-яё-]+){0,3}\s+области"
    r")\s+(?P=fragment)",
    re.IGNORECASE,
)
DOUBLE_RAYON_PATTERN = re.compile(
    r"(?P<fragment>[А-ЯЁа-яё-]+ского\s+района)\s+района",
    re.IGNORECASE,
)
MISSING_MUNICIPAL_DISTRICT_PATTERN = re.compile(
    r"(?<!муниципального\s)\b(?!муниципального\b)(?P<district>[А-ЯЁа-яё-]+(?:ского|цкого|ого))\s+района\b",
    re.IGNORECASE,
)
BAD_UPPERCASE_WORDS = {
    "ДОГОВОР": "Договор",
    "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ": "Коммерческое предложение",
}
PROTECTED_UPPERCASE_HEADINGS = {
    "ДОГОВОР",
    "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ",
}
PROTECTED_ATTACHMENT_HEADING_PATTERN = re.compile(r"^Приложение\s+№\s+\d+\b", re.IGNORECASE)
LEGAL_TERM_CASE_REPLACEMENTS = {
    "Выполнение Работ": "Выполнение работ",
    "нахождения Исполнителя": "нахождения исполнителя",
    "настоящим Договором": "настоящим договором",
    "Техническим заданием": "техническим заданием",
    "Календарным планом": "календарным планом",
    "Приложение №": "приложение №",
    "к Договору": "к договору",
    "настоящего Договора": "настоящего договора",
}


@dataclass
class ReviewIssue:
    source: str
    location: str
    fragment: str
    issue: str
    suggestion: str
    severity: str = "warning"


def _safe_text(value) -> str:
    if value is None:
        return ""
    return " ".join(str(value).split())


EDITORIAL_SUGGESTION_PREFIXES = (
    "заменить ",
    "исправить ",
    "нужно ",
    "следует ",
    "проверить ",
    "убрать ",
)


def _looks_like_editorial_instruction(text: str) -> bool:
    normalized = _safe_text(text).strip().lower()
    if not normalized:
        return False
    if any(normalized.startswith(prefix) for prefix in EDITORIAL_SUGGESTION_PREFIXES):
        return True
    if "заменить" in normalized and (" на " in normalized or '"' in normalized or "«" in normalized):
        return True
    return False


def _is_overlong_replacement(fragment: str, suggestion: str) -> bool:
    fragment_text = _safe_text(fragment)
    suggestion_text = _safe_text(suggestion)
    if not fragment_text or not suggestion_text:
        return False
    if len(suggestion_text) > 400:
        return True
    if len(fragment_text) <= 80 and len(suggestion_text) > len(fragment_text) * 3:
        return True
    return False


def _population_with_unit(number_text: str) -> str:
    number = int(number_text)
    mod100 = number % 100
    mod10 = number % 10
    if 11 <= mod100 <= 14:
        unit = "человек"
    elif mod10 == 1:
        unit = "человек"
    elif mod10 in (2, 3, 4):
        unit = "человека"
    else:
        unit = "человек"
    return f"{number_text} {unit}"


def _title_case_locality_name(name: str) -> str:
    parts = [part for part in name.split("-") if part]
    return "-".join(part[:1].upper() + part[1:] for part in parts)


def _is_protected_uppercase_heading(text: str, source: str) -> bool:
    normalized = _safe_text(text).strip()
    return normalized == source and normalized in PROTECTED_UPPERCASE_HEADINGS


def _is_protected_legal_term_context(text: str, source: str) -> bool:
    normalized = _safe_text(text).strip()
    if source == "Приложение №" and PROTECTED_ATTACHMENT_HEADING_PATTERN.search(normalized):
        return True
    return False


def _iter_docx_blocks(doc: Document) -> Iterable[tuple[str, str]]:
    for index, paragraph in enumerate(doc.paragraphs, 1):
        text = _safe_text(paragraph.text)
        if text:
            yield (f"paragraph:{index}", text)

    for table_index, table in enumerate(doc.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            for cell_index, cell in enumerate(row.cells, 1):
                for paragraph_index, paragraph in enumerate(cell.paragraphs, 1):
                    text = _safe_text(paragraph.text.replace("\n", " / "))
                    if text:
                        yield (
                            f"table:{table_index}:row:{row_index}:cell:{cell_index}:paragraph:{paragraph_index}",
                            text,
                        )


def _add_local_issue(
    issues: list[ReviewIssue],
    *,
    location: str,
    fragment: str,
    issue: str,
    suggestion: str,
    severity: str = "warning",
) -> None:
    issues.append(
        ReviewIssue(
            source="local",
            location=location,
            fragment=fragment,
            issue=issue,
            suggestion=suggestion,
            severity=severity,
        )
    )


def _add_local_replacement_issue(
    issues: list[ReviewIssue],
    *,
    location: str,
    fragment: str,
    replacement: str,
    issue: str,
    severity: str = "warning",
) -> None:
    _add_local_issue(
        issues,
        location=location,
        fragment=fragment,
        issue=issue,
        suggestion=replacement,
        severity=severity,
    )


def _run_local_checks(blocks: Iterable[tuple[str, str]]) -> list[ReviewIssue]:
    issues: list[ReviewIssue] = []
    for location, text in blocks:
        placeholder_match = PLACEHOLDER_PATTERN.search(text)
        if placeholder_match:
            placeholder = placeholder_match.group(0)
            _add_local_issue(
                issues,
                location=location,
                fragment=text,
                issue=f"В документе остался незаменённый плейсхолдер `{placeholder}`.",
                suggestion="Проверить подстановку данных и шаблон.",
                severity="error",
            )

        for match in POPULATION_PATTERN.finditer(text):
            prefix, number_text, unit = match.groups()
            expected = _population_with_unit(number_text)
            current = f"{number_text} {unit}"
            if current != expected:
                _add_local_issue(
                    issues,
                    location=location,
                    fragment=text,
                    issue=f"После числа используется неверная форма слова: `{current}`.",
                    suggestion=f"{prefix}{expected}.",
                    severity="warning",
                )

        if "  " in text:
            _add_local_issue(
                issues,
                location=location,
                fragment=text,
                issue="Обнаружены двойные пробелы.",
                suggestion="Убрать лишние пробелы.",
                severity="info",
            )

        for source, replacement in BAD_UPPERCASE_WORDS.items():
            if source not in text:
                continue
            if _is_protected_uppercase_heading(text, source):
                continue
            _add_local_replacement_issue(
                issues,
                location=location,
                fragment=source,
                replacement=replacement,
                issue="Есть неудачное использование капса в текстовом блоке.",
                severity="warning",
            )

        for source, replacement in LEGAL_TERM_CASE_REPLACEMENTS.items():
            if source not in text:
                continue
            if _is_protected_legal_term_context(text, source):
                continue
            _add_local_replacement_issue(
                issues,
                location=location,
                fragment=source,
                replacement=replacement,
                issue="В середине предложения используется лишняя заглавная буква.",
                severity="warning",
            )

        for match in LOWERCASE_LOCALITY_PATTERN.finditer(text):
            prefix = match.group("prefix")
            locality_name = match.group("name")
            titled_locality = _title_case_locality_name(locality_name)
            if locality_name == titled_locality:
                continue
            fragment = match.group(0)
            replacement = f"{prefix} {titled_locality}"
            _add_local_replacement_issue(
                issues,
                location=location,
                fragment=fragment,
                replacement=replacement,
                issue="Название населённого пункта в официальной конструкции должно начинаться с заглавной буквы.",
                severity="warning",
            )

        repeated_geo_match = REPEATED_GEO_FRAGMENT_PATTERN.search(text)
        if repeated_geo_match:
            fragment = repeated_geo_match.group("fragment")
            _add_local_replacement_issue(
                issues,
                location=location,
                fragment=repeated_geo_match.group(0),
                replacement=fragment,
                issue="Повторяется один и тот же географический хвост (район/область).",
                severity="error",
            )

        double_rayon_match = DOUBLE_RAYON_PATTERN.search(text)
        if double_rayon_match:
            fragment = double_rayon_match.group(0)
            replacement = double_rayon_match.group("fragment")
            _add_local_replacement_issue(
                issues,
                location=location,
                fragment=fragment,
                replacement=replacement,
                issue="Дублируется слово 'района' в официальной конструкции.",
                severity="error",
            )

        for match in MISSING_MUNICIPAL_DISTRICT_PATTERN.finditer(text):
            fragment = match.group(0)
            replacement = f"{match.group('district')} муниципального района"
            _add_local_replacement_issue(
                issues,
                location=location,
                fragment=fragment,
                replacement=replacement,
                issue="В официальной конструкции района пропущено слово 'муниципального'.",
                severity="error",
            )

        if "обоснование предмета нормирование" in text:
            _add_local_replacement_issue(
                issues,
                location=location,
                fragment="обоснование предмета нормирование",
                replacement="обоснование предмета нормирования",
                issue="Нарушена падежная форма в словосочетании 'предмета нормирование'.",
                severity="error",
            )

        if "утвержденный постановление" in text:
            _add_local_replacement_issue(
                issues,
                location=location,
                fragment="утвержденный постановление",
                replacement="утвержденный постановлением",
                issue="Нарушено управление в конструкции 'утвержденный постановление'.",
                severity="error",
            )

        if "подписанный в 2 (двух) экземплярах акта сдачи-приемки работ" in text:
            _add_local_replacement_issue(
                issues,
                location=location,
                fragment="подписанный в 2 (двух) экземплярах акта сдачи-приемки работ",
                replacement="подписанный в 2 (двух) экземплярах акт сдачи-приемки работ",
                issue="Нарушено согласование в конструкции с актом сдачи-приемки работ.",
                severity="error",
            )

        if "дней, следующих за днем поступления Заказчику акта сдачи-приемки работ Заказчик подписывает" in text:
            _add_local_replacement_issue(
                issues,
                location=location,
                fragment="дней, следующих за днем поступления Заказчику акта сдачи-приемки работ Заказчик подписывает",
                replacement="дней, следующих за днем поступления Заказчику акта сдачи-приемки работ, Заказчик подписывает",
                issue="В предложении есть лишний повтор слова 'Заказчик', который ломает синтаксис.",
                severity="warning",
            )

        if "В стоимость работ включены консультационное сопровождение" in text:
            _add_local_replacement_issue(
                issues,
                location=location,
                fragment="В стоимость работ включены консультационное сопровождение",
                replacement="В стоимость работ включено консультационное сопровождение",
                issue="Нарушено согласование сказуемого со словом 'сопровождение'.",
                severity="warning",
            )
    return issues


def _build_ai_prompt(blocks: list[tuple[str, str]]) -> str:
    payload = [
        {"location": location, "text": text}
        for location, text in blocks
    ]
    rules = find_relevant_rules(" ".join(text for _, text in blocks), limit=5)
    rules_context = format_rules_context(rules)
    return (
        "Проверь текст фрагментов договора/КП на грамматику, падежи, согласование и канцелярский стиль. "
        "Сначала опирайся на локальную базу правил русского языка, затем на сам текст фрагмента. "
        "Не придумывай новый текст целиком. Ищи только реальные ошибки русского языка или шаблонные огрехи. "
        "Если нашлась ошибка, location должен указывать точное место, а fragment должен быть МИНИМАЛЬНЫМ точным куском исходного текста, где есть ошибка. "
        "Если нашлась ошибка, в поле suggestion верни только ИСПРАВЛЕННУЮ ЗАМЕНУ для fragment, а не весь абзац. "
        "suggestion должен быть готовым текстом документа, а не комментарием редактора. "
        "Нельзя писать 'Заменить ...', 'Исправить ...', 'Нужно ...', 'Следует ...', 'Проверить ...'. "
        "Если для фрагмента нельзя дать безопасную готовую замену, оставь suggestion пустым, но опиши ошибку в issue. "
        "Не меняй смысл документа и не переписывай абзац шире, чем нужно для исправления ошибки. "
        "Верни только JSON-объект вида "
        '{"issues":[{"location":"...","fragment":"...","issue":"...","suggestion":"...","severity":"warning|error|info"}]}. '
        "Если ошибок нет, верни {\"issues\": []}. Без markdown и пояснений вне JSON.\n\n"
        f"Локальная база правил:\n{rules_context}\n\n"
        f"{json.dumps(payload, ensure_ascii=False, indent=2)}"
    )


def _run_ai_review(blocks: list[tuple[str, str]], *, ai_enabled: bool = True) -> list[ReviewIssue]:
    if not ai_enabled or not ENABLE_DOCUMENT_REVIEW_AI or not OpenAI:
        return []

    api_key = _resolve_openai_api_key()
    if not api_key:
        return []

    client_kwargs = {"api_key": api_key, "max_retries": 0}
    base_url = _resolve_openai_base_url()
    if base_url:
        client_kwargs["base_url"] = base_url
    if httpx:
        client_kwargs["http_client"] = httpx.Client(
            http2=False,
            timeout=httpx.Timeout(connect=10, read=30, write=30, pool=30),
            trust_env=False,
        )
    client = OpenAI(**client_kwargs)

    request_kwargs = {
        "model": DOCUMENT_REVIEW_MODEL,
        "messages": [{"role": "user", "content": _build_ai_prompt(blocks)}],
    }
    if not base_url:
        request_kwargs["response_format"] = {"type": "json_object"}

    response = client.chat.completions.create(**request_kwargs)
    content = response.choices[0].message.content or "{}"
    parsed = json.loads(_extract_json_payload(content))
    issues = parsed.get("issues") if isinstance(parsed, dict) else []
    result: list[ReviewIssue] = []
    for item in issues or []:
        if not isinstance(item, dict):
            continue
        fragment = _safe_text(item.get("fragment"))
        suggestion = _safe_text(item.get("suggestion"))
        if _looks_like_editorial_instruction(suggestion) or _is_overlong_replacement(fragment, suggestion):
            suggestion = ""
        result.append(
            ReviewIssue(
                source="ai",
                location=_safe_text(item.get("location")),
                fragment=fragment,
                issue=_safe_text(item.get("issue")),
                suggestion=suggestion,
                severity=_safe_text(item.get("severity")) or "warning",
            )
        )
    return result


def review_docx(document_path: Path, *, ai_enabled: bool = True) -> dict:
    doc = Document(document_path)
    blocks = list(_iter_docx_blocks(doc))
    local_issues = _run_local_checks(blocks)
    ai_issues: list[ReviewIssue] = []
    ai_error = None
    try:
        ai_issues = _run_ai_review(blocks, ai_enabled=ai_enabled)
    except Exception as exc:  # pragma: no cover
        ai_error = f"{type(exc).__name__}: {exc}"

    issues = local_issues + ai_issues
    return {
        "document": str(document_path),
        "issue_count": len(issues),
        "local_issue_count": len(local_issues),
        "ai_issue_count": len(ai_issues),
        "ai_enabled": bool(ai_enabled and ENABLE_DOCUMENT_REVIEW_AI),
        "ai_model": DOCUMENT_REVIEW_MODEL,
        "ai_error": ai_error,
        "issues": [asdict(issue) for issue in issues],
    }


def _main(argv: list[str]) -> int:
    if len(argv) < 2:
        print("Usage: python document_review_agent.py <path-to-docx> [json-output-path]")
        return 1

    document_path = Path(argv[1])
    result = review_docx(document_path)
    payload = json.dumps(result, ensure_ascii=False, indent=2)
    print(payload)

    if len(argv) >= 3:
        output_path = Path(argv[2])
        output_path.write_text(payload, encoding="utf-8")

    return 0


if __name__ == "__main__":
    raise SystemExit(_main(sys.argv))

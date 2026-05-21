from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from time import perf_counter
from typing import Any, Iterable

from docx import Document
from docx.oxml.ns import qn

from src.generator.orchestration.agent_handoff import (
    append_agent_event,
    count_tasks_for_agent,
    create_task,
    get_recent_events,
    get_tasks_for_agent,
    mark_tasks_in_progress,
    set_task_statuses,
)
from src.generator.knowledge.agent_memory import build_quarantine_items, save_agent_report, save_learning_memory
from src.generator.inflection.ai_case_agent import (
    OpenAI,
    _extract_json_payload,
    _resolve_openai_api_key,
    _resolve_openai_base_url,
)
from src.generator.generation.config_generator import (
    DOCUMENT_REVIEW_MODEL,
    OUTPUT_DIR,
    PHILOLOGIST_CONTEXT_LLM,
    PHILOLOGIST_CONTEXT_LLM_BATCH_SIZE,
    PHILOLOGIST_CONTEXT_LLM_MAX_ITEMS,
    PHILOLOGIST_CONTEXT_LLM_MIN_CONFIDENCE,
    PHILOLOGIST_LLM_FIX_STRATEGY,
    PHILOLOGIST_LLM_ROUTER,
    PHILOLOGIST_MODE,
    PHILOLOGIST_REBUILD_PDF,
)
from src.generator.philologist.document_review_agent import review_docx
from src.generator.inflection.inflection_report import format_inflection_report, load_inflection_log
from src.generator.philologist.philologist_decisions import AUTO_FIX, NEEDS_HUMAN, QUARANTINE, decide_issue_fix
from src.generator.philologist.philologist_executor import PhilologistAgentLoop, merge_plan_execution
from src.generator.philologist.philologist_planner import build_philologist_plan
from src.generator.philologist.philologist_rag import explain_fix_decision_with_rag
from src.generator.knowledge.philology_knowledge import find_relevant_rules, format_rules_context
from src.generator.philologist.philologist_tools import PhilologistToolRunner, build_philologist_tool_manifest
from src.generator.generation.pdf_converter import convert_docx_batch
from src.generator.orchestration.responsibility_matrix import diagnose_responsibility
from src.jobs import load_agent_state, save_agent_state
from src.jobs.storage import resolve_job_paths
from src.utils.logger import logger
from src.utils.config import settings

try:
    import httpx  # type: ignore
except ImportError:  # pragma: no cover
    httpx = None


PHILOLOGIST_STATE: dict[str, Any] = {
    "status": "idle",
    "started_at": None,
    "completed_at": None,
    "total_documents": 0,
    "processed_documents": 0,
    "fixed_documents": 0,
    "documents_with_issues": 0,
    "documents": [],
    "summary_text": "Агент-филолог ещё не запускался.",
    "tool_manifest": build_philologist_tool_manifest(),
    "tool_trace": [],
    "plan": None,
    "agent_loop": None,
    "task_stats": {"total": 0, "pending": 0, "in_progress": 0, "done": 0, "blocked": 0},
    "tasks": [],
    "recent_events": [],
    "stop_requested": False,
    "stop_requested_at": None,
    "template_memory": {},
}


def _load_philologist_state(job_id: str | None = None) -> dict[str, Any]:
    return load_agent_state("philologist", PHILOLOGIST_STATE, job_id)


def _save_philologist_state(state: dict[str, Any], job_id: str | None = None) -> dict[str, Any]:
    return save_agent_state("philologist", state, job_id)


def request_philologist_stop(job_id: str | None = None) -> dict[str, Any]:
    state = _load_philologist_state(job_id)
    state["stop_requested"] = True
    state["stop_requested_at"] = datetime.now().isoformat(timespec="seconds")
    if state.get("status") == "running":
        state["summary_text"] = (
            "Получен запрос на остановку. Филолог завершит текущий документ и остановится."
        )
    _save_philologist_state(state, job_id)
    return get_philologist_status(job_id)


def clear_philologist_stop_request(job_id: str | None = None) -> None:
    state = _load_philologist_state(job_id)
    state["stop_requested"] = False
    state["stop_requested_at"] = None
    _save_philologist_state(state, job_id)


def _refresh_philologist_stop_flag(state: dict[str, Any], job_id: str | None = None) -> dict[str, Any]:
    persisted = _load_philologist_state(job_id)
    state["stop_requested"] = bool(persisted.get("stop_requested", False))
    state["stop_requested_at"] = persisted.get("stop_requested_at")
    return state


def _sorted_locations(locations: set[str]) -> list[str]:
    return sorted(location for location in locations if _safe_text(location))


def _update_template_memory(
    template_memory: dict[str, Any],
    *,
    document_name: str,
    review_result: dict[str, Any],
    fix_result: dict[str, Any],
) -> dict[str, Any]:
    template_fingerprint = _safe_text(review_result.get("template_fingerprint"))
    if not template_fingerprint:
        return template_memory

    entry = template_memory.get(template_fingerprint)
    if not isinstance(entry, dict):
        entry = {}

    risky_locations = {
        _safe_text(item.get("location"))
        for item in review_result.get("issues", [])
        if isinstance(item, dict) and _safe_text(item.get("location"))
    }
    risky_locations.update(
        _safe_text(item.get("location"))
        for item in fix_result.get("applied_fixes", [])
        if isinstance(item, dict) and _safe_text(item.get("location"))
    )
    risky_locations.update(
        _safe_text(item.get("location"))
        for item in fix_result.get("skipped_fixes", [])
        if isinstance(item, dict) and _safe_text(item.get("location"))
    )
    preserved_locations = {
        _safe_text(location)
        for location in entry.get("risky_locations", [])
        if _safe_text(location)
    }
    merged_locations = preserved_locations | risky_locations

    entry.update(
        {
            "documents_seen": int(entry.get("documents_seen", 0) or 0) + 1,
            "documents_with_issues": int(entry.get("documents_with_issues", 0) or 0)
            + (1 if int(review_result.get("issue_count", 0) or 0) > 0 else 0),
            "risky_locations": _sorted_locations(merged_locations)[:200],
            "last_document_name": document_name,
            "last_review_mode": _safe_text(review_result.get("review_mode")) or "full",
            "total_block_count": int(review_result.get("total_block_count", 0) or 0),
            "last_reviewed_block_count": int(review_result.get("reviewed_block_count", 0) or 0),
        }
    )
    template_memory[template_fingerprint] = entry
    return template_memory


def _build_llm_client():
    if not OpenAI:
        return None
    api_key = _resolve_openai_api_key()
    if not api_key:
        return None
    client_kwargs: dict[str, Any] = {"api_key": api_key, "max_retries": 0}
    base_url = _resolve_openai_base_url()
    if base_url:
        client_kwargs["base_url"] = base_url
    if httpx:
        client_kwargs["http_client"] = httpx.Client(
            http2=False,
            timeout=httpx.Timeout(connect=10, read=30, write=30, pool=30),
            trust_env=False,
        )
    return OpenAI(**client_kwargs)


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    return " ".join(str(value).split())


def _is_llm_auth_error(value: Any) -> bool:
    normalized = _safe_text(value).lower()
    return "401" in normalized or "unauthorized" in normalized or "incorrect api key" in normalized


def _normalize_philologist_mode(mode: str | None) -> str:
    normalized = _safe_text(mode or PHILOLOGIST_MODE).lower()
    if normalized in {"deep", "llm", "full"}:
        return "deep"
    return "fast"


def _is_fio_context_suspicious(row: dict[str, Any]) -> bool:
    field = _safe_text(row.get("field"))
    if not field.startswith("HEAD_FIO_"):
        return False

    source_parts = [_safe_text(part) for part in _safe_text(row.get("source_value")).split() if _safe_text(part)]
    result_parts = [_safe_text(part) for part in _safe_text(row.get("result_value")).split() if _safe_text(part)]
    if len(source_parts) != 3 or len(result_parts) != 3:
        return True

    source_patronymic = source_parts[2].lower()
    result_patronymic = result_parts[2].lower()
    target_case = _safe_text(row.get("target_case")).lower()

    female_patronymic = source_patronymic.endswith(("вна", "чна", "шна"))
    male_patronymic = source_patronymic.endswith(("ич", "оглы"))
    if target_case == "genitive":
        if female_patronymic and result_patronymic.endswith(("вна", "чна", "шна")):
            return True
        if male_patronymic and result_patronymic == source_patronymic:
            return True
    if target_case == "dative":
        if female_patronymic and not result_patronymic.endswith(("вне", "чне", "шне")):
            return True
        if male_patronymic and result_patronymic == source_patronymic:
            return True
    return False


def _inflection_context_is_suspicious(row: dict[str, Any]) -> bool:
    warning = _safe_text(row.get("warning"))
    method = _safe_text(row.get("method")).lower()
    confidence = _safe_text(row.get("confidence")).lower()
    source_value = _safe_text(row.get("source_value"))
    result_value = _safe_text(row.get("result_value"))
    target_case = _safe_text(row.get("target_case")).lower()
    filled_sentence = _safe_text(row.get("filled_sentence"))

    if not filled_sentence or not source_value or not result_value:
        return False
    if warning:
        return True
    if method == "fallback" or confidence in {"empty", "no_morph", "low"}:
        return True
    if target_case in {"genitive", "dative", "prepositional", "project_genitive"} and source_value == result_value:
        return True
    if _is_fio_context_suspicious(row):
        return True
    return False


def _select_inflection_contexts_for_llm(
    rows: list[dict[str, Any]],
    *,
    row_ids: set[str],
    limit: int,
) -> list[dict[str, Any]]:
    selected: list[dict[str, Any]] = []
    seen: set[tuple[str, str, str]] = set()
    for row in rows:
        row_id = _safe_text(row.get("row_id"))
        if row_ids and row_id not in row_ids:
            continue
        if not _inflection_context_is_suspicious(row):
            continue
        key = (row_id, _safe_text(row.get("field")), _safe_text(row.get("result_value")))
        if key in seen:
            continue
        seen.add(key)
        selected.append(row)
        if limit and len(selected) >= limit:
            break
    return selected


def _build_inflection_context_prompt(rows: list[dict[str, Any]]) -> str:
    payload = []
    for index, row in enumerate(rows, start=1):
        payload.append(
            {
                "id": index,
                "row_id": row.get("row_id"),
                "field": row.get("field"),
                "source_value": row.get("source_value"),
                "generated_value": row.get("result_value"),
                "target_case": row.get("target_case"),
                "context_sentence": row.get("context_sentence"),
                "filled_sentence": row.get("filled_sentence"),
                "slot_instruction": row.get("slot_instruction"),
                "warning": row.get("warning"),
            }
        )
    return (
        "Ты агент-филолог. Проверь только подстановку generated_value внутри filled_sentence. "
        "Не переписывай всё предложение. Если generated_value в этой фразе корректен, верни status='ok'. "
        "Если неверен падеж, род, число или согласование, верни status='fix' и corrected_value — только исправленную подстановку для [SLOT], не всё предложение. "
        "Если недостаточно уверенности, верни status='needs_review'. "
        "Не меняй название организации/муниципалитета по смыслу. Не придумывай новые данные. "
        "confidence должен быть от 0.0 до 1.0. "
        "Верни только JSON-объект вида "
        '{"items":[{"id":1,"status":"ok|fix|needs_review","corrected_value":"...","confidence":0.0,"comment":"..."}]} '
        "без markdown.\n\n"
        f"Данные:\n{json.dumps(payload, ensure_ascii=False, indent=2)}"
    )


def _call_inflection_context_llm(
    rows: list[dict[str, Any]],
    *,
    client: Any,
) -> list[dict[str, Any]]:
    if not rows or client is None:
        return []
    request_kwargs = {
        "model": DOCUMENT_REVIEW_MODEL,
        "messages": [{"role": "user", "content": _build_inflection_context_prompt(rows)}],
    }
    if not _resolve_openai_base_url():
        request_kwargs["response_format"] = {"type": "json_object"}
    response = client.chat.completions.create(**request_kwargs)
    content = response.choices[0].message.content or "{}"
    parsed = json.loads(_extract_json_payload(content))
    if isinstance(parsed, list):
        items = parsed
    elif isinstance(parsed, dict):
        items = parsed.get("items") or parsed.get("results") or parsed.get("result") or []
    else:
        items = []
    return [item for item in items if isinstance(item, dict)]


def _review_inflection_contexts_with_llm(
    rows: list[dict[str, Any]],
    *,
    row_ids: set[str],
    tool_runner: PhilologistToolRunner,
) -> dict[str, Any]:
    if not PHILOLOGIST_CONTEXT_LLM:
        return {"enabled": False, "checked_count": 0, "items": [], "error": ""}
    selected = _select_inflection_contexts_for_llm(
        rows,
        row_ids=row_ids,
        limit=PHILOLOGIST_CONTEXT_LLM_MAX_ITEMS,
    )
    if not selected:
        return {"enabled": True, "checked_count": 0, "items": [], "error": ""}

    client = _build_llm_client()
    if client is None:
        return {
            "enabled": True,
            "checked_count": 0,
            "items": [],
            "error": "LLM client is unavailable.",
            "selected_count": len(selected),
        }

    normalized_items: list[dict[str, Any]] = []
    error = ""
    for offset in range(0, len(selected), PHILOLOGIST_CONTEXT_LLM_BATCH_SIZE):
        batch = selected[offset : offset + PHILOLOGIST_CONTEXT_LLM_BATCH_SIZE]
        try:
            raw_items = tool_runner.call(
                "review_inflection_contexts",
                {"count": len(batch)},
                lambda batch=batch, client=client: _call_inflection_context_llm(batch, client=client),
            )
        except Exception as exc:  # pragma: no cover
            error = f"{type(exc).__name__}: {exc}"
            break

        batch_by_id = {index: row for index, row in enumerate(batch, start=1)}
        for item in raw_items:
            try:
                item_id = int(item.get("id"))
            except (TypeError, ValueError):
                continue
            source_row = batch_by_id.get(item_id)
            if not source_row:
                continue
            status = _safe_text(item.get("status")).lower() or "needs_review"
            if status not in {"ok", "fix", "needs_review"}:
                status = "needs_review"
            corrected_value = _safe_text(item.get("corrected_value")) or _safe_text(source_row.get("result_value"))
            try:
                confidence = float(item.get("confidence", 0.0) or 0.0)
            except (TypeError, ValueError):
                confidence = 0.0
            normalized_items.append(
                {
                    "row_id": source_row.get("row_id"),
                    "field": source_row.get("field"),
                    "source_value": source_row.get("source_value"),
                    "generated_value": source_row.get("result_value"),
                    "corrected_value": corrected_value,
                    "status": status,
                    "confidence": confidence,
                    "comment": _safe_text(item.get("comment")),
                    "filled_sentence": source_row.get("filled_sentence"),
                    "context_sentence": source_row.get("context_sentence"),
                }
            )

    return {
        "enabled": True,
        "selected_count": len(selected),
        "checked_count": len(normalized_items),
        "items": normalized_items,
        "error": error,
    }


def _apply_inflection_context_corrections(
    docx_paths: list[Path],
    context_review: dict[str, Any],
) -> dict[str, Any]:
    corrections = []
    for item in context_review.get("items") or []:
        if not isinstance(item, dict):
            continue
        try:
            confidence = float(item.get("confidence", 0.0) or 0.0)
        except (TypeError, ValueError):
            confidence = 0.0
        generated_value = _safe_text(item.get("generated_value"))
        corrected_value = _safe_text(item.get("corrected_value"))
        if (
            _safe_text(item.get("status")).lower() == "fix"
            and confidence >= PHILOLOGIST_CONTEXT_LLM_MIN_CONFIDENCE
            and generated_value
            and corrected_value
            and generated_value != corrected_value
        ):
            corrections.append({**item, "confidence": confidence})

    if not corrections:
        return {"applied_count": 0, "documents": [], "corrections": []}

    by_row_id: dict[str, list[dict[str, Any]]] = {}
    for correction in corrections:
        by_row_id.setdefault(_safe_text(correction.get("row_id")), []).append(correction)

    applied_records: list[dict[str, Any]] = []
    changed_documents: set[str] = set()
    for docx_path in docx_paths:
        row_id = _extract_row_id_from_docx_path(docx_path)
        row_corrections = by_row_id.get(row_id)
        if not row_corrections:
            continue

        doc = Document(docx_path)
        document_changed = False
        for correction in row_corrections:
            generated_value = _safe_text(correction.get("generated_value"))
            corrected_value = _safe_text(correction.get("corrected_value"))
            applied_here = False
            for _location, paragraph in _iter_document_paragraphs(doc):
                if generated_value not in _paragraph_text(paragraph):
                    continue
                if _replace_fragment_in_paragraph(paragraph, generated_value, corrected_value):
                    applied_here = True
                    document_changed = True
            if applied_here:
                applied_records.append(
                    {
                        "document": str(docx_path),
                        "row_id": row_id,
                        "field": correction.get("field"),
                        "fragment": generated_value,
                        "suggestion": corrected_value,
                        "confidence": correction.get("confidence"),
                        "comment": correction.get("comment"),
                    }
                )
        if document_changed:
            doc.save(docx_path)
            changed_documents.add(str(docx_path))

    return {
        "applied_count": len(applied_records),
        "documents": sorted(changed_documents),
        "corrections": applied_records,
    }


def _iter_document_paragraphs(doc: Document) -> Iterable[tuple[str, Any]]:
    for index, paragraph in enumerate(doc.paragraphs, 1):
        yield (f"paragraph:{index}", paragraph)

    for table_index, table in enumerate(doc.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            for cell_index, cell in enumerate(row.cells, 1):
                for paragraph_index, paragraph in enumerate(cell.paragraphs, 1):
                    yield (
                        f"table:{table_index}:row:{row_index}:cell:{cell_index}:paragraph:{paragraph_index}",
                        paragraph,
                    )


def _replace_paragraph_text(paragraph, new_text: str) -> bool:
    current_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
    if current_text == new_text:
        return False

    non_empty_runs = [run for run in paragraph.runs if run.text]
    if len(non_empty_runs) > 1:
        return False

    if paragraph.runs:
        target_run = non_empty_runs[0] if non_empty_runs else paragraph.runs[0]
        target_run.text = new_text
        for run in paragraph.runs:
            if run is not target_run:
                run.text = ""
    else:
        paragraph.add_run(new_text)
    return True


def _run_format_signature(run) -> tuple[Any, ...]:
    color = None
    if run.font.color is not None and run.font.color.rgb is not None:
        color = str(run.font.color.rgb)
    size = run.font.size.pt if run.font.size is not None else None
    return (
        run.bold,
        run.italic,
        run.underline,
        run.font.name,
        size,
        color,
    )


def _paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text


def _paragraph_style_fingerprint(paragraph) -> dict[str, Any]:
    non_empty_runs = [run for run in paragraph.runs if run.text]
    run_signatures = sorted(repr(_run_format_signature(run)) for run in non_empty_runs)
    style_name = ""
    alignment = ""
    try:
        style_name = paragraph.style.name if paragraph.style is not None else ""
    except Exception:
        style_name = ""
    try:
        alignment = str(paragraph.alignment)
    except Exception:
        p_pr = getattr(getattr(paragraph, "_p", None), "pPr", None)
        jc = getattr(p_pr, "jc", None) if p_pr is not None else None
        alignment = jc.get(qn("w:val")) if jc is not None else ""
    return {
        "paragraph_style": style_name,
        "alignment": alignment,
        "run_styles": run_signatures,
    }


def _collect_docx_snapshot(docx_path: Path) -> dict[str, Any]:
    doc = Document(docx_path)
    locations: dict[str, dict[str, Any]] = {}
    for location, paragraph in _iter_document_paragraphs(doc):
        locations[location] = {
            "text": _paragraph_text(paragraph),
            "style": _paragraph_style_fingerprint(paragraph),
        }
    return {
        "path": str(docx_path),
        "paragraph_count": len(locations),
        "locations": locations,
    }


def _verify_safe_fixes(
    docx_path: Path,
    before_snapshot: dict[str, Any],
    fix_result: dict[str, Any],
) -> dict[str, Any]:
    after_snapshot = _collect_docx_snapshot(docx_path)
    before_locations = before_snapshot.get("locations") or {}
    after_locations = after_snapshot.get("locations") or {}
    applied_fixes = [fix for fix in fix_result.get("applied_fixes", []) if isinstance(fix, dict)]
    expected_locations = {_safe_text(fix.get("location")) for fix in applied_fixes if _safe_text(fix.get("location"))}
    changed_locations = {
        location
        for location, before_item in before_locations.items()
        if (after_locations.get(location) or {}).get("text") != before_item.get("text")
    }
    warnings: list[dict[str, str]] = []

    for location in sorted(changed_locations - expected_locations):
        warnings.append(
            {
                "location": location,
                "reason": "unexpected_text_change",
                "message": "Текст изменился вне списка применённых автоправок.",
            }
        )

    for location in sorted(expected_locations - changed_locations):
        warnings.append(
            {
                "location": location,
                "reason": "expected_change_missing",
                "message": "Автоправка заявлена, но текст в этой локации не изменился.",
            }
        )

    for location in sorted(changed_locations):
        before_style = (before_locations.get(location) or {}).get("style")
        after_style = (after_locations.get(location) or {}).get("style")
        if before_style != after_style:
            warnings.append(
                {
                    "location": location,
                    "reason": "style_fingerprint_changed",
                    "message": "После автоправки изменился отпечаток стиля. Нужна ручная проверка.",
                }
            )

    for fix in applied_fixes:
        location = _safe_text(fix.get("location"))
        suggestion = _safe_text(fix.get("suggestion"))
        after_text = _safe_text((after_locations.get(location) or {}).get("text"))
        if suggestion and suggestion not in after_text:
            warnings.append(
                {
                    "location": location,
                    "reason": "suggestion_not_found_after_fix",
                    "message": "Предложенный текст не найден после применения правки.",
                }
            )

    return {
        "verified": not warnings,
        "warning_count": len(warnings),
        "warnings": warnings,
        "changed_locations": sorted(changed_locations),
        "expected_locations": sorted(expected_locations),
        "after_paragraph_count": after_snapshot.get("paragraph_count", 0),
    }


PHILOLOGIST_REACT_ACTIONS = (
    "review_docx",
    "snapshot_docx",
    "apply_safe_fixes",
    "verify_safe_fixes",
    "rebuild_pdf",
    "finish_document",
)

PHILOLOGIST_FIX_STRATEGIES = (
    AUTO_FIX,
    QUARANTINE,
    NEEDS_HUMAN,
    "skip",
)


def _allowed_fix_strategy_actions(base_decision: dict[str, Any]) -> list[str]:
    base_action = _safe_text(base_decision.get("action"))
    if base_action == AUTO_FIX:
        return [AUTO_FIX, QUARANTINE, NEEDS_HUMAN, "skip"]
    if base_action == QUARANTINE:
        return [QUARANTINE, NEEDS_HUMAN, "skip"]
    if base_action == NEEDS_HUMAN:
        return [NEEDS_HUMAN, QUARANTINE, "skip"]
    return ["skip", NEEDS_HUMAN, QUARANTINE]


def _fallback_fix_strategy(
    base_decision: dict[str, Any],
    rag: dict[str, Any],
    *,
    reason: str,
) -> dict[str, Any]:
    action = _safe_text(base_decision.get("action")) or "skip"
    allowed = _allowed_fix_strategy_actions(base_decision)
    if action not in allowed:
        action = allowed[0]
    return {
        "action": action,
        "thought": "Использую безопасное базовое решение, потому что агентная стратегия недоступна.",
        "reason": reason or _safe_text(base_decision.get("reason")),
        "confidence": float(base_decision.get("confidence") or 0.0),
        "source": "fallback_strategy",
        "allowed_actions": allowed,
        "rag_recommendation": _safe_text(rag.get("recommendation")),
        "rag_support_score": int(rag.get("support_score") or 0),
    }


def _react_decide_fix_strategy(
    *,
    client: Any,
    issue: dict[str, Any],
    current_text: str,
    base_decision: dict[str, Any],
    rag: dict[str, Any],
    use_llm: bool = True,
) -> dict[str, Any]:
    allowed = _allowed_fix_strategy_actions(base_decision)
    if not use_llm or not client:
        return _fallback_fix_strategy(
            base_decision,
            rag,
            reason="Стратегия выбрана по безопасным правилам без LLM-вызова.",
        )

    compact_rules = [
        {
            "id": rule.get("id", ""),
            "title": rule.get("title", ""),
            "score": rule.get("score", 0),
            "matched_terms": rule.get("matched_terms", ""),
            "rule": rule.get("rule", ""),
        }
        for rule in (rag.get("rules") or [])[:3]
        if isinstance(rule, dict)
    ]
    prompt = (
        "Ты ReAct-стратег агента-филолога. "
        "Выбери, что делать с конкретной правкой DOCX. "
        "Можно выбрать только action из allowed_actions. "
        "auto_fix разрешен только если он есть в allowed_actions; если есть риск сломать стиль, выбирай quarantine или needs_human. "
        "Ответ верни только JSON без markdown:\n"
        "{\n"
        '  "thought": "краткое рассуждение",\n'
        '  "action": "auto_fix|quarantine|needs_human|skip",\n'
        '  "reason": "почему выбран этот вариант",\n'
        '  "confidence": 0.0\n'
        "}\n\n"
        f"allowed_actions: {json.dumps(allowed, ensure_ascii=False)}\n"
        f"issue: {json.dumps(issue, ensure_ascii=False)}\n"
        f"current_text: {json.dumps(_safe_text(current_text), ensure_ascii=False)}\n"
        f"base_decision: {json.dumps(base_decision, ensure_ascii=False)}\n"
        f"rag: {json.dumps({'support_score': rag.get('support_score'), 'recommendation': rag.get('recommendation'), 'reason': rag.get('reason'), 'rules': compact_rules}, ensure_ascii=False)}"
    )
    try:
        response = client.chat.completions.create(
            model=DOCUMENT_REVIEW_MODEL,
            messages=[{"role": "user", "content": prompt}],
        )
        payload = _extract_json_payload(_safe_text(response.choices[0].message.content))
        parsed = json.loads(payload) if payload else {}
    except Exception as exc:
        decision = _fallback_fix_strategy(base_decision, rag, reason=f"LLM-стратегия недоступна: {_safe_text(exc)}")
        decision["source"] = "fallback_strategy_after_llm_error"
        return decision

    action = _safe_text(parsed.get("action"))
    if action not in allowed:
        decision = _fallback_fix_strategy(
            base_decision,
            rag,
            reason=f"LLM выбрала недопустимую стратегию `{action}`, применен безопасный fallback.",
        )
        decision["source"] = "fallback_invalid_strategy"
        decision["llm_action"] = action
        return decision

    try:
        confidence = float(parsed.get("confidence"))
    except (TypeError, ValueError):
        confidence = float(base_decision.get("confidence") or 0.0)
    return {
        "action": action,
        "thought": _safe_text(parsed.get("thought")),
        "reason": _safe_text(parsed.get("reason")) or _safe_text(base_decision.get("reason")),
        "confidence": max(0.0, min(1.0, confidence)),
        "source": "llm_strategy",
        "allowed_actions": allowed,
        "rag_recommendation": _safe_text(rag.get("recommendation")),
        "rag_support_score": int(rag.get("support_score") or 0),
    }


def _quarantine_issue_record(issue: dict[str, Any], decision: dict[str, Any]) -> dict[str, str]:
    return {
        "status": "quarantine",
        "location": _safe_text(issue.get("location")),
        "issue": _safe_text(issue.get("issue")),
        "fragment": _safe_text(issue.get("fragment")),
        "suggestion": _safe_text(issue.get("suggestion")),
        "action": _safe_text(decision.get("action")),
        "reason": _safe_text(decision.get("reason")),
        "source": _safe_text(decision.get("strategy_source") or decision.get("source")),
    }


def _available_react_actions(context: dict[str, Any]) -> list[str]:
    if not context.get("review_done"):
        return ["review_docx"]
    if int(context.get("issue_count", 0) or 0) <= 0:
        return ["finish_document"]
    if not context.get("snapshot_done"):
        return ["snapshot_docx"]
    if not context.get("fix_done"):
        return ["apply_safe_fixes"]
    if int(context.get("applied_fix_count", 0) or 0) <= 0:
        return ["finish_document"]
    if not context.get("verify_done"):
        return ["verify_safe_fixes"]
    if (
        context.get("rebuild_pdf_enabled")
        and context.get("applied_fix_count", 0) > 0
        and not context.get("pdf_rebuilt")
    ):
        return ["rebuild_pdf"]
    return ["finish_document"]


def _default_react_decision(context: dict[str, Any], *, reason: str) -> dict[str, Any]:
    available = _available_react_actions(context)
    action = available[0]
    return {
        "action": action,
        "thought": "Выбираю следующий безопасный обязательный шаг.",
        "reason": reason,
        "source": "fallback",
        "available_actions": available,
    }


def _compact_react_context(context: dict[str, Any]) -> dict[str, Any]:
    return {
        "document": context.get("document"),
        "review_done": bool(context.get("review_done")),
        "snapshot_done": bool(context.get("snapshot_done")),
        "fix_done": bool(context.get("fix_done")),
        "verify_done": bool(context.get("verify_done")),
        "pdf_rebuilt": bool(context.get("pdf_rebuilt")),
        "issue_count": int(context.get("issue_count", 0) or 0),
        "applied_fix_count": int(context.get("applied_fix_count", 0) or 0),
        "skipped_fix_count": int(context.get("skipped_fix_count", 0) or 0),
        "verification_warning_count": int(context.get("verification_warning_count", 0) or 0),
    }


def _react_decide_next_action(
    *,
    client: Any,
    context: dict[str, Any],
    trace: list[dict[str, Any]],
    use_llm: bool = True,
) -> dict[str, Any]:
    available = _available_react_actions(context)
    if not use_llm or not client:
        return _default_react_decision(context, reason="Использую быстрый безопасный порядок без LLM-вызова.")

    prompt = (
        "Ты ReAct-контроллер агента-филолога. "
        "Выбери ровно один следующий action из available_actions. "
        "Нельзя придумывать новые tools. Нельзя пропускать самопроверку после правок. "
        "Если документ уже проверен, правки применены и самопроверка завершена, выбирай finish_document. "
        "Ответ верни только JSON без markdown:\n"
        "{\n"
        '  "thought": "краткое рассуждение",\n'
        '  "action": "одно значение из available_actions",\n'
        '  "reason": "почему выбран этот шаг"\n'
        "}\n\n"
        f"available_actions: {json.dumps(available, ensure_ascii=False)}\n"
        f"document_state: {json.dumps(_compact_react_context(context), ensure_ascii=False)}\n"
        f"recent_trace: {json.dumps(trace[-6:], ensure_ascii=False)}"
    )
    try:
        response = client.chat.completions.create(
            model=DOCUMENT_REVIEW_MODEL,
            messages=[{"role": "user", "content": prompt}],
        )
        payload = _extract_json_payload(_safe_text(response.choices[0].message.content))
        parsed = json.loads(payload) if payload else {}
    except Exception as exc:
        decision = _default_react_decision(context, reason=f"LLM-решение недоступно: {_safe_text(exc)}")
        decision["source"] = "fallback_after_llm_error"
        return decision

    action = _safe_text(parsed.get("action"))
    if action not in available:
        decision = _default_react_decision(
            context,
            reason=f"LLM выбрала недопустимый action `{action}`, применен безопасный fallback.",
        )
        decision["source"] = "fallback_invalid_llm_action"
        decision["llm_action"] = action
        return decision

    return {
        "action": action,
        "thought": _safe_text(parsed.get("thought")),
        "reason": _safe_text(parsed.get("reason")),
        "source": "llm",
        "available_actions": available,
    }


def _run_docx_react_loop(
    *,
    docx_path: Path,
    ai_enabled: bool,
    tool_runner: PhilologistToolRunner,
    client: Any,
    use_llm_router: bool = True,
    use_llm_fix_strategy: bool = True,
    rebuild_pdf_enabled: bool = False,
    use_rag_decisions: bool = True,
    template_memory: dict[str, Any] | None = None,
    max_iterations: int = 8,
) -> dict[str, Any]:
    context: dict[str, Any] = {
        "document": docx_path.name,
        "review_done": False,
        "snapshot_done": False,
        "fix_done": False,
        "verify_done": False,
        "pdf_rebuilt": False,
        "issue_count": 0,
        "applied_fix_count": 0,
        "skipped_fix_count": 0,
        "verification_warning_count": 0,
        "rebuild_pdf_enabled": bool(rebuild_pdf_enabled),
    }
    trace: list[dict[str, Any]] = []
    review_result: dict[str, Any] = {"issues": [], "issue_count": 0, "local_issue_count": 0, "ai_issue_count": 0}
    before_snapshot: dict[str, Any] = {}
    fix_result: dict[str, Any] = {
        "applied_fix_count": 0,
        "applied_fixes": [],
        "skipped_fix_count": 0,
        "skipped_fixes": [],
        "decision_count": 0,
        "fix_decisions": [],
    }
    verification_result: dict[str, Any] = {
        "verified": True,
        "warning_count": 0,
        "warnings": [],
    }
    pdf_path: str | None = None

    for iteration in range(1, max_iterations + 1):
        decision = tool_runner.call(
            "react_decide_next_action",
            {"available_actions": _available_react_actions(context), "document_state": _compact_react_context(context)},
            lambda context=context, trace=trace: _react_decide_next_action(
                client=client,
                context=context,
                trace=trace,
                use_llm=use_llm_router,
            ),
        )
        if decision.get("source") == "fallback_after_llm_error":
            client = None
        action = decision["action"]
        observation = ""

        if action == "review_docx":
            review_result = tool_runner.call(
                "review_docx",
                {"path": str(docx_path), "ai_enabled": ai_enabled},
                lambda docx_path=docx_path, template_memory=template_memory: review_docx(
                    docx_path,
                    ai_enabled=ai_enabled,
                    template_memory=template_memory,
                ),
            )
            context["review_done"] = True
            context["issue_count"] = int(review_result.get("issue_count", 0) or 0)
            review_mode = _safe_text(review_result.get("review_mode")) or "full"
            reviewed_block_count = int(review_result.get("reviewed_block_count", 0) or 0)
            total_block_count = int(review_result.get("total_block_count", 0) or 0)
            observation = (
                f"Найдено замечаний: {context['issue_count']}. "
                f"Режим проверки: {'прицельный' if review_mode == 'targeted' else 'полный'} "
                f"({reviewed_block_count}/{total_block_count} блоков)."
            )
        elif action == "snapshot_docx":
            before_snapshot = tool_runner.call(
                "snapshot_docx",
                {"path": str(docx_path)},
                lambda docx_path=docx_path: _collect_docx_snapshot(docx_path),
            )
            context["snapshot_done"] = True
            observation = f"Snapshot создан, локаций: {before_snapshot.get('paragraph_count', 0)}."
        elif action == "apply_safe_fixes":
            fix_result = tool_runner.call(
                "apply_safe_fixes",
                {"path": str(docx_path), "issue_count": int(review_result.get("issue_count", 0))},
                lambda docx_path=docx_path, review_result=review_result, client=client, tool_runner=tool_runner: _auto_fix_docx(
                    docx_path,
                    review_result,
                    client=client,
                    tool_runner=tool_runner,
                    use_llm_strategy=use_llm_fix_strategy,
                    use_rag_decisions=use_rag_decisions,
                ),
            )
            context["fix_done"] = True
            context["applied_fix_count"] = int(fix_result.get("applied_fix_count", 0) or 0)
            context["skipped_fix_count"] = int(fix_result.get("skipped_fix_count", 0) or 0)
            observation = (
                f"Применено правок: {context['applied_fix_count']}; "
                f"отложено: {context['skipped_fix_count']}."
            )
        elif action == "verify_safe_fixes":
            verification_result = tool_runner.call(
                "verify_safe_fixes",
                {"path": str(docx_path), "applied_fix_count": int(fix_result.get("applied_fix_count", 0))},
                lambda docx_path=docx_path, before_snapshot=before_snapshot, fix_result=fix_result: _verify_safe_fixes(
                    docx_path,
                    before_snapshot,
                    fix_result,
                ),
            )
            context["verify_done"] = True
            context["verification_warning_count"] = int(verification_result.get("warning_count", 0) or 0)
            observation = f"Самопроверка завершена, предупреждений: {context['verification_warning_count']}."
        elif action == "rebuild_pdf":
            pdf_path = tool_runner.call(
                "rebuild_pdf",
                {"path": str(docx_path)},
                lambda docx_path=docx_path: _rebuild_pdf_for_docx(docx_path),
            )
            context["pdf_rebuilt"] = True
            observation = f"PDF пересобран: {bool(pdf_path)}."
        elif action == "finish_document":
            observation = "Документальный цикл завершен."
            trace.append(
                {
                    "iteration": iteration,
                    **decision,
                    "observation": observation,
                    "state": _compact_react_context(context),
                }
            )
            break

        trace.append(
            {
                "iteration": iteration,
                **decision,
                "observation": observation,
                "state": _compact_react_context(context),
            }
        )
    else:
        trace.append(
            {
                "iteration": max_iterations,
                "action": "max_iterations",
                "thought": "Цикл достиг лимита итераций.",
                "reason": "Нужна ручная проверка, чтобы не зациклиться.",
                "source": "system",
                "observation": "Документ не был завершен через finish_document.",
                "state": _compact_react_context(context),
            }
        )

    return {
        "review_result": review_result,
        "fix_result": fix_result,
        "verification_result": verification_result,
        "pdf_path": pdf_path,
        "react_trace": trace,
        "react_context": _compact_react_context(context),
    }


def _paragraph_is_safe_for_text_rewrite(paragraph) -> bool:
    non_empty_runs = [run for run in paragraph.runs if run.text]
    if len(non_empty_runs) <= 1:
        return True
    signatures = {_run_format_signature(run) for run in non_empty_runs}
    return len(signatures) == 1


def _replace_fragment_in_paragraph(paragraph, fragment: str, replacement: str) -> bool:
    current_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
    if not fragment or not replacement or fragment not in current_text:
        return False
    if _replace_fragment_inside_single_run(paragraph, fragment, replacement):
        return True
    if not _paragraph_is_safe_for_text_rewrite(paragraph):
        return False
    new_text = current_text.replace(fragment, replacement, 1)
    if new_text == current_text:
        return False
    return _replace_paragraph_text(paragraph, new_text)


def _replace_all_safe_fragments_in_paragraph(paragraph, fragment: str, replacement: str) -> bool:
    if fragment and fragment in replacement:
        return _replace_fragment_in_paragraph(paragraph, fragment, replacement)

    changed = False
    replacement_count = 0
    while True:
        current_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
        if fragment not in current_text:
            break
        if not _replace_fragment_in_paragraph(paragraph, fragment, replacement):
            break
        changed = True
        replacement_count += 1
        if replacement_count > 100:
            break
    return changed


def _replace_fragment_inside_single_run(paragraph, fragment: str, replacement: str) -> bool:
    for run in paragraph.runs:
        if fragment not in run.text:
            continue
        new_text = run.text.replace(fragment, replacement, 1)
        if new_text == run.text:
            return False
        run.text = new_text
        return True
    return False


def _fragment_exists_inside_single_run(paragraph, fragment: str) -> bool:
    return any(fragment and fragment in run.text for run in paragraph.runs)


def _normalize_double_spaces_in_runs(paragraph) -> bool:
    changed = False
    for run in paragraph.runs:
        normalized = _normalize_double_spaces(run.text)
        if normalized != run.text:
            run.text = normalized
            changed = True
    return changed


def _normalize_double_spaces(text: str) -> str:
    while "  " in text:
        text = text.replace("  ", " ")
    return text


EDITORIAL_SUGGESTION_PREFIXES = (
    "заменить ",
    "исправить ",
    "нужно ",
    "следует ",
    "проверить ",
    "убрать ",
)


def _looks_like_editorial_instruction(text: str) -> bool:
    normalized = _safe_text(text).strip().lower()
    if not normalized:
        return False
    if any(normalized.startswith(prefix) for prefix in EDITORIAL_SUGGESTION_PREFIXES):
        return True
    if "заменить" in normalized and (" на " in normalized or '"' in normalized or "«" in normalized):
        return True
    return False


def _apply_issue_to_document(location_map: dict[str, Any], issue: dict[str, Any]) -> dict[str, Any]:
    location = _safe_text(issue.get("location"))
    paragraph = location_map.get(location)
    if paragraph is None:
        return {"applied": False, "reason": "location_not_found"}

    issue_text = _safe_text(issue.get("issue")).lower()
    fragment = _safe_text(issue.get("fragment"))
    suggestion = _safe_text(issue.get("suggestion"))
    current_text = paragraph.text

    if "двойные пробелы" in issue_text:
        if _normalize_double_spaces_in_runs(paragraph):
            return {"applied": True, "mode": "run_normalization"}
        normalized = _normalize_double_spaces(current_text)
        applied = _replace_paragraph_text(paragraph, normalized)
        return {"applied": applied, "mode": "paragraph_rewrite" if applied else "", "reason": "" if applied else "unsafe_multirun"}

    if _looks_like_editorial_instruction(suggestion):
        return {"applied": False, "reason": "editorial_instruction"}

    if issue.get("source") == "local" and suggestion:
        if fragment and fragment in current_text:
            if not _paragraph_is_safe_for_text_rewrite(paragraph) and not _fragment_exists_inside_single_run(paragraph, fragment):
                return {"applied": False, "reason": "unsafe_formatting"}
            applied = _replace_all_safe_fragments_in_paragraph(paragraph, fragment, suggestion)
            return {"applied": applied, "mode": "fragment_replace" if applied else "", "reason": "" if applied else "unsafe_or_missing_fragment"}
        if current_text.strip() == fragment.strip():
            if not _paragraph_is_safe_for_text_rewrite(paragraph):
                return {"applied": False, "reason": "unsafe_formatting"}
            applied = _replace_paragraph_text(paragraph, suggestion)
            return {"applied": applied, "mode": "paragraph_rewrite" if applied else "", "reason": "" if applied else "unsafe_multirun"}

    if "после числа используется неверная форма слова" in issue_text and suggestion:
        if not _paragraph_is_safe_for_text_rewrite(paragraph):
            return {"applied": False, "reason": "unsafe_formatting"}
        applied = _replace_paragraph_text(paragraph, suggestion)
        return {"applied": applied, "mode": "paragraph_rewrite" if applied else "", "reason": "" if applied else "unsafe_multirun"}

    if issue.get("source") == "ai" and suggestion and fragment:
        if fragment in current_text:
            if not _paragraph_is_safe_for_text_rewrite(paragraph) and not _fragment_exists_inside_single_run(paragraph, fragment):
                return {"applied": False, "reason": "unsafe_formatting"}
            applied = _replace_all_safe_fragments_in_paragraph(paragraph, fragment, suggestion)
            return {"applied": applied, "mode": "fragment_replace" if applied else "", "reason": "" if applied else "unsafe_or_missing_fragment"}
        if current_text.strip() == fragment.strip():
            if not _paragraph_is_safe_for_text_rewrite(paragraph):
                return {"applied": False, "reason": "unsafe_formatting"}
            applied = _replace_paragraph_text(paragraph, suggestion)
            return {"applied": applied, "mode": "paragraph_rewrite" if applied else "", "reason": "" if applied else "unsafe_multirun"}
        if issue.get("severity") == "error":
            if not _paragraph_is_safe_for_text_rewrite(paragraph):
                return {"applied": False, "reason": "unsafe_formatting"}
            applied = _replace_paragraph_text(paragraph, current_text.replace(fragment, suggestion))
            return {"applied": applied, "mode": "paragraph_rewrite" if applied else "", "reason": "" if applied else "unsafe_multirun"}

    if issue.get("source") == "ai" and suggestion and not fragment and issue.get("severity") == "error":
        return {"applied": False, "reason": "ai_full_paragraph_rewrite_blocked"}

    return {"applied": False, "reason": "no_safe_fix"}


def _format_skipped_fix(
    issue: dict[str, Any],
    result: dict[str, Any],
    decision: dict[str, Any] | None = None,
) -> dict[str, str]:
    decision = decision or {}
    return {
        "location": _safe_text(issue.get("location")),
        "fragment": _safe_text(issue.get("fragment")),
        "issue": _safe_text(issue.get("issue")),
        "suggestion": _safe_text(issue.get("suggestion")),
        "reason": _safe_text(result.get("reason")) or "not_safe_to_apply",
        "decision_action": _safe_text(decision.get("action")),
        "decision_reason": _safe_text(decision.get("reason")),
        "decision_confidence": _safe_text(decision.get("confidence")),
        "strategy_source": _safe_text(decision.get("strategy_source")),
        "strategy_reason": _safe_text(decision.get("strategy_reason")),
        "rag_recommendation": _safe_text((decision.get("rag") or {}).get("recommendation")),
        "rag_reason": _safe_text((decision.get("rag") or {}).get("reason")),
        "rag_support_score": _safe_text((decision.get("rag") or {}).get("support_score")),
    }


def _auto_fix_docx(
    docx_path: Path,
    review_result: dict[str, Any],
    *,
    client: Any = None,
    tool_runner: PhilologistToolRunner | None = None,
    use_llm_strategy: bool = True,
    use_rag_decisions: bool = True,
) -> dict[str, Any]:
    doc = Document(docx_path)
    location_map = {location: paragraph for location, paragraph in _iter_document_paragraphs(doc)}
    applied_fixes: list[dict[str, str]] = []
    skipped_fixes: list[dict[str, str]] = []
    fix_decisions: list[dict[str, Any]] = []

    for issue in review_result.get("issues", []):
        if not isinstance(issue, dict):
            continue
        location = _safe_text(issue.get("location"))
        paragraph = location_map.get(location)
        current_text = paragraph.text if paragraph is not None else ""
        base_decision = decide_issue_fix(issue, current_text=current_text).to_dict()
        if use_rag_decisions and tool_runner:
            rag = tool_runner.call(
                "lookup_rag_rule",
                {"issue": issue, "base_decision": base_decision},
                lambda base_decision=base_decision: explain_fix_decision_with_rag(base_decision),
            )
            strategy = tool_runner.call(
                "decide_fix_strategy",
                {
                    "issue": issue,
                    "rag": rag,
                    "allowed_actions": _allowed_fix_strategy_actions(base_decision),
                },
                lambda issue=issue, current_text=current_text, base_decision=base_decision, rag=rag: _react_decide_fix_strategy(
                    client=client,
                    issue=issue,
                    current_text=current_text,
                    base_decision=base_decision,
                    rag=rag,
                    use_llm=use_llm_strategy,
                ),
            )
        elif use_rag_decisions:
            rag = explain_fix_decision_with_rag(base_decision)
            strategy = _react_decide_fix_strategy(
                client=client,
                issue=issue,
                current_text=current_text,
                base_decision=base_decision,
                rag=rag,
                use_llm=use_llm_strategy,
            )
        else:
            rag = {
                "query": "",
                "support_score": 0,
                "recommendation": "skipped_in_fast_mode",
                "reason": "RAG-объяснение отключено в быстром массовом режиме.",
                "rules": [],
            }
            strategy = _fallback_fix_strategy(
                base_decision,
                rag,
                reason="Быстрый режим: стратегия выбрана локальным правилом без RAG/LLM.",
            )
        decision = dict(base_decision)
        decision["base_action"] = base_decision.get("action")
        decision["base_reason"] = base_decision.get("reason")
        decision["action"] = strategy.get("action") or base_decision.get("action")
        decision["reason"] = strategy.get("reason") or base_decision.get("reason")
        decision["confidence"] = strategy.get("confidence", base_decision.get("confidence"))
        decision["strategy_source"] = strategy.get("source", "")
        decision["strategy_reason"] = strategy.get("reason", "")
        decision["strategy_thought"] = strategy.get("thought", "")
        decision["rag"] = rag
        fix_decisions.append(decision)
        if decision.get("action") != AUTO_FIX:
            if issue.get("suggestion"):
                quarantine_record = (
                    tool_runner.call(
                        "quarantine_issue",
                        {"issue": issue, "decision": decision},
                        lambda issue=issue, decision=decision: _quarantine_issue_record(issue, decision),
                    )
                    if tool_runner
                    else _quarantine_issue_record(issue, decision)
                )
                skipped_fixes.append(
                    _format_skipped_fix(
                        issue,
                        {"reason": quarantine_record.get("reason") or decision.get("action") or "not_auto_fix"},
                        decision,
                    )
                )
            continue
        apply_result = _apply_issue_to_document(location_map, issue)
        if apply_result.get("applied"):
            applied_fixes.append(
                {
                    "location": _safe_text(issue.get("location")),
                    "fragment": _safe_text(issue.get("fragment")),
                    "issue": _safe_text(issue.get("issue")),
                    "suggestion": _safe_text(issue.get("suggestion")),
                    "mode": _safe_text(apply_result.get("mode")),
                    "decision_action": _safe_text(decision.get("action")),
                    "decision_reason": _safe_text(decision.get("reason")),
                    "decision_confidence": _safe_text(decision.get("confidence")),
                    "strategy_source": _safe_text(decision.get("strategy_source")),
                    "strategy_reason": _safe_text(decision.get("strategy_reason")),
                    "rag_recommendation": _safe_text((decision.get("rag") or {}).get("recommendation")),
                    "rag_reason": _safe_text((decision.get("rag") or {}).get("reason")),
                    "rag_support_score": _safe_text((decision.get("rag") or {}).get("support_score")),
                }
            )
        elif issue.get("suggestion"):
            skipped_fixes.append(_format_skipped_fix(issue, apply_result, decision))

    if applied_fixes:
        doc.save(docx_path)

    return {
        "applied_fix_count": len(applied_fixes),
        "applied_fixes": applied_fixes,
        "skipped_fix_count": len(skipped_fixes),
        "skipped_fixes": skipped_fixes,
        "decision_count": len(fix_decisions),
        "fix_decisions": fix_decisions,
    }


def _rebuild_pdf_for_docx(docx_path: Path) -> str | None:
    pdf_map = convert_docx_batch([docx_path], docx_path.parent, chunk_size=1, worker_count=1)
    pdf_path = pdf_map.get(docx_path)
    return str(pdf_path) if pdf_path and pdf_path.exists() else None


def _format_summary(documents: list[dict[str, Any]]) -> str:
    total = len(documents)
    with_issues = sum(1 for item in documents if item.get("issue_count", 0) > 0)
    fixed = sum(1 for item in documents if item.get("applied_fix_count", 0) > 0)
    if total == 0:
        return "Готовые документы для проверки пока не найдены."
    if with_issues == 0:
        return f"Проверил {total} документов. Явных языковых ошибок не нашёл."
    return (
        f"Проверил {total} документов. "
        f"Замечания нашёл в {with_issues}, автоматически исправил {fixed}. "
        "Если хочешь, могу показать, что именно исправил и где остались спорные места."
    )


def _format_run_summary(documents: list[dict[str, Any]], *, sender_handoffs: int = 0) -> str:
    base = _format_summary(documents)
    if sender_handoffs > 0:
        return base + f" Передал отправщику задач на дополнительную проверку перед отправкой: {sender_handoffs}."
    return base


def _format_fixed_details(documents: list[dict[str, Any]], limit: int = 5) -> str:
    lines: list[str] = []
    for item in documents:
        fixes = item.get("applied_fixes") or []
        if not fixes:
            continue
        lines.append(f"{item['name']}:")
        for fix in fixes[:2]:
            issue = _safe_text(fix.get("issue")) or "языковая правка"
            suggestion = _safe_text(fix.get("suggestion"))
            details = issue if not suggestion else f"{issue} -> {suggestion}"
            lines.append(f"- {details}")
        if len(lines) >= limit * 2:
            break
    if not lines:
        return "Автоматических исправлений пока не было. Я только нашёл замечания."
    return "Исправил:\n" + "\n".join(lines[: limit * 2])


def _format_issue_details(documents: list[dict[str, Any]], limit: int = 5) -> str:
    lines: list[str] = []
    for item in documents:
        issues = item.get("issues") or []
        if not issues:
            continue
        lines.append(f"{item['name']}:")
        for issue in issues[:2]:
            issue_text = _safe_text(issue.get("issue")) or "есть замечание"
            suggestion = _safe_text(issue.get("suggestion"))
            details = issue_text if not suggestion else f"{issue_text} -> {suggestion}"
            lines.append(f"- {details}")
        if len(lines) >= limit * 2:
            break
    if not lines:
        return "Явных языковых проблем в проверенных документах я не нашёл."
    return "Остались замечания:\n" + "\n".join(lines[: limit * 2])


def _format_skipped_fix_details(documents: list[dict[str, Any]], limit: int = 5) -> str:
    lines: list[str] = []
    for item in documents:
        skipped = item.get("skipped_fixes") or []
        if not skipped:
            continue
        lines.append(f"{item['name']}:")
        for fix in skipped[:2]:
            issue = _safe_text(fix.get("issue")) or "правка пропущена"
            reason = _safe_text(fix.get("reason")) or "небезопасно применять автоматически"
            lines.append(f"- {issue} ({reason})")
        if len(lines) >= limit * 2:
            break
    if not lines:
        return ""
    return "Не применил автоматически:\n" + "\n".join(lines[: limit * 2])


def _format_verification_details(documents: list[dict[str, Any]], limit: int = 5) -> str:
    lines: list[str] = []
    for item in documents:
        warnings = item.get("verification_warnings") or []
        if not warnings:
            continue
        lines.append(f"{item['name']}:")
        for warning in warnings[:2]:
            reason = _safe_text(warning.get("reason")) or "verification_warning"
            message = _safe_text(warning.get("message"))
            location = _safe_text(warning.get("location"))
            details = f"{reason}: {message}" if message else reason
            if location:
                details = f"{location} - {details}"
            lines.append(f"- {details}")
        if len(lines) >= limit * 2:
            break
    if not lines:
        return ""
    return "Самопроверка после правок:\n" + "\n".join(lines[: limit * 2])


def _format_fix_decision_details(documents: list[dict[str, Any]], limit: int = 6) -> str:
    lines: list[str] = []
    for item in documents:
        decisions = item.get("fix_decisions") or []
        if not decisions:
            continue
        lines.append(f"{item['name']}:")
        for decision in decisions[:3]:
            action = _safe_text(decision.get("action")) or "unknown"
            reason = _safe_text(decision.get("reason"))
            issue = _safe_text(decision.get("issue")) or "правка"
            strategy_source = _safe_text(decision.get("strategy_source"))
            rag = decision.get("rag") or {}
            rag_text = ""
            if isinstance(rag, dict) and rag.get("recommendation"):
                rag_text = f"; RAG: {rag.get('recommendation')}, score={rag.get('support_score', 0)}"
            strategy_text = f"; strategy={strategy_source}" if strategy_source else ""
            lines.append(f"- {action}: {issue} ({reason}{strategy_text}{rag_text})")
        if len(lines) >= limit:
            break
    if not lines:
        return ""
    return "Решения по правкам:\n" + "\n".join(lines[:limit])


def _format_react_trace_details(documents: list[dict[str, Any]], limit: int = 5) -> str:
    lines: list[str] = []
    for item in documents:
        trace = item.get("react_trace") or []
        if not trace:
            continue
        actions = [
            f"{_safe_text(step.get('action'))}:{_safe_text(step.get('source'))}"
            for step in trace
            if _safe_text(step.get("action"))
        ]
        lines.append(f"{item['name']}: " + " -> ".join(actions[:8]))
        if len(lines) >= limit:
            break
    if not lines:
        return ""
    return "ReAct-цикл:\n" + "\n".join(lines)


def _format_rule_details(message: str) -> str:
    rules = find_relevant_rules(message, limit=4)
    if not rules:
        return "В локальной базе правил я не нашёл точного совпадения по этому вопросу. Могу всё равно проверить формулировку по найденным замечаниям."
    lines = ["Вот на какие правила я могу опереться:"]
    for rule in rules:
        lines.append(
            f"- {_safe_text(rule.get('title'))}: {_safe_text(rule.get('rule'))} "
            f"(источник: {_safe_text(rule.get('source'))})"
        )
    return "\n".join(lines)


def _format_tool_trace(tool_trace: list[dict[str, Any]], limit: int = 8) -> str:
    if not tool_trace:
        return ""
    lines = ["Инструменты:"]
    for record in tool_trace[-limit:]:
        name = _safe_text(record.get("name")) or "tool"
        status = _safe_text(record.get("status")) or "unknown"
        elapsed = record.get("elapsed_seconds")
        elapsed_text = f", {elapsed} сек." if isinstance(elapsed, (int, float)) else ""
        lines.append(f"- {name}: {status}{elapsed_text}")
    return "\n".join(lines)


def _format_plan_summary(plan: dict[str, Any] | None, limit: int = 7) -> str:
    if not plan:
        return ""
    lines = [
        f"План агента: {plan.get('status', 'unknown')}",
        f"Цель: {plan.get('goal', '')}",
    ]
    execution = plan.get("execution") or {}
    if execution:
        lines.append(f"Цикл исполнения: {execution.get('status', 'unknown')}")
    if plan.get("next_step"):
        lines.append(f"Следующий шаг: {plan.get('next_step')}")
    for step in (plan.get("steps") or [])[:limit]:
        lines.append(
            "- "
            f"{step.get('id')}: {step.get('status')} "
            f"({step.get('tool')}) - {step.get('reason')}"
        )
    return "\n".join(lines)


def _format_philologist_structured_reply(message: str, state: dict[str, Any]) -> str:
    documents = state.get("documents") or []
    if not documents:
        parts = [
            "Сводка:\nФилолог ещё не запускался или не нашёл готовых документов для проверки.",
            _format_plan_summary(state.get("plan")),
            _format_rule_details(message),
        ]
        return "\n\n".join(part for part in parts if part)

    issues = [item for item in documents if item.get("issue_count", 0) > 0]
    fixed = [item for item in documents if item.get("applied_fix_count", 0) > 0]
    parts = [
        "Сводка:\n" + (state.get("summary_text") or "Проверка завершена, но краткая сводка пока недоступна.")
    ]
    inflection_report = _safe_text(state.get("inflection_report"))
    if inflection_report:
        parts.append(inflection_report)
    tool_trace = state.get("tool_trace") or []
    if isinstance(tool_trace, list):
        parts.append(_format_tool_trace(tool_trace))
    plan_summary = _format_plan_summary(state.get("plan"))
    if plan_summary:
        parts.append(plan_summary)
    decision_details = _format_fix_decision_details(documents, limit=6)
    if decision_details:
        parts.append(decision_details)
    react_details = _format_react_trace_details(documents, limit=5)
    if react_details:
        parts.append(react_details)
    parts.append(_format_fixed_details(fixed, limit=5) if fixed else "Исправил:\nАвтоматических исправлений не было.")
    skipped_details = _format_skipped_fix_details(documents, limit=5)
    if skipped_details:
        parts.append(skipped_details)
    verification_details = _format_verification_details(documents, limit=5)
    if verification_details:
        parts.append(verification_details)
    parts.append(_format_issue_details(issues, limit=5) if issues else "Остались замечания:\nКритичных языковых замечаний не осталось.")
    parts.append(_format_rule_details(message))
    return "\n\n".join(part for part in parts if part)


def _extract_row_id_from_docx_path(docx_path: Path) -> str:
    folder_name = docx_path.parent.name
    return folder_name.split("_", 1)[0].strip()


def _extract_mun_name_from_docx_path(docx_path: Path) -> str:
    folder_name = docx_path.parent.name
    if "_" not in folder_name:
        return folder_name
    return folder_name.split("_", 1)[1].strip()


def run_philologist(
    *,
    output_dir: Path | None = None,
    ai_enabled: bool = True,
    row_ids: list[str] | None = None,
    job_id: str | None = None,
    mode: str | None = None,
) -> dict[str, Any]:
    run_mode = _normalize_philologist_mode(mode)
    job_paths = resolve_job_paths(job_id)
    target_dir = output_dir or (job_paths.output_dir if not job_paths.uses_legacy_layout else OUTPUT_DIR)
    docx_paths = sorted(target_dir.rglob("*.docx"))
    requested_row_ids = {str(item).strip() for item in (row_ids or []) if str(item).strip()}
    if requested_row_ids:
        docx_paths = [
            path for path in docx_paths
            if _extract_row_id_from_docx_path(path) in requested_row_ids
        ]
    claimed_tasks = mark_tasks_in_progress(
        "philologist",
        row_ids=sorted(requested_row_ids) if requested_row_ids else None,
        job_id=job_id,
    )
    tool_runner = PhilologistToolRunner()
    plan = build_philologist_plan(job_id, output_dir=target_dir)
    agent_loop = PhilologistAgentLoop(plan)
    agent_loop.start(
        f"К исполнению принято документов: {len(docx_paths)}; "
        f"фильтр строк: {len(requested_row_ids) if requested_row_ids else 'все'}."
    )
    agent_loop.mark_step(
        "inspect_job",
        "done",
        "Проверены входные условия запуска филолога.",
        data={"docx_count": len(docx_paths), "requested_row_ids": len(requested_row_ids)},
    )
    inflection_rows = tool_runner.call(
        "read_inflection_log",
        {"job_id": job_id},
        lambda: load_inflection_log(job_id),
    )
    agent_loop.mark_step(
        "read_inflection_log",
        "done" if inflection_rows else "skipped",
        (
            f"Прочитан журнал склонений: {len(inflection_rows)} записей."
            if inflection_rows
            else "Журнал склонений пуст, филолог продолжает проверку по тексту документов."
        ),
        data={"inflection_log_rows": len(inflection_rows)},
    )

    state = _load_philologist_state(job_id)
    was_stopped = str(state.get("status") or "") == "stopped"
    processed_documents: list[dict[str, Any]] = list(state.get("documents") or []) if was_stopped else []
    processed_paths = {str(item.get("path")) for item in processed_documents if str(item.get("path") or "").strip()}
    template_memory: dict[str, Any] = (
        dict(state.get("template_memory") or {}) if isinstance(state.get("template_memory"), dict) else {}
    )
    if was_stopped:
        docx_paths = [path for path in docx_paths if str(path) not in processed_paths]
    state.update(
        {
            "status": "running",
            "started_at": state.get("started_at") if was_stopped and state.get("started_at") else datetime.now().isoformat(timespec="seconds"),
            "completed_at": None,
            "total_documents": (len(processed_documents) + len(docx_paths)) if was_stopped else len(docx_paths),
            "processed_documents": len(processed_documents) if was_stopped else 0,
            "fixed_documents": sum(1 for item in processed_documents if item.get("applied_fix_count", 0) > 0) if was_stopped else 0,
            "documents_with_issues": sum(1 for item in processed_documents if item.get("issue_count", 0) > 0) if was_stopped else 0,
            "documents": processed_documents,
            "mode": run_mode,
            "ai_review_enabled": bool(ai_enabled and run_mode == "deep"),
            "summary_text": (
                (
                    f"Агент-филолог продолжает проверку с сохраненного места. Уже обработано {len(processed_documents)} документов."
                    if was_stopped
                    else
                    "Агент-филолог начал быструю массовую проверку документов: "
                    "локальные правила, безопасные автоправки и карантин спорных мест."
                    if run_mode == "fast"
                    else "Агент-филолог начал глубокую LLM-проверку документов."
                )
                if not claimed_tasks
                else f"Агент-филолог начал проверку документов и принял {len(claimed_tasks)} внутренних задач."
            ),
            "inflection_report": format_inflection_report(inflection_rows, limit=8) if inflection_rows else "",
            "inflection_log_count": len(inflection_rows),
            "tool_manifest": build_philologist_tool_manifest(),
            "tool_trace": tool_runner.as_state(),
            "plan": agent_loop.as_plan(),
            "agent_loop": agent_loop.as_plan().get("execution"),
            "task_stats": count_tasks_for_agent("philologist", job_id),
            "tasks": get_tasks_for_agent("philologist", job_id)[:20],
            "recent_events": get_recent_events(agent_name="philologist", limit=20, job_id=job_id),
            "stop_requested": False,
            "stop_requested_at": None,
            "template_memory": template_memory,
        }
    )
    _save_philologist_state(state, job_id)

    if ai_enabled:
        context_review = _review_inflection_contexts_with_llm(
            inflection_rows,
            row_ids=requested_row_ids,
            tool_runner=tool_runner,
        )
    else:
        context_review = {"enabled": False, "checked_count": 0, "items": [], "error": ""}
    context_apply_result = _apply_inflection_context_corrections(docx_paths, context_review)
    agent_loop.observe(
        "inflection_context_review",
        (
            "LLM проверила подозрительные контексты подстановок: "
            f"выбрано {context_review.get('selected_count', 0)}, "
            f"проверено {context_review.get('checked_count', 0)}, "
            f"исправлений применено {context_apply_result.get('applied_count', 0)}."
            if context_review.get("enabled")
            else "LLM-проверка контекстов подстановок отключена."
        ),
        step_id="read_inflection_log",
        data={
            "selected_count": context_review.get("selected_count", 0),
            "checked_count": context_review.get("checked_count", 0),
            "applied_count": context_apply_result.get("applied_count", 0),
            "error": context_review.get("error", ""),
        },
    )
    state["inflection_context_review"] = context_review
    state["inflection_context_corrections"] = context_apply_result
    state["tool_trace"] = tool_runner.as_state()
    state["plan"] = agent_loop.as_plan()
    state["agent_loop"] = state["plan"].get("execution")
    _save_philologist_state(state, job_id)

    started_at = perf_counter()
    react_client = _build_llm_client() if settings.orchestrator_mode == "agentic" else None
    effective_ai_enabled = bool(ai_enabled and run_mode == "deep")
    for index, docx_path in enumerate(docx_paths, start=1):
        _refresh_philologist_stop_flag(state, job_id)
        if state.get("stop_requested"):
            state["status"] = "stopped"
            state["completed_at"] = None
            state["summary_text"] = "Проверка остановлена. Можно продолжить с сохраненного места."
            state["current_document"] = None
            state["elapsed_seconds"] = round(perf_counter() - started_at, 2)
            state["tool_trace"] = tool_runner.as_state()
            state["plan"] = agent_loop.as_plan()
            state["agent_loop"] = state["plan"].get("execution")
            _save_philologist_state(state, job_id)
            return dict(state)
        agent_loop.observe(
            "document_start",
            f"Проверяю документ {len(processed_documents) + 1} из {state.get('total_documents', len(docx_paths))}: {docx_path.name}.",
            step_id="review_docx",
            data={"index": len(processed_documents) + 1, "path": str(docx_path)},
        )
        state["current_document"] = {
            "index": len(processed_documents) + 1,
            "total": state.get("total_documents", len(docx_paths)),
            "name": docx_path.name,
            "path": str(docx_path),
        }
        state["processed_documents"] = len(processed_documents)
        state["summary_text"] = (
            f"Агент-филолог проверяет документ {len(processed_documents) + 1} из {state.get('total_documents', len(docx_paths))}: {docx_path.name}. "
            f"Режим: {'быстрый' if run_mode == 'fast' else 'глубокий LLM'}."
        )
        state["tool_trace"] = tool_runner.as_state()
        state["plan"] = agent_loop.as_plan()
        state["agent_loop"] = state["plan"].get("execution")
        _save_philologist_state(state, job_id)

        try:
            react_result = _run_docx_react_loop(
                docx_path=docx_path,
                ai_enabled=effective_ai_enabled,
                tool_runner=tool_runner,
                client=react_client,
                use_llm_router=PHILOLOGIST_LLM_ROUTER,
                use_llm_fix_strategy=PHILOLOGIST_LLM_FIX_STRATEGY,
                rebuild_pdf_enabled=PHILOLOGIST_REBUILD_PDF,
                use_rag_decisions=run_mode == "deep",
                template_memory=template_memory,
            )
        except Exception as exc:
            logger.exception("philologist_document_failed", job_id=job_id, path=str(docx_path))
            react_result = {
                "review_result": {
                    "issues": [],
                    "issue_count": 1,
                    "local_issue_count": 0,
                    "ai_issue_count": 0,
                    "ai_error": f"{type(exc).__name__}: {exc}",
                },
                "fix_result": {
                    "applied_fix_count": 0,
                    "applied_fixes": [],
                    "skipped_fix_count": 1,
                    "skipped_fixes": [
                        {
                            "location": "",
                            "fragment": docx_path.name,
                            "issue": "Документ не удалось проверить.",
                            "suggestion": "",
                            "reason": f"{type(exc).__name__}: {exc}",
                        }
                    ],
                    "decision_count": 0,
                    "fix_decisions": [],
                },
                "verification_result": {"verified": False, "warning_count": 1, "warnings": [str(exc)]},
                "pdf_path": None,
                "react_trace": [],
                "react_context": {},
            }
        review_result = react_result["review_result"]
        fix_result = react_result["fix_result"]
        verification_result = react_result["verification_result"]
        pdf_path = react_result["pdf_path"]
        template_memory = _update_template_memory(
            template_memory,
            document_name=docx_path.name,
            review_result=review_result,
            fix_result=fix_result,
        )
        if effective_ai_enabled and _is_llm_auth_error(review_result.get("ai_error")):
            effective_ai_enabled = False
            react_client = None
            agent_loop.observe(
                "ai_review_disabled",
                "LLM-проверка отключена до конца запуска из-за ошибки авторизации API; продолжаю локальными правилами.",
                step_id="review_docx",
                data={"ai_error": review_result.get("ai_error")},
            )

        document_entry = {
            "index": index,
            "name": docx_path.name,
            "path": str(docx_path),
            "folder": str(docx_path.parent),
            "row_id": _extract_row_id_from_docx_path(docx_path),
            "mun_name": _extract_mun_name_from_docx_path(docx_path),
            "issue_count": int(review_result.get("issue_count", 0)),
            "local_issue_count": int(review_result.get("local_issue_count", 0)),
            "ai_issue_count": int(review_result.get("ai_issue_count", 0)),
            "ai_error": review_result.get("ai_error"),
            "template_fingerprint": review_result.get("template_fingerprint"),
            "review_mode": review_result.get("review_mode"),
            "reviewed_block_count": int(review_result.get("reviewed_block_count", 0)),
            "total_block_count": int(review_result.get("total_block_count", 0)),
            "focused_locations": review_result.get("focus_locations", []),
            "issues": review_result.get("issues", []),
            "applied_fix_count": fix_result["applied_fix_count"],
            "applied_fixes": fix_result["applied_fixes"],
            "skipped_fix_count": fix_result.get("skipped_fix_count", 0),
            "skipped_fixes": fix_result.get("skipped_fixes", []),
            "decision_count": fix_result.get("decision_count", 0),
            "fix_decisions": fix_result.get("fix_decisions", []),
            "verification": verification_result,
            "verification_warning_count": int(verification_result.get("warning_count", 0)),
            "verification_warnings": verification_result.get("warnings", []),
            "react_trace": react_result.get("react_trace", []),
            "react_context": react_result.get("react_context", {}),
            "updated_pdf": pdf_path,
        }
        processed_documents.append(document_entry)

        state["processed_documents"] = len(processed_documents)
        state["documents"] = processed_documents
        state["fixed_documents"] = sum(1 for item in processed_documents if item.get("applied_fix_count", 0) > 0)
        state["documents_with_issues"] = sum(1 for item in processed_documents if item.get("issue_count", 0) > 0)
        state["summary_text"] = _format_summary(processed_documents)
        state["current_document"] = None
        state["template_memory"] = template_memory
        state["tool_trace"] = tool_runner.as_state()
        state["plan"] = agent_loop.as_plan()
        state["agent_loop"] = state["plan"].get("execution")
        _save_philologist_state(state, job_id)

    if docx_paths:
        agent_loop.mark_step(
            "review_docx",
            "done",
            f"Проверено документов: {len(processed_documents)}.",
            data={
                "processed_documents": len(processed_documents),
                "documents_with_issues": sum(1 for item in processed_documents if item.get("issue_count", 0) > 0),
            },
        )
        total_applied = sum(int(item.get("applied_fix_count", 0) or 0) for item in processed_documents)
        total_skipped = sum(int(item.get("skipped_fix_count", 0) or 0) for item in processed_documents)
        total_decisions = sum(int(item.get("decision_count", 0) or 0) for item in processed_documents)
        agent_loop.mark_step(
            "apply_safe_fixes",
            "done",
            (
                f"Решений по правкам принято: {total_decisions}; "
                f"безопасных правок применено: {total_applied}; "
                f"отложено в ручную проверку: {total_skipped}."
            ),
            data={
                "decisions": total_decisions,
                "applied_fixes": total_applied,
                "skipped_fixes": total_skipped,
            },
        )
        verification_warnings = sum(int(item.get("verification_warning_count", 0) or 0) for item in processed_documents)
        agent_loop.mark_step(
            "verify_safe_fixes",
            "done" if verification_warnings == 0 else "blocked",
            (
                "Самопроверка автоправок не нашла повреждений стилей или неожиданных изменений."
                if verification_warnings == 0
                else f"Самопроверка нашла предупреждения после автоправок: {verification_warnings}."
            ),
            data={"verification_warnings": verification_warnings},
        )
        rebuilt_pdfs = sum(1 for item in processed_documents if item.get("updated_pdf"))
        agent_loop.mark_step(
            "rebuild_pdf",
            "done" if rebuilt_pdfs else "skipped",
            (
                f"PDF пересобраны для документов с правками: {rebuilt_pdfs}."
                if rebuilt_pdfs
                else "Правок, требующих пересборки PDF, не было."
            ),
            data={"rebuilt_pdfs": rebuilt_pdfs},
        )
    else:
        agent_loop.mark_step("review_docx", "blocked", "DOCX-файлы не найдены.")
        agent_loop.mark_step("apply_safe_fixes", "blocked", "Нет документов для применения правок.")
        agent_loop.mark_step("verify_safe_fixes", "blocked", "Нет документов для самопроверки правок.")
        agent_loop.mark_step("rebuild_pdf", "blocked", "Нет документов для пересборки PDF.")

    row_rollups: dict[str, dict[str, Any]] = {}
    for item in processed_documents:
        row_id = _safe_text(item.get("row_id"))
        if not row_id:
            continue
        row_entry = row_rollups.setdefault(
            row_id,
            {
                "row_id": row_id,
                "mun_name": _safe_text(item.get("mun_name")),
                "issue_count": 0,
                "applied_fix_count": 0,
                "verification_warning_count": 0,
            },
        )
        row_entry["issue_count"] += int(item.get("issue_count", 0))
        row_entry["applied_fix_count"] += int(item.get("applied_fix_count", 0))
        row_entry["verification_warning_count"] += int(item.get("verification_warning_count", 0))

    sender_handoffs = 0
    for row_entry in row_rollups.values():
        row_id = row_entry["row_id"]
        mun_name = row_entry["mun_name"]
        issue_count = int(row_entry["issue_count"])
        applied_fix_count = int(row_entry["applied_fix_count"])
        verification_warning_count = int(row_entry.get("verification_warning_count", 0))
        unresolved_issue_count = max(0, issue_count - applied_fix_count, verification_warning_count)
        note = (
            "Филолог завершил проверку документов."
            if unresolved_issue_count == 0
            else f"Филолог нашёл {unresolved_issue_count} нерешённых замечаний."
        )
        set_task_statuses(
            "philologist",
            row_id=row_id,
            task_type="review_generated_documents",
            new_status="done",
            note=note,
            resolution_summary="Филолог завершил проверку документов по строке.",
            job_id=job_id,
        )
        if not settings.inter_agent_handoffs_enabled:
            continue
        if unresolved_issue_count > 0:
            diagnosis = diagnose_responsibility(
                symptom="philology_review_block",
                context={"unresolved_issue_count": unresolved_issue_count},
            )
            create_task(
                source_agent="philologist",
                target_agent=diagnosis["owner_agent"],
                owner_agent=diagnosis["owner_agent"],
                task_type="review_before_send",
                problem_type=diagnosis["problem_type"],
                symptom="philology_review_block",
                root_cause=diagnosis["root_cause"],
                priority=diagnosis["priority"],
                blocking=diagnosis["blocking"],
                can_retry_after=diagnosis["can_retry_after"],
                row_id=row_id,
                mun_name=mun_name,
                details={
                    "issue_count": issue_count,
                    "applied_fix_count": applied_fix_count,
                    "verification_warning_count": verification_warning_count,
                    "unresolved_issue_count": unresolved_issue_count,
                    "reason": "Перед отправкой нужно учесть замечания филолога.",
                },
                job_id=job_id,
            )
            sender_handoffs += 1
            append_agent_event(
                source_agent="philologist",
                target_agent="sender",
                event_type="review_flagged",
                message=f"Филолог нашёл замечания по комплекту документов для строки {row_id}.",
                row_id=row_id,
                mun_name=mun_name,
                details={
                    "issue_count": issue_count,
                    "unresolved_issue_count": unresolved_issue_count,
                },
                job_id=job_id,
            )
        else:
            create_task(
                source_agent="philologist",
                target_agent="sender",
                owner_agent="sender",
                task_type="resume_send_readiness",
                problem_type="delivery_blocked",
                symptom="documents_ready_after_review",
                root_cause="Филолог завершил проверку и снял языковой блокер, можно повторно проверить готовность к отправке.",
                priority="medium",
                blocking=False,
                can_retry_after=True,
                row_id=row_id,
                mun_name=mun_name,
                details={
                    "issue_count": issue_count,
                    "applied_fix_count": applied_fix_count,
                    "verification_warning_count": verification_warning_count,
                    "reason": "После филологической проверки строку можно снова прогнать через отправщика.",
                },
                job_id=job_id,
            )
            set_task_statuses(
                "sender",
                row_id=row_id,
                task_type="review_before_send",
                new_status="done",
                note="Филолог подтвердил, что критичных замечаний перед отправкой не осталось.",
                resolution_summary="Критичных замечаний перед отправкой не осталось.",
                job_id=job_id,
            )
            append_agent_event(
                source_agent="philologist",
                target_agent="sender",
                event_type="review_completed",
                message=f"Филолог проверил документы для строки {row_id}; критичных замечаний не осталось.",
                row_id=row_id,
                mun_name=mun_name,
                details={
                    "issue_count": issue_count,
                    "applied_fix_count": applied_fix_count,
                },
                job_id=job_id,
            )

    state["status"] = "finalizing"
    state["current_document"] = None
    state["elapsed_seconds"] = round(perf_counter() - started_at, 2)
    state["task_stats"] = count_tasks_for_agent("philologist", job_id)
    state["tasks"] = get_tasks_for_agent("philologist", job_id)[:20]
    state["recent_events"] = get_recent_events(agent_name="philologist", limit=20, job_id=job_id)
    state["summary_text"] = "Документы проверены. Формирую журнал исправлений и итоговый отчет."
    state["tool_trace"] = tool_runner.as_state()
    state["plan"] = agent_loop.as_plan()
    state["agent_loop"] = state["plan"].get("execution")
    _save_philologist_state(state, job_id)

    state["status"] = "completed"
    state["completed_at"] = datetime.now().isoformat(timespec="seconds")
    state["elapsed_seconds"] = round(perf_counter() - started_at, 2)
    state["task_stats"] = count_tasks_for_agent("philologist", job_id)
    state["tasks"] = get_tasks_for_agent("philologist", job_id)[:20]
    state["recent_events"] = get_recent_events(agent_name="philologist", limit=20, job_id=job_id)
    state["summary_text"] = _format_run_summary(processed_documents, sender_handoffs=sender_handoffs)
    state["inflection_report"] = format_inflection_report(inflection_rows, limit=8) if inflection_rows else ""
    state["inflection_log_count"] = len(inflection_rows)
    tool_runner.call(
        "write_report",
        {"document_count": len(processed_documents)},
        lambda: {"status": "completed"},
    )
    candidates = tool_runner.call(
        "save_learning_memory",
        {"job_id": job_id},
        lambda: save_learning_memory(job_id),
    )
    quarantine = build_quarantine_items(job_id)
    agent_loop.mark_step(
        "update_memory",
        "done",
        f"Память обновлена: кандидатов {len(candidates)}, в карантине {len(quarantine)}.",
        data={"candidates": len(candidates), "quarantine_items": len(quarantine)},
    )
    agent_loop.mark_step(
        "finalize_report",
        "done",
        "Итоговый отчет агента подготовлен.",
        data={"document_count": len(processed_documents)},
    )
    agent_loop.complete(
        "needs_review" if quarantine else "completed",
        (
            f"Цикл филолога завершен, ручной проверки требуют {len(quarantine)} решений."
            if quarantine
            else "Цикл филолога завершен без карантина."
        ),
    )
    state["tool_manifest"] = build_philologist_tool_manifest()
    state["tool_trace"] = tool_runner.as_state()
    state["plan"] = agent_loop.as_plan()
    state["agent_loop"] = state["plan"].get("execution")
    _save_philologist_state(state, job_id)
    tool_runner.call(
        "save_agent_report",
        {"job_id": job_id},
        lambda: str(save_agent_report(job_id)),
    )
    state["tool_trace"] = tool_runner.as_state()
    _save_philologist_state(state, job_id)
    return dict(state)


def get_philologist_status(job_id: str | None = None) -> dict[str, Any]:
    state = _load_philologist_state(job_id)
    job_paths = resolve_job_paths(job_id)
    target_dir = job_paths.output_dir if not job_paths.uses_legacy_layout else OUTPUT_DIR
    state["task_stats"] = count_tasks_for_agent("philologist", job_id)
    state["tasks"] = get_tasks_for_agent("philologist", job_id)[:20]
    state["recent_events"] = get_recent_events(agent_name="philologist", limit=20, job_id=job_id)
    inflection_rows = load_inflection_log(job_id)
    state["inflection_report"] = format_inflection_report(inflection_rows, limit=8) if inflection_rows else ""
    state["inflection_log_count"] = len(inflection_rows)
    state["tool_manifest"] = build_philologist_tool_manifest()
    state["tool_trace"] = state.get("tool_trace") or []
    state["plan"] = merge_plan_execution(
        build_philologist_plan(job_id, output_dir=target_dir, current_state=state),
        state.get("plan") if isinstance(state.get("plan"), dict) else None,
    )
    state["agent_loop"] = (state.get("plan") or {}).get("execution")
    if state.get("status") == "idle":
        state["total_documents"] = len(list(target_dir.rglob("*.docx"))) if target_dir.exists() else 0
    return state


def _fallback_chat_answer(message: str, state: dict[str, Any]) -> str:
    return _format_philologist_structured_reply(message, state)


def chat_with_philologist(message: str, *, job_id: str | None = None) -> dict[str, Any]:
    state = get_philologist_status(job_id)
    client = _build_llm_client()
    relevant_rules = find_relevant_rules(message, limit=4)
    rules_context = format_rules_context(relevant_rules)
    if not client:
        return {"reply": _fallback_chat_answer(message, state), "state": state}

    compact_documents = []
    for item in (state.get("documents") or [])[:20]:
        compact_documents.append(
            {
                "name": item.get("name"),
                "issue_count": item.get("issue_count"),
                "applied_fix_count": item.get("applied_fix_count"),
                "ai_error": item.get("ai_error"),
                "issues": (item.get("issues") or [])[:5],
            }
        )

    prompt = (
        "Ты агент-филолог для проверки коммерческих предложений и договоров. "
        "Ты уже проверил документы и должен отвечать пользователю по результатам этой проверки. "
        "Отвечай по-русски, кратко, структурированно и по делу. "
        "Если пользователь спрашивает об ошибках или исправлениях, сначала опирайся на локальную базу правил русского языка, "
        "а потом на переданные данные по документам. "
        "Если правило найдено, кратко ссылайся на него и объясняй исправление. "
        "Не придумывай факты, которых нет в сводке. "
        "Если вопрос про исправления, используй разделы: «Сводка:», «Исправил:», «Остались замечания:», «Правила:». "
        "В разделах про исправления и замечания указывай конкретные документы и не более 1-2 примеров на документ.\n\n"
        f"Локальная база правил:\n{rules_context}\n\n"
        f"Состояние:\n{json.dumps({'summary_text': state.get('summary_text'), 'status': state.get('status'), 'documents': compact_documents}, ensure_ascii=False, indent=2)}\n\n"
        f"Вопрос пользователя:\n{message}"
    )

    request_kwargs = {
        "model": DOCUMENT_REVIEW_MODEL,
        "messages": [{"role": "user", "content": prompt}],
    }
    if not _resolve_openai_base_url():
        request_kwargs["response_format"] = {"type": "text"}

    try:
        response = client.chat.completions.create(**request_kwargs)
        reply = _safe_text(response.choices[0].message.content)
        if not reply:
            reply = _fallback_chat_answer(message, state)
    except Exception:
        reply = _fallback_chat_answer(message, state)

    return {"reply": reply, "state": state}

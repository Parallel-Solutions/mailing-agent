"""AI-assisted template generation for the library."""

from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document

from src.campaigns import template_service
from src.infra.llm_pricing import LlmUsage, estimate_llm_cost_usd, usage_from_response
from src.infra.llm_pricing import usage_from_response as _usage_from_response
from src.infra.spend_ledger import record_llm_usage
from src.utils.config import settings

try:
    from src.generator.inflection.ai_case_agent import (
        OpenAI,
        _build_openai_http_client,
        _resolve_openai_api_key,
        _resolve_openai_base_url,
    )
except ImportError:  # pragma: no cover
    OpenAI = None  # type: ignore

    def _build_openai_http_client():  # type: ignore
        return None

    def _resolve_openai_api_key():  # type: ignore
        return None

    def _resolve_openai_base_url():  # type: ignore
        return None


_JSON_FENCE_RE = re.compile(r"```(?:json)?\s*([\s\S]*?)\s*```", re.IGNORECASE)


@dataclass(frozen=True)
class VisionLlmResult:
    payload: dict[str, Any]
    usage: LlmUsage = field(default_factory=LlmUsage)
    estimated_cost_usd: float = 0.0
    model: str = ""


def list_models() -> list[dict[str, str]]:
    raw = str(getattr(settings, "template_ai_models", "") or "").strip()
    if raw:
        models = [part.strip() for part in raw.split(",") if part.strip()]
    else:
        models = ["gpt-4o-mini", "gpt-4.1"]
    default = models[0] if models else "gpt-4o-mini"
    return [{"id": model, "label": model, "default": model == default} for model in models]


def _build_client():
    if not OpenAI:
        return None
    api_key = _resolve_openai_api_key()
    if not api_key:
        return None
    kwargs: dict[str, Any] = {"api_key": api_key}
    base_url = _resolve_openai_base_url()
    if base_url:
        kwargs["base_url"] = base_url
    http_client = _build_openai_http_client()
    if http_client is not None:
        kwargs["http_client"] = http_client
    return OpenAI(**kwargs)


def _parse_json_payload(text: str) -> dict[str, Any]:
    raw = (text or "").strip()
    fence = _JSON_FENCE_RE.search(raw)
    if fence:
        raw = fence.group(1).strip()
    return json.loads(raw)


def _llm_error_message(exc: BaseException, *, model: str) -> str:
    raw = str(exc) or exc.__class__.__name__
    lowered = raw.lower()
    if "no healthy deployments" in lowered or "model_not_found" in lowered or "does not exist" in lowered:
        return f"Модель {model} сейчас недоступна. Выберите другую модель."
    if "insufficient_quota" in lowered or "rate limit" in lowered or "429" in lowered:
        return "AI временно недоступен из‑за лимитов. Попробуйте позже или другую модель."
    if "401" in lowered or "invalid api key" in lowered or "authentication" in lowered:
        return "AI недоступен: ошибка авторизации у провайдера моделей"
    # Keep gateway detail short for operators, but avoid dumping huge JSON.
    snippet = raw.replace("\n", " ").strip()
    if len(snippet) > 220:
        snippet = snippet[:217] + "..."
    return f"AI не смог сгенерировать шаблон: {snippet}"


def _attachments_context(files: list[tuple[str, bytes]]) -> str:
    chunks: list[str] = []
    for filename, data in files:
        try:
            text = template_service._file_text(filename, data)  # noqa: SLF001
        except Exception:
            text = ""
        if text.strip():
            chunks.append(f"--- файл: {filename} ---\n{text[:8000]}")
    return "\n\n".join(chunks)


def _resolve_model(model: str) -> str:
    allowed = {item["id"] for item in list_models()}
    return model.strip() if model.strip() in allowed else next(iter(allowed), "gpt-4o-mini")


def _call_llm(model: str, system: str, user: str) -> dict[str, Any]:
    client = _build_client()
    if client is None:
        raise RuntimeError("AI недоступен: не настроен OpenAI API ключ")
    resolved_model = _resolve_model(model)
    try:
        response = client.chat.completions.create(
            model=resolved_model,
            temperature=0.4,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
        )
    except Exception as exc:
        raise RuntimeError(_llm_error_message(exc, model=resolved_model)) from exc
    record_llm_usage(
        service="openai",
        model=resolved_model,
        operation="template_generate",
        usage=_usage_from_response(response),
    )
    content = response.choices[0].message.content or ""
    try:
        return _parse_json_payload(content)
    except json.JSONDecodeError as exc:
        raise RuntimeError("Модель вернула некорректный JSON") from exc


def _call_vision_llm(
    model: str,
    system: str,
    user_text: str,
    image_data_urls: list[str] | None = None,
    *,
    max_tokens: int = 8000,
) -> dict[str, Any]:
    return _call_vision_llm_tracked(
        model,
        system,
        user_text,
        image_data_urls,
        max_tokens=max_tokens,
    ).payload


def _call_vision_llm_tracked(
    model: str,
    system: str,
    user_text: str,
    image_data_urls: list[str] | None = None,
    *,
    max_tokens: int = 8000,
) -> VisionLlmResult:
    client = _build_client()
    if client is None:
        raise RuntimeError("AI недоступен: не настроен OpenAI API ключ")
    resolved_model = _resolve_model(model)
    image_urls = [url for url in (image_data_urls or []) if url]
    content: list[dict[str, Any]] = [{"type": "text", "text": user_text}]
    for url in image_urls:
        content.append({"type": "image_url", "image_url": {"url": url}})
    try:
        response = client.chat.completions.create(
            model=resolved_model,
            temperature=0.2,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": content},
            ],
            max_tokens=max(1000, int(max_tokens or 8000)),
        )
    except Exception as exc:
        raise RuntimeError(_llm_error_message(exc, model=resolved_model)) from exc
    raw = response.choices[0].message.content or ""
    try:
        payload = _parse_json_payload(raw)
    except json.JSONDecodeError as exc:
        raise RuntimeError("Модель вернула некорректный JSON") from exc
    usage = _usage_from_response(response, image_count=len(image_urls))
    cost = estimate_llm_cost_usd(
        model=resolved_model,
        prompt_tokens=usage.prompt_tokens,
        completion_tokens=usage.completion_tokens,
        image_count=usage.image_count,
    )
    record_llm_usage(
        service="openai",
        model=resolved_model,
        operation="template_import_vision",
        usage=usage,
    )
    return VisionLlmResult(payload=payload, usage=usage, estimated_cost_usd=cost, model=resolved_model)


def _docx_from_paragraphs(title: str, paragraphs: list[str]) -> bytes:
    document = Document()
    document.add_heading(title or "Документ", level=1)
    for paragraph in paragraphs:
        document.add_paragraph(str(paragraph))
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _create_email_from_files(owner_username: str, files: list[tuple[str, bytes]]) -> dict[str, Any]:
    if not files:
        raise ValueError("Нужен промпт или файл")
    filename, data = files[0]
    suffix = Path(filename).suffix.lower()
    text = template_service._file_text(filename, data)  # noqa: SLF001
    if not text.strip():
        raise ValueError("Не удалось извлечь текст из файла")
    if suffix in {".html", ".htm"}:
        body_html = text
    else:
        body_html = "".join(f"<p>{line}</p>" for line in text.splitlines() if line.strip()) or f"<p>{text}</p>"
    return template_service.create_template(
        owner_username,
        name=Path(filename).stem or "Шаблон письма",
        template_type="email",
        subject="Тема письма",
        body_html=body_html,
        body_text=text,
    )


def generate_template(
    owner_username: str,
    *,
    template_type: str,
    prompt: str = "",
    model: str = "",
    files: list[tuple[str, bytes]] | None = None,
) -> dict[str, Any]:
    normalized_type = template_service.normalize_file_template_type(template_type)
    if normalized_type not in {"email", "document"}:
        raise ValueError("Поддерживаются типы email и document")

    attachments = list(files or [])
    prompt_text = (prompt or "").strip()
    if not prompt_text and not attachments:
        raise ValueError("Укажите описание или приложите файлы")

    if not prompt_text:
        if normalized_type == "document":
            filename, data = attachments[0]
            return template_service.upload_file_version(
                owner_username,
                name=Path(filename).stem or "Документ",
                template_type="document",
                filename=Path(filename).name,
                data=data,
            )
        return _create_email_from_files(owner_username, attachments)

    context = _attachments_context(attachments)
    if normalized_type == "email":
        payload = _call_llm(
            model,
            system=(
                "Ты помощник по шаблонам деловых писем на русском. "
                "Верни только JSON: "
                '{"name":"...","subject":"...","body_html":"..."} '
                "Используй плейсхолдеры {{company}}, {{contact_name}}, {{email}}, {{region}} где уместно. "
                "body_html — простой HTML с тегами <p>."
            ),
            user=f"Запрос:\n{prompt_text}\n\nКонтекст файлов:\n{context or '(нет)'}",
        )
        return template_service.create_template(
            owner_username,
            name=str(payload.get("name") or "Шаблон письма"),
            template_type="email",
            subject=str(payload.get("subject") or "Тема письма"),
            body_html=str(payload.get("body_html") or "<p></p>"),
            body_text="",
            tags=["ai"],
        )

    payload = _call_llm(
        model,
        system=(
            "Ты помощник по шаблонам деловых документов на русском. "
            "Верни только JSON: "
            '{"name":"...","title":"...","paragraphs":["..."]} '
            "Используй плейсхолдеры {{company}}, {{contact_name}}, {{email}}, {{region}} где уместно. "
            "paragraphs — список абзацев документа."
        ),
        user=f"Запрос:\n{prompt_text}\n\nКонтекст файлов:\n{context or '(нет)'}",
    )
    paragraphs = payload.get("paragraphs") or []
    if not isinstance(paragraphs, list):
        paragraphs = [str(paragraphs)]
    doc_bytes = _docx_from_paragraphs(
        str(payload.get("title") or payload.get("name") or "Документ"),
        [str(item) for item in paragraphs],
    )
    safe_name = str(payload.get("name") or "Документ")
    return template_service.upload_file_version(
        owner_username,
        name=safe_name,
        template_type="document",
        filename=f"{Path(safe_name).stem or 'document'}.docx",
        data=doc_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

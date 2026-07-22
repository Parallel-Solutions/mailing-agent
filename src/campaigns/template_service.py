"""Mail / document templates with versions."""

from __future__ import annotations

import re
import uuid
from datetime import datetime, timezone
from html import escape
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any
from urllib.parse import quote

from docx import Document
from pypdf import PdfReader
from sqlalchemy import or_, select

from src.infra.db import session_scope
from src.infra.models import MailTemplate, TemplateVersion
from src.infra.object_store import delete as delete_object
from src.infra.object_store import get_bytes, put_bytes
from src.security.company_access import apply_owner_filter, can_access_owner

_VISIBILITY_NOT_SET = object()

VARIABLE_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_а-яА-ЯёЁ]+)\s*\}\}")


def _owns_template(row: MailTemplate | None, owner_username: str, visible_owners: Any = _VISIBILITY_NOT_SET) -> bool:
    if row is None:
        return False
    if visible_owners is not _VISIBILITY_NOT_SET:
        return can_access_owner(visible_owners, row.owner_username)
    return row.owner_username == owner_username

FILE_TEMPLATE_TYPE_ALIASES = {
    "kp": "document",
    "contract": "document",
}
FILE_TEMPLATE_EXTENSIONS = {
    "document": {".docx", ".pdf", ".html", ".htm"},
}
LEGACY_DOCUMENT_TYPES = ("kp", "contract")


def normalize_file_template_type(template_type: str) -> str:
    normalized = str(template_type or "").strip().lower()
    return FILE_TEMPLATE_TYPE_ALIASES.get(normalized, normalized)


def normalize_template_type_filter(template_type: str | None) -> str | None:
    if not template_type:
        return None
    return normalize_file_template_type(template_type)

def _now() -> datetime:
    return datetime.now(timezone.utc)


def _new_id() -> str:
    return str(uuid.uuid4())


def _extract_variables(text: str) -> list[dict[str, Any]]:
    names = sorted(set(VARIABLE_RE.findall(text or "")))
    return [
        {
            "name": name,
            "source": "recipient" if name in {"company", "contact_name", "email", "region"} else "user_input",
            "label": name,
        }
        for name in names
    ]


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1251"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _file_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".html", ".htm", ".txt"}:
        return _decode_text(data)
    if suffix == ".docx":
        document = Document(BytesIO(data))
        chunks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                chunks.extend(cell.text for cell in row.cells)
        for section in document.sections:
            chunks.extend(paragraph.text for paragraph in section.header.paragraphs)
            chunks.extend(paragraph.text for paragraph in section.footer.paragraphs)
        return "\n".join(chunks)
    if suffix == ".pdf":
        reader = PdfReader(BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    return ""


def _is_file_document_template(template_type: str) -> bool:
    return normalize_file_template_type(template_type) == "document"


def _validate_file_template_type(template_type: str, filename: str) -> str:
    normalized_type = normalize_file_template_type(template_type)
    allowed = FILE_TEMPLATE_EXTENSIONS.get(normalized_type)
    if not allowed:
        raise ValueError("Файловая загрузка доступна только для документов")
    suffix = Path(filename).suffix.lower()
    if suffix not in allowed:
        formats = ", ".join(sorted(allowed))
        raise ValueError(f"Для этого типа шаблона доступны форматы: {formats}")
    return normalized_type


def _template_types_compatible(existing: str, incoming: str) -> bool:
    return normalize_file_template_type(existing) == normalize_file_template_type(incoming)

def _default_is_template(*, template_type: str, variables: list[dict[str, Any]]) -> bool:
    if str(template_type or "").strip().lower() == "email":
        return bool(variables)
    return bool(variables)


def template_to_dict(row: MailTemplate, version: TemplateVersion | None = None) -> dict[str, Any]:
    payload = {
        "id": row.id,
        "name": row.name,
        "template_type": row.template_type,
        "status": row.status,
        "active_version_id": row.active_version_id,
        "tags": list(row.tags or []),
        "archived": bool(row.archived),
        "is_template": bool(row.is_template),
        "created_at": row.created_at.isoformat() if row.created_at else None,
        "updated_at": row.updated_at.isoformat() if row.updated_at else None,
    }
    if version is not None:
        payload["version"] = version_to_dict(version)
    return payload


def version_to_dict(row: TemplateVersion) -> dict[str, Any]:
    return {
        "id": row.id,
        "template_id": row.template_id,
        "version_number": row.version_number,
        "subject": row.subject,
        "body_html": row.body_html,
        "body_text": row.body_text,
        "variables": list(row.variables or []),
        "storage_key": row.storage_key,
        "filename": row.filename,
        "rendered_pdf_storage_key": row.rendered_pdf_storage_key,
        "rendered_pdf_filename": row.rendered_pdf_filename,
        "editor_state": row.editor_state,
        "artifacts": {
            "source": {"filename": row.filename, "storage_key": row.storage_key} if row.filename else None,
            "delivery_pdf": {
                "filename": row.rendered_pdf_filename,
                "storage_key": row.rendered_pdf_storage_key,
            } if row.rendered_pdf_filename else None,
        },
        "created_by": row.created_by,
        "created_at": row.created_at.isoformat() if row.created_at else None,
    }


def create_template(
    owner_username: str,
    *,
    name: str,
    template_type: str,
    subject: str = "",
    body_html: str = "",
    body_text: str = "",
    tags: list[str] | None = None,
    editor_state: dict[str, Any] | None = None,
) -> dict[str, Any]:
    template_id = _new_id()
    version_id = _new_id()
    combined = f"{subject}\n{body_html}\n{body_text}"
    variables = _extract_variables(combined)
    with session_scope() as session:
        tmpl = MailTemplate(
            id=template_id,
            owner_username=owner_username,
            name=name or "Шаблон",
            template_type=template_type,
            status="ready",
            active_version_id=version_id,
            tags=list(tags or []),
            is_template=_default_is_template(template_type=template_type, variables=variables),
        )
        session.add(tmpl)
        session.add(
            TemplateVersion(
                id=version_id,
                template_id=template_id,
                version_number=1,
                subject=subject,
                body_html=body_html,
                body_text=body_text,
                variables=variables,
                editor_state=editor_state,
                created_by=owner_username,
            )
        )
        session.flush()
        return template_to_dict(tmpl, session.get(TemplateVersion, version_id))


EMAIL_ASSET_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
EMAIL_ASSET_CONTENT_TYPES = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".gif": "image/gif",
    ".webp": "image/webp",
}


def _template_asset_storage_key(template_id: str, asset_id: str) -> str:
    return f"template-library/{template_id}/assets/{asset_id}"


def upload_template_asset(
    template_id: str,
    owner_username: str,
    *,
    filename: str,
    data: bytes,
    content_type: str | None = None,
) -> dict[str, str] | None:
    safe_filename = Path(filename).name
    suffix = Path(safe_filename).suffix.lower()
    if suffix not in EMAIL_ASSET_EXTENSIONS:
        raise ValueError("Для письма можно загрузить только изображение JPG, PNG, GIF или WEBP")
    if not data:
        raise ValueError("Файл изображения пуст")

    with session_scope() as session:
        tmpl = session.get(MailTemplate, template_id)
        if not _owns_template(tmpl, owner_username):
            return None
        if tmpl.template_type != "email":
            raise ValueError("Изображения можно загружать только для email-шаблонов")

    asset_id = f"{_new_id()}-{safe_filename}"
    storage_key = _template_asset_storage_key(template_id, asset_id)
    resolved_content_type = content_type or EMAIL_ASSET_CONTENT_TYPES.get(suffix, "application/octet-stream")
    put_bytes(storage_key, data, content_type=resolved_content_type)
    return {
        "asset_id": asset_id,
        "url": f"/api/v1/templates/{template_id}/assets/{quote(asset_id, safe='')}",
    }


def get_template_asset(template_id: str, asset_id: str, owner_username: str) -> dict[str, Any] | None:
    safe_asset_id = Path(asset_id).name
    if not safe_asset_id or safe_asset_id != asset_id:
        return None
    with session_scope() as session:
        tmpl = session.get(MailTemplate, template_id)
        if not _owns_template(tmpl, owner_username):
            return None
    storage_key = _template_asset_storage_key(template_id, safe_asset_id)
    try:
        data = get_bytes(storage_key)
    except Exception:
        return None
    suffix = Path(safe_asset_id).suffix.lower()
    content_type = EMAIL_ASSET_CONTENT_TYPES.get(suffix, "application/octet-stream")
    return {"content": data, "content_type": content_type}


def _build_kp_pdf_artifact(filename: str, data: bytes) -> tuple[bytes, str]:
    from src.generator.generation.kp_one_page_fitter import KpLayoutError, fit_docx_to_one_page_pdf

    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return data, f"{Path(filename).stem}.pdf"
    if suffix != ".docx":
        raise ValueError("Исходником КП должен быть DOCX или PDF")
    with TemporaryDirectory(prefix="kp-template-pdf-") as temp_dir:
        root = Path(temp_dir)
        source_path = root / Path(filename).name
        output_pdf = root / "converted" / f"{Path(filename).stem}.pdf"
        source_path.write_bytes(data)
        from src.generator.generation.pdf_safe import is_kp_docx

        file_kind = "kp" if is_kp_docx(source_path) else None
        if file_kind == "kp":
            fit_docx_to_one_page_pdf(
                source_path,
                output_pdf,
                file_kind=file_kind,
                template_docx=source_path,
            )
        else:
            from src.generator.generation.template_preview import convert_docx_to_delivery_pdf

            convert_docx_to_delivery_pdf(
                source_path,
                output_pdf,
                file_kind=file_kind,
                template_docx=source_path,
            )
        return output_pdf.read_bytes(), output_pdf.name


def upload_file_version(
    owner_username: str,
    *,
    name: str,
    template_type: str,
    filename: str,
    data: bytes,
    content_type: str | None = None,
    template_id: str | None = None,
) -> dict[str, Any]:
    requested_name = str(name or "").strip()
    safe_filename = Path(filename).name
    normalized_type = _validate_file_template_type(template_type, safe_filename)
    if not data:
        raise ValueError("Файл шаблона пуст")

    try:
        source_text = _file_text(safe_filename, data)
        variables = _extract_variables(source_text)
        from src.campaigns.substitution_ai import normalize_placeholders

        for item in normalize_placeholders(source_text):
            placeholder_name = str(item.get("name") or "").strip()
            if placeholder_name and placeholder_name not in {str(v.get("name")) for v in variables}:
                variables.append(
                    {
                        "name": placeholder_name,
                        "label": str(item.get("label") or placeholder_name),
                        "source": str(item.get("source") or "recipient"),
                    }
                )
        rendered_pdf_data: bytes | None = None
        rendered_pdf_filename: str | None = None
        suffix = Path(safe_filename).suffix.lower()
        inferred_delivery_filename: str | None = None
        if _is_file_document_template(normalized_type):
            from src.campaigns.delivery_filename_service import infer_static_delivery_filename

            inferred_delivery_filename = infer_static_delivery_filename(
                text=source_text,
                upload_filename=safe_filename,
            )
            inferred_stem = Path(inferred_delivery_filename).stem
            if inferred_stem:
                safe_filename = f"{inferred_stem}{suffix}"
        if _is_file_document_template(normalized_type) and suffix == ".docx":
            rendered_pdf_data, _artifact_name = _build_kp_pdf_artifact(safe_filename, data)
            rendered_pdf_filename = inferred_delivery_filename or _artifact_name
        elif _is_file_document_template(normalized_type) and suffix == ".pdf":
            rendered_pdf_data = data
            rendered_pdf_filename = inferred_delivery_filename or f"{Path(safe_filename).stem}.pdf"
    except (ValueError, RuntimeError):
        raise
    except Exception as exc:
        raise ValueError("Не удалось прочитать содержимое шаблона") from exc

    editor_state = None
    if _is_file_document_template(normalized_type) and suffix == ".pdf":
        from src.campaigns.pdf_overlay_service import analyze_pdf

        editor_state = analyze_pdf(data)
    if rendered_pdf_data is not None and suffix == ".docx":
        from src.generator.generation.document_builder import DOCUMENT_RENDERER_VERSION

        state = dict(editor_state or {})
        state["delivery_renderer_version"] = DOCUMENT_RENDERER_VERSION
        editor_state = state

    upload_stem = Path(filename).name
    upload_stem_value = Path(upload_stem).stem
    resolved_template_name = requested_name
    if _is_file_document_template(normalized_type) and rendered_pdf_filename:
        from src.campaigns.delivery_filename_service import infer_template_display_name

        inferred_template_name = infer_template_display_name(rendered_pdf_filename)
        if not resolved_template_name or resolved_template_name == upload_stem_value:
            resolved_template_name = inferred_template_name

    version_id = _new_id()
    resolved_template_id = template_id or _new_id()
    storage_key = f"template-library/{resolved_template_id}/{version_id}/source/{safe_filename}"
    rendered_pdf_storage_key = None
    if rendered_pdf_data is not None and rendered_pdf_filename:
        rendered_pdf_storage_key = (
            storage_key
            if Path(safe_filename).suffix.lower() == ".pdf"
            else f"template-library/{resolved_template_id}/{version_id}/delivery/{rendered_pdf_filename}"
        )

    stored_keys: list[str] = []
    try:
        put_bytes(storage_key, data, content_type=content_type)
        stored_keys.append(storage_key)
        if rendered_pdf_storage_key and rendered_pdf_storage_key != storage_key and rendered_pdf_data is not None:
            put_bytes(rendered_pdf_storage_key, rendered_pdf_data, content_type="application/pdf")
            stored_keys.append(rendered_pdf_storage_key)

        with session_scope() as session:
            if template_id:
                tmpl = session.get(MailTemplate, template_id)
                if not _owns_template(tmpl, owner_username):
                    raise FileNotFoundError("Шаблон не найден")
                if not _template_types_compatible(tmpl.template_type, normalized_type):
                    raise ValueError("Тип загружаемого файла не совпадает с типом шаблона")
                if tmpl.template_type in LEGACY_DOCUMENT_TYPES:
                    tmpl.template_type = "document"
                current = session.get(TemplateVersion, tmpl.active_version_id) if tmpl.active_version_id else None
                version_number = (current.version_number + 1) if current else 1
                if resolved_template_name:
                    tmpl.name = resolved_template_name
            else:
                version_number = 1
                tmpl = MailTemplate(
                    id=resolved_template_id,
                    owner_username=owner_username,
                    name=resolved_template_name or Path(safe_filename).stem or "Шаблон",
                    template_type=normalized_type,
                    status="ready",
                    active_version_id=version_id,
                    tags=[],
                    is_template=_default_is_template(template_type=normalized_type, variables=variables),
                )
                session.add(tmpl)

            version = TemplateVersion(
                id=version_id,
                template_id=resolved_template_id,
                version_number=version_number,
                subject="",
                body_html="",
                body_text="",
                variables=variables,
                storage_key=storage_key,
                filename=safe_filename,
                rendered_pdf_storage_key=rendered_pdf_storage_key,
                rendered_pdf_filename=rendered_pdf_filename,
                editor_state=editor_state,
                created_by=owner_username,
            )
            session.add(version)
            tmpl.active_version_id = version_id
            tmpl.status = "ready"
            tmpl.archived = False
            tmpl.updated_at = _now()
            session.flush()
            return template_to_dict(tmpl, version)
    except BaseException:
        for key in reversed(stored_keys):
            delete_object(key)
        raise


def get_template_file(template_id: str, owner_username: str) -> dict[str, Any] | None:
    with session_scope() as session:
        tmpl = session.get(MailTemplate, template_id)
        if tmpl is None or tmpl.owner_username != owner_username or not tmpl.active_version_id:
            return None
        version = session.get(TemplateVersion, tmpl.active_version_id)
        if version is None or not version.storage_key or not version.filename:
            return None
        storage_key = version.storage_key
        filename = version.filename
        template_type = tmpl.template_type
    return {
        "content": get_bytes(storage_key),
        "filename": filename,
        "template_type": template_type,
        "storage_key": storage_key,
    }


def get_template_delivery_file(template_id: str, owner_username: str) -> dict[str, Any] | None:
    from src.generator.generation.document_builder import DOCUMENT_RENDERER_VERSION

    with session_scope() as session:
        tmpl = session.get(MailTemplate, template_id)
        if tmpl is None or tmpl.owner_username != owner_username or not tmpl.active_version_id:
            return None
        version = session.get(TemplateVersion, tmpl.active_version_id)
        if version is None:
            return None
        version_id = version.id
        rendered_key = version.rendered_pdf_storage_key
        rendered_name = version.rendered_pdf_filename
        source_key = version.storage_key
        source_name = version.filename
        template_type = tmpl.template_type
        editor_state = dict(version.editor_state or {}) if isinstance(version.editor_state, dict) else {}
        delivery_renderer_version = str(editor_state.get("delivery_renderer_version") or "")
    if rendered_key and rendered_name and delivery_renderer_version == DOCUMENT_RENDERER_VERSION:
        return {
            "content": get_bytes(rendered_key),
            "filename": rendered_name,
            "media_type": "application/pdf",
            "template_type": template_type,
        }
    if not _is_file_document_template(template_type) or not source_key or not source_name:
        return None

    source_data = get_bytes(source_key)
    suffix = Path(source_name).suffix.lower()
    if suffix == ".pdf":
        pdf_data = source_data
        pdf_name = f"{Path(source_name).stem}.pdf"
        pdf_key = source_key
    elif suffix == ".docx":
        pdf_data, pdf_name = _build_kp_pdf_artifact(source_name, source_data)
        pdf_key = f"template-library/{template_id}/{version_id}/delivery/{pdf_name}"
        put_bytes(pdf_key, pdf_data, content_type="application/pdf")
    else:
        return None

    with session_scope() as session:
        current = session.get(TemplateVersion, version_id)
        if current is not None and current.template_id == template_id:
            current.rendered_pdf_storage_key = pdf_key
            current.rendered_pdf_filename = pdf_name
            state = dict(current.editor_state or {}) if isinstance(current.editor_state, dict) else {}
            state["delivery_renderer_version"] = DOCUMENT_RENDERER_VERSION
            current.editor_state = state
            session.flush()
    return {
        "content": pdf_data,
        "filename": pdf_name,
        "media_type": "application/pdf",
        "template_type": template_type,
    }


def get_template_version_file(template_id: str, version_id: str, owner_username: str) -> dict[str, Any]:
    with session_scope() as session:
        tmpl = session.get(MailTemplate, template_id)
        if not _owns_template(tmpl, owner_username):
            raise FileNotFoundError("Шаблон не найден")
        version = session.get(TemplateVersion, version_id)
        if version is None or version.template_id != template_id or not version.storage_key or not version.filename:
            raise FileNotFoundError("Версия документа не найдена")
        storage_key = version.storage_key
        filename = version.filename
    return {
        "content": get_bytes(storage_key),
        "filename": filename,
        "media_type": (
            "application/pdf"
            if Path(filename).suffix.lower() == ".pdf"
            else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    }


def _render_kp_html_pdf_bytes(body_html: str, *, sample_values: bool) -> bytes:
    html = str(body_html or "").strip()
    if not html:
        raise ValueError("В шаблоне КП отсутствует редактируемый макет")
    if sample_values:
        from src.generator.templates.certification import certification_context

        names = tuple(sorted(set(VARIABLE_RE.findall(html))))
        values = certification_context(names, "normal")
        html = VARIABLE_RE.sub(lambda match: escape(values.get(match.group(1), match.group(0))), html)
    with TemporaryDirectory(prefix="kp-editor-render-") as temp_dir:
        output_path = Path(temp_dir) / "preview.pdf"
        from src.generator.generation.pdf_converter import convert_html_to_pdf

        converted = convert_html_to_pdf(html, output_path, filename="index.html")
        if converted is None or not converted.exists():
            raise RuntimeError("Не удалось собрать PDF из макета КП")
        return converted.read_bytes()


def build_kp_pdf_preview(
    template_id: str,
    owner_username: str,
    *,
    body_html: str | None = None,
) -> dict[str, Any]:
    template = get_template(template_id, owner_username)
    if not template or template.get("template_type") != "kp":
        raise FileNotFoundError("Шаблон КП не найден")
    version = template.get("version") or {}
    html = body_html if body_html is not None else str(version.get("body_html") or "")
    return {
        "content": _render_kp_html_pdf_bytes(html, sample_values=True),
        "filename": f"{template['name']}_preview.pdf",
        "media_type": "application/pdf",
    }


def save_kp_html_version(
    template_id: str,
    owner_username: str,
    *,
    body_html: str,
    name: str | None = None,
) -> dict[str, Any]:
    template = get_template(template_id, owner_username)
    if not template or template.get("template_type") != "kp":
        raise FileNotFoundError("Шаблон КП не найден")
    html = str(body_html or "").strip()
    if not html:
        raise ValueError("Макет КП не может быть пустым")
    pdf_data = _render_kp_html_pdf_bytes(html, sample_values=False)
    version_id = _new_id()
    filename = f"{Path(name or template['name']).stem}.pdf"
    storage_key = f"template-library/{template_id}/{version_id}/{filename}"
    put_bytes(storage_key, pdf_data, content_type="application/pdf")
    try:
        with session_scope() as session:
            tmpl = session.get(MailTemplate, template_id)
            if tmpl is None or tmpl.owner_username != owner_username or tmpl.template_type != "kp":
                raise FileNotFoundError("Шаблон КП не найден")
            current = session.get(TemplateVersion, tmpl.active_version_id) if tmpl.active_version_id else None
            version = TemplateVersion(
                id=version_id,
                template_id=template_id,
                version_number=(current.version_number + 1) if current else 1,
                subject="",
                body_html=html,
                body_text="",
                variables=_extract_variables(html),
                storage_key=storage_key,
                filename=filename,
                rendered_pdf_storage_key=storage_key,
                rendered_pdf_filename=filename,
                created_by=owner_username,
            )
            session.add(version)
            tmpl.active_version_id = version_id
            tmpl.status = "ready"
            tmpl.archived = False
            if name is not None:
                tmpl.name = name
            tmpl.updated_at = _now()
            session.flush()
            return template_to_dict(tmpl, version)
    except BaseException:
        delete_object(storage_key)
        raise


def save_docx_editor_version(template_id: str, owner_username: str, data: bytes) -> dict[str, Any]:
    template = get_template(template_id, owner_username)
    if not template or not _is_file_document_template(str(template.get("template_type") or "")):
        raise FileNotFoundError("Шаблон документа не найден")
    version = template.get("version") or {}
    filename = str(version.get("filename") or f"{template['name']}.docx")
    if not filename.lower().endswith(".docx"):
        raise ValueError("Редактируемый исходник должен иметь формат DOCX")
    Document(BytesIO(data))
    if version.get("storage_key") and get_bytes(str(version["storage_key"])) == data:
        return template
    return upload_file_version(
        owner_username,
        name=str(template["name"]),
        template_type=str(template["template_type"]),
        filename=filename,
        data=data,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        template_id=template_id,
    )


def build_file_preview(template_id: str, owner_username: str) -> dict[str, Any] | None:
    template = get_template(template_id, owner_username)
    if template and _is_file_document_template(str(template.get("template_type") or "")):
        delivery = get_template_delivery_file(template_id, owner_username)
        if delivery is not None:
            return delivery
    item = get_template_file(template_id, owner_username)
    if item is None:
        return None
    data = item["content"]
    filename = str(item["filename"])
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return {"content": data, "filename": f"{Path(filename).stem}.pdf", "media_type": "application/pdf"}

    with TemporaryDirectory(prefix="template-preview-") as temp_dir:
        root = Path(temp_dir)
        preview_path = root / "preview.pdf"
        if suffix == ".docx":
            source_path = root / filename
            source_path.write_bytes(data)
            from src.generator.generation.template_preview import _convert_preview_docx_to_pdf

            converted = _convert_preview_docx_to_pdf(source_path, root / "converted")
        elif suffix in {".html", ".htm"}:
            from src.generator.generation.pdf_converter import convert_html_to_pdf

            converted = convert_html_to_pdf(_decode_text(data), preview_path, filename=filename)
        else:
            converted = None
        if converted is None or not converted.exists():
            raise RuntimeError("Не удалось сформировать PDF-предпросмотр шаблона")
        return {
            "content": converted.read_bytes(),
            "filename": f"{Path(filename).stem}.pdf",
            "media_type": "application/pdf",
        }


def list_templates(
    owner_username: str,
    *,
    template_type: str | None = None,
    q: str | None = None,
    include_archived: bool = False,
    visible_owners: frozenset[str] | None = _VISIBILITY_NOT_SET,
) -> list[dict[str, Any]]:
    with session_scope() as session:
        stmt = select(MailTemplate)
        if visible_owners is not _VISIBILITY_NOT_SET:
            stmt = apply_owner_filter(stmt, MailTemplate.owner_username, visible_owners)
        else:
            stmt = stmt.where(MailTemplate.owner_username == owner_username)
        if not include_archived:
            stmt = stmt.where(MailTemplate.archived.is_(False))
        normalized_filter = normalize_template_type_filter(template_type)
        if normalized_filter == "document":
            stmt = stmt.where(
                or_(
                    MailTemplate.template_type == "document",
                    MailTemplate.template_type.in_(LEGACY_DOCUMENT_TYPES),
                )
            )
        elif normalized_filter:
            stmt = stmt.where(MailTemplate.template_type == normalized_filter)
        if q:
            stmt = stmt.where(MailTemplate.name.ilike(f"%{q.strip()}%"))
        rows = session.scalars(stmt.order_by(MailTemplate.updated_at.desc())).all()
        result = []
        for row in rows:
            version = session.get(TemplateVersion, row.active_version_id) if row.active_version_id else None
            result.append(template_to_dict(row, version))
        return result


def get_template(
    template_id: str,
    owner_username: str,
    *,
    visible_owners: frozenset[str] | None = _VISIBILITY_NOT_SET,
) -> dict[str, Any] | None:
    with session_scope() as session:
        row = session.get(MailTemplate, template_id)
        if not _owns_template(row, owner_username, visible_owners):
            return None
        version = session.get(TemplateVersion, row.active_version_id) if row.active_version_id else None
        return template_to_dict(row, version)


def save_version(
    template_id: str,
    owner_username: str,
    *,
    subject: str | None = None,
    body_html: str | None = None,
    body_text: str | None = None,
    variables: list[dict[str, Any]] | None = None,
    name: str | None = None,
    editor_state: dict[str, Any] | None = None,
    is_template: bool | None = None,
    rendered_pdf_filename: str | None = None,
) -> dict[str, Any] | None:
    with session_scope() as session:
        tmpl = session.get(MailTemplate, template_id)
        if not _owns_template(tmpl, owner_username):
            return None
        content_changed = any(
            value is not None
            for value in (subject, body_html, body_text, variables, editor_state)
        )
        metadata_changed = is_template is not None or rendered_pdf_filename is not None
        if not content_changed and name is None and metadata_changed:
            current = session.get(TemplateVersion, tmpl.active_version_id) if tmpl.active_version_id else None
            if rendered_pdf_filename is not None:
                if current is None:
                    raise ValueError("У шаблона нет активной версии")
                if not _is_file_document_template(str(tmpl.template_type or "")):
                    raise ValueError("Имя вложения можно задать только для документов")
                from src.campaigns.delivery_filename_service import normalize_delivery_filename

                current.rendered_pdf_filename = normalize_delivery_filename(rendered_pdf_filename)
            if is_template is not None:
                tmpl.is_template = bool(is_template)
            tmpl.updated_at = _now()
            session.flush()
            return template_to_dict(tmpl, current)
        if is_template is not None:
            tmpl.is_template = bool(is_template)
            tmpl.updated_at = _now()
        if not content_changed and name is None:
            return None
        current = session.get(TemplateVersion, tmpl.active_version_id) if tmpl.active_version_id else None
        next_number = (current.version_number + 1) if current else 1
        subj = subject if subject is not None else (current.subject if current else "")
        html = body_html if body_html is not None else (current.body_html if current else "")
        text = body_text if body_text is not None else (current.body_text if current else "")
        vars_list = variables if variables is not None else _extract_variables(f"{subj}\n{html}\n{text}")
        if is_template is None and variables is not None:
            tmpl.is_template = _default_is_template(template_type=tmpl.template_type, variables=vars_list)
        resolved_editor_state = editor_state if editor_state is not None else (current.editor_state if current else None)
        version_id = _new_id()
        session.add(
            TemplateVersion(
                id=version_id,
                template_id=template_id,
                version_number=next_number,
                subject=subj,
                body_html=html,
                body_text=text,
                variables=vars_list,
                storage_key=current.storage_key if current else None,
                filename=current.filename if current else None,
                rendered_pdf_storage_key=current.rendered_pdf_storage_key if current else None,
                rendered_pdf_filename=current.rendered_pdf_filename if current else None,
                editor_state=resolved_editor_state,
                created_by=owner_username,
            )
        )
        tmpl.active_version_id = version_id
        tmpl.status = "ready"
        if name is not None:
            tmpl.name = name
        tmpl.updated_at = _now()
        session.flush()
        return template_to_dict(tmpl, session.get(TemplateVersion, version_id))


def duplicate_template(template_id: str, owner_username: str) -> dict[str, Any] | None:
    source = get_template(template_id, owner_username)
    if not source:
        return None
    version = source.get("version") or {}
    storage_key = str(version.get("storage_key") or "")
    filename = str(version.get("filename") or "")
    if storage_key and filename:
        file_item = get_template_file(template_id, owner_username)
        if file_item is None:
            return None
        return upload_file_version(
            owner_username,
            name=f"{source['name']} (копия)",
            template_type=source["template_type"],
            filename=filename,
            data=file_item["content"],
            content_type=None,
        )
    return create_template(
        owner_username,
        name=f"{source['name']} (копия)",
        template_type=source["template_type"],
        subject=str(version.get("subject") or ""),
        body_html=str(version.get("body_html") or ""),
        body_text=str(version.get("body_text") or ""),
        tags=list(source.get("tags") or []),
        editor_state=version.get("editor_state") if isinstance(version.get("editor_state"), dict) else None,
    )


def archive_template(template_id: str, owner_username: str) -> dict[str, Any] | None:
    with session_scope() as session:
        tmpl = session.get(MailTemplate, template_id)
        if not _owns_template(tmpl, owner_username):
            return None
        tmpl.archived = True
        tmpl.updated_at = _now()
        session.flush()
        return template_to_dict(tmpl)


def list_versions(template_id: str, owner_username: str) -> list[dict[str, Any]]:
    with session_scope() as session:
        tmpl = session.get(MailTemplate, template_id)
        if not _owns_template(tmpl, owner_username):
            return []
        rows = session.scalars(
            select(TemplateVersion)
            .where(TemplateVersion.template_id == template_id)
            .order_by(TemplateVersion.version_number.desc())
        ).all()
        return [version_to_dict(r) for r in rows]


def preview_template(template_id: str, owner_username: str, sample: dict[str, Any] | None = None) -> dict[str, Any] | None:
    tmpl = get_template(template_id, owner_username)
    if not tmpl:
        return None
    version = tmpl.get("version") or {}
    sample = sample or {
        "company": "ООО Пример",
        "contact_name": "Иван Иванов",
        "email": "ivan@example.com",
        "region": "Москва",
        "campaign_name": "Тестовая рассылка",
    }
    html = str(version.get("body_html") or "")
    subject = str(version.get("subject") or "")
    for key, value in sample.items():
        html = html.replace("{{" + key + "}}", str(value))
        subject = subject.replace("{{" + key + "}}", str(value))
    from src.campaigns.chain_template_utils import substitute_chain_buttons_preview

    html = substitute_chain_buttons_preview(html)
    return {"subject": subject, "body_html": html, "sample": sample}

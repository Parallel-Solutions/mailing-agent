"""Surgical find/replace edits for existing DOCX files (preserve structure)."""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document


def _paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs) if paragraph.runs else (paragraph.text or "")


def _run_format_signature(run) -> tuple[Any, ...]:
    color = None
    if run.font.color is not None and run.font.color.rgb is not None:
        color = str(run.font.color.rgb)
    size = run.font.size.pt if run.font.size is not None else None
    return (
        run.bold,
        run.italic,
        run.underline,
        run.font.name,
        size,
        color,
    )


def _paragraph_is_safe_for_text_rewrite(paragraph) -> bool:
    non_empty_runs = [run for run in paragraph.runs if run.text]
    if len(non_empty_runs) <= 1:
        return True
    signatures = {_run_format_signature(run) for run in non_empty_runs}
    return len(signatures) == 1


def _force_replace_paragraph_text(paragraph, new_text: str) -> bool:
    """Put full paragraph text into the first run; clear the rest."""
    current_text = _paragraph_text(paragraph)
    if current_text == new_text:
        return False
    runs = list(paragraph.runs)
    if runs:
        non_empty_runs = [run for run in runs if run.text]
        target_run = non_empty_runs[0] if non_empty_runs else runs[0]
        # python-docx rebuilds Run wrappers on each .runs access — compare by XML node.
        target_node = target_run._r  # noqa: SLF001
        target_run.text = new_text
        for run in runs:
            if run._r is not target_node:  # noqa: SLF001
                run.text = ""
    else:
        paragraph.add_run(new_text)
    return True


def _replace_paragraph_text(paragraph, new_text: str) -> bool:
    current_text = _paragraph_text(paragraph)
    if current_text == new_text:
        return False
    non_empty_runs = [run for run in paragraph.runs if run.text]
    if len(non_empty_runs) > 1 and not _paragraph_is_safe_for_text_rewrite(paragraph):
        return False
    return _force_replace_paragraph_text(paragraph, new_text)


def _replace_fragment_inside_single_run(paragraph, fragment: str, replacement: str) -> bool:
    for run in paragraph.runs:
        if fragment not in run.text:
            continue
        new_text = run.text.replace(fragment, replacement, 1)
        if new_text == run.text:
            return False
        run.text = new_text
        return True
    return False


def _replace_fragment_in_paragraph(paragraph, fragment: str, replacement: str) -> bool:
    current_text = _paragraph_text(paragraph)
    if not fragment or fragment not in current_text:
        return False
    if _replace_fragment_inside_single_run(paragraph, fragment, replacement):
        return True
    new_text = current_text.replace(fragment, replacement, 1)
    if new_text == current_text:
        return False
    # Prefer format-safe rewrite; if runs differ, still apply so assistant edits land.
    if _paragraph_is_safe_for_text_rewrite(paragraph):
        return _replace_paragraph_text(paragraph, new_text)
    return _force_replace_paragraph_text(paragraph, new_text)


def _replace_all_safe_fragments_in_paragraph(paragraph, fragment: str, replacement: str) -> int:
    """Apply replacements in one paragraph. Returns number of successful replacements."""
    if fragment and fragment in replacement:
        return 1 if _replace_fragment_in_paragraph(paragraph, fragment, replacement) else 0

    count = 0
    for _ in range(100):
        current_text = _paragraph_text(paragraph)
        if fragment not in current_text:
            break
        if not _replace_fragment_in_paragraph(paragraph, fragment, replacement):
            break
        count += 1
    return count


def iter_document_paragraphs(document):
    """Yield unique body and table paragraphs (stable order)."""
    seen: set[int] = set()
    for paragraph in document.paragraphs:
        key = id(paragraph._p)  # noqa: SLF001
        if key in seen:
            continue
        seen.add(key)
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    key = id(paragraph._p)  # noqa: SLF001
                    if key in seen:
                        continue
                    seen.add(key)
                    yield paragraph


def extract_plain_text(docx_bytes: bytes, *, limit: int = 12000) -> str:
    """Return plain text with paragraph markers like ``[p0] …``."""
    document = Document(BytesIO(docx_bytes))
    lines: list[str] = []
    for index, paragraph in enumerate(iter_document_paragraphs(document)):
        text = _paragraph_text(paragraph)
        lines.append(f"[p{index}] {text}")
    joined = "\n".join(lines)
    if len(joined) <= limit:
        return joined
    return joined[: limit - 1] + "…"


def apply_text_replacements(
    docx_bytes: bytes,
    edits: list[dict[str, Any]],
) -> tuple[bytes, list[dict[str, Any]]]:
    """
    Apply exact find→replace edits to an existing DOCX.

    Each edit: ``{"find": str, "replace": str, "replace_all": bool?}``.
    Returns ``(new_bytes, report)`` where report items have status
    ``applied`` / ``not_found`` / ``unsafe`` / ``skipped``.
    """
    document = Document(BytesIO(docx_bytes))
    paragraphs = list(iter_document_paragraphs(document))
    report: list[dict[str, Any]] = []

    for raw in edits:
        find = str((raw or {}).get("find") or "")
        replace = str((raw or {}).get("replace") if (raw or {}).get("replace") is not None else "")
        replace_all = bool((raw or {}).get("replace_all"))
        if not find:
            report.append(
                {
                    "find": find,
                    "replace": replace,
                    "status": "skipped",
                    "reason": "empty_find",
                    "count": 0,
                }
            )
            continue
        if find == replace:
            report.append(
                {
                    "find": find,
                    "replace": replace,
                    "status": "skipped",
                    "reason": "noop",
                    "count": 0,
                }
            )
            continue

        total = 0
        saw_fragment = False
        hit_unsafe = False
        for paragraph in paragraphs:
            text = _paragraph_text(paragraph)
            if find not in text:
                continue
            saw_fragment = True
            if replace_all:
                applied = _replace_all_safe_fragments_in_paragraph(paragraph, find, replace)
            else:
                applied = 1 if _replace_fragment_in_paragraph(paragraph, find, replace) else 0
            if applied:
                total += applied
                if not replace_all:
                    break
            elif find in _paragraph_text(paragraph):
                hit_unsafe = True

        if total > 0:
            status = "applied"
            reason = ""
        elif not saw_fragment:
            status = "not_found"
            reason = "fragment_not_in_document"
        elif hit_unsafe:
            status = "unsafe"
            reason = "unsafe_formatting"
        else:
            status = "not_found"
            reason = "fragment_not_in_document"

        report.append(
            {
                "find": find,
                "replace": replace,
                "status": status,
                "reason": reason,
                "count": total,
            }
        )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue(), report

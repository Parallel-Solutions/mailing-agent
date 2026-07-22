from __future__ import annotations

import unittest
import uuid
from io import BytesIO
from unittest.mock import patch

import fitz
from docx import Document

from src.campaigns import template_import_service
from src.generator.generation.docxjs_converter import DocxJsHtmlResult
from src.generator.generation.import_visual_qa import PickResult
from src.security.user_store import create_user
from tests.bootstrap import bootstrap_test_runtime


def _mini_png(color: tuple[int, int, int] = (255, 0, 0)) -> bytes:
    document = fitz.open()
    page = document.new_page(width=120, height=160)
    page.draw_rect(page.rect, color=tuple(c / 255 for c in color), fill=tuple(c / 255 for c in color))
    png = page.get_pixmap(alpha=False).tobytes("png")
    document.close()
    return png


class TemplateImportServiceTests(unittest.TestCase):
    def setUp(self) -> None:
        bootstrap_test_runtime(reset_db=True)
        self.username = f"import{uuid.uuid4().hex[:8]}"
        create_user(self.username, "Pass12345!")

    def test_normalize_email_html_wraps_fragment(self) -> None:
        normalized = template_import_service.normalize_email_html("<p>Hello</p>")
        self.assertIn("<table", normalized.lower())
        self.assertIn("Hello", normalized)
        self.assertIn("max-width:640px", normalized)

    def test_normalize_email_html_uses_custom_max_width(self) -> None:
        normalized = template_import_service.normalize_email_html(
            "<p>Hello</p>",
            max_width=794,
        )
        self.assertIn("max-width:794px", normalized)
        self.assertIn('data-content-width="794"', normalized)

    def test_normalize_email_html_preserves_rich_layout(self) -> None:
        normalized = template_import_service.normalize_email_html(
            '<div class="docx-wrapper"><p>Hello</p></div>',
            preserve_inner_layout=True,
            max_width=720,
        )
        self.assertIn("docx-wrapper", normalized)
        self.assertIn("main-body", normalized)
        self.assertIn("max-width:720px", normalized)

    def test_normalize_email_html_strips_scripts(self) -> None:
        normalized = template_import_service.normalize_email_html(
            '<p>Hi</p><script>alert(1)</script>'
        )
        self.assertNotIn("<script", normalized.lower())
        self.assertIn("Hi", normalized)

    def test_import_html_creates_visual_template(self) -> None:
        html = b"<table><tr><td><h1>Promo</h1><p>Hello {{company}}</p></td></tr></table>"
        result = template_import_service.import_visual_email_template(
            self.username,
            filename="promo.html",
            data=html,
        )
        self.assertEqual(result["template_type"], "email")
        self.assertEqual(result["version"]["editor_state"]["email_format"], "visual")
        self.assertIn("{{company}}", result["version"]["body_html"])
        self.assertIn("import", result["tags"])
        refinement = result["version"]["editor_state"].get("import_refinement")
        self.assertIsInstance(refinement, dict)
        self.assertEqual(refinement.get("selected_source"), "html")
        self.assertFalse(refinement.get("available"))

    def test_import_txt_creates_paragraphs(self) -> None:
        result = template_import_service.import_visual_email_template(
            self.username,
            filename="letter.txt",
            data="Line one\nLine two".encode("utf-8"),
        )
        body_html = result["version"]["body_html"]
        self.assertIn("Line one", body_html)
        self.assertIn("Line two", body_html)
        self.assertEqual(result["version"]["editor_state"]["import_source"], "txt")

    def test_import_docx_uses_docxjs_candidate(self) -> None:
        buffer = BytesIO()
        document = Document()
        document.add_heading("Коммерческое предложение", level=1)
        document.add_paragraph("Уважаемый {{contact_name}}!")
        document.save(buffer)

        rich_html = (
            '<div class="docx-wrapper" data-content-width="794">'
            "<h1>Коммерческое предложение</h1><p>Уважаемый {{contact_name}}!</p></div>"
        )
        with (
            patch(
                "src.campaigns.template_import_service.convert_docx_to_html_result",
                return_value=DocxJsHtmlResult(html=rich_html, content_width=794),
            ),
            patch(
                "src.campaigns.template_import_service.convert_docx_to_pdf_bytes",
                return_value=None,
            ),
        ):
            result = template_import_service.import_visual_email_template(
                self.username,
                filename="offer.docx",
                data=buffer.getvalue(),
            )

        body_html = result["version"]["body_html"]
        self.assertEqual(result["version"]["editor_state"]["import_source"], "docxjs")
        self.assertIn("{{contact_name}}", body_html)

    def test_import_docx_picks_best_candidate(self) -> None:
        buffer = BytesIO()
        document = Document()
        document.add_paragraph("Draft {{company}}")
        document.save(buffer)

        docx_html = '<div data-content-width="640"><p>Draft {{company}}</p></div>'
        fixed_html = '<div data-layout="fixed"><span>Draft {{company}}</span></div>'
        with (
            patch(
                "src.campaigns.template_import_service.convert_docx_to_html_result",
                return_value=DocxJsHtmlResult(html=docx_html, content_width=640),
            ),
            patch(
                "src.campaigns.template_import_service.convert_docx_to_pdf_bytes",
                return_value=b"%PDF-1.4",
            ),
            patch(
                "src.campaigns.template_import_service.extract_fixed_layout",
                return_value=(fixed_html, "Draft {{company}}", [_mini_png()], 640),
            ),
            patch(
                "src.campaigns.template_import_service.pick_best_candidate",
                return_value=PickResult(
                    html=fixed_html,
                    name="fixed_layout",
                    score=0.95,
                    scores={"docxjs": 0.7, "fixed_layout": 0.95},
                ),
            ),
        ):
            result = template_import_service.import_visual_email_template(
                self.username,
                filename="offer.docx",
                data=buffer.getvalue(),
            )

        self.assertEqual(result["version"]["editor_state"]["import_source"], "fixed_layout")
        refinement = result["version"]["editor_state"]["import_refinement"]
        self.assertEqual(refinement.get("qa", {}).get("winner"), "fixed_layout")

    def test_import_pdf_uses_fixed_layout(self) -> None:
        fixture = (
            __import__("pathlib").Path(__file__).resolve().parent / "fixtures" / "pismo_sl.pdf"
        )
        if not fixture.is_file():
            self.skipTest("tests/fixtures/pismo_sl.pdf missing")

        result = template_import_service.import_visual_email_template(
            self.username,
            filename="pismo_sl.pdf",
            data=fixture.read_bytes(),
        )
        self.assertEqual(result["version"]["editor_state"]["import_source"], "fixed_layout")
        self.assertIn('data-layout="fixed"', result["version"]["body_html"])

    def test_prepare_import_html_preserves_fixed_layout(self) -> None:
        fragment = (
            '<style>@font-face{font-family:TestFont;src:url(data:font/woff;base64,AA==)}</style>'
            '<div data-layout="fixed" style="font-size:12px"><span>Hi</span></div>'
        )
        prepared = template_import_service._prepare_import_html(fragment, content_width=640)
        self.assertIn('data-layout="fixed"', prepared)
        self.assertIn("@font-face", prepared)

    def test_regenerate_imported_template_requires_import_tag(self) -> None:
        created = template_import_service.import_visual_email_template(
            self.username,
            filename="promo.html",
            data=b"<p>Hello</p>",
        )
        template_id = created["id"]
        with patch.object(
            template_import_service,
            "_refine_import_iteratively",
            return_value=(
                {"body_html": "<p>Updated</p>", "name": "promo", "subject": "Тема"},
                {"rounds": 1, "best_score": 0.9, "spent_usd": 0.1, "stop_reason": "plan_done", "source": "vision_iterative", "trace": []},
            ),
        ):
            updated = template_import_service.regenerate_imported_template(
                self.username,
                template_id,
            )
        self.assertEqual(updated["version"]["editor_state"]["import_source"], "vision_iterative")
        self.assertIn("Updated", updated["version"]["body_html"])


if __name__ == "__main__":
    unittest.main()

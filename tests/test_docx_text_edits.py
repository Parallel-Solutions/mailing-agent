from __future__ import annotations

import unittest
from io import BytesIO

from docx import Document

from src.campaigns import docx_text_edits


def _docx_bytes(*paragraphs: str) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class DocxTextEditsTests(unittest.TestCase):
    def test_extract_plain_text_marks_paragraphs(self) -> None:
        data = _docx_bytes("Первый", "Второй {{company}}")
        text = docx_text_edits.extract_plain_text(data)
        self.assertIn("[p0] Первый", text)
        self.assertIn("[p1] Второй {{company}}", text)

    def test_replace_one_phrase_keeps_other_paragraphs(self) -> None:
        data = _docx_bytes(
            "Уважаемый клиент!",
            "Мы предлагаем услугу.",
            "С уважением, {{contact_name}}",
        )
        new_bytes, report = docx_text_edits.apply_text_replacements(
            data,
            [{"find": "услугу", "replace": "поддержку"}],
        )
        self.assertEqual(report[0]["status"], "applied")
        self.assertEqual(report[0]["count"], 1)
        text = docx_text_edits.extract_plain_text(new_bytes)
        self.assertIn("поддержку", text)
        self.assertNotIn("услугу", text)
        self.assertIn("Уважаемый клиент!", text)
        self.assertIn("{{contact_name}}", text)

    def test_replace_missing_fragment(self) -> None:
        data = _docx_bytes("Один абзац")
        _new_bytes, report = docx_text_edits.apply_text_replacements(
            data,
            [{"find": "нет такого", "replace": "другое"}],
        )
        self.assertEqual(report[0]["status"], "not_found")
        self.assertEqual(report[0]["count"], 0)

    def test_replace_all_occurrences(self) -> None:
        data = _docx_bytes("тест и ещё тест")
        new_bytes, report = docx_text_edits.apply_text_replacements(
            data,
            [{"find": "тест", "replace": "проверка", "replace_all": True}],
        )
        self.assertEqual(report[0]["status"], "applied")
        self.assertEqual(report[0]["count"], 2)
        text = docx_text_edits.extract_plain_text(new_bytes)
        self.assertEqual(text.count("проверка"), 2)
        self.assertNotIn("тест", text)

    def test_replace_across_mixed_format_runs(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        normal = paragraph.add_run("Привет, ")
        normal.bold = False
        bold = paragraph.add_run("мир")
        bold.bold = True
        buffer = BytesIO()
        document.save(buffer)
        new_bytes, report = docx_text_edits.apply_text_replacements(
            buffer.getvalue(),
            [{"find": "Привет, мир", "replace": "Здравствуй, коллега"}],
        )
        self.assertEqual(report[0]["status"], "applied")
        text = docx_text_edits.extract_plain_text(new_bytes)
        self.assertIn("Здравствуй, коллега", text)
        self.assertNotIn("Привет, мир", text)

    def test_replace_inside_single_run_preserves_bold(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        run = paragraph.add_run("важный текст здесь")
        run.bold = True
        buffer = BytesIO()
        document.save(buffer)
        new_bytes, report = docx_text_edits.apply_text_replacements(
            buffer.getvalue(),
            [{"find": "текст", "replace": "фрагмент"}],
        )
        self.assertEqual(report[0]["status"], "applied")
        updated = Document(BytesIO(new_bytes))
        target = next(
            paragraph
            for paragraph in updated.paragraphs
            if any(run.text for run in paragraph.runs)
        )
        bold_run = next(run for run in target.runs if run.text)
        self.assertTrue(bold_run.bold)
        self.assertEqual(bold_run.text, "важный фрагмент здесь")


if __name__ == "__main__":
    unittest.main()

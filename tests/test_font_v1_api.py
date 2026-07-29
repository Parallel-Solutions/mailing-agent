from __future__ import annotations

import unittest
import uuid
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

from src.campaigns.font_service import inspect_font_bytes
from src.security.auth import Principal
from src.security.user_store import create_user
from src.web.v1_router import create_v1_router
from tests.bootstrap import bootstrap_test_runtime


def _installed_font() -> Path | None:
    candidates = (
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    )
    return next((path for path in candidates if path.is_file()), None)


class FontV1ApiTests(unittest.TestCase):
    def setUp(self) -> None:
        bootstrap_test_runtime(reset_db=True)
        self.username = f"font{uuid.uuid4().hex[:8]}"
        create_user(self.username, "Pass12345!")
        app = FastAPI()
        app.include_router(
            create_v1_router(check_auth=lambda: Principal(self.username, "t1", "user"))
        )
        self.client = TestClient(app)

    def test_upload_font_resolves_matching_docx_requirement(self) -> None:
        font_path = _installed_font()
        if font_path is None:
            self.skipTest("No known system TTF is installed")
        font_data = font_path.read_bytes()
        metadata = inspect_font_bytes(font_path.name, font_data)

        missing_license = self.client.post(
            "/api/v1/fonts/upload",
            data={"license_confirmed": "false"},
            files={"file": (font_path.name, font_data, "font/ttf")},
        )
        self.assertEqual(missing_license.status_code, 400, missing_license.text)

        uploaded = self.client.post(
            "/api/v1/fonts/upload",
            data={"license_confirmed": "true"},
            files={"file": (font_path.name, font_data, "font/ttf")},
        )
        self.assertEqual(uploaded.status_code, 200, uploaded.text)
        self.assertEqual(uploaded.json()["result"]["family"], metadata.family)

        document = Document()
        run = document.add_paragraph().add_run("Проверка шрифта")
        run.font.name = metadata.family
        payload = BytesIO()
        document.save(payload)
        with patch(
            "src.campaigns.template_service._build_document_pdf_artifact",
            return_value=(b"%PDF-1.4 test", "font-test.pdf"),
        ):
            template_response = self.client.post(
                "/api/v1/templates/upload",
                data={"template_type": "document", "name": "Font test"},
                files={
                    "file": (
                        "font-test.docx",
                        payload.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        self.assertEqual(template_response.status_code, 200, template_response.text)
        template_id = template_response.json()["result"]["id"]

        analysis = self.client.get(f"/api/v1/templates/{template_id}/fonts")

        self.assertEqual(analysis.status_code, 200, analysis.text)
        requirements = analysis.json()["result"]["requirements"]
        matching = next(
            item
            for item in requirements
            if item["family_normalized"] == metadata.family_normalized
        )
        self.assertEqual(matching["status"], "resolved")
        self.assertEqual(matching["font_asset"]["id"], uploaded.json()["result"]["id"])


if __name__ == "__main__":
    unittest.main()

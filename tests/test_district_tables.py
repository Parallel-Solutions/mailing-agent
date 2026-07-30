import shutil
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from src.generator.generation.document_builder import (
    build_contract_filename,
    build_contract_replacements,
    build_kp_replacements,
    render_docx,
)
from src.generator.generation.excel_io import load_rows
from src.generator.generation.transforms import build_document_context, build_output_folder_name
from src.generator.generation.work_types import WORK_TYPE_STP_MO
from src.generator.verification.municipality_name_verifier import verify_municipality_names_in_workbook


class DistrictTableTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp_dir = Path("tmp_test_district_tables")
        self.tmp_dir.mkdir(exist_ok=True)

    def tearDown(self) -> None:
        if self.tmp_dir.exists():
            shutil.rmtree(self.tmp_dir)

    def _district_xlsx(self) -> Path:
        path = self.tmp_dir / "district.xlsx"
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.cell(row=1, column=1).value = "N"
        headers = [
            "ID",
            "SUB_RF",
            "MUN_R_NAME",
            "MUN_NAME",
            "ADM_NAME",
            "ADRES",
            "HEAD_FIO",
            "POPULATION",
            "EMAIL_OSN",
            "REQUISITES_OKTNO",
            "STATUS",
        ]
        for column_index, header in enumerate(headers, start=1):
            worksheet.cell(row=2, column=column_index).value = header
        values = [
            1,
            "Краснодарский край",
            "Новокубанский район",
            None,
            "АДМИНИСТРАЦИЯ МУНИЦИПАЛЬНОГО ОБРАЗОВАНИЯ НОВОКУБАНСКИЙ РАЙОН",
            "352240, Краснодарский край, Новокубанский район",
            "Гомодин Александр Владимирович",
            81734,
            "omsiknk@mail.ru",
            "3634101001",
            "ОК",
        ]
        for column_index, value in enumerate(values, start=1):
            worksheet.cell(row=3, column=column_index).value = value
        source_values = ["Источники:", 'Файл "БАЗА МО 2025 год"', 'Файл "БАЗА МО 2025 год"']
        for column_index, value in enumerate(source_values, start=1):
            worksheet.cell(row=4, column=column_index).value = value
        workbook.save(path)
        workbook.close()
        return path

    def test_load_rows_skips_source_row_and_keeps_oktmo_alias(self) -> None:
        _, _, rows = load_rows(self._district_xlsx())

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["MUN_R_NAME"], "Новокубанский район")
        self.assertEqual(rows[0]["REQUISITES_OKTMO"], "3634101001")

    def test_district_table_verification_is_accepted_without_mun_name(self) -> None:
        result = verify_municipality_names_in_workbook(self._district_xlsx(), use_oktmo=True)

        self.assertEqual(result["status"], "ok")
        self.assertEqual(result["table_mode"], "district")
        self.assertEqual(result["total_rows"], 1)
        self.assertEqual(result["kept_rows"], 1)

    def test_district_context_uses_mun_r_name_as_main_entity(self) -> None:
        _, _, rows = load_rows(self._district_xlsx())
        context = build_document_context(rows[0], outgoing_number=101)
        replacements = dict(build_kp_replacements(context))
        contract_replacements = dict(build_contract_replacements(context))

        self.assertEqual(context["DOCUMENT_ENTITY_TYPE"], "district")
        self.assertEqual(context["MUN_NAME"], "Новокубанский муниципальный район")
        self.assertEqual(context["MUN_R_NAME"], "Новокубанский район")
        self.assertIn("Новокубанского", context["MUN_R_NAME_1"])
        self.assertIn("Новокубанского", context["WORK_SCOPE_FRAGMENT"])
        self.assertEqual(replacements["MUN_R_NAME"], "Новокубанский район")
        self.assertEqual(replacements["SUB_RF"], "Краснодарский край")
        self.assertEqual(contract_replacements["Глава MUN_NAME"], "Глава Новокубанского района")
        self.assertIn("Новокубанский", build_contract_filename(rows[0]))
        self.assertIn("Новокубанский", build_output_folder_name(rows[0]))

    def test_district_context_canonicalizes_administration_company_name(self) -> None:
        row = {
            "ID": 1,
            "SUB_RF": "Брянская область",
            "MUN_R_NAME": "Дятьковский район",
            "MUN_NAME": "Администрация Дятьковского района",
            "ADM_NAME": (
                'Администрация муниципального образования '
                '«Администрация Дятьковского района»'
            ),
            "HEAD_FIO": "Иванов Иван Иванович",
        }

        context = build_document_context(row, outgoing_number=101)

        self.assertEqual(context["DOCUMENT_ENTITY_TYPE"], "district")
        self.assertEqual(context["MUN_NAME"], "Дятьковский муниципальный район")
        self.assertEqual(
            context["ADM_NAME"],
            "администрация Дятьковского муниципального района",
        )
        self.assertEqual(
            context["ADM_NAME_1"],
            "администрации Дятьковского муниципального района",
        )
        self.assertNotIn("«Администрация", context["ADM_NAME_1"])

    def test_kp_replacements_inflect_district_scope_with_extra_placeholder_space(self) -> None:
        row = {
            "ID": 1,
            "SUB_RF": "Иркутская область",
            "MUN_R_NAME": "Жигаловский муниципальный округ",
            "MUN_NAME": None,
            "ADM_NAME": "",
            "HEAD_FIO": "Иванов Иван Иванович",
        }
        context = build_document_context(row, outgoing_number=101, work_type=WORK_TYPE_STP_MO)
        template_path = self.tmp_dir / "kp-extra-space-template.docx"
        output_path = self.tmp_dir / "kp-extra-space-output.docx"
        document = Document()
        table = document.add_table(rows=1, cols=1)
        paragraph = table.cell(0, 0).paragraphs[0]
        paragraph.add_run("Выполнение работ по схеме территориального планирования (СТП) ")
        paragraph.add_run("MUN_R_NAME ")
        paragraph.add_run(" ")
        paragraph.add_run("SUB_RF")
        document.save(template_path)

        render_docx(template_path, build_kp_replacements(context), output_path, context)

        rendered_text = Document(output_path).tables[0].cell(0, 0).text
        self.assertIn("Жигаловского муниципального округа Иркутской области", rendered_text)
        self.assertNotIn("Жигаловский муниципальный округ Иркутская область", rendered_text)

    def test_kp_render_inflects_single_district_placeholder_after_work_title(self) -> None:
        row = {
            "ID": 1,
            "SUB_RF": "Республика Адыгея",
            "MUN_R_NAME": "Тахтамукайский муниципальный район",
            "MUN_NAME": "Яблоновское городское поселение",
            "ADM_NAME": "Администрация Яблоновского городского поселения",
            "HEAD_FIO": "Иванов Иван Иванович",
        }
        context = build_document_context(row, outgoing_number=186)
        template_path = self.tmp_dir / "uploaded-single-district-template.docx"
        output_path = self.tmp_dir / "uploaded-single-district-output.docx"
        document = Document()
        document.add_paragraph("Район: MUN_R_NAME")
        table = document.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Вид работ"
        paragraph = table.rows[1].cells[0].paragraphs[0]
        paragraph.add_run(
            "Выполнение работ по разработке проекта местных нормативов "
            "градостроительного проектирования "
        )
        district_run = paragraph.add_run("MUN_R_NAME")
        district_run.bold = False
        document.save(template_path)

        render_docx(
            template_path,
            build_kp_replacements(context),
            output_path,
            context,
        )

        rendered = Document(output_path)
        self.assertEqual(
            rendered.paragraphs[0].text,
            "Район: Тахтамукайский муниципальный район",
        )
        rendered_work = rendered.tables[0].rows[1].cells[0].paragraphs[0]
        self.assertIn(
            "Тахтамукайского муниципального района",
            rendered_work.text,
        )
        self.assertNotIn(
            "Тахтамукайский муниципальный район",
            rendered_work.text,
        )
        self.assertFalse(rendered_work.runs[-1].bold)

    def test_municipality_contract_replacements_inflect_district_scope(self) -> None:
        row = {
            "ID": 1,
            "SUB_RF": "Иркутская область",
            "MUN_R_NAME": "Усть-Кутский район",
            "MUN_NAME": "Нийское сельское поселение",
            "ADM_NAME": 'АДМИНИСТРАЦИЯ НИЙСКОГО СЕЛЬСКОГО ПОСЕЛЕНИЯ УСТЬ-КУТСКОГО МУНИЦИПАЛЬНОГО РАЙОНА ИРКУТСКОЙ ОБЛАСТИ',
            "HEAD_FIO": "Дудник Евгения Викторовна",
        }
        context = build_document_context(row, outgoing_number=101)
        replacements = dict(build_contract_replacements(context))

        self.assertEqual(context["MUN_R_NAME_1"], "Усть-Кутского района")
        self.assertEqual(context["SUB_RF_1"], "Иркутской области")
        self.assertEqual(replacements["MUN_R_NAME SUB_RF"], "Усть-Кутского муниципального района Иркутской области")
        self.assertEqual(replacements["Глава MUN_NAME"], "Глава Нийского сельского поселения")


if __name__ == "__main__":
    unittest.main()

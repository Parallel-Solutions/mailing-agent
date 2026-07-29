from __future__ import annotations

import unittest
import uuid
from io import BytesIO
from unittest.mock import patch

from docx import Document

from src.campaigns.validation_auto_fix_service import _apply_philologist_docx_fixes
from src.security.auth import Principal
from src.security.user_store import create_user
from src.web.v1_router import create_v1_router
from tests.bootstrap import bootstrap_test_runtime
from fastapi import FastAPI
from fastapi.testclient import TestClient


class ValidationAutoFixDocxTests(unittest.TestCase):
    def setUp(self) -> None:
        bootstrap_test_runtime(reset_db=True)
        self.username = f"afd{uuid.uuid4().hex[:8]}"
        create_user(self.username, "Pass12345!")
        app = FastAPI()
        app.include_router(create_v1_router(check_auth=lambda: Principal(self.username, "t1", "user")))
        self.client = TestClient(app)

    def test_apply_philologist_docx_fixes_updates_source_template(self) -> None:
        docx_buffer = BytesIO()
        document = Document()
        document.add_paragraph(
            "1.2. Выполнение Работ осуществляется по месту нахождения Исполнителя "
            "на условиях и в сроки, установленные настоящим Договором."
        )
        document.save(docx_buffer)

        with patch(
            "src.campaigns.template_service._build_document_pdf_artifact",
            return_value=(b"%PDF-1.4 test", "legal.pdf"),
        ):
            uploaded = self.client.post(
                "/api/v1/templates/upload",
                data={"template_type": "document", "name": "Legal doc"},
                files={
                    "file": (
                        "legal.docx",
                        docx_buffer.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        self.assertEqual(uploaded.status_code, 200, uploaded.text)
        document_id = uploaded.json()["result"]["id"]

        with patch(
            "src.campaigns.template_service._build_document_pdf_artifact",
            return_value=(b"%PDF-1.4 test", "legal.pdf"),
        ):
            applied, skipped = _apply_philologist_docx_fixes(document_id, self.username)
        self.assertTrue(applied, msg=f"skipped={skipped}")
        self.assertTrue(any(item.get("kind") == "document" for item in applied))

        file_response = self.client.get(f"/api/v1/templates/{document_id}/file")
        self.assertEqual(file_response.status_code, 200, file_response.text)
        saved = Document(BytesIO(file_response.content))
        text = "\n".join(paragraph.text for paragraph in saved.paragraphs)
        self.assertIn("Выполнение работ", text)
        self.assertIn("нахождения исполнителя", text)


if __name__ == "__main__":
    unittest.main()

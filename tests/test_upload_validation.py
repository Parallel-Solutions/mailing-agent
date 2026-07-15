from __future__ import annotations

import shutil
import unittest
import uuid
import zipfile
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from fastapi import FastAPI
from fastapi.testclient import TestClient
from docx import Document

from src.security.auth import Principal
from src.web.jobs_router import JobsWebController
from src.web.upload_validation import validate_uploaded_file


PROJECT_ROOT = Path(__file__).resolve().parents[1]


def _workspace_temp_root(prefix: str) -> Path:
    root = PROJECT_ROOT / "tmp" / f"{prefix}-{uuid.uuid4().hex}"
    root.mkdir(parents=True)
    return root


def _cleanup(path: Path) -> None:
    shutil.rmtree(path, ignore_errors=True)


def _job_paths(root: Path, job_id: str | None) -> SimpleNamespace:
    base = root / ("legacy" if not job_id else f"jobs/{job_id}")

    def ensure_dirs() -> None:
        base.mkdir(parents=True, exist_ok=True)

    return SimpleNamespace(
        job_id=job_id,
        root_dir=base,
        data_xlsx=base / "input" / "data.xlsx",
        base_xlsx=base / "input" / "base.xlsx",
        templates_dir=base / "templates",
        output_dir=base / "output",
        consents_dir=base / "consents",
        sent_mail_log_path=base / "sent_mail_log.jsonl",
        uses_legacy_layout=job_id is None,
        ensure_dirs=ensure_dirs,
    )


def _ooxml_bytes(*members: str) -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        for member in members:
            archive.writestr(member, "<xml/>")
    return buffer.getvalue()


def _legacy_docx_template_bytes() -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph("№ 101-КП от 12.05.2026")
    document.add_paragraph("ADM_NAME")
    document.add_paragraph("Работы для MUN_R_NAME SUB_RF")
    document.save(buffer)
    return buffer.getvalue()

class UploadValidationTests(unittest.TestCase):
    def _client(self, root: Path, verification_calls: list[dict]) -> TestClient:
        resolver = lambda job_id=None: _job_paths(root, job_id)
        controller = JobsWebController(
            check_auth=lambda: Principal("admin", "root", "admin"),
            settings=SimpleNamespace(upload_data_max_bytes=4096, upload_template_max_bytes=128 * 1024),
            logger=SimpleNamespace(exception=lambda *args, **kwargs: None, info=lambda *args, **kwargs: None),
            prefer_existing_file=lambda primary, fallback: primary if primary.exists() else fallback,
            validate_uploaded_file=validate_uploaded_file,
            cached_excel_row_count=lambda path: 0,
            cached_tree_file_count=lambda path, pattern: 0,
            safe_int=lambda value: int(value or 0),
            create_job_id=lambda: "job-upload",
            resolve_job_paths=resolver,
            jobs_dir=root / "jobs",
            create_documents_load_test_job=lambda **kwargs: {},
            start_parser_verification_process=lambda **kwargs: verification_calls.append(kwargs),
            get_parser_status=lambda job_id: {},
            get_generator_status=lambda job_id: {},
            get_philologist_status=lambda job_id, include_details=False: {},
            get_sender_status=lambda job_id: {},
            run_parser_municipality_verification=lambda *args, **kwargs: {},
        )
        app = FastAPI()
        app.include_router(controller.router)
        return TestClient(app)

    def test_upload_data_rejects_renamed_text_file_before_save(self) -> None:
        root = _workspace_temp_root("upload-validation-data-invalid")
        try:
            verification_calls: list[dict] = []
            client = self._client(root, verification_calls)

            response = client.post(
                "/api/upload/data",
                files={
                    "file": (
                        "data.xlsx",
                        b"plain text, not a workbook",
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    )
                },
            )

            self.assertEqual(response.status_code, 400)
            self.assertIn("не соответствует формату .xlsx", response.json()["detail"])
            self.assertFalse(_job_paths(root, None).data_xlsx.exists())
            self.assertEqual(verification_calls, [])
        finally:
            _cleanup(root)

    def test_upload_template_rejects_docx_zip_without_document_member(self) -> None:
        root = _workspace_temp_root("upload-validation-template-invalid")
        try:
            verification_calls: list[dict] = []
            client = self._client(root, verification_calls)
            wrong_ooxml = _ooxml_bytes("xl/workbook.xml")

            response = client.post(
                "/api/upload/template",
                data={"template_kind": "kp"},
                files={
                    "file": (
                        "kp.docx",
                        wrong_ooxml,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )

            self.assertEqual(response.status_code, 400)
            self.assertIn("не соответствует формату .docx", response.json()["detail"])
            self.assertFalse((_job_paths(root, None).templates_dir / "kp_template_source.docx").exists())
            self.assertEqual(verification_calls, [])
        finally:
            _cleanup(root)

    def test_upload_data_accepts_xlsx_ooxml_signature(self) -> None:
        root = _workspace_temp_root("upload-validation-data-valid")
        try:
            verification_calls: list[dict] = []
            client = self._client(root, verification_calls)
            workbook = _ooxml_bytes("xl/workbook.xml")

            with patch("src.web.jobs_router.append_audit_event", lambda **kwargs: None):
                response = client.post(
                    "/api/upload/data",
                    files={
                        "file": (
                            "data.xlsx",
                            workbook,
                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        )
                    },
                )

            self.assertEqual(response.status_code, 200)
            self.assertTrue(_job_paths(root, None).data_xlsx.exists())
            self.assertEqual(verification_calls, [{"job_id": None, "filename": "data.xlsx", "source": "upload"}])
        finally:
            _cleanup(root)

    def test_upload_template_accepts_docx_ooxml_signature(self) -> None:
        root = _workspace_temp_root("upload-validation-template-valid")
        try:
            verification_calls: list[dict] = []
            client = self._client(root, verification_calls)
            document = _ooxml_bytes("word/document.xml")

            with patch("src.web.jobs_router.append_audit_event", lambda **kwargs: None):
                response = client.post(
                    "/api/upload/template",
                    data={"template_kind": "contract"},
                    files={
                        "file": (
                            "contract.docx",
                            document,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    },
                )

            dest = _job_paths(root, None).templates_dir / "contract_template_source.docx"
            self.assertEqual(response.status_code, 200)
            self.assertTrue(dest.exists())
            self.assertEqual(verification_calls, [])
        finally:
            _cleanup(root)


    def test_upload_kp_template_accepts_pdf_and_replaces_docx(self) -> None:
        root = _workspace_temp_root("upload-validation-kp-pdf")
        try:
            verification_calls: list[dict] = []
            client = self._client(root, verification_calls)
            stale_docx = _job_paths(root, None).templates_dir / "kp_template_source.docx"
            stale_docx.parent.mkdir(parents=True, exist_ok=True)
            stale_docx.write_bytes(_ooxml_bytes("word/document.xml"))

            with (
                patch("src.web.jobs_router.append_audit_event", lambda **kwargs: None),
                patch("src.generator.generation.config_generator.KP_ADAPTIVE_TEMPLATE_ENGINE", False),
            ):
                response = client.post(
                    "/api/upload/template",
                    data={"template_kind": "kp"},
                    files={"file": ("kp.pdf", b"%PDF-1.4\n% test", "application/pdf")},
                )

            dest = _job_paths(root, None).templates_dir / "kp_template_source.pdf"
            self.assertEqual(response.status_code, 200)
            self.assertTrue(dest.exists())
            self.assertFalse(stale_docx.exists())
            self.assertEqual(response.json()["result"]["stored_as"], "kp_template_source.pdf")
        finally:
            _cleanup(root)

    def test_upload_kp_template_auto_discovers_fields_without_manual_markup(self) -> None:
        root = _workspace_temp_root("upload-validation-kp-auto")
        try:
            verification_calls: list[dict] = []
            client = self._client(root, verification_calls)

            with (
                patch("src.generator.generation.config_generator.KP_ADAPTIVE_TEMPLATE_ENGINE", True),
                patch("src.web.jobs_router.append_audit_event", lambda **kwargs: None),
                patch("src.workers.task_queue.enqueue_task", return_value=({"task_id": "certify-1"}, True)),
            ):
                response = client.post(
                    "/api/upload/template",
                    data={"template_kind": "kp"},
                    files={
                        "file": (
                            "kp.docx",
                            _legacy_docx_template_bytes(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    },
                )

            payload = response.json()
            self.assertEqual(response.status_code, 200, payload)
            self.assertEqual(
                payload["result"]["adaptive_template"]["fields"],
                ["ADM_NAME_1", "DATE", "MUN_R_SCOPE_FRAGMENT", "OUTGOING_NUMBER"],
            )
            self.assertEqual(payload["result"]["adaptive_task"]["task_id"], "certify-1")
            self.assertTrue((_job_paths(root, None).templates_dir / "kp_template_source.docx").exists())
        finally:
            _cleanup(root)
if __name__ == "__main__":
    unittest.main()

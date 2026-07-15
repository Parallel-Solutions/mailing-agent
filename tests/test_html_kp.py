from __future__ import annotations

import shutil
import sys
import types
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.shared import RGBColor

if "alembic" not in sys.modules:
    alembic_stub = types.ModuleType("alembic")
    alembic_config_stub = types.ModuleType("alembic.config")
    alembic_command_stub = types.ModuleType("alembic.command")
    alembic_config_stub.Config = object
    alembic_stub.command = alembic_command_stub
    sys.modules["alembic"] = alembic_stub
    sys.modules["alembic.config"] = alembic_config_stub
    sys.modules["alembic.command"] = alembic_command_stub
from src.generator.generation import html_kp, pdf_converter

try:
    from src.generator.generation import document_builder
except ModuleNotFoundError as exc:
    document_builder = None
    DOCUMENT_BUILDER_IMPORT_ERROR = exc
else:
    DOCUMENT_BUILDER_IMPORT_ERROR = None
from src.generator.generation.transforms import build_document_context


class HtmlKPTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp_dir = Path("tmp_test_html_kp")
        shutil.rmtree(self.tmp_dir, ignore_errors=True)
        self.tmp_dir.mkdir(parents=True, exist_ok=True)

    def tearDown(self) -> None:
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def _row(self) -> dict:
        return {
            "ID": "1",
            "SUB_RF": "Test region",
            "MUN_R_NAME": "Test district",
            "MUN_NAME": "Test settlement",
            "ADM_NAME": "Test administration",
            "HEAD_FIO": "Ivanov Ivan Ivanovich",
            "EMAIL_OSN": "test@example.com",
            "TEL_OSN": "+70000000000",
        }

    def test_build_kp_html_keeps_images_behind_text(self) -> None:
        context = build_document_context(self._row(), 101)

        html = html_kp.build_kp_html(context)

        self.assertIn(".kp-background", html)
        self.assertIn("z-index: 0", html)
        self.assertIn("pointer-events: none", html)
        self.assertIn(".kp-content", html)
        self.assertIn("z-index: 1", html)
        self.assertLess(html.index("kp-background"), html.index("kp-content"))

    def test_build_kp_html_exposes_overflow_to_page_count_validation(self) -> None:
        context = build_document_context(self._row(), 101)

        html = html_kp.build_kp_html(context)

        self.assertIn("min-height: 297mm", html)
        self.assertIn("height: auto", html)
        self.assertIn("overflow: visible", html)
        self.assertNotIn("height: 297mm;\n    overflow: hidden", html)

    def test_html_kp_has_readable_minimum_density_fallback(self) -> None:
        minimum = html_kp.HTML_KP_DENSITIES[-1]

        self.assertEqual(minimum.name, "minimum")
        self.assertGreaterEqual(minimum.font_size_pt, 8.5)
        self.assertGreaterEqual(minimum.table_font_size_pt, 8.0)

    def test_adaptive_density_expands_above_template_font_and_descends(self) -> None:
        with patch.object(
            html_kp,
            "analyze_docx_style_profile",
            return_value={"body_font_size_pt": 11.6},
        ):
            densities = html_kp._adaptive_density_candidates(None)

        self.assertEqual(densities[0].font_size_pt, 13.6)
        self.assertEqual(densities[1].font_size_pt, 13.4)
        self.assertEqual(densities[-1].name, "minimum")
        self.assertTrue(all(
            left.font_size_pt > right.font_size_pt
            for left, right in zip(densities, densities[1:])
        ))

    def test_template_images_are_assigned_to_semantic_slots(self) -> None:
        contact = "\u0418\u0441\u043f. \u041a\u0440\u0430\u0448\u0435\u043d\u0438\u043d\u043d\u0438\u043a\u043e\u0432  \u0442\u0435\u043b. +7 921 409-45-61  ks@parresh.ru"
        assets = [
            html_kp.DocxImageAsset("phone", 3.1, 3.4, False, "word/document.xml", contact),
            html_kp.DocxImageAsset("email", 3.2, 2.2, False, "word/document.xml", contact),
            html_kp.DocxImageAsset("stamp", 45.5, 51.6, False, "word/document.xml"),
            html_kp.DocxImageAsset("decoration", 119.0, 80.0, True, "word/document.xml"),
            html_kp.DocxImageAsset("logo", 98.0, 12.4, False, "word/header1.xml"),
        ]

        classified = html_kp._classify_template_assets(assets)

        self.assertEqual(classified.logo, "logo")
        self.assertEqual(classified.phone_icon, "phone")
        self.assertEqual(classified.email_icon, "email")
        self.assertEqual(classified.stamp, "stamp")
        self.assertEqual(classified.decorations, ("decoration",))
        self.assertEqual(len(classified.contact_lines), 3)

    def test_stamp_is_rendered_inside_signature_after_validity_date(self) -> None:
        context = build_document_context(self._row(), 101)
        assets = html_kp.HtmlKPTemplateAssets(
            logo="logo",
            phone_icon="phone",
            email_icon="email",
            stamp="stamp",
            decorations=("decoration",),
            contact_lines=("executor", "tel. +7 921 409-45-61", "ks@parresh.ru"),
        )

        with patch.object(html_kp, "_extract_docx_template_assets", return_value=assets):
            rendered = html_kp.build_kp_html(context)

        self.assertIn('class="kp-logo"', rendered)
        self.assertIn('class="kp-contact-icon"', rendered)
        self.assertIn('class="kp-stamp"', rendered)
        self.assertIn('class="kp-decoration kp-decoration-1"', rendered)
        self.assertLess(rendered.index('class="kp-validity"'), rendered.index('class="kp-stamp-slot"'))

    def test_build_kp_html_uses_template_style_profile(self) -> None:
        context = build_document_context(self._row(), 101)
        template_path = self.tmp_dir / "template.docx"
        doc = Document()
        run = doc.add_paragraph().add_run("ADM_NAME")
        run.font.name = "Courier New"
        run.font.color.rgb = RGBColor(12, 34, 56)
        doc.save(template_path)

        html = html_kp.build_kp_html(context, template_path=template_path)

        self.assertIn('--kp-font-family: "Courier New"', html)
        self.assertIn("--kp-primary-color: #0C2238", html)

    def test_neutral_body_color_is_not_used_as_template_primary_color(self) -> None:
        with patch.object(
            html_kp,
            "analyze_docx_style_profile",
            return_value={"font_family": "Tahoma", "primary_color": "#595959"},
        ):
            css = html_kp._template_css_vars(None)

        self.assertIn("--kp-primary-color: #232E50", css)
        self.assertIn("--kp-muted-color: #595959", css)

    def test_build_kp_html_preserves_template_colors_and_bold_fragments(self) -> None:
        context = build_document_context(self._row(), 101)
        context["KP_PRICE_RUBLES"] = "150000"
        assets = html_kp.HtmlKPTemplateAssets(
            company_lines=("Company legal name", "Company short name", "Company address"),
        )

        with patch.object(html_kp, "_extract_docx_template_assets", return_value=assets):
            rendered = html_kp.build_kp_html(context)

        self.assertIn("--kp-muted-color: #595959", rendered)
        self.assertIn("--kp-table-header-bg: #D9D9D9", rendered)
        self.assertIn("color: #404040", rendered)
        self.assertIn("color: var(--kp-primary-color); font-weight: 400", rendered)
        self.assertIn("font-size: clamp(8.2pt, calc(var(--kp-font-size) - 3.5pt), 9.5pt)", rendered)
        self.assertNotIn("font-size: 7.6pt", rendered)
        self.assertIn(f"<strong>{context['WORK_TITLE']}</strong>", rendered)
        self.assertIn("<strong>150 000</strong>", rendered)
        self.assertNotIn("<strong>Company legal name</strong>", rendered)

    def test_variable_blocks_have_independent_browser_fitting(self) -> None:
        context = build_document_context(self._row(), 101)
        assets = html_kp.HtmlKPTemplateAssets(
            contact_lines=("executor", "tel. +7 921 409-45-61", "ks@parresh.ru"),
        )

        with patch.object(html_kp, "_extract_docx_template_assets", return_value=assets):
            rendered = html_kp.build_kp_html(context)

        self.assertIn('class="kp-recipient" data-fit-lines="3"', rendered)
        self.assertIn('class="kp-intro" data-fit-lines="5"', rendered)
        self.assertIn('class="kp-included-services" data-fit-lines="5"', rendered)
        self.assertIn('class="kp-work-title" data-fit-lines="4"', rendered)
        self.assertIn('class="kp-price-note" data-fit-lines="3"', rendered)
        self.assertIn('class="kp-contact-row" data-fit-lines="1"', rendered)
        self.assertIn('querySelectorAll("[data-fit-lines]")', rendered)
        self.assertIn('element.style.fontSize = currentPt.toFixed(2) + "pt"', rendered)
        self.assertNotIn('}`n  [data-fit-lines]', rendered)

    def test_html_conversion_uses_separate_gotenberg_html_urls(self) -> None:
        output_path = self.tmp_dir / "html.pdf"
        calls = []

        class FakeResponse:
            content = b"%PDF ok"
            headers = {"content-type": "application/pdf"}

            def raise_for_status(self) -> None:
                return None

        class FakeClient:
            def __init__(self, *args, **kwargs) -> None:
                pass

            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, tb) -> bool:
                return False

            def post(self, endpoint, files, data=None):
                calls.append(
                    {
                        "endpoint": endpoint,
                        "filename": files["files"][0],
                        "data": data or {},
                    }
                )
                return FakeResponse()

        with (
            patch.object(pdf_converter, "GOTENBERG_BASE_URLS", ("http://libreoffice",)),
            patch.object(pdf_converter, "GOTENBERG_HTML_BASE_URLS", ("http://html",)),
            patch.object(pdf_converter, "_healthy_gotenberg_base_urls", return_value=("http://html",)) as healthy,
            patch.object(pdf_converter.httpx, "Client", FakeClient),
        ):
            result = pdf_converter.convert_html_to_pdf("<html></html>", output_path)

        self.assertEqual(result, output_path)
        self.assertEqual(output_path.read_bytes(), b"%PDF ok")
        healthy.assert_called_once_with(("http://html",))
        self.assertEqual(calls[0]["endpoint"], "http://html/forms/chromium/convert/html")
        self.assertEqual(calls[0]["filename"], "index.html")
        self.assertEqual(calls[0]["data"].get("printBackground"), "true")
        self.assertEqual(calls[0]["data"].get("preferCssPageSize"), "true")

    def test_render_html_kp_pdf_retries_until_one_page(self) -> None:
        context = build_document_context(self._row(), 101)
        output_path = self.tmp_dir / "kp.pdf"
        attempts = []

        def fake_convert(html: str, candidate: Path, **kwargs):
            attempts.append(html)
            candidate.write_bytes(b"%PDF fake")
            return candidate

        def fake_validate(path: Path):
            if len(attempts) == 1:
                return {"ok": False, "reason": "page_count", "message": "two pages", "page_count": 2}
            return {"ok": True, "reason": "", "message": "", "page_count": 1}

        with (
            patch.object(html_kp, "convert_html_to_pdf", side_effect=fake_convert),
            patch.object(html_kp, "validate_kp_pdf", side_effect=fake_validate),
        ):
            result = html_kp.render_html_kp_pdf(context, output_path)

        self.assertEqual(result, output_path)
        self.assertTrue(output_path.exists())
        self.assertEqual(len(attempts), 2)
        self.assertIn("--kp-font-size: 14.0pt", attempts[0])
        self.assertIn("--kp-font-size: 13.8pt", attempts[1])
        self.assertFalse(any(output_path.parent.glob("kp.html_try_*.pdf")))

    @unittest.skipIf(document_builder is None, f"document_builder dependencies unavailable: {DOCUMENT_BUILDER_IMPORT_ERROR}")
    def test_generate_documents_for_row_uses_html_direct_pdf_engine(self) -> None:
        row = self._row()
        context = build_document_context(row, 101)
        output_dir = self.tmp_dir / "output"
        batch_dir = self.tmp_dir / "batch"
        templates_dir = self.tmp_dir / "templates"
        templates_dir.mkdir()

        def fake_render(ctx, output_path, **kwargs):
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_bytes(b"%PDF fake")
            return output_path

        with (
            patch.object(document_builder, "KP_ADAPTIVE_TEMPLATE_ENGINE", False),
            patch.object(document_builder, "KP_GENERATION_ENGINE", "html"),
            patch.object(document_builder, "render_html_kp_pdf", side_effect=fake_render),
        ):
            generated = document_builder.generate_documents_for_row(
                row,
                context,
                output_dir=output_dir,
                batch_docx_dir=batch_dir,
                templates_dir=templates_dir,
                document_mode="kp",
            )

        self.assertIn("kp_pdf", generated)
        self.assertIn("kp_final_pdf", generated)
        self.assertNotIn("kp", generated)
        self.assertNotIn("kp_final_docx", generated)
        self.assertTrue(generated["kp_pdf"].exists())

    @unittest.skipIf(document_builder is None, f"document_builder dependencies unavailable: {DOCUMENT_BUILDER_IMPORT_ERROR}")
    def test_pending_adaptive_template_never_falls_back_to_html(self) -> None:
        row = self._row()
        context = build_document_context(row, 101)
        output_dir = self.tmp_dir / "pending-output"
        batch_dir = self.tmp_dir / "pending-batch"
        templates_dir = self.tmp_dir / "pending-templates"
        templates_dir.mkdir()

        pending_state = {
            "latest_template_id": "latest-template",
            "active_template_id": None,
            "certification_status": "pending",
            "certification_error": "",
            "ready": False,
        }
        with (
            patch.object(document_builder, "KP_ADAPTIVE_TEMPLATE_ENGINE", True),
            patch.object(document_builder, "KP_GENERATION_ENGINE", "html"),
            patch(
                "src.generator.templates.store.AdaptiveTemplateStore.activation_state",
                return_value=pending_state,
            ),
            patch.object(document_builder, "render_html_kp_pdf") as html_renderer,
        ):
            with self.assertRaisesRegex(ValueError, "не активирован"):
                document_builder.generate_documents_for_row(
                    row,
                    context,
                    output_dir=output_dir,
                    batch_docx_dir=batch_dir,
                    templates_dir=templates_dir,
                    document_mode="kp",
                )

        html_renderer.assert_not_called()

    @unittest.skipIf(document_builder is None, f"document_builder dependencies unavailable: {DOCUMENT_BUILDER_IMPORT_ERROR}")
    def test_active_latest_template_has_priority_over_html(self) -> None:
        row = self._row()
        context = build_document_context(row, 101)
        output_dir = self.tmp_dir / "active-output"
        batch_dir = self.tmp_dir / "active-batch"
        templates_dir = self.tmp_dir / "active-templates"
        templates_dir.mkdir()

        active_state = {
            "latest_template_id": "latest-template",
            "active_template_id": "latest-template",
            "certification_status": "passed",
            "certification_error": "",
            "ready": True,
        }

        def fake_adaptive_render(_templates_dir, _context, output_path, **_kwargs):
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_bytes(b"%PDF adaptive")
            return output_path

        with (
            patch.object(document_builder, "KP_ADAPTIVE_TEMPLATE_ENGINE", True),
            patch.object(document_builder, "KP_GENERATION_ENGINE", "html"),
            patch(
                "src.generator.templates.store.AdaptiveTemplateStore.activation_state",
                return_value=active_state,
            ),
            patch(
                "src.generator.templates.renderer.render_active_template",
                side_effect=fake_adaptive_render,
            ) as adaptive_renderer,
            patch.object(document_builder, "render_html_kp_pdf") as html_renderer,
        ):
            generated = document_builder.generate_documents_for_row(
                row,
                context,
                output_dir=output_dir,
                batch_docx_dir=batch_dir,
                templates_dir=templates_dir,
                document_mode="kp",
            )

        self.assertTrue(generated["kp_pdf"].exists())
        adaptive_renderer.assert_called_once()
        html_renderer.assert_not_called()


if __name__ == "__main__":
    unittest.main()

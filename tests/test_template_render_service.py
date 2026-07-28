from __future__ import annotations

import unittest
import uuid
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from sqlalchemy import select

from src.campaigns import template_render_service, template_service
from src.campaigns.service import create_campaign, replace_recipients, validate_campaign_for_launch
from src.infra.db import session_scope
from src.infra.models import Campaign, CampaignRecipient, MailTemplate, TemplateVersion
from src.security.auth import Principal
from src.security.user_store import create_user
from src.web.v1_router import create_v1_router
from tests.bootstrap import bootstrap_test_runtime


class TemplateRenderServiceTests(unittest.TestCase):
    def setUp(self) -> None:
        bootstrap_test_runtime(reset_db=True)
        self.username = f"tr{uuid.uuid4().hex[:8]}"
        create_user(self.username, "Pass12345!")
        self.campaign = create_campaign(self.username, {"name": "Render docs"})
        self.campaign_id = self.campaign["id"]
        self.job_id = self.campaign["job_id"]
        replace_recipients(
            self.campaign_id,
            self.username,
            recipients=[{"company": "A", "contact_name": "A", "email": "a@example.com"}],
        )

        app = FastAPI()
        app.include_router(create_v1_router(check_auth=lambda: Principal(self.username, "t1", "user")))
        self.client = TestClient(app)

        pdf_buffer = BytesIO()
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        writer.write(pdf_buffer)
        uploaded = self.client.post(
            "/api/v1/templates/upload",
            data={"template_type": "document", "name": "Doc"},
            files={"file": ("doc.pdf", pdf_buffer.getvalue(), "application/pdf")},
        )
        self.assertEqual(uploaded.status_code, 200, uploaded.text)
        self.template_id = uploaded.json()["result"]["id"]
        self.client.patch(
            f"/api/v1/templates/{self.template_id}",
            json={"is_template": True},
        )

        with session_scope() as session:
            camp = session.get(Campaign, self.campaign_id)
            assert camp is not None
            camp.draft_payload = {
                "email_chain": {
                    "version": 1,
                    "root_node_id": "n1",
                    "nodes": [
                        {
                            "id": "n1",
                            "name": "Root",
                            "kind": "email",
                            "email_template_id": None,
                            "document_template_ids": [self.template_id],
                        }
                    ],
                    "edges": [],
                }
            }
            session.flush()
            self.recipient_id = session.scalar(
                select(CampaignRecipient.id).where(CampaignRecipient.campaign_id == self.campaign_id)
            )
        assert self.recipient_id is not None

    def test_launch_validation_uses_persisted_pdf_text(self) -> None:
        with session_scope() as session:
            template = session.get(MailTemplate, self.template_id)
            assert template is not None and template.active_version_id
            version = session.get(TemplateVersion, template.active_version_id)
            assert version is not None
            self.assertEqual(version.source_text, "")
            self.assertEqual(version.text_extraction_status, "ready")
            self.assertEqual(len(str(version.source_sha256 or "")), 64)

        with patch.object(
            template_service,
            "_file_text",
            side_effect=AssertionError("launch validation must not parse the source PDF"),
        ):
            result = validate_campaign_for_launch(self.campaign_id, self.username)

        self.assertIn("errors", result)

    def test_launch_validation_reports_pending_legacy_pdf_cache(self) -> None:
        with session_scope() as session:
            template = session.get(MailTemplate, self.template_id)
            assert template is not None and template.active_version_id
            version = session.get(TemplateVersion, template.active_version_id)
            assert version is not None
            version.source_text = None
            version.text_extraction_status = "pending"
            session.flush()

        with patch.object(
            template_service,
            "_file_text",
            side_effect=AssertionError("launch validation must not parse a pending source PDF"),
        ):
            result = validate_campaign_for_launch(self.campaign_id, self.username)

        self.assertTrue(any("\u043e\u0431\u0440\u0430\u0431\u0430\u0442\u044b\u0432\u0430\u0435\u0442\u0441\u044f" in error for error in result["errors"]))

    def test_background_task_backfills_legacy_pdf_text_once(self) -> None:
        with session_scope() as session:
            template = session.get(MailTemplate, self.template_id)
            assert template is not None and template.active_version_id
            version_id = template.active_version_id
            version = session.get(TemplateVersion, version_id)
            assert version is not None
            version.source_text = None
            version.source_sha256 = None
            version.text_extraction_status = "pending"
            session.flush()

        from src.campaigns.template_text_cache_service import run_template_text_extraction

        run_template_text_extraction({"version_id": version_id})

        with session_scope() as session:
            version = session.get(TemplateVersion, version_id)
            assert version is not None
            self.assertEqual(version.source_text, "")
            self.assertEqual(version.text_extraction_status, "ready")
            self.assertEqual(len(str(version.source_sha256 or "")), 64)

    def test_pending_pdf_backfill_tasks_are_deduplicated(self) -> None:
        with session_scope() as session:
            template = session.get(MailTemplate, self.template_id)
            assert template is not None and template.active_version_id
            version = session.get(TemplateVersion, template.active_version_id)
            assert version is not None
            version.source_text = None
            version.text_extraction_status = "pending"
            session.flush()

        from src.campaigns.template_text_cache_service import enqueue_pending_template_text_extractions

        self.assertEqual(enqueue_pending_template_text_extractions(), 1)
        self.assertEqual(enqueue_pending_template_text_extractions(), 0)

    def test_static_attachment_when_not_template(self) -> None:
        with session_scope() as session:
            tmpl = session.get(MailTemplate, self.template_id)
            assert tmpl is not None
            tmpl.is_template = False
            session.flush()
            recipient = session.get(CampaignRecipient, int(self.recipient_id))
            campaign = session.get(Campaign, self.campaign_id)
            assert recipient is not None and campaign is not None
            session.expunge(recipient)
            session.expunge(campaign)

        filename, data = template_render_service.render_document_template_for_recipient(
            template_id=self.template_id,
            recipient=recipient,
            campaign=campaign,
            job_id=self.job_id,
        )
        self.assertTrue(filename)
        self.assertTrue(data)

    def test_personalized_pdf_is_cached(self) -> None:
        with session_scope() as session:
            recipient = session.get(CampaignRecipient, int(self.recipient_id))
            campaign = session.get(Campaign, self.campaign_id)
            assert recipient is not None and campaign is not None
            session.expunge(recipient)
            session.expunge(campaign)

        fetched = self.client.get(f"/api/v1/templates/{self.template_id}")
        self.assertEqual(fetched.status_code, 200, fetched.text)
        delivery_name = fetched.json()["result"]["version"]["rendered_pdf_filename"]
        self.assertTrue(delivery_name)

        first_filename, first_data = template_render_service.render_document_template_for_recipient(
            template_id=self.template_id,
            recipient=recipient,
            campaign=campaign,
            job_id=self.job_id,
        )
        self.assertEqual(first_filename, delivery_name)
        self.assertNotEqual(first_filename, f"{self.template_id}.pdf")
        self.assertTrue(first_data)

        cache_path = template_render_service._cache_path(
            self.job_id,
            int(self.recipient_id),
            self.template_id,
            ".pdf",
        )
        self.assertTrue(cache_path.exists())
        self.assertEqual(cache_path.name, f"{self.template_id}.pdf")

        with patch(
            "src.campaigns.template_render_service.get_bytes",
            side_effect=AssertionError("should use cache"),
        ):
            cached_filename, cached_data = template_render_service.render_document_template_for_recipient(
                template_id=self.template_id,
                recipient=recipient,
                campaign=campaign,
                job_id=self.job_id,
            )
        self.assertEqual(cached_filename, delivery_name)
        self.assertNotEqual(cached_filename, f"{self.template_id}.pdf")
        self.assertTrue(cached_data)

    def test_pre_generate_batch_creates_manifest(self) -> None:
        result = template_render_service.pre_generate_batch_templates(
            campaign_id=self.campaign_id,
            recipient_ids=[int(self.recipient_id)],
        )
        self.assertIn("template_ids", result)
        manifest = template_render_service.load_manifest(self.job_id)
        self.assertEqual(manifest.get("template_ids"), [self.template_id])

    def test_launch_schedules_pre_generate_task(self) -> None:
        from cryptography.fernet import Fernet
        from src.generator.delivery.smtp_mailboxes import create_mailbox
        from src.infra.models import BackgroundTask
        from src.utils.config import settings

        with patch.object(settings, "smtp_credentials_key", Fernet.generate_key().decode("ascii")):
            mailbox = create_mailbox(
                owner_username=self.username,
                provider="custom",
                email="sender@example.com",
                password="x",
                host="mailpit",
                port=1025,
                use_ssl=False,
                use_starttls=False,
                make_default=True,
            )
            self.client.patch(
                f"/api/v1/campaigns/{self.campaign_id}",
                json={
                    "smtp_mailbox_id": mailbox["id"],
                    "transport": "smtp",
                    "mail_subject": "Hello",
                },
            )
            self.client.put(
                f"/api/v1/campaigns/{self.campaign_id}/schedule",
                json={"send_immediately": True, "batch_size": 1, "interval_seconds": 60},
            )
            launched = self.client.post(f"/api/v1/campaigns/{self.campaign_id}/launch?force_now=true")
            self.assertEqual(launched.status_code, 200, launched.text)

        from sqlalchemy import func

        with session_scope() as session:
            count = session.scalar(
                select(func.count())
                .select_from(BackgroundTask)
                .where(
                    BackgroundTask.task_type == "campaign_pre_generate",
                    BackgroundTask.payload["campaign_id"].as_string() == self.campaign_id,
                )
            )
        self.assertGreaterEqual(int(count or 0), 1)

    def test_build_replacements_merges_kp_specific_tokens(self) -> None:
        from pathlib import Path

        from src.generator.generation.transforms import build_document_context

        row = {
            "ID": "1",
            "SUB_RF": "Иркутская область",
            "MUN_R_NAME": "Усть-Кутский муниципальный район",
            "MUN_NAME": "Нийское сельское поселение",
            "ADM_NAME": "Администрация Нийского сельского поселения",
            "HEAD_FIO": "Иванов Иван Иванович",
        }
        context = build_document_context(row, 101)
        string_context = {key: str(value) for key, value in context.items() if value is not None}
        pairs = dict(
            template_render_service._build_replacements(
                string_context,
                "Уважаемый (ая) HEAD_FIO  ! MUN_R_NAME SUB_RF",
                source_path=Path("КП_test.docx"),
            )
        )
        self.assertIn("Уважаемый (ая) HEAD_FIO  !", pairs)
        self.assertIn("MUN_R_NAME SUB_RF", pairs)
        self.assertNotIn("HEAD_FIO", pairs.get("Уважаемый (ая) HEAD_FIO  !", ""))

    def test_convert_docx_to_pdf_uses_delivery_pipeline(self) -> None:
        from pathlib import Path
        from tempfile import TemporaryDirectory

        with TemporaryDirectory(prefix="template-render-test-") as temp_dir:
            root = Path(temp_dir)
            docx_path = root / "rendered.docx"
            pdf_path = root / "rendered.pdf"
            docx_path.write_bytes(b"docx")
            with patch(
                "src.generator.generation.template_preview.convert_docx_to_delivery_pdf",
                return_value=pdf_path,
            ) as convert_mock:
                result = template_render_service._convert_docx_to_pdf(docx_path, pdf_path, file_kind="kp")
            convert_mock.assert_called_once_with(
                docx_path,
                pdf_path,
                file_kind="kp",
                template_docx=docx_path,
            )
            self.assertEqual(result, pdf_path)

    @patch("src.generator.generation.pdf_safe.is_kp_docx", return_value=True)
    @patch("src.campaigns.template_render_service._convert_kp_docx_to_pdf")
    @patch("src.campaigns.template_render_service.render_docx")
    @patch("src.campaigns.template_service._build_document_pdf_artifact")
    def test_docx_render_returns_stored_delivery_filename(
        self,
        mock_build_pdf,
        mock_render_docx,
        mock_convert_kp,
        _mock_is_kp_docx,
    ) -> None:
        mock_build_pdf.return_value = (b"%PDF-1.4 delivery copy", "ignored.pdf")

        def _write_pdf(docx_path: Path, pdf_path: Path, **kwargs: object) -> None:
            pdf_path.write_bytes(b"%PDF-1.4 rendered")

        mock_convert_kp.side_effect = _write_pdf

        def _fake_render(source: Path, replacements: list[tuple[str, str]], output: Path, context: dict) -> Path:
            output.write_bytes(source.read_bytes())
            return output

        mock_render_docx.side_effect = _fake_render

        source_docx = Document()
        source_docx.add_paragraph("Коммерческое предложение")
        source_docx.add_paragraph(
            "на разработку схемы территориального планирования муниципального образования"
        )
        payload = BytesIO()
        source_docx.save(payload)
        uploaded = self.client.post(
            "/api/v1/templates/upload",
            data={"template_type": "document", "name": "Docx render name"},
            files={
                "file": (
                    "КП_СТП_районы (1) (1).docx",
                    payload.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        self.assertEqual(uploaded.status_code, 200, uploaded.text)
        template_id = uploaded.json()["result"]["id"]
        delivery_name = uploaded.json()["result"]["version"]["rendered_pdf_filename"]
        self.assertEqual(delivery_name, "КП_СТП_районы.pdf")
        updated = self.client.patch(
            f"/api/v1/templates/{template_id}",
            json={
                "is_template": True,
                "attachment_output_format": "pdf",
            },
        )
        self.assertEqual(updated.status_code, 200, updated.text)

        with session_scope() as session:
            recipient = session.get(CampaignRecipient, int(self.recipient_id))
            campaign = session.get(Campaign, self.campaign_id)
            assert recipient is not None and campaign is not None
            session.expunge(recipient)
            session.expunge(campaign)

        filename, data = template_render_service.render_document_template_for_recipient(
            template_id=template_id,
            recipient=recipient,
            campaign=campaign,
            job_id=self.job_id,
            force=True,
        )
        self.assertEqual(filename, delivery_name)
        self.assertNotEqual(filename, f"{template_id}.pdf")
        self.assertTrue(data.startswith(b"%PDF"))

    def test_docx_attachment_keeps_original_format_by_default(self) -> None:
        source_docx = Document()
        source_docx.add_paragraph("Документ без конвертации")
        payload = BytesIO()
        source_docx.save(payload)
        uploaded = self.client.post(
            "/api/v1/templates/upload",
            data={"template_type": "document", "name": "Original DOCX"},
            files={
                "file": (
                    "original.docx",
                    payload.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        self.assertEqual(uploaded.status_code, 200, uploaded.text)
        item = uploaded.json()["result"]
        self.assertEqual(item["attachment_output_format"], "original")

        with session_scope() as session:
            recipient = session.get(CampaignRecipient, int(self.recipient_id))
            campaign = session.get(Campaign, self.campaign_id)
            assert recipient is not None and campaign is not None
            session.expunge(recipient)
            session.expunge(campaign)

        filename, data = template_render_service.render_document_template_for_recipient(
            template_id=item["id"],
            recipient=recipient,
            campaign=campaign,
            job_id=self.job_id,
            force=True,
        )
        self.assertEqual(filename, "original.docx")
        self.assertTrue(data.startswith(b"PK"))

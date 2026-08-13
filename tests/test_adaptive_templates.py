from __future__ import annotations

import json
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt
from pypdf import PdfReader, PdfWriter
from pypdf.generic import ArrayObject, DictionaryObject, NameObject, NumberObject, TextStringObject

from src.generator.templates.adapters.docx import DocxTemplateAdapter
from src.generator.templates.adapters.html import HtmlTemplateAdapter
from src.generator.templates.adapters.pdf import PdfAcroFormAdapter, _canonical_form_field_name
from src.generator.templates.auto_compiler import _find_pdf_regions, _font_path, _table_cell_right_boundary
from src.generator.templates.base import TemplateCompileError
from src.generator.templates.certification import _pdf_layout_issues
from src.generator.templates.compiler import compile_template
from src.generator.templates.models import CertificationResult
from src.generator.templates.pdf_semantic import build_semantic_pdf_html
from src.generator.templates.protocol import TemplateProtocolError, find_fields, replace_placeholders
from src.generator.templates.store import AdaptiveTemplateStore


def _docx_text(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def test_protocol_requires_strict_fields_and_complete_context() -> None:
    assert find_fields("To {{ ADM_NAME }} at {{DATE}}") == ("ADM_NAME", "DATE")
    assert replace_placeholders("{{ADM_NAME}} / {{DATE}}", {"adm_name": "A", "date": "D"}) == "A / D"
    with pytest.raises(TemplateProtocolError, match="Missing template values"):
        replace_placeholders("{{ADM_NAME}} {{DATE}}", {"ADM_NAME": "A"})
    with pytest.raises(TemplateProtocolError, match="Invalid placeholder syntax"):
        replace_placeholders("{{bad field}}", {})


def test_docx_split_run_placeholder_is_detected_and_replaced(tmp_path: Path) -> None:
    source = tmp_path / "template.docx"
    output = tmp_path / "rendered.docx"
    document = Document()
    paragraph = document.add_paragraph()
    first = paragraph.add_run("Кому: {{ADM")
    first.bold = True
    first.font.size = Pt(11)
    paragraph.add_run("_NAME}}")
    document.save(source)

    adapter = DocxTemplateAdapter()
    occurrences, capabilities, _ = adapter.inspect(source)
    assert [item.field_name for item in occurrences] == ["ADM_NAME"]
    assert capabilities["preserves_ooxml"] is True

    adapter.render(source, {"ADM_NAME": "Администрация города"}, output)
    xml = _docx_text(output)
    assert "Администрация города" in xml
    assert "{{ADM" not in xml
    rendered = Document(output)
    assert rendered.paragraphs[0].runs[0].bold is True
    assert rendered.paragraphs[0].runs[0].font.size.pt == pytest.approx(11)


def test_html_adapter_replaces_fields_and_blocks_active_content(tmp_path: Path) -> None:
    source = tmp_path / "template.html"
    source.write_text("<html><body><div style='width:100px'>{{ADM_NAME}}</div></body></html>", encoding="utf-8")
    output = tmp_path / "rendered.html"
    adapter = HtmlTemplateAdapter()
    adapter.render(source, {"ADM_NAME": "Администрация"}, output)
    rendered = output.read_text(encoding="utf-8")
    assert 'data-adaptive-field="ADM_NAME"' in rendered
    assert "Администрация" in rendered
    assert "dataset.adaptiveFontSize" in rendered

    unsafe = tmp_path / "unsafe.html"
    unsafe.write_text("<script>alert(1)</script>{{ADM_NAME}}", encoding="utf-8")
    with pytest.raises(TemplateCompileError, match="blocked active elements"):
        adapter.inspect(unsafe)


def test_compiler_versions_package_and_activation_requires_certification(tmp_path: Path) -> None:
    templates_dir = tmp_path / "templates"
    source = tmp_path / "template.html"
    source.write_text("<html><body>{{ADM_NAME}}</body></html>", encoding="utf-8")
    package = compile_template(source, templates_dir)
    store = AdaptiveTemplateStore(templates_dir)

    assert store.load_package(package.template_id) == package
    assert store.latest_template_id() == package.template_id
    assert store.source_path(package.template_id).read_bytes() == source.read_bytes()
    with pytest.raises(ValueError, match="certified"):
        store.activate(package.template_id)

    certification = CertificationResult(
        template_id=package.template_id,
        status="passed",
        created_at="2026-07-13T12:00:00+00:00",
        checks=({"name": "test", "status": "passed"},),
    )
    store.save_certification(certification)
    store.activate(package.template_id)
    assert store.load_active() == package
    active_payload = json.loads(store.active_path.read_text(encoding="utf-8"))
    assert active_payload["template_id"] == package.template_id

def test_pdf_acroform_is_autosized_and_flattened(tmp_path: Path) -> None:
    source = tmp_path / "form.pdf"
    output = tmp_path / "filled.pdf"
    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=300)
    font_ref = writer._add_object(
        DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
    )
    fields = ArrayObject()
    acro_form = DictionaryObject(
        {
            NameObject("/Fields"): fields,
            NameObject("/DR"): DictionaryObject(
                {NameObject("/Font"): DictionaryObject({NameObject("/Helv"): font_ref})}
            ),
            NameObject("/DA"): TextStringObject("/Helv 0 Tf 0 g"),
        }
    )
    writer._root_object[NameObject("/AcroForm")] = writer._add_object(acro_form)
    widget = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Annot"),
            NameObject("/Subtype"): NameObject("/Widget"),
            NameObject("/FT"): NameObject("/Tx"),
            NameObject("/T"): TextStringObject("ADM_NAME"),
            NameObject("/Rect"): ArrayObject(
                [NumberObject(40), NumberObject(200), NumberObject(260), NumberObject(235)]
            ),
            NameObject("/DA"): TextStringObject("/Helv 0 Tf 0 g"),
            NameObject("/P"): page.indirect_reference,
        }
    )
    widget_ref = writer._add_object(widget)
    fields.append(widget_ref)
    page[NameObject("/Annots")] = ArrayObject([widget_ref])
    with source.open("wb") as handle:
        writer.write(handle)

    adapter = PdfAcroFormAdapter()
    assert adapter.probe(source) is True
    occurrences, capabilities, _ = adapter.inspect(source)
    assert [item.field_name for item in occurrences] == ["ADM_NAME"]
    assert capabilities["flattened"] is True

    adapter.render(
        source,
        {"ADM_NAME": "Very Long Municipal Administration Name"},
        output,
    )
    reader = PdfReader(str(output))
    assert len(reader.pages) == 1
    rendered_text = (reader.pages[0].extract_text() or "").replace("\x00", "")
    assert "Municipal Administration" in rendered_text
    assert output.read_bytes().startswith(b"%PDF")


def test_compiler_auto_discovers_legacy_docx_markers(tmp_path: Path) -> None:
    source = tmp_path / "legacy.docx"
    document = Document()
    document.add_paragraph("№ 101-КП от 12.05.2026")
    document.add_paragraph("ADM_NAME")
    document.add_paragraph("Работы для MUN_R_NAME SUB_RF")
    document.save(source)

    package = compile_template(source, tmp_path / "templates")

    assert package.adapter == "docx-ooxml-v1"
    assert package.fields == ("ADM_NAME_1", "DATE", "MUN_R_SCOPE_FRAGMENT", "OUTGOING_NUMBER")
    assert package.capabilities["auto_discovery"]["mode"] == "automatic"
    stored_source = AdaptiveTemplateStore(tmp_path / "templates").source_path(package.template_id)
    stored_xml = _docx_text(stored_source)
    assert "{{ADM_NAME_1}}" in stored_xml
    assert "{{MUN_R_SCOPE_FRAGMENT}}" in stored_xml
    assert (stored_source.parent / "original.docx").exists()


def test_pdf_auto_detection_uses_outgoing_line_not_company_details() -> None:
    fitz = pytest.importorskip("fitz")

    class SearchablePage:
        rect = fitz.Rect(0, 0, 595, 842)

        def get_text(self, kind: str) -> str:
            assert kind == "text"
            return "ИНН 101 5038\n№ 101-КП от 12.05.2026\n"

        def search_for(self, value: str) -> list:
            if value == "101":
                return [fitz.Rect(330, 70, 350, 82), fitz.Rect(58, 160, 80, 173)]
            if value == "12.05.2026":
                return [fitz.Rect(116, 160, 178, 173)]
            return []

    document = [SearchablePage()]

    regions = _find_pdf_regions(document, {})
    number_region = next(item for item in regions if item.field_name == "OUTGOING_NUMBER")

    assert number_region.source_box[1] > 150
    assert number_region.field_box[2] <= number_region.source_box[2]


def test_pdf_scope_field_stops_at_table_column_boundary() -> None:
    fitz = pytest.importorskip("fitz")
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    page.draw_line((474, 370), (474, 473), width=1)
    page.draw_line((561, 370), (561, 473), width=1)

    boundary = _table_cell_right_boundary(page, fitz.Rect(245, 420, 425, 446))

    assert boundary == pytest.approx(474)

def test_pdf_certification_detects_image_over_validity_date(tmp_path: Path) -> None:
    fitz = pytest.importorskip("fitz")
    pdf_path = tmp_path / "overlap.pdf"
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    page.insert_text(
        (50, 300),
        "Срок действия коммерческого предложения: до 31.07.2026.",
        fontname="validity-font",
        fontfile=str(_font_path()),
        fontsize=11,
    )
    pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 50, 50), False)
    pixmap.clear_with(0x3366CC)
    page.insert_image(fitz.Rect(250, 285, 400, 430), stream=pixmap.tobytes("png"))
    document.save(pdf_path)
    document.close()

    issues = _pdf_layout_issues(pdf_path)

    assert issues[0]["type"] == "image_overlaps_validity_date"

def test_pdf_background_sampling_preserves_white() -> None:
    fitz = pytest.importorskip("fitz")
    from src.generator.templates.auto_compiler import _sample_background_color

    document = fitz.open()
    page = document.new_page(width=200, height=100)
    color = _sample_background_color(page, fitz.Rect(20, 20, 100, 40))

    assert color == (1.0, 1.0, 1.0)


def test_pdf_form_instance_names_map_to_one_context_field() -> None:
    assert _canonical_form_field_name("MUN_R_SCOPE_FRAGMENT") == "MUN_R_SCOPE_FRAGMENT"
    assert _canonical_form_field_name("MUN_R_SCOPE_FRAGMENT__2") == "MUN_R_SCOPE_FRAGMENT"

def test_plain_pdf_is_rebuilt_as_one_semantic_html_flow(tmp_path: Path) -> None:
    fitz = pytest.importorskip("fitz")
    source = tmp_path / "ordinary.pdf"
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    font_path = str(_font_path())

    def put(x: float, y: float, text: str, size: float = 10) -> None:
        page.insert_text((x, y), text, fontname="semantic-test", fontfile=font_path, fontsize=size)

    work_title = "разработке проекта местных нормативов градостроительного проектирования"
    scope = "Тестового муниципального района Московской области"
    put(330, 55, "Общество с ограниченной ответственностью")
    put(45, 165, "№ 101-КП от 12.05.2026", 11)
    put(330, 165, "Администрации Тестового района", 11)
    put(180, 210, "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ", 12)
    put(45, 240, f"ООО «Компания» предлагает выполнить работы по {work_title} {scope}.")
    put(45, 280, "В стоимость работ включено консультационное сопровождение Заказчика.")
    put(45, 340, f"Выполнение работ по {work_title} {scope}.")
    put(470, 340, "150 000,00")
    put(45, 410, "Стоимость выполнения работ составляет 150 000 рублей 00 копеек.")
    put(45, 455, "ООО «Компания» специализируется на комплексной разработке документов.")
    put(45, 510, "Просим Вас ознакомиться с проектом договора и направить обратную связь.")
    put(45, 590, "Срок действия коммерческого предложения: до 31.07.2026.")
    put(45, 650, "Исп. Иванов Иван\nтел. +7 921 000-00-00\ntest@example.ru")
    put(410, 660, "И.И. Иванов")
    document.save(source)
    document.close()

    html, report = build_semantic_pdf_html(
        source,
        {"WORK_TITLE": work_title},
        field_names=("OUTGOING_NUMBER", "DATE", "ADM_NAME_1", "WORK_TITLE", "MUN_R_SCOPE_FRAGMENT"),
    )

    assert report["mode"] == "semantic_html"
    assert html.count("{{WORK_TITLE}}") == 2
    assert html.count("{{MUN_R_SCOPE_FRAGMENT}}") == 2
    assert '<strong class="work-phrase">{{WORK_TITLE}} {{MUN_R_SCOPE_FRAGMENT}}</strong>' in html
    assert "contact-row" in html
    assert "AcroForm" not in html

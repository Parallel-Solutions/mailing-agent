import json
import shutil
import unittest
from pathlib import Path
from types import SimpleNamespace

from docx import Document

from src.generator.philologist_agent import (
    _auto_fix_docx,
    _collect_docx_snapshot,
    _react_decide_fix_strategy,
    _react_decide_next_action,
    _run_docx_react_loop,
    _verify_safe_fixes,
)
from src.generator.philologist_decisions import AUTO_FIX, QUARANTINE
from src.generator.philologist_tools import PhilologistToolRunner


def _ru(value: str) -> str:
    return value.encode("ascii").decode("unicode_escape")


def _fake_client(payload: dict) -> object:
    class FakeCompletions:
        def create(self, **kwargs):
            message = SimpleNamespace(content=json.dumps(payload, ensure_ascii=False))
            return SimpleNamespace(choices=[SimpleNamespace(message=message)])

    return SimpleNamespace(chat=SimpleNamespace(completions=FakeCompletions()))


class PhilologistAgenticTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp_dir = Path("tmp_test_philologist_agentic")
        self.tmp_dir.mkdir(exist_ok=True)

    def tearDown(self) -> None:
        if self.tmp_dir.exists():
            shutil.rmtree(self.tmp_dir)

    def _docx(self, name: str, text: str) -> Path:
        path = self.tmp_dir / name
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        doc.save(path)
        return path

    def test_react_tool_controller_rejects_unknown_tool(self) -> None:
        client = _fake_client(
            {
                "thought": "try unsafe action",
                "action": "delete_file",
                "reason": "not allowed",
            }
        )
        decision = _react_decide_next_action(
            client=client,
            context={"document": "x.docx", "review_done": False},
            trace=[],
        )

        self.assertEqual(decision["action"], "review_docx")
        self.assertEqual(decision["source"], "fallback_invalid_llm_action")
        self.assertEqual(decision["llm_action"], "delete_file")

    def test_fix_strategy_cannot_escalate_quarantine_to_auto_fix(self) -> None:
        client = _fake_client(
            {
                "thought": "force auto",
                "action": AUTO_FIX,
                "reason": "unsafe override",
                "confidence": 1.0,
            }
        )
        base_decision = {
            "action": QUARANTINE,
            "reason": "AI offered a rewrite without exact fragment.",
            "confidence": 0.9,
            "source": "ai",
        }
        decision = _react_decide_fix_strategy(
            client=client,
            issue={"source": "ai", "issue": "rewrite", "fragment": "", "suggestion": "new text"},
            current_text="old text",
            base_decision=base_decision,
            rag={"support_score": 10, "recommendation": "candidate_for_human_approval", "rules": []},
        )

        self.assertEqual(decision["action"], QUARANTINE)
        self.assertEqual(decision["source"], "fallback_invalid_strategy")
        self.assertEqual(decision["llm_action"], AUTO_FIX)

    def test_safe_fix_is_verified_without_style_warning(self) -> None:
        before_text = _ru(
            "\\u0422\\u0435\\u0441\\u0442  \\u0434\\u0432\\u043e\\u0439\\u043d\\u043e\\u0433\\u043e "
            "\\u043f\\u0440\\u043e\\u0431\\u0435\\u043b\\u0430"
        )
        after_text = _ru(
            "\\u0422\\u0435\\u0441\\u0442 \\u0434\\u0432\\u043e\\u0439\\u043d\\u043e\\u0433\\u043e "
            "\\u043f\\u0440\\u043e\\u0431\\u0435\\u043b\\u0430"
        )
        path = self._docx("safe_fix.docx", before_text)
        before_snapshot = _collect_docx_snapshot(path)
        review = {
            "issues": [
                {
                    "source": "local",
                    "location": "paragraph:1",
                    "fragment": before_text,
                    "issue": _ru(
                        "\\u041e\\u0431\\u043d\\u0430\\u0440\\u0443\\u0436\\u0435\\u043d\\u044b "
                        "\\u0434\\u0432\\u043e\\u0439\\u043d\\u044b\\u0435 \\u043f\\u0440\\u043e\\u0431\\u0435\\u043b\\u044b."
                    ),
                    "suggestion": after_text,
                    "severity": "info",
                }
            ],
            "issue_count": 1,
        }

        fix_result = _auto_fix_docx(path, review, client=None, tool_runner=PhilologistToolRunner())
        verification = _verify_safe_fixes(path, before_snapshot, fix_result)

        self.assertEqual(fix_result["applied_fix_count"], 1)
        self.assertTrue(verification["verified"])
        self.assertEqual(verification["warning_count"], 0)

    def test_apply_safe_fixes_records_quarantine_tool_for_unsafe_ai_rewrite(self) -> None:
        path = self._docx("unsafe_ai.docx", "old text")
        runner = PhilologistToolRunner()
        review = {
            "issues": [
                {
                    "source": "ai",
                    "location": "paragraph:1",
                    "fragment": "",
                    "issue": "AI rewrite without fragment",
                    "suggestion": "new paragraph text",
                    "severity": "error",
                }
            ],
            "issue_count": 1,
        }

        result = _auto_fix_docx(path, review, client=None, tool_runner=runner)
        tool_names = [record["name"] for record in runner.as_state()]

        self.assertEqual(result["applied_fix_count"], 0)
        self.assertEqual(result["skipped_fix_count"], 1)
        self.assertIn("lookup_rag_rule", tool_names)
        self.assertIn("decide_fix_strategy", tool_names)
        self.assertIn("quarantine_issue", tool_names)

    def test_document_react_loop_records_observations_and_finishes(self) -> None:
        path = self._docx("react_loop.docx", "clean text")
        runner = PhilologistToolRunner()

        result = _run_docx_react_loop(
            docx_path=path,
            ai_enabled=False,
            tool_runner=runner,
            client=None,
        )
        actions = [step["action"] for step in result["react_trace"]]

        self.assertEqual(
            actions,
            ["review_docx", "snapshot_docx", "apply_safe_fixes", "verify_safe_fixes", "finish_document"],
        )
        self.assertTrue(all(step.get("observation") for step in result["react_trace"]))


if __name__ == "__main__":
    unittest.main()

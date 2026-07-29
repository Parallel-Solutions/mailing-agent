from __future__ import annotations

import unittest
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from docx import Document

from src.campaigns import font_service
from src.generator.generation import pdf_converter


class FontServiceTests(unittest.TestCase):
    def test_analyze_docx_reports_family_style_and_weight(self) -> None:
        document = Document()
        run = document.add_paragraph().add_run("Тестовый текст")
        run.font.name = "Tahoma"
        run.bold = True
        run.italic = True
        payload = BytesIO()
        document.save(payload)

        requirements = font_service.analyze_docx_fonts(payload.getvalue())

        tahoma = next(
            item
            for item in requirements
            if item["family_normalized"] == font_service.normalize_font_family("Tahoma")
        )
        self.assertEqual(tahoma["weight"], 700)
        self.assertTrue(tahoma["italic"])
        self.assertIn("word/document.xml", tahoma["source_parts"])

    def test_rejects_non_font_payload(self) -> None:
        with self.assertRaises(ValueError):
            font_service.inspect_font_bytes("not-a-font.ttf", b"not a font")

    def test_reads_metadata_from_installed_font(self) -> None:
        candidates = (
            Path("C:/Windows/Fonts/arial.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
        )
        font_path = next((path for path in candidates if path.is_file()), None)
        if font_path is None:
            self.skipTest("No known system TTF is installed")

        metadata = font_service.inspect_font_bytes(font_path.name, font_path.read_bytes())

        self.assertTrue(metadata.family)
        self.assertTrue(metadata.family_normalized)
        self.assertGreater(metadata.glyph_coverage["glyph_count"], 100)

    def test_trusted_download_rejects_unapproved_host(self) -> None:
        with self.assertRaises(ValueError):
            font_service._trusted_download(  # noqa: SLF001
                object(),  # type: ignore[arg-type]
                "https://example.com/font.ttf",
                max_bytes=1024,
            )

    @patch.object(pdf_converter, "find_soffice", return_value="/usr/bin/soffice")
    def test_uploaded_fonts_prefer_local_libreoffice(self, _find_soffice) -> None:
        self.assertEqual(
            pdf_converter._backend_sequence(prefer_local=True)[:2],  # noqa: SLF001
            ["libreoffice", "gotenberg"],
        )


if __name__ == "__main__":
    unittest.main()

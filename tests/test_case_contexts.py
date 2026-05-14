import unittest
from pathlib import Path
import shutil

from docx import Document

from src.generator.case_engine import build_inflected_fields_with_trace
from src.generator.inflection.inflect import inflect_fio_dative
from src.generator.inflection.inflect import inflect_fio_genitive
from src.generator.philologist.philologist_agent import (
    _apply_inflection_context_corrections,
    _inflection_context_is_suspicious,
)


class CaseContextTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp_dir = Path("tmp_test_case_contexts")
        self.tmp_dir.mkdir(exist_ok=True)

    def tearDown(self) -> None:
        if self.tmp_dir.exists():
            shutil.rmtree(self.tmp_dir)

    def test_feminine_indeclinable_fio_parts_are_preserved(self) -> None:
        self.assertEqual(
            inflect_fio_genitive("Гучетль Аминет Валерьевна").value,
            "Гучетль Аминет Валерьевны",
        )
        self.assertEqual(
            inflect_fio_dative("Гучетль Аминет Валерьевна").value,
            "Гучетль Аминет Валерьевне",
        )

    def test_fio_overrides_keep_correct_case_forms(self) -> None:
        female_fields, _ = build_inflected_fields_with_trace({"HEAD_FIO": "Гучетль Аминет Валерьевна"})
        self.assertEqual(female_fields["HEAD_FIO_1"], "Гучетль Аминет Валерьевны")
        self.assertEqual(female_fields["HEAD_FIO_2"], "Гучетль Аминет Валерьевне")

        male_fields, _ = build_inflected_fields_with_trace({"HEAD_FIO": "Пшизов Муртаз Сальбиевич"})
        self.assertEqual(male_fields["HEAD_FIO_1"], "Пшизова Муртаза Сальбиевича")
        self.assertEqual(male_fields["HEAD_FIO_2"], "Пшизову Муртазу Сальбиевичу")

        adyghe_fields, _ = build_inflected_fields_with_trace({"HEAD_FIO": "Емтыль Асланбий Рамазанович"})
        self.assertEqual(adyghe_fields["HEAD_FIO_1"], "Емтыля Асланбия Рамазановича")
        self.assertEqual(adyghe_fields["HEAD_FIO_2"], "Емтылю Асланбию Рамазановичу")

        indeclinable_surname_fields, _ = build_inflected_fields_with_trace({"HEAD_FIO": "Евтых Мурат Аскерович"})
        self.assertEqual(indeclinable_surname_fields["HEAD_FIO_1"], "Евтых Мурата Аскеровича")
        self.assertEqual(indeclinable_surname_fields["HEAD_FIO_2"], "Евтых Мурату Аскеровичу")

        rare_name_fields, _ = build_inflected_fields_with_trace({"HEAD_FIO": "Тадыев Судур Борисович"})
        self.assertEqual(rare_name_fields["HEAD_FIO_1"], "Тадыева Судура Борисовича")
        self.assertEqual(rare_name_fields["HEAD_FIO_2"], "Тадыеву Судуру Борисовичу")

        indeclinable_rare_name_fields, _ = build_inflected_fields_with_trace(
            {"HEAD_FIO": "Ундулганов Адучы Станиславович"}
        )
        self.assertEqual(indeclinable_rare_name_fields["HEAD_FIO_1"], "Ундулганова Адучы Станиславовича")
        self.assertEqual(indeclinable_rare_name_fields["HEAD_FIO_2"], "Ундулганову Адучы Станиславовичу")

        female_consonant_surname_fields, _ = build_inflected_fields_with_trace(
            {"HEAD_FIO": "Ткач Урсула Викторовна"}
        )
        self.assertEqual(female_consonant_surname_fields["HEAD_FIO_1"], "Ткач Урсулы Викторовны")
        self.assertEqual(female_consonant_surname_fields["HEAD_FIO_2"], "Ткач Урсуле Викторовне")

        plural_like_surname_fields, _ = build_inflected_fields_with_trace(
            {"HEAD_FIO": "Хоробрых Елена Николаевна"}
        )
        self.assertEqual(plural_like_surname_fields["HEAD_FIO_1"], "Хоробрых Елены Николаевны")
        self.assertEqual(plural_like_surname_fields["HEAD_FIO_2"], "Хоробрых Елене Николаевне")

        feminine_indeclinable_name_fields, _ = build_inflected_fields_with_trace(
            {"HEAD_FIO": "Сынгизова Гульюзум Хадисовна"}
        )
        self.assertEqual(feminine_indeclinable_name_fields["HEAD_FIO_1"], "Сынгизовой Гульюзум Хадисовны")
        self.assertEqual(feminine_indeclinable_name_fields["HEAD_FIO_2"], "Сынгизовой Гульюзум Хадисовне")

        rare_masculine_name_fields, _ = build_inflected_fields_with_trace({"HEAD_FIO": "Шакуров Дауыт Уралович"})
        self.assertEqual(rare_masculine_name_fields["HEAD_FIO_1"], "Шакурова Дауыта Ураловича")
        self.assertEqual(rare_masculine_name_fields["HEAD_FIO_2"], "Шакурову Дауыту Ураловичу")

    def test_rare_masculine_names_do_not_become_plural_dative(self) -> None:
        self.assertEqual(
            inflect_fio_dative("Баяндинов Мирас Бауржанович").value,
            "Баяндинову Мирасу Бауржановичу",
        )
        self.assertEqual(
            inflect_fio_dative("Мазитов Файм Наилович").value,
            "Мазитову Файму Наиловичу",
        )

    def test_fio_override_is_checked_in_context_sentence(self) -> None:
        row = {
            "HEAD_FIO": "Юнусова Лилия Римовна",
            "MUN_NAME": "Городское поселение Приютово",
            "MUN_R_NAME": "Белебеевский район",
            "SUB_RF": "Республика Башкортостан",
            "ADM_NAME": 'Администрация муниципального образования "Городское поселение Приютово"',
        }

        fields, decisions = build_inflected_fields_with_trace(row)
        by_field = {decision.field: decision for decision in decisions}

        self.assertEqual(fields["HEAD_FIO_1"], "Юнусовой Лилии Римовны")
        self.assertIn("в лице главы", by_field["HEAD_FIO_1"].context_sentence)
        self.assertIn("Юнусовой Лилии Римовны", by_field["HEAD_FIO_1"].filled_sentence)
        self.assertIn("[SLOT]", by_field["HEAD_FIO_1"].context_sentence)
        self.assertIn("Проверь ФИО", by_field["HEAD_FIO_1"].slot_instruction)

    def test_bad_fio_context_is_suspicious(self) -> None:
        self.assertTrue(
            _inflection_context_is_suspicious(
                {
                    "field": "HEAD_FIO_1",
                    "source_value": "Юнусова Лилия Римовна",
                    "result_value": "Юнусовой Лилии Римовна",
                    "target_case": "genitive",
                    "method": "override",
                    "confidence": "high",
                    "filled_sentence": "в лице главы поселения Юнусовой Лилии Римовна, действующего",
                }
            )
        )

    def test_context_correction_is_applied_to_docx(self) -> None:
        row_dir = self.tmp_dir / "1_Городское поселение Приютово"
        row_dir.mkdir()
        docx_path = row_dir / "contract.docx"
        doc = Document()
        doc.add_paragraph("в лице главы поселения Юнусовой Лилии Римовна, действующего")
        doc.save(docx_path)

        result = _apply_inflection_context_corrections(
            [docx_path],
            {
                "items": [
                    {
                        "row_id": "1",
                        "field": "HEAD_FIO_1",
                        "status": "fix",
                        "generated_value": "Юнусовой Лилии Римовна",
                        "corrected_value": "Юнусовой Лилии Римовны",
                        "confidence": 0.96,
                        "comment": "Родительный падеж отчества.",
                    }
                ]
            },
        )

        self.assertEqual(result["applied_count"], 1)
        fixed_doc = Document(docx_path)
        self.assertIn("Юнусовой Лилии Римовны", fixed_doc.paragraphs[0].text)


if __name__ == "__main__":
    unittest.main()

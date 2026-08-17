from __future__ import annotations

import unittest
from datetime import datetime
import uuid
import zipfile
from io import BytesIO
from pathlib import Path

from cryptography.fernet import Fernet
import fitz
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from unittest.mock import patch

from src.security.auth import Principal
from src.security.user_store import create_user
from src.web.v1_router import create_v1_router
from tests.bootstrap import bootstrap_test_runtime


class CampaignV1ApiTests(unittest.TestCase):
    def setUp(self) -> None:
        bootstrap_test_runtime(reset_db=True)
        from src.utils.config import settings

        self._key = patch.object(settings, "smtp_credentials_key", Fernet.generate_key().decode("ascii"))
        self._key.start()
        self.addCleanup(self._key.stop)
        self.username = f"c{uuid.uuid4().hex[:8]}"
        create_user(self.username, "Pass12345!")
        app = FastAPI()
        app.include_router(create_v1_router(check_auth=lambda: Principal(self.username, "t1", "user")))
        self.client = TestClient(app)

    def _attach_required_chain(self, campaign_id: str) -> str:
        created = self.client.post(
            "/api/v1/chains",
            json={"name": f"Required chain {campaign_id}"},
        )
        self.assertEqual(created.status_code, 200, created.text)
        chain_id = str(created.json()["result"]["id"])
        linked = self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={"email_chain_id": chain_id},
        )
        self.assertEqual(linked.status_code, 200, linked.text)
        return chain_id

    def test_setting_email_chain_scenario_without_a_chain_is_rejected(self) -> None:
        """Migration 0040_repair_detached_campaign_chains had to bulk-repair
        drafts stuck in send_scenario='email_chain' with no email_chain_id.
        The old guard only covered the explicit-detach direction (clearing
        email_chain_id while in chain mode) — a PATCH setting the scenario
        to email_chain alone, with no chain ever attached, must be rejected
        instead of reproducing the same bad state."""
        created = self.client.post("/api/v1/campaigns", json={"name": "Chainless"})
        self.assertEqual(created.status_code, 200)
        campaign_id = created.json()["result"]["id"]

        rejected = self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={"send_scenario": "email_chain"},
        )
        self.assertEqual(rejected.status_code, 400, rejected.text)

    def test_setting_email_chain_scenario_with_a_chain_still_succeeds(self) -> None:
        created = self.client.post("/api/v1/campaigns", json={"name": "Chained"})
        self.assertEqual(created.status_code, 200)
        campaign_id = created.json()["result"]["id"]
        chain_id = self._attach_required_chain(campaign_id)

        patched = self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={"send_scenario": "email_chain", "email_chain_id": chain_id},
        )
        self.assertEqual(patched.status_code, 200, patched.text)
        self.assertEqual(patched.json()["result"]["send_scenario"], "email_chain")

    def test_create_update_schedule_launch_pause(self) -> None:
        created = self.client.post("/api/v1/campaigns", json={"name": "API Campaign", "mail_subject": "Hello"})
        self.assertEqual(created.status_code, 200)
        campaign_id = created.json()["result"]["id"]
        self._attach_required_chain(campaign_id)

        patched = self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={"mail_subject": "Hello 2", "document_mode": "kp"},
        )
        self.assertEqual(patched.status_code, 200)
        self.assertEqual(patched.json()["result"]["mail_subject"], "Hello 2")

        recipients = self.client.put(
            f"/api/v1/campaigns/{campaign_id}/recipients",
            json={
                "recipients": [
                    {"company": "A", "contact_name": "A", "email": "a@example.com"},
                    {"company": "B", "contact_name": "B", "email": "b@example.com"},
                ]
            },
        )
        self.assertEqual(recipients.status_code, 200)
        self.assertEqual(recipients.json()["result"]["total"], 2)

        from src.generator.delivery.smtp_mailboxes import create_mailbox

        mailbox = create_mailbox(
            owner_username=self.username,
            provider="custom",
            email="sender@mailpit.local",
            password="x",
            host="mailpit",
            port=1025,
            use_ssl=False,
            use_starttls=False,
            make_default=True,
        )
        self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={"smtp_mailbox_id": mailbox["id"], "transport": "smtp"},
        )

        schedule = self.client.put(
            f"/api/v1/campaigns/{campaign_id}/schedule",
            json={"send_immediately": True, "batch_size": 1, "interval_seconds": 60},
        )
        self.assertEqual(schedule.status_code, 200)
        self.assertGreaterEqual(schedule.json()["result"]["preview"]["batch_count"], 1)

        blocked_without_template = self.client.post(
            f"/api/v1/campaigns/{campaign_id}/launch?force_now=true"
        )
        self.assertEqual(blocked_without_template.status_code, 200, blocked_without_template.text)
        self.assertGreaterEqual(len(blocked_without_template.json()["result"]["batches"]), 1)

        paused = self.client.post(f"/api/v1/campaigns/{campaign_id}/pause")
        self.assertEqual(paused.status_code, 200)
        self.assertEqual(paused.json()["result"]["status"], "paused")

    def test_late_launch_keeps_campaign_interval_instead_of_filling_connection_limit(self) -> None:
        from src.campaigns.connection_service import create_connection

        connection = create_connection(
            self.username,
            {
                "transport": "rusender",
                "email": "sender@example.com",
                "sending_key_id": 42,
                "max_per_hour": 350,
            },
        )
        created = self.client.post(
            "/api/v1/campaigns",
            json={
                "name": "Late paced campaign",
                "mail_subject": "Hello",
                "smtp_mailbox_id": connection["id"],
                "connection_ids": [connection["id"]],
                "transport": "rusender",
            },
        )
        self.assertEqual(created.status_code, 200, created.text)
        campaign_id = created.json()["result"]["id"]
        self._attach_required_chain(campaign_id)

        recipients = [
            {
                "company": f"Company {index}",
                "contact_name": "Contact",
                "email": f"user{index}@example.com",
            }
            for index in range(351)
        ]
        uploaded = self.client.put(
            f"/api/v1/campaigns/{campaign_id}/recipients",
            json={"recipients": recipients},
        )
        self.assertEqual(uploaded.status_code, 200, uploaded.text)
        scheduled = self.client.put(
            f"/api/v1/campaigns/{campaign_id}/schedule",
            json={
                "send_immediately": False,
                "start_at": "2020-01-01T10:00:00+00:00",
                "timezone": "Europe/Moscow",
                "weekdays": list(range(7)),
                "time_windows": [{"start": "00:00", "end": "23:59"}],
                "batch_size": 250,
                "interval_seconds": 3600,
            },
        )
        self.assertEqual(scheduled.status_code, 200, scheduled.text)

        launched = self.client.post(f"/api/v1/campaigns/{campaign_id}/launch")
        self.assertEqual(launched.status_code, 200, launched.text)
        batches = launched.json()["result"]["batches"]
        self.assertEqual([item["size"] for item in batches], [250, 101])
        first = datetime.fromisoformat(batches[0]["scheduled_at"])
        second = datetime.fromisoformat(batches[1]["scheduled_at"])
        self.assertGreaterEqual((second - first).total_seconds(), 3599)

    def test_update_campaign_with_connection_ids(self) -> None:
        from src.generator.delivery.smtp_mailboxes import create_mailbox

        created = self.client.post("/api/v1/campaigns", json={"name": "Multi sender"})
        self.assertEqual(created.status_code, 200)
        campaign_id = created.json()["result"]["id"]

        mailbox_a = create_mailbox(
            owner_username=self.username,
            provider="custom",
            email="sender-a@mailpit.local",
            password="x",
            host="mailpit",
            port=1025,
            use_ssl=False,
            use_starttls=False,
            make_default=True,
        )
        mailbox_b = create_mailbox(
            owner_username=self.username,
            provider="custom",
            email="sender-b@mailpit.local",
            password="x",
            host="mailpit",
            port=1025,
            use_ssl=False,
            use_starttls=False,
            make_default=False,
        )

        patched = self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={"connection_ids": [mailbox_a["id"], mailbox_b["id"]]},
        )
        self.assertEqual(patched.status_code, 200, patched.text)
        result = patched.json()["result"]
        self.assertEqual(result["connection_ids"], [mailbox_a["id"], mailbox_b["id"]])
        self.assertEqual(result["smtp_mailbox_id"], mailbox_a["id"])
        self.assertEqual(result["transport"], "smtp")

        loaded = self.client.get(f"/api/v1/campaigns/{campaign_id}")
        self.assertEqual(loaded.status_code, 200)
        self.assertEqual(loaded.json()["result"]["connection_ids"], [mailbox_a["id"], mailbox_b["id"]])

        validate = self.client.get(f"/api/v1/campaigns/{campaign_id}/validate")
        self.assertEqual(validate.status_code, 200)
        self.assertNotIn("подключение отправителя", " ".join(validate.json()["result"]["errors"]).lower())

    def test_reset_campaign_draft_clears_fields_and_recipients(self) -> None:
        created = self.client.post("/api/v1/campaigns", json={"name": "Reset me", "mail_subject": "Hello"})
        self.assertEqual(created.status_code, 200)
        campaign_id = created.json()["result"]["id"]

        recipients = self.client.put(
            f"/api/v1/campaigns/{campaign_id}/recipients",
            json={
                "recipients": [
                    {"company": "A", "contact_name": "A", "email": "a@example.com"},
                ]
            },
        )
        self.assertEqual(recipients.status_code, 200)
        self.assertEqual(recipients.json()["result"]["total"], 1)

        from src.generator.delivery.smtp_mailboxes import create_mailbox

        mailbox = create_mailbox(
            owner_username=self.username,
            provider="custom",
            email="sender@mailpit.local",
            password="x",
            host="mailpit",
            port=1025,
            use_ssl=False,
            use_starttls=False,
            make_default=True,
        )
        patched = self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={"smtp_mailbox_id": mailbox["id"], "mail_subject": "Changed"},
        )
        self.assertEqual(patched.status_code, 200)

        reset = self.client.post(f"/api/v1/campaigns/{campaign_id}/reset")
        self.assertEqual(reset.status_code, 200, reset.text)
        payload = reset.json()["result"]
        self.assertEqual(payload["name"], "Черновик рассылки")
        self.assertEqual(payload["mail_subject"], "")
        self.assertIsNone(payload.get("smtp_mailbox_id"))
        self.assertIsNone(payload.get("email_chain_id"))
        self.assertEqual(payload.get("draft_payload") or {}, {})

        listed = self.client.get(f"/api/v1/campaigns/{campaign_id}/recipients")
        self.assertEqual(listed.status_code, 200)
        self.assertEqual(listed.json()["result"]["total"], 0)

        schedule = self.client.get(f"/api/v1/campaigns/{campaign_id}/schedule")
        self.assertEqual(schedule.status_code, 200)
        self.assertEqual(schedule.json()["result"]["batch_size"], 25)
        self.assertEqual(schedule.json()["result"]["interval_seconds"], 300)

    def test_duplicate_completed_campaign_creates_complete_draft(self) -> None:
        created = self.client.post(
            "/api/v1/campaigns",
            json={"name": "Completed source", "mail_subject": "Original subject"},
        )
        self.assertEqual(created.status_code, 200, created.text)
        campaign_id = created.json()["result"]["id"]

        template = self.client.post(
            "/api/v1/templates",
            json={
                "name": "Duplicate email",
                "template_type": "email",
                "subject": "Template subject",
                "body_html": "<p>Template body</p>",
                "body_text": "Template body",
            },
        )
        self.assertEqual(template.status_code, 200, template.text)
        template_id = template.json()["result"]["id"]

        audience = self.client.post(
            "/api/v1/audiences",
            json={"name": "Duplicate audience"},
        )
        self.assertEqual(audience.status_code, 200, audience.text)
        audience_id = audience.json()["result"]["id"]

        chain = self.client.post(
            "/api/v1/chains",
            json={"name": "Duplicate chain"},
        )
        self.assertEqual(chain.status_code, 200, chain.text)
        chain_id = chain.json()["result"]["id"]

        patched = self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={
                "description": "Preserve me",
                "email_template_id": template_id,
                "audience_id": audience_id,
                "email_chain_id": chain_id,
                "send_scenario": "email_chain",
                "draft_payload": {"email_body": "<p>Draft body</p>"},
            },
        )
        self.assertEqual(patched.status_code, 200, patched.text)

        recipients = self.client.put(
            f"/api/v1/campaigns/{campaign_id}/recipients",
            json={
                "recipients": [
                    {"company": "A", "contact_name": "A", "email": "a@example.com"},
                    {"company": "B", "contact_name": "B", "email": "b@example.com"},
                ]
            },
        )
        self.assertEqual(recipients.status_code, 200, recipients.text)

        schedule = self.client.put(
            f"/api/v1/campaigns/{campaign_id}/schedule",
            json={
                "send_immediately": False,
                "start_at": "2030-08-05T12:00:00+00:00",
                "timezone": "Europe/Moscow",
                "weekdays": [1, 3],
                "time_windows": [{"start": "10:00", "end": "16:00"}],
                "batch_size": 7,
                "interval_seconds": 45,
                "max_per_hour": 70,
                "max_per_day": 700,
                "on_error": "pause",
                "max_retries": 5,
            },
        )
        self.assertEqual(schedule.status_code, 200, schedule.text)

        from src.infra.db import session_scope
        from src.infra.models import Campaign

        with session_scope() as session:
            source = session.get(Campaign, campaign_id)
            assert source is not None
            source.status = "completed"

        duplicated = self.client.post(f"/api/v1/campaigns/{campaign_id}/duplicate")
        self.assertEqual(duplicated.status_code, 200, duplicated.text)
        copy = duplicated.json()["result"]
        self.assertNotEqual(copy["id"], campaign_id)
        self.assertEqual(copy["status"], "draft")
        self.assertEqual(copy["email_template_id"], template_id)
        self.assertEqual(copy["audience_id"], audience_id)
        self.assertEqual(copy["email_chain_id"], chain_id)
        self.assertEqual(copy["draft_payload"]["email_body"], "<p>Draft body</p>")

        copied_recipients = self.client.get(f"/api/v1/campaigns/{copy['id']}/recipients")
        self.assertEqual(copied_recipients.json()["result"]["total"], 2)
        self.assertTrue(
            all(item["send_status"] == "pending" for item in copied_recipients.json()["result"]["items"])
        )

        copied_schedule = self.client.get(f"/api/v1/campaigns/{copy['id']}/schedule")
        self.assertEqual(copied_schedule.status_code, 200, copied_schedule.text)
        copied_schedule_payload = copied_schedule.json()["result"]
        self.assertEqual(copied_schedule_payload["batch_size"], 7)
        self.assertEqual(copied_schedule_payload["interval_seconds"], 45)
        self.assertEqual(copied_schedule_payload["weekdays"], [1, 3])
        self.assertEqual(copied_schedule_payload["max_per_hour"], 70)
        self.assertEqual(copied_schedule_payload["max_per_day"], 700)
        self.assertEqual(copied_schedule_payload["max_retries"], 5)

    def test_replace_recipients_twice_does_not_500(self) -> None:
        created = self.client.post("/api/v1/campaigns", json={"name": "Reimport Campaign"})
        self.assertEqual(created.status_code, 200)
        campaign_id = created.json()["result"]["id"]
        payload = {
            "recipients": [
                {"company": "A", "contact_name": "A", "email": "a@example.com"},
                {"company": "B", "contact_name": "B", "email": "b@example.com"},
            ]
        }

        first = self.client.put(f"/api/v1/campaigns/{campaign_id}/recipients", json=payload)
        self.assertEqual(first.status_code, 200, first.text)
        self.assertEqual(first.json()["result"]["total"], 2)

        second = self.client.put(f"/api/v1/campaigns/{campaign_id}/recipients", json=payload)
        self.assertEqual(second.status_code, 200, second.text)
        self.assertEqual(second.json()["result"]["total"], 2)

        csv_body = "company,contact_name,email\nC,C,c@example.com\nD,D,d@example.com\n"
        imported = self.client.post(
            f"/api/v1/campaigns/{campaign_id}/recipients/import",
            files={"file": ("recipients.csv", csv_body.encode("utf-8"), "text/csv")},
        )
        self.assertEqual(imported.status_code, 200, imported.text)
        self.assertEqual(imported.json()["result"]["import"]["total"], 2)

        reimported = self.client.post(
            f"/api/v1/campaigns/{campaign_id}/recipients/import",
            files={"file": ("recipients.csv", csv_body.encode("utf-8"), "text/csv")},
        )
        self.assertEqual(reimported.status_code, 200, reimported.text)
        self.assertEqual(reimported.json()["result"]["import"]["total"], 2)

    def test_import_recipients_preserves_long_fallback_email_list(self) -> None:
        from openpyxl import Workbook

        created = self.client.post("/api/v1/campaigns", json={"name": "Long fallback emails"})
        self.assertEqual(created.status_code, 200)
        campaign_id = created.json()["result"]["id"]
        fallback_emails = ", ".join(
            f"fallback-{index:02d}@example.com" for index in range(20)
        )
        self.assertGreater(len(fallback_emails), 320)

        workbook = Workbook()
        worksheet = workbook.active
        worksheet.append(["Organization", "Primary email", "Fallback emails"])
        worksheet.append(["ADM_NAME", "EMAIL_OSN", "EMAIL_DOP"])
        worksheet.append(["Administration", "primary@example.com", fallback_emails])
        buffer = BytesIO()
        workbook.save(buffer)

        imported = self.client.post(
            f"/api/v1/campaigns/{campaign_id}/recipients/import",
            files={
                "file": (
                    "recipients.xlsx",
                    buffer.getvalue(),
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )
        self.assertEqual(imported.status_code, 200, imported.text)
        self.assertEqual(imported.json()["result"]["import"]["total"], 1)

        listed = self.client.get(f"/api/v1/campaigns/{campaign_id}/recipients")
        self.assertEqual(listed.status_code, 200, listed.text)
        recipient = listed.json()["result"]["items"][0]
        self.assertEqual(recipient["email"], "primary@example.com")
        self.assertEqual(recipient["email_fallback"], fallback_emails)

    def test_rusender_connection_requires_sending_key_id(self) -> None:
        missing_id = self.client.post(
            "/api/v1/connections",
            json={
                "transport": "rusender",
                "email": "missing-id@example.com",
                "api_token": "must-not-be-used",
            },
        )
        self.assertEqual(missing_id.status_code, 400, missing_id.text)
        self.assertIn("ID ключа отправки", missing_id.text)

        created = self.client.post(
            "/api/v1/connections",
            json={
                "transport": "rusender",
                "email": "configured@example.com",
                "api_token": "must-not-be-used",
                "sending_key_id": 42,
            },
        )
        self.assertEqual(created.status_code, 200, created.text)
        item = created.json()["result"]
        self.assertEqual(item["transport"], "rusender")
        self.assertEqual(item["sending_key_id"], 42)
        self.assertFalse(item["has_secret"])
        from src.campaigns.connection_service import resolve_connection

        resolved = resolve_connection(item["id"], self.username)
        self.assertEqual(resolved.secret, "")


    def test_create_list_and_test_provider_connections(self) -> None:
        rusender = self.client.post(
            "/api/v1/connections",
            json={
                "transport": "rusender",
                "email": "verified@example.com",
                "sender_name": "Sales",
                "sending_key_id": 42,
            },
        )
        self.assertEqual(rusender.status_code, 200, rusender.text)
        rusender_item = rusender.json()["result"]
        self.assertEqual(rusender_item["transport"], "rusender")
        self.assertEqual(rusender_item["sending_key_id"], 42)
        self.assertNotIn("api_token", rusender_item)
        self.assertNotIn("password_encrypted", rusender_item)

        mailopost = self.client.post(
            "/api/v1/connections",
            json={
                "transport": "mailopost",
                "email": "mailopost@example.com",
                "api_token": "mailopost-secret",
            },
        )
        self.assertEqual(mailopost.status_code, 200, mailopost.text)
        self.assertEqual(mailopost.json()["result"]["transport"], "mailopost")

        listed = self.client.get("/api/v1/connections")
        self.assertEqual(listed.status_code, 200, listed.text)
        self.assertEqual({item["transport"] for item in listed.json()["result"]}, {"rusender", "mailopost"})
        self.assertTrue(all("api_token" not in item for item in listed.json()["result"]))

        updated = self.client.patch(
            f"/api/v1/connections/{rusender_item['id']}",
            json={
                "transport": "rusender",
                "email": "updated@example.com",
                "sender_name": "Updated Sales",
                "sending_key_id": 84,
            },
        )
        self.assertEqual(updated.status_code, 200, updated.text)
        self.assertEqual(updated.json()["result"]["email"], "updated@example.com")
        self.assertEqual(updated.json()["result"]["sender_name"], "Updated Sales")
        self.assertEqual(updated.json()["result"]["sending_key_id"], 84)

        with patch(
            "src.generator.delivery.sender_agent._send_via_rusender",
            return_value={"message_id": "message-1"},
        ) as sender:
            checked = self.client.post(f"/api/v1/connections/{rusender_item['id']}/test")
        self.assertEqual(checked.status_code, 200, checked.text)
        self.assertIn("тестовое письмо", checked.json()["result"]["message"])
        self.assertNotIn("credential_api_key", sender.call_args.kwargs)
        self.assertEqual(sender.call_args.kwargs["credential_sending_key_id"], 84)
        self.assertEqual(sender.call_args.kwargs["sender_email"], "updated@example.com")

        deleted = self.client.delete(f"/api/v1/connections/{mailopost.json()['result']['id']}")
        self.assertEqual(deleted.status_code, 200, deleted.text)


    def test_mailru_connection_is_verified_and_uses_fixed_settings(self) -> None:
        with patch("src.campaigns.connection_service.verify_smtp_credentials") as verify:
            created = self.client.post(
                "/api/v1/connections",
                json={
                    "transport": "smtp",
                    "provider": "mailru",
                    "email": "Sender@Mail.ru",
                    "sender_name": "Sales",
                    "password": "external-app-password",
                    "host": "untrusted.example",
                    "port": 25,
                },
            )

        self.assertEqual(created.status_code, 200, created.text)
        item = created.json()["result"]
        self.assertEqual(item["provider"], "mailru")
        self.assertEqual(item["email"], "sender@mail.ru")
        self.assertEqual(item["host"], "smtp.mail.ru")
        self.assertEqual(item["port"], 465)
        self.assertTrue(item["use_ssl"])
        self.assertFalse(item["use_starttls"])
        credentials = verify.call_args.args[0]
        self.assertEqual(credentials.smtp_username, "sender@mail.ru")
        self.assertEqual(credentials.host, "smtp.mail.ru")
        self.assertEqual(credentials.port, 465)

        with patch("src.campaigns.connection_service.verify_smtp_credentials") as reverify:
            updated = self.client.patch(
                f"/api/v1/connections/{item['id']}",
                json={
                    "transport": "smtp",
                    "email": "updated@inbox.ru",
                    "password": "new-external-app-password",
                },
            )
        self.assertEqual(updated.status_code, 200, updated.text)
        self.assertEqual(updated.json()["result"]["email"], "updated@inbox.ru")
        self.assertEqual(reverify.call_args.args[0].smtp_username, "updated@inbox.ru")

    def test_mailru_connection_accepts_custom_email_domain(self) -> None:
        with patch("src.campaigns.connection_service.verify_smtp_credentials") as verify:
            response = self.client.post(
                "/api/v1/connections",
                json={
                    "transport": "smtp",
                    "provider": "mailru",
                    "email": "personal.offer@parresh.ru",
                    "password": "external-app-password",
                },
            )

        self.assertEqual(response.status_code, 200, response.text)
        item = response.json()["result"]
        self.assertEqual(item["email"], "personal.offer@parresh.ru")
        self.assertEqual(item["host"], "smtp.mail.ru")
        self.assertEqual(verify.call_args.args[0].smtp_username, "personal.offer@parresh.ru")

    def test_mailru_connection_does_not_save_invalid_credentials(self) -> None:
        import smtplib

        error = smtplib.SMTPAuthenticationError(535, b"Authentication failed")
        with patch("src.campaigns.connection_service.verify_smtp_credentials", side_effect=error):
            response = self.client.post(
                "/api/v1/connections",
                json={
                    "transport": "smtp",
                    "provider": "mailru",
                    "email": "sender@bk.ru",
                    "password": "wrong-password",
                },
            )

        self.assertEqual(response.status_code, 400, response.text)
        self.assertIn("пароль для внешнего приложения", response.json()["detail"])
        self.assertEqual(self.client.get("/api/v1/connections").json()["result"], [])

    def test_work_types_list_create_and_reject_duplicate(self) -> None:
        listed = self.client.get("/api/v1/work-types")
        self.assertEqual(listed.status_code, 200, listed.text)
        system_items = listed.json()["result"]
        self.assertGreaterEqual(len(system_items), 5)
        self.assertTrue(all(item["mail_subject"] for item in system_items))
        self.assertTrue(all(item["is_system"] for item in system_items))

        created = self.client.post(
            "/api/v1/work-types",
            json={
                "name": "Градостроительный аудит",
                "mail_subject": "Предложение по градостроительному аудиту",
            },
        )
        self.assertEqual(created.status_code, 200, created.text)
        item = created.json()["result"]
        self.assertFalse(item["is_system"])
        self.assertTrue(item["key"].startswith("custom_"))

        repeated = self.client.get("/api/v1/work-types")
        self.assertIn(item, repeated.json()["result"])
        duplicate = self.client.post(
            "/api/v1/work-types",
            json={"name": "градостроительный аудит", "mail_subject": "Другая тема"},
        )
        self.assertEqual(duplicate.status_code, 400)

    def _fake_docx_pdf_artifact(
        self,
        filename: str,
        data: bytes,
        *,
        owner_username: str | None = None,
        file_kind: str | None = None,
    ) -> tuple[bytes, str]:
        del data, owner_username, file_kind
        return (b"%PDF-1.4 delivery copy", f"{Path(filename).stem}.pdf")

    def _pptx_payload(self, text: str) -> bytes:
        slide_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
       xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:cSld>
    <p:spTree>
      <p:sp>
        <p:txBody>
          <a:p><a:r><a:t>{text}</a:t></a:r></a:p>
        </p:txBody>
      </p:sp>
    </p:spTree>
  </p:cSld>
</p:sld>"""
        payload = BytesIO()
        with zipfile.ZipFile(payload, "w", zipfile.ZIP_DEFLATED) as archive:
            archive.writestr(
                "[Content_Types].xml",
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"/>',
            )
            archive.writestr(
                "ppt/presentation.xml",
                (
                    '<p:presentation '
                    'xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"/>'
                ),
            )
            archive.writestr("ppt/slides/slide1.xml", slide_xml)
        return payload.getvalue()

    def test_upload_pptx_keeps_original_delivery_and_extracts_text(self) -> None:
        pptx_payload = self._pptx_payload("Разработка для администрация района.")
        created = self.client.post(
            "/api/v1/templates/upload",
            data={"template_type": "document", "name": "Static deck"},
            files={
                "file": (
                    "deck.pptx",
                    pptx_payload,
                    "application/vnd.openxmlformats-officedocument.presentationml.presentation",
                )
            },
        )

        self.assertEqual(created.status_code, 200, created.text)
        template = created.json()["result"]
        self.assertEqual(template["version"]["filename"], "deck.pptx")
        self.assertIsNone(template["version"]["rendered_pdf_filename"])
        self.assertEqual(template["version"]["text_extraction_status"], "ready")
        self.assertFalse(template["is_template"])
        self.assertEqual(template["attachment_output_format"], "original")

        source_download = self.client.get(f"/api/v1/templates/{template['id']}/file")
        self.assertEqual(source_download.status_code, 200, source_download.text)
        self.assertEqual(source_download.content, pptx_payload)

        from src.campaigns import template_service

        self.assertIn("Разработка", template_service._file_text("deck.pptx", pptx_payload))

        rejected_pdf = self.client.patch(
            f"/api/v1/templates/{template['id']}",
            json={"attachment_output_format": "pdf"},
        )
        self.assertEqual(rejected_pdf.status_code, 400, rejected_pdf.text)
        self.assertIn("исходном формате", rejected_pdf.text)

    @patch("src.campaigns.template_service._build_document_pdf_artifact")
    def test_upload_file_template_and_download_active_version(self, mock_build_pdf) -> None:
        mock_build_pdf.side_effect = self._fake_docx_pdf_artifact
        source_docx = Document()
        source_docx.add_paragraph("Offer for {{company}} and {{contact_name}}")
        source_payload = BytesIO()
        source_docx.save(source_payload)
        delivery_pdf = b"%PDF-1.4 delivery copy"
        created = self.client.post(
            "/api/v1/templates/upload",
            data={"template_type": "document", "name": "Own document"},
            files={
                "file": (
                    "offer.docx",
                    source_payload.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        self.assertEqual(created.status_code, 200, created.text)
        template = created.json()["result"]
        self.assertEqual(template["template_type"], "document")
        self.assertEqual(template["version"]["filename"], "offer.docx")
        self.assertEqual(template["version"]["rendered_pdf_filename"], "offer.pdf")
        self.assertTrue(template["enforce_one_page"])
        self.assertEqual(
            {item["name"] for item in template["version"]["variables"]},
            {"company", "contact_name"},
        )

        source_download = self.client.get(f"/api/v1/templates/{template['id']}/file")
        self.assertEqual(source_download.status_code, 200, source_download.text)
        self.assertEqual(source_download.content, source_payload.getvalue())
        delivery_download = self.client.get(f"/api/v1/templates/{template['id']}/delivery-file")
        self.assertEqual(delivery_download.status_code, 200, delivery_download.text)
        self.assertEqual(delivery_download.content, delivery_pdf)

        page_mode = self.client.patch(
            f"/api/v1/templates/{template['id']}",
            json={"enforce_one_page": False},
        )
        self.assertEqual(page_mode.status_code, 200, page_mode.text)
        self.assertFalse(page_mode.json()["result"]["enforce_one_page"])

        legacy = self.client.post(
            "/api/v1/templates/upload",
            data={"template_type": "kp", "name": "Legacy alias"},
            files={
                "file": (
                    "legacy.docx",
                    source_payload.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        self.assertEqual(legacy.status_code, 200, legacy.text)
        self.assertEqual(legacy.json()["result"]["template_type"], "document")

        replacement_pdf = PdfWriter()
        replacement_pdf.add_blank_page(width=595, height=842)
        replacement = BytesIO()
        replacement_pdf.write(replacement)
        updated = self.client.post(
            "/api/v1/templates/upload",
            data={"template_type": "document", "template_id": template["id"]},
            files={"file": ("offer-v2.pdf", replacement.getvalue(), "application/pdf")},
        )
        self.assertEqual(updated.status_code, 200, updated.text)
        active = updated.json()["result"]
        self.assertEqual(active["version"]["version_number"], 2)
        self.assertEqual(active["version"]["filename"], "offer-v2.pdf")
        self.assertEqual(active["version"]["rendered_pdf_filename"], "offer-v2.pdf")
        downloaded = self.client.get(f"/api/v1/templates/{template['id']}/file")
        self.assertEqual(downloaded.status_code, 200, downloaded.text)
        self.assertEqual(downloaded.content, replacement.getvalue())
        preview = self.client.get(f"/api/v1/templates/{template['id']}/preview-file")
        self.assertEqual(preview.content, replacement.getvalue())

        document = Document()
        document.add_paragraph("Договор оказания услуг по разработке МНГП для {{company}}")
        document_payload = BytesIO()
        document.save(document_payload)
        contract = self.client.post(
            "/api/v1/templates/upload",
            data={"template_type": "contract", "name": "Own document"},
            files={
                "file": (
                    "contract.docx",
                    document_payload.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        self.assertEqual(contract.status_code, 200, contract.text)
        contract_template = contract.json()["result"]
        self.assertEqual(contract_template["template_type"], "document")
        self.assertEqual(contract_template["name"], "Own document")
        self.assertEqual(contract_template["version"]["filename"], "contract.docx")
        self.assertEqual(contract_template["version"]["rendered_pdf_filename"], "Договор_МНГП.pdf")

    @patch("src.campaigns.template_service._build_document_pdf_artifact")
    def test_upload_incidental_kp_phrase_uses_generic_document_builder(self, mock_build_pdf) -> None:
        mock_build_pdf.return_value = (b"%PDF-1.4 delivery copy", "document.pdf")
        source_docx = Document()
        source_docx.add_paragraph(
            "В случае заинтересованности готовы оперативно представить коммерческое предложение."
        )
        source_docx.add_paragraph("Подготовка границ территориальных зон.")
        source_payload = BytesIO()
        source_docx.save(source_payload)

        created = self.client.post(
            "/api/v1/templates/upload",
            data={"template_type": "document"},
            files={
                "file": (
                    "Шаблон письма Администрациям.docx",
                    source_payload.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        self.assertEqual(created.status_code, 200, created.text)
        template = created.json()["result"]
        self.assertEqual(template["version"]["filename"], "Шаблон письма Администрациям.docx")
        self.assertEqual(template["version"]["rendered_pdf_filename"], "КП_Территориальные_зоны.pdf")
        mock_build_pdf.assert_called_once_with(
            "Шаблон письма Администрациям.docx",
            source_payload.getvalue(),
            owner_username=self.username,
            file_kind=None,
        )

    def test_upload_conversion_error_returns_structured_422(self) -> None:
        from src.campaigns import template_service

        source_docx = Document()
        source_docx.add_paragraph("Документ")
        source_payload = BytesIO()
        source_docx.save(source_payload)
        before = self.client.get("/api/v1/templates", params={"template_type": "document"}).json()["result"]

        with patch(
            "src.campaigns.template_service._build_document_pdf_artifact",
            side_effect=template_service.DocumentConversionError(),
        ):
            response = self.client.post(
                "/api/v1/templates/upload",
                data={"template_type": "document"},
                files={
                    "file": (
                        "document.docx",
                        source_payload.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )

        self.assertEqual(response.status_code, 422, response.text)
        self.assertEqual(
            response.json()["detail"],
            {
                "code": "document_conversion_failed",
                "title": "Не удалось преобразовать документ",
                "message": "Документ не удалось преобразовать в PDF.",
                "hint": "Проверьте, что файл открывается корректно, и повторите загрузку.",
            },
        )
        after = self.client.get("/api/v1/templates", params={"template_type": "document"}).json()["result"]
        self.assertEqual(len(after), len(before))

    def test_delivery_conversion_errors_return_structured_422(self) -> None:
        from src.campaigns import template_service

        expected_detail = template_service.DocumentConversionError().to_detail()
        cases = (
            (
                "src.campaigns.template_service.get_template_delivery_file",
                "/api/v1/templates/document-id/delivery-file",
            ),
            (
                "src.campaigns.template_service.build_file_preview",
                "/api/v1/templates/document-id/preview-file",
            ),
        )

        for target, endpoint in cases:
            with self.subTest(endpoint=endpoint), patch(
                target,
                side_effect=template_service.DocumentConversionError(),
            ):
                response = self.client.get(endpoint)

            self.assertEqual(response.status_code, 422, response.text)
            self.assertEqual(response.json()["detail"], expected_detail)

    @patch("src.campaigns.template_service._build_document_pdf_artifact")
    def test_upload_infers_delivery_filename_from_document_content(self, mock_build_pdf) -> None:
        mock_build_pdf.side_effect = self._fake_docx_pdf_artifact
        source_docx = Document()
        source_docx.add_paragraph("Коммерческое предложение")
        source_docx.add_paragraph(
            "на разработку схемы территориального планирования муниципального образования"
        )
        source_payload = BytesIO()
        source_docx.save(source_payload)
        created = self.client.post(
            "/api/v1/templates/upload",
            data={"template_type": "document"},
            files={
                "file": (
                    "КП_СТП_районы (1) (1).docx",
                    source_payload.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        self.assertEqual(created.status_code, 200, created.text)
        template = created.json()["result"]
        self.assertEqual(template["version"]["filename"], "КП_СТП_районы (1) (1).docx")
        self.assertEqual(template["version"]["rendered_pdf_filename"], "КП_СТП_районы.pdf")
        self.assertEqual(template["name"], "КП СТП районы")
        self.assertEqual(template["version"]["editor_state"]["document_file_kind"], "kp")
        self.assertEqual(mock_build_pdf.call_args.kwargs["file_kind"], "kp")

    @patch("src.campaigns.template_service._build_document_pdf_artifact")
    def test_patch_delivery_filename_without_new_version(self, mock_build_pdf) -> None:
        mock_build_pdf.side_effect = self._fake_docx_pdf_artifact
        source_docx = Document()
        source_docx.add_paragraph("Коммерческое предложение")
        source_payload = BytesIO()
        source_docx.save(source_payload)
        created = self.client.post(
            "/api/v1/templates/upload",
            data={"template_type": "document", "name": "Delivery name doc"},
            files={
                "file": (
                    "offer.docx",
                    source_payload.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        self.assertEqual(created.status_code, 200, created.text)
        template = created.json()["result"]
        template_id = template["id"]
        version_number = template["version"]["version_number"]
        version_id = template["version"]["id"]

        patched = self.client.patch(
            f"/api/v1/templates/{template_id}",
            json={"rendered_pdf_filename": "КП_СТП_районы.pdf"},
        )
        self.assertEqual(patched.status_code, 200, patched.text)
        updated = patched.json()["result"]
        self.assertEqual(updated["version"]["rendered_pdf_filename"], "КП_СТП_районы.pdf")
        self.assertEqual(updated["version"]["version_number"], version_number)
        self.assertEqual(updated["version"]["id"], version_id)

        fetched = self.client.get(f"/api/v1/templates/{template_id}")
        self.assertEqual(fetched.status_code, 200, fetched.text)
        self.assertEqual(
            fetched.json()["result"]["version"]["rendered_pdf_filename"],
            "КП_СТП_районы.pdf",
        )

        empty = self.client.patch(
            f"/api/v1/templates/{template_id}",
            json={"rendered_pdf_filename": "   "},
        )
        self.assertEqual(empty.status_code, 400, empty.text)

    @patch("src.campaigns.template_service._build_document_pdf_artifact")
    def test_patch_name_without_new_version(self, mock_build_pdf) -> None:
        mock_build_pdf.side_effect = self._fake_docx_pdf_artifact
        source_docx = Document()
        source_docx.add_paragraph("Коммерческое предложение")
        source_payload = BytesIO()
        source_docx.save(source_payload)
        created = self.client.post(
            "/api/v1/templates/upload",
            data={"template_type": "document", "name": "КП СТП районы"},
            files={
                "file": (
                    "offer.docx",
                    source_payload.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        self.assertEqual(created.status_code, 200, created.text)
        template = created.json()["result"]
        template_id = template["id"]
        version_number = template["version"]["version_number"]
        version_id = template["version"]["id"]

        patched = self.client.patch(
            f"/api/v1/templates/{template_id}",
            json={"name": "КП СТП районы (новое название)"},
        )
        self.assertEqual(patched.status_code, 200, patched.text)
        updated = patched.json()["result"]
        self.assertEqual(updated["name"], "КП СТП районы (новое название)")
        self.assertEqual(updated["version"]["version_number"], version_number)
        self.assertEqual(updated["version"]["id"], version_id)

        fetched = self.client.get(f"/api/v1/templates/{template_id}")
        self.assertEqual(fetched.status_code, 200, fetched.text)
        self.assertEqual(fetched.json()["result"]["name"], "КП СТП районы (новое название)")
        self.assertEqual(fetched.json()["result"]["version"]["version_number"], version_number)

        empty = self.client.patch(
            f"/api/v1/templates/{template_id}",
            json={"name": "   "},
        )
        self.assertEqual(empty.status_code, 400, empty.text)

    @patch("src.campaigns.template_service._build_document_pdf_artifact")
    def test_template_starters_and_models(self, mock_build_pdf) -> None:
        mock_build_pdf.side_effect = self._fake_docx_pdf_artifact
        models = self.client.get("/api/v1/templates/models")
        self.assertEqual(models.status_code, 200, models.text)
        model_ids = {item["id"] for item in models.json()["result"]}
        self.assertIn("gpt-4o-mini", model_ids)

        starters = self.client.get("/api/v1/templates/starters", params={"template_type": "email"})
        self.assertEqual(starters.status_code, 200, starters.text)
        email_starters = starters.json()["result"]
        self.assertGreaterEqual(len(email_starters), 3)
        used = self.client.post(f"/api/v1/templates/starters/{email_starters[0]['id']}/use")
        self.assertEqual(used.status_code, 200, used.text)
        self.assertEqual(used.json()["result"]["template_type"], "email")

        doc_starters = self.client.get("/api/v1/templates/starters", params={"template_type": "document"})
        self.assertEqual(doc_starters.status_code, 200, doc_starters.text)
        doc_used = self.client.post(f"/api/v1/templates/starters/{doc_starters.json()['result'][0]['id']}/use")
        self.assertEqual(doc_used.status_code, 200, doc_used.text)
        self.assertEqual(doc_used.json()["result"]["template_type"], "document")

        materials = self.client.post("/api/v1/templates/starters/email-materials/use")
        self.assertEqual(materials.status_code, 200, materials.text)
        materials_html = materials.json()["result"]["version"]["body_html"]
        self.assertNotIn("Контакт: {{email}}", materials_html)
        self.assertNotIn("{{campaign_name}}", materials_html)

    def test_generate_template_files_only_and_ai(self) -> None:
        html = b"<p>Hello {{contact_name}} from {{company}}</p>"
        created = self.client.post(
            "/api/v1/templates/generate",
            data={"template_type": "email", "prompt": ""},
            files={"files": ("mail.html", html, "text/html")},
        )
        self.assertEqual(created.status_code, 200, created.text)
        self.assertEqual(created.json()["result"]["template_type"], "email")
        self.assertIn("{{company}}", created.json()["result"]["version"]["body_html"])

        with patch("src.campaigns.template_ai._call_llm") as mock_llm:
            mock_llm.return_value = {
                "name": "AI Letter",
                "subject": "Hi {{company}}",
                "body_html": "<p>Hello {{contact_name}}</p>",
            }
            ai = self.client.post(
                "/api/v1/templates/generate",
                data={
                    "template_type": "email",
                    "prompt": "Сделай короткое письмо",
                    "model": "gpt-4o-mini",
                },
            )
        self.assertEqual(ai.status_code, 200, ai.text)
        self.assertEqual(ai.json()["result"]["name"], "AI Letter")
        mock_llm.assert_called_once()

    def test_archive_template_hides_it_but_preserves_history(self) -> None:
        created = self.client.post(
            "/api/v1/templates",
            json={
                "name": "Disposable template",
                "template_type": "email",
                "subject": "Archive me",
                "body_html": "<p>Archive me</p>",
                "body_text": "Archive me",
            },
        )
        self.assertEqual(created.status_code, 200, created.text)
        template_id = created.json()["result"]["id"]

        archived = self.client.post(f"/api/v1/templates/{template_id}/archive")
        self.assertEqual(archived.status_code, 200, archived.text)
        self.assertTrue(archived.json()["result"]["archived"])

        listed = self.client.get("/api/v1/templates", params={"template_type": "email"})
        self.assertEqual(listed.status_code, 200, listed.text)
        self.assertFalse(any(item["id"] == template_id for item in listed.json()["result"]))

        preserved = self.client.get(f"/api/v1/templates/{template_id}")
        self.assertEqual(preserved.status_code, 200, preserved.text)
        self.assertEqual(preserved.json()["result"]["version"]["body_text"], "Archive me")

    def test_visual_email_template_editor_state_and_assets(self) -> None:
        created = self.client.post(
            "/api/v1/templates",
            json={
                "name": "Visual email",
                "template_type": "email",
                "subject": "Hello {{company}}",
                "body_html": "<table><tr><td>Hi {{contact_name}}</td></tr></table>",
                "body_text": "Hi {{contact_name}}",
                "editor_state": {
                    "email_format": "visual",
                    "grapesjs_project": {"pages": [{"frames": []}]},
                },
            },
        )
        self.assertEqual(created.status_code, 200, created.text)
        template = created.json()["result"]
        self.assertEqual(template["version"]["editor_state"]["email_format"], "visual")
        template_id = template["id"]

        saved = self.client.patch(
            f"/api/v1/templates/{template_id}",
            json={
                "body_html": "<table><tr><td>Updated {{company}}</td></tr></table>",
                "body_text": "Updated {{company}}",
                "editor_state": {
                    "email_format": "visual",
                    "grapesjs_project": {"pages": [{"frames": [{"component": {"type": "wrapper"}}]}]},
                },
            },
        )
        self.assertEqual(saved.status_code, 200, saved.text)
        saved_version = saved.json()["result"]["version"]
        self.assertEqual(saved_version["body_text"], "Updated {{company}}")
        self.assertIn("frames", saved_version["editor_state"]["grapesjs_project"]["pages"][0])

        png_bytes = (
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00"
            b"\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
        )
        uploaded = self.client.post(
            f"/api/v1/templates/{template_id}/assets",
            files={"file": ("logo.png", png_bytes, "image/png")},
        )
        self.assertEqual(uploaded.status_code, 200, uploaded.text)
        asset_src = uploaded.json()["result"]["data"][0]["src"]
        self.assertIn(f"/api/v1/templates/{template_id}/assets/", asset_src)

        asset_id = asset_src.rsplit("/", 1)[-1]
        fetched = self.client.get(f"/api/v1/templates/{template_id}/assets/{asset_id}")
        self.assertEqual(fetched.status_code, 200, fetched.text)
        self.assertEqual(fetched.content[:8], png_bytes[:8])

        visual_starters = [
            item for item in self.client.get("/api/v1/templates/starters", params={"template_type": "email"}).json()["result"]
            if item.get("email_format") == "visual"
        ]
        self.assertGreaterEqual(len(visual_starters), 2)
        used = self.client.post(f"/api/v1/templates/starters/{visual_starters[0]['id']}/use")
        self.assertEqual(used.status_code, 200, used.text)
        self.assertEqual(used.json()["result"]["version"]["editor_state"]["email_format"], "visual")

    def test_import_template_endpoint_html_and_docx(self) -> None:
        html = b"<table><tr><td><p>Hello {{company}}</p></td></tr></table>"
        imported_html = self.client.post(
            "/api/v1/templates/import",
            data={"template_type": "email"},
            files={"file": ("mail.html", html, "text/html")},
        )
        self.assertEqual(imported_html.status_code, 200, imported_html.text)
        payload = imported_html.json()["result"]
        self.assertEqual(payload["template_type"], "email")
        self.assertEqual(payload["version"]["editor_state"]["email_format"], "visual")

        buffer = BytesIO()
        document = Document()
        document.add_heading("Письмо", level=1)
        document.add_paragraph("Для {{company}}")
        document.save(buffer)
        from src.generator.generation.docxjs_converter import DocxJsHtmlResult

        with (
            patch(
                "src.campaigns.template_import_service.convert_docx_to_html_result",
                return_value=DocxJsHtmlResult(
                    html='<div data-content-width="640"><p>Для {{company}}</p></div>',
                    content_width=640,
                ),
            ),
            patch(
                "src.campaigns.template_import_service.convert_docx_to_pdf_bytes",
                return_value=None,
            ),
        ):
            imported_docx = self.client.post(
                "/api/v1/templates/import",
                data={"template_type": "email"},
                files={
                    "file": (
                        "mail.docx",
                        buffer.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        self.assertEqual(imported_docx.status_code, 200, imported_docx.text)
        self.assertIn("{{company}}", imported_docx.json()["result"]["version"]["body_html"])

    def test_generate_template_llm_failure_returns_503(self) -> None:
        with patch("src.campaigns.template_ai._build_client") as mock_client:
            mock_client.return_value.chat.completions.create.side_effect = RuntimeError(
                "Error code: 400 - no healthy deployments for this model"
            )
            response = self.client.post(
                "/api/v1/templates/generate",
                data={
                    "template_type": "email",
                    "prompt": "Сделай короткое письмо",
                    "model": "gpt-4o-mini",
                },
            )
        self.assertEqual(response.status_code, 503, response.text)
        self.assertIn("недоступна", response.json()["detail"].lower())

    def _highlighted_pdf(self) -> bytes:
        document = fitz.open()
        page = document.new_page(width=595, height=842)
        page.draw_rect(fitz.Rect(100, 100, 172, 118), color=None, fill=(1, 1, 0), overlay=True)
        page.insert_text(fitz.Point(101, 114), "ADM_NAME", fontsize=11)
        data = document.tobytes()
        document.close()
        return data

    def test_pdf_editor_for_document_template(self) -> None:
        created = self.client.post(
            "/api/v1/templates/upload",
            data={"template_type": "document", "name": "PDF overlay"},
            files={"file": ("offer.pdf", self._highlighted_pdf(), "application/pdf")},
        )
        self.assertEqual(created.status_code, 200, created.text)
        template = created.json()["result"]
        self.assertEqual(template["template_type"], "document")
        self.assertEqual(template["version"]["filename"], "offer.pdf")
        self.assertEqual(len(template["version"]["editor_state"]["fields"]), 1)

        editor = self.client.get(f"/api/v1/templates/{template['id']}/pdf-editor")
        self.assertEqual(editor.status_code, 200, editor.text)
        fields = editor.json()["result"]["fields"]
        self.assertEqual(len(fields), 1)
        self.assertEqual(fields[0]["variable"], "ADM_NAME")

        saved = self.client.patch(
            f"/api/v1/templates/{template['id']}/pdf-editor",
            json={"fields": [{"id": fields[0]["id"], "value": "Ivanov I.I.", "font_size": fields[0]["font_size"]}]},
        )
        self.assertEqual(saved.status_code, 200, saved.text)
        self.assertEqual(saved.json()["result"]["version"]["version_number"], 2)
        self.assertTrue(saved.json()["result"]["version"]["rendered_pdf_storage_key"])

        email = self.client.post(
            "/api/v1/templates",
            json={
                "name": "Email only",
                "template_type": "email",
                "subject": "Hello",
                "body_html": "<p>Hi {{contact_name}}</p>",
            },
        )
        self.assertEqual(email.status_code, 200, email.text)
        email_id = email.json()["result"]["id"]
        denied = self.client.get(f"/api/v1/templates/{email_id}/pdf-editor")
        self.assertEqual(denied.status_code, 404, denied.text)

    def test_variable_mapping_import_suggest_save_and_launch_gate(self) -> None:
        from io import BytesIO

        from openpyxl import Workbook

        created = self.client.post(
            "/api/v1/campaigns",
            json={"name": "Mapping Campaign", "mail_subject": "Hello {{company}}"},
        )
        self.assertEqual(created.status_code, 200)
        campaign_id = created.json()["result"]["id"]
        self._attach_required_chain(campaign_id)

        email = self.client.post(
            "/api/v1/templates",
            json={
                "name": "Email with vars",
                "template_type": "email",
                "subject": "Hello {{company}}",
                "body_html": "<p>Dear {{contact_name}} from {{ADM_NAME}}</p>",
            },
        )
        self.assertEqual(email.status_code, 200)
        email_id = email.json()["result"]["id"]
        kp = self.client.post(
            "/api/v1/templates",
            json={"name": "KP gate", "template_type": "kp"},
        )
        self.assertEqual(kp.status_code, 200, kp.text)
        kp_id = kp.json()["result"]["id"]
        self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={
                "email_template_id": email_id,
                "kp_template_id": kp_id,
            },
        )

        wb = Workbook()
        ws = wb.active
        ws.append(["A", "B", "C", "D"])
        ws.append(["ADM_NAME", "HEAD_FIO", "EMAIL_OSN", "SUB_RF"])
        ws.append(["Administration A", "Ivanov I.I.", "a@example.com", "Region"])
        buffer = BytesIO()
        wb.save(buffer)
        imported = self.client.post(
            f"/api/v1/campaigns/{campaign_id}/recipients/import",
            files={"file": ("recipients.xlsx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.sheetml.main")},
        )
        self.assertEqual(imported.status_code, 200, imported.text)
        self.assertIn("adm_name", imported.json()["result"]["import"]["recipient_columns"])

        from src.generator.delivery.smtp_mailboxes import create_mailbox

        mailbox = create_mailbox(
            owner_username=self.username,
            provider="custom",
            email="sender@mailpit.local",
            password="x",
            host="mailpit",
            port=1025,
            use_ssl=False,
            use_starttls=False,
            make_default=True,
        )
        self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={"smtp_mailbox_id": mailbox["id"], "transport": "smtp"},
        )
        self.client.put(
            f"/api/v1/campaigns/{campaign_id}/schedule",
            json={"send_immediately": True, "batch_size": 1, "interval_seconds": 60},
        )

        validate_before = self.client.get(f"/api/v1/campaigns/{campaign_id}/validate")
        self.assertEqual(validate_before.status_code, 200)
        self.assertFalse(validate_before.json()["result"]["mapping_confirmed"])
        self.assertEqual(validate_before.json()["result"]["template_issues"], [])
        self.assertIn("сопоставление переменных", " ".join(validate_before.json()["result"]["errors"]).lower())

        suggest = self.client.post(f"/api/v1/campaigns/{campaign_id}/variable-mapping/suggest")
        self.assertEqual(suggest.status_code, 200, suggest.text)
        self.assertEqual(suggest.json()["result"]["status"], "complete")
        self.assertEqual(suggest.json()["result"]["suggested_mapping"].get("ADM_NAME"), "adm_name")

        launch_blocked = self.client.post(f"/api/v1/campaigns/{campaign_id}/launch?force_now=true")
        self.assertEqual(launch_blocked.status_code, 400)

        saved = self.client.put(
            f"/api/v1/campaigns/{campaign_id}/variable-mapping",
            json={"mapping": suggest.json()["result"]["suggested_mapping"]},
        )
        self.assertEqual(saved.status_code, 200, saved.text)
        self.assertTrue(saved.json()["result"]["mapping_confirmed"])

        stale_autosave = self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={
                "description": "autosaved after mapping",
                "email_template_id": email_id,
                "kp_template_id": kp_id,
                "draft_payload": {
                    "description": "autosaved after mapping",
                    "mapping_confirmed": False,
                    "mapping_confirmed_at": None,
                    "variable_mapping": {},
                    "system_variables": {},
                    "recipient_columns": [],
                },
            },
        )
        self.assertEqual(stale_autosave.status_code, 200, stale_autosave.text)
        stale_draft = stale_autosave.json()["result"]["draft_payload"]
        self.assertTrue(stale_draft["mapping_confirmed"])
        self.assertEqual(stale_draft["variable_mapping"].get("ADM_NAME"), "adm_name")

        repeated_import = self.client.post(
            f"/api/v1/campaigns/{campaign_id}/recipients/import",
            files={
                "file": (
                    "recipients.xlsx",
                    buffer.getvalue(),
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )
        self.assertEqual(repeated_import.status_code, 200, repeated_import.text)

        validate_after = self.client.get(f"/api/v1/campaigns/{campaign_id}/validate")
        self.assertEqual(validate_after.status_code, 200)
        self.assertTrue(validate_after.json()["result"]["mapping_confirmed"])
        self.assertNotIn(
            "сопоставление переменных",
            " ".join(validate_after.json()["result"]["errors"]).lower(),
        )

        launched = self.client.post(f"/api/v1/campaigns/{campaign_id}/launch?force_now=true")
        self.assertEqual(launched.status_code, 200, launched.text)

    def test_validate_blocks_launch_on_unresolved_template_artifact(self) -> None:
        created = self.client.post("/api/v1/campaigns", json={"name": "Artifact Campaign"})
        self.assertEqual(created.status_code, 200)
        campaign_id = created.json()["result"]["id"]

        email = self.client.post(
            "/api/v1/templates",
            json={
                "name": "Artifact email",
                "template_type": "email",
                "body_html": "<p>на {{ стp }} для {{company}}.</p>",
            },
        )
        self.assertEqual(email.status_code, 200, email.text)
        email_id = email.json()["result"]["id"]

        loaded = self.client.get(f"/api/v1/campaigns/{campaign_id}/email-chain")
        self.assertEqual(loaded.status_code, 200)
        chain = loaded.json()["result"]["chain"]
        chain["nodes"][0]["email_template_id"] = email_id
        saved = self.client.put(f"/api/v1/campaigns/{campaign_id}/email-chain", json=chain)
        self.assertEqual(saved.status_code, 200, saved.text)
        published = self.client.post(f"/api/v1/campaigns/{campaign_id}/email-chain/publish")
        self.assertEqual(published.status_code, 200, published.text)

        recipients = self.client.put(
            f"/api/v1/campaigns/{campaign_id}/recipients",
            json={
                "recipients": [
                    {"company": "A", "contact_name": "Ivan", "email": "a@example.com"},
                ]
            },
        )
        self.assertEqual(recipients.status_code, 200, recipients.text)

        from src.generator.delivery.smtp_mailboxes import create_mailbox

        mailbox = create_mailbox(
            owner_username=self.username,
            provider="custom",
            email="sender@mailpit.local",
            password="x",
            host="mailpit",
            port=1025,
            use_ssl=False,
            use_starttls=False,
            make_default=True,
        )
        self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={"smtp_mailbox_id": mailbox["id"], "transport": "smtp"},
        )
        self.client.put(
            f"/api/v1/campaigns/{campaign_id}/schedule",
            json={"send_immediately": True, "batch_size": 1, "interval_seconds": 60},
        )

        mapping = self.client.put(
            f"/api/v1/campaigns/{campaign_id}/variable-mapping",
            json={"mapping": {"company": "company"}},
        )
        self.assertEqual(mapping.status_code, 200, mapping.text)
        self.assertTrue(mapping.json()["result"]["mapping_confirmed"])

        validate_fast = self.client.get(f"/api/v1/campaigns/{campaign_id}/validate")
        self.assertEqual(validate_fast.status_code, 200, validate_fast.text)
        payload = validate_fast.json()["result"]
        self.assertTrue(any(issue.get("kind") == "artifact" for issue in payload["template_issues"]))
        self.assertTrue(
            any("артефакт" in error.lower() for error in payload["errors"]),
            payload,
        )

        preview = self.client.post(f"/api/v1/campaigns/{campaign_id}/email-chain/preview")
        self.assertEqual(preview.status_code, 200, preview.text)
        preview_issues = preview.json()["result"]["items"][0]["issues"]
        self.assertTrue(any(issue.get("kind") == "artifact" for issue in preview_issues))
        self.assertTrue(
            any("артефакт" in str(issue.get("message") or "").lower() for issue in preview_issues),
        )

        launch = self.client.post(f"/api/v1/campaigns/{campaign_id}/launch?force_now=true")
        self.assertEqual(launch.status_code, 400, launch.text)

    def test_email_chain_crud_and_publish(self) -> None:
        created = self.client.post("/api/v1/campaigns", json={"name": "Chain API"})
        self.assertEqual(created.status_code, 200)
        campaign_id = created.json()["result"]["id"]

        loaded = self.client.get(f"/api/v1/campaigns/{campaign_id}/email-chain")
        self.assertEqual(loaded.status_code, 200)
        chain = loaded.json()["result"]["chain"]
        self.assertTrue(chain["nodes"])

        chain["nodes"][0]["email_template_id"] = "email-template-1"
        saved = self.client.put(f"/api/v1/campaigns/{campaign_id}/email-chain", json=chain)
        self.assertEqual(saved.status_code, 200)
        self.assertIn("validation", saved.json()["result"])

        published = self.client.post(f"/api/v1/campaigns/{campaign_id}/email-chain/publish")
        self.assertEqual(published.status_code, 200, published.text)
        self.assertTrue(published.json()["result"]["published"])

        camp = self.client.get(f"/api/v1/campaigns/{campaign_id}")
        self.assertEqual(camp.json()["result"]["send_scenario"], "email_chain")

    def test_standalone_chains_api(self) -> None:
        created = self.client.post("/api/v1/chains", json={"name": "Standalone API chain"})
        self.assertEqual(created.status_code, 200, created.text)
        chain_id = created.json()["result"]["id"]

        listed = self.client.get("/api/v1/chains")
        self.assertEqual(listed.status_code, 200)
        self.assertTrue(any(item["id"] == chain_id for item in listed.json()["result"]["items"]))

        loaded = self.client.get(f"/api/v1/chains/{chain_id}")
        self.assertEqual(loaded.status_code, 200)
        chain = loaded.json()["result"]["chain"]
        chain["nodes"][0]["email_template_id"] = "email-template-1"

        saved = self.client.put(f"/api/v1/chains/{chain_id}", json=chain)
        self.assertEqual(saved.status_code, 200, saved.text)

        published = self.client.post(f"/api/v1/chains/{chain_id}/publish")
        self.assertEqual(published.status_code, 200, published.text)
        self.assertTrue(published.json()["result"]["published"])

        campaign = self.client.post("/api/v1/campaigns", json={"name": "Uses standalone chain"})
        campaign_id = campaign.json()["result"]["id"]
        missing_chain = self.client.get(f"/api/v1/campaigns/{campaign_id}/validate")
        self.assertEqual(missing_chain.status_code, 200, missing_chain.text)
        self.assertIn(
            "Выберите цепочку писем",
            missing_chain.json()["result"]["errors"],
        )
        blocked_launch = self.client.post(
            f"/api/v1/campaigns/{campaign_id}/launch?force_now=true"
        )
        self.assertEqual(blocked_launch.status_code, 400, blocked_launch.text)
        self.assertIn("Выберите цепочку писем", blocked_launch.text)
        linked = self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={"send_scenario": "email_chain", "email_chain_id": chain_id},
        )
        self.assertEqual(linked.status_code, 200, linked.text)
        self.assertEqual(linked.json()["result"]["email_chain_id"], chain_id)

        cleared = self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={"email_chain_id": None},
        )
        self.assertEqual(cleared.status_code, 200, cleared.text)
        cleared_payload = cleared.json()["result"]
        self.assertIsNone(cleared_payload["email_chain_id"])
        self.assertEqual(cleared_payload["send_scenario"], "consent_then_materials")
        self.assertIsNone(cleared_payload["draft_payload"]["email_chain_id"])

        relinked = self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={"send_scenario": "email_chain", "email_chain_id": chain_id},
        )
        self.assertEqual(relinked.status_code, 200, relinked.text)

        deleted = self.client.delete(f"/api/v1/chains/{chain_id}")
        self.assertEqual(deleted.status_code, 200, deleted.text)
        self.assertEqual(deleted.json()["result"], {"deleted": True, "id": chain_id})
        self.assertEqual(self.client.get(f"/api/v1/chains/{chain_id}").status_code, 404)
        self.assertEqual(self.client.delete(f"/api/v1/chains/{chain_id}").status_code, 404)

        detached = self.client.get(f"/api/v1/campaigns/{campaign_id}")
        self.assertEqual(detached.status_code, 200, detached.text)
        self.assertIsNone(detached.json()["result"]["email_chain_id"])
        self.assertEqual(
            detached.json()["result"]["send_scenario"],
            "consent_then_materials",
        )
        self.assertIsNone(detached.json()["result"]["draft_payload"]["email_chain_id"])

    def test_standalone_chain_rename(self) -> None:
        created = self.client.post("/api/v1/chains", json={"name": "Старое название"})
        self.assertEqual(created.status_code, 200, created.text)
        chain_id = created.json()["result"]["id"]

        renamed = self.client.patch(f"/api/v1/chains/{chain_id}", json={"name": "  Новое название  "})
        self.assertEqual(renamed.status_code, 200, renamed.text)
        self.assertEqual(renamed.json()["result"]["name"], "Новое название")

        loaded = self.client.get(f"/api/v1/chains/{chain_id}")
        self.assertEqual(loaded.status_code, 200)
        self.assertEqual(loaded.json()["result"]["name"], "Новое название")

        listed = self.client.get("/api/v1/chains")
        self.assertEqual(listed.status_code, 200)
        item = next(row for row in listed.json()["result"]["items"] if row["id"] == chain_id)
        self.assertEqual(item["name"], "Новое название")

        empty = self.client.patch(f"/api/v1/chains/{chain_id}", json={"name": "   "})
        self.assertEqual(empty.status_code, 400, empty.text)

    def test_email_chain_with_link_nodes_and_consent_stats(self) -> None:
        created = self.client.post("/api/v1/campaigns", json={"name": "Chain links API"})
        self.assertEqual(created.status_code, 200)
        campaign_id = created.json()["result"]["id"]

        loaded = self.client.get(f"/api/v1/campaigns/{campaign_id}/email-chain")
        chain = loaded.json()["result"]["chain"]
        root = chain["root_node_id"]
        chain["nodes"][0]["email_template_id"] = "email-template-1"
        chain["nodes"].extend(
            [
                {
                    "id": "node-sub",
                    "name": "Подписаться",
                    "kind": "link",
                    "link_kind": "subscribe",
                },
                {
                    "id": "node-unsub",
                    "name": "Отписаться",
                    "kind": "link",
                    "link_kind": "unsubscribe",
                },
            ]
        )
        chain["edges"] = [
            {"id": "e1", "source_id": root, "target_id": "node-sub", "button_label": "Подписаться"},
            {"id": "e2", "source_id": root, "target_id": "node-unsub", "button_label": "Отписаться"},
        ]

        saved = self.client.put(f"/api/v1/campaigns/{campaign_id}/email-chain", json=chain)
        self.assertEqual(saved.status_code, 200, saved.text)
        saved_chain = saved.json()["result"]["chain"]
        self.assertEqual(saved_chain["nodes"][1]["kind"], "link")

        published = self.client.post(f"/api/v1/campaigns/{campaign_id}/email-chain/publish")
        self.assertEqual(published.status_code, 200, published.text)

        from src.campaigns.chain_consent_service import record_subscribe, record_unsubscribe

        record_subscribe(
            campaign_id=campaign_id,
            recipient_id=1,
            email="sub@example.com",
            node_id="node-sub",
            edge_id="e1",
            token=str(uuid.uuid4()),
        )
        record_unsubscribe(
            campaign_id=campaign_id,
            recipient_id=2,
            email="unsub@example.com",
            node_id="node-unsub",
            edge_id="e2",
            token=str(uuid.uuid4()),
        )

        stats = self.client.get(f"/api/v1/campaigns/{campaign_id}/email-chain/stats")
        self.assertEqual(stats.status_code, 200)
        consents = stats.json()["result"]["consents"]
        self.assertEqual(consents["subscribe"]["count"], 1)
        self.assertEqual(consents["unsubscribe"]["count"], 1)

    def test_validation_auto_fix_applies_variable_mapping(self) -> None:
        created = self.client.post("/api/v1/campaigns", json={"name": "AutoFix Campaign"})
        self.assertEqual(created.status_code, 200)
        campaign_id = created.json()["result"]["id"]

        email = self.client.post(
            "/api/v1/templates",
            json={
                "name": "AutoFix email",
                "template_type": "email",
                "body_html": "<p>Hello {{ADM_NAME}} from {{company}}</p>",
            },
        )
        self.assertEqual(email.status_code, 200, email.text)
        email_id = email.json()["result"]["id"]

        loaded = self.client.get(f"/api/v1/campaigns/{campaign_id}/email-chain")
        chain = loaded.json()["result"]["chain"]
        chain["nodes"][0]["email_template_id"] = email_id
        self.client.put(f"/api/v1/campaigns/{campaign_id}/email-chain", json=chain)
        self.client.post(f"/api/v1/campaigns/{campaign_id}/email-chain/publish")

        recipients = self.client.put(
            f"/api/v1/campaigns/{campaign_id}/recipients",
            json={
                "recipients": [
                    {
                        "company": "A",
                        "contact_name": "Ivan",
                        "email": "a@example.com",
                        "adm_name": "Ivan Admin",
                    },
                ]
            },
        )
        self.assertEqual(recipients.status_code, 200, recipients.text)

        auto_fix = self.client.post(f"/api/v1/campaigns/{campaign_id}/validation/auto-fix")
        self.assertEqual(auto_fix.status_code, 200, auto_fix.text)
        payload = auto_fix.json()["result"]
        self.assertTrue(payload["applied"])
        self.assertTrue(payload["validation"]["mapping_confirmed"])

    def test_validation_auto_fix_applies_philologist_docx_fixes(self) -> None:
        docx_buffer = BytesIO()
        document = Document()
        document.add_paragraph(
            "1.2. Выполнение Работ осуществляется по месту нахождения Исполнителя "
            "на условиях и в сроки, установленные настоящим Договором."
        )
        document.save(docx_buffer)

        created = self.client.post("/api/v1/campaigns", json={"name": "DocFix Campaign"})
        self.assertEqual(created.status_code, 200)
        campaign_id = created.json()["result"]["id"]

        with patch(
            "src.campaigns.template_service._build_document_pdf_artifact",
            return_value=(b"%PDF-1.4 test", "legal.pdf"),
        ):
            uploaded = self.client.post(
                "/api/v1/templates/upload",
                data={"template_type": "document", "name": "Legal doc"},
                files={
                    "file": (
                        "legal.docx",
                        docx_buffer.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
            self.assertEqual(uploaded.status_code, 200, uploaded.text)
            document_id = uploaded.json()["result"]["id"]

            email = self.client.post(
                "/api/v1/templates",
                json={
                    "name": "DocFix email",
                    "template_type": "email",
                    "body_html": "<p>Hello {{company}}</p>",
                },
            )
            self.assertEqual(email.status_code, 200, email.text)
            email_id = email.json()["result"]["id"]

            loaded = self.client.get(f"/api/v1/campaigns/{campaign_id}/email-chain")
            chain = loaded.json()["result"]["chain"]
            chain["nodes"][0]["email_template_id"] = email_id
            chain["nodes"][0]["document_template_ids"] = [document_id]
            self.client.put(f"/api/v1/campaigns/{campaign_id}/email-chain", json=chain)
            self.client.post(f"/api/v1/campaigns/{campaign_id}/email-chain/publish")

            recipients = self.client.put(
                f"/api/v1/campaigns/{campaign_id}/recipients",
                json={
                    "recipients": [
                        {
                            "company": "A",
                            "contact_name": "Ivan",
                            "email": "a@example.com",
                        },
                    ]
                },
            )
            self.assertEqual(recipients.status_code, 200, recipients.text)

            auto_fix = self.client.post(f"/api/v1/campaigns/{campaign_id}/validation/auto-fix")
            self.assertEqual(auto_fix.status_code, 200, auto_fix.text)
            payload = auto_fix.json()["result"]
            self.assertTrue(
                any(item.get("kind") == "document" for item in payload["applied"]),
                msg=f"applied={payload['applied']} skipped={payload['skipped']}",
            )

            file_response = self.client.get(f"/api/v1/templates/{document_id}/file")
            self.assertEqual(file_response.status_code, 200, file_response.text)
            saved = Document(BytesIO(file_response.content))
            text = "\n".join(paragraph.text for paragraph in saved.paragraphs)
            self.assertIn("Выполнение работ", text)
            self.assertIn("нахождения исполнителя", text)

    def _prepare_campaign_with_email_template(
        self,
        *,
        body_html: str,
        subject: str = "",
    ) -> tuple[str, str]:
        created = self.client.post("/api/v1/campaigns", json={"name": "AutoFix Email Campaign"})
        self.assertEqual(created.status_code, 200)
        campaign_id = created.json()["result"]["id"]

        email = self.client.post(
            "/api/v1/templates",
            json={
                "name": "AutoFix email template",
                "template_type": "email",
                "subject": subject,
                "body_html": body_html,
            },
        )
        self.assertEqual(email.status_code, 200, email.text)
        email_id = email.json()["result"]["id"]

        loaded = self.client.get(f"/api/v1/campaigns/{campaign_id}/email-chain")
        self.assertEqual(loaded.status_code, 200)
        chain = loaded.json()["result"]["chain"]
        chain["nodes"][0]["email_template_id"] = email_id
        saved = self.client.put(f"/api/v1/campaigns/{campaign_id}/email-chain", json=chain)
        self.assertEqual(saved.status_code, 200, saved.text)
        published = self.client.post(f"/api/v1/campaigns/{campaign_id}/email-chain/publish")
        self.assertEqual(published.status_code, 200, published.text)

        recipients = self.client.put(
            f"/api/v1/campaigns/{campaign_id}/recipients",
            json={
                "recipients": [
                    {
                        "company": "A",
                        "contact_name": "Ivan",
                        "email": "a@example.com",
                    },
                ]
            },
        )
        self.assertEqual(recipients.status_code, 200, recipients.text)
        return campaign_id, email_id

    def test_validation_auto_fix_repairs_triple_brace_placeholder(self) -> None:
        campaign_id, email_id = self._prepare_campaign_with_email_template(
            body_html="<p>Работы по {{{WORK_TITLE}}} для {{company}}.</p>",
        )

        preview_before = self.client.post(f"/api/v1/campaigns/{campaign_id}/email-chain/preview")
        self.assertEqual(preview_before.status_code, 200, preview_before.text)
        issues_before = preview_before.json()["result"]["items"][0]["issues"]
        self.assertTrue(any(issue.get("kind") == "malformed" for issue in issues_before))

        auto_fix = self.client.post(f"/api/v1/campaigns/{campaign_id}/validation/auto-fix")
        self.assertEqual(auto_fix.status_code, 200, auto_fix.text)
        payload = auto_fix.json()["result"]
        self.assertTrue(
            any(item.get("kind") == "placeholder" for item in payload["applied"]),
            msg=f"applied={payload['applied']} skipped={payload['skipped']}",
        )

        template = self.client.get(f"/api/v1/templates/{email_id}")
        self.assertEqual(template.status_code, 200, template.text)
        version = template.json()["result"]["version"]
        self.assertIn("{{WORK_TITLE}}", version["body_html"])
        self.assertNotIn("{{{WORK_TITLE}}}", version["body_html"])

        preview_after = self.client.post(f"/api/v1/campaigns/{campaign_id}/email-chain/preview")
        self.assertEqual(preview_after.status_code, 200, preview_after.text)
        issues_after = preview_after.json()["result"]["items"][0]["issues"]
        self.assertFalse(any(issue.get("kind") == "malformed" for issue in issues_after))

    def test_validation_auto_fix_saves_artifact_mapping_without_template_edit(self) -> None:
        campaign_id, email_id = self._prepare_campaign_with_email_template(
            body_html="<p>на {{вид работ}} для {{company}}.</p>",
        )
        patch_resp = self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={"work_type": "stp_mo"},
        )
        self.assertEqual(patch_resp.status_code, 200, patch_resp.text)

        auto_fix = self.client.post(f"/api/v1/campaigns/{campaign_id}/validation/auto-fix")
        self.assertEqual(auto_fix.status_code, 200, auto_fix.text)
        payload = auto_fix.json()["result"]
        self.assertTrue(
            any(item.get("kind") == "mapping" for item in payload["applied"]),
            msg=f"applied={payload['applied']} skipped={payload['skipped']}",
        )

        template = self.client.get(f"/api/v1/templates/{email_id}")
        self.assertEqual(template.status_code, 200, template.text)
        version = template.json()["result"]["version"]
        self.assertIn("{{вид работ}}", version["body_html"])

        campaign = self.client.get(f"/api/v1/campaigns/{campaign_id}")
        self.assertEqual(campaign.status_code, 200, campaign.text)
        system_variables = campaign.json()["result"]["draft_payload"].get("system_variables") or {}
        self.assertEqual(system_variables.get("вид работ"), "WORK_TITLE")

        preview_after = self.client.post(f"/api/v1/campaigns/{campaign_id}/email-chain/preview")
        self.assertEqual(preview_after.status_code, 200, preview_after.text)
        item = preview_after.json()["result"]["items"][0]
        self.assertNotIn("{{вид работ}}", item["body_html"])
        self.assertFalse(any(issue.get("kind") == "artifact" for issue in item["issues"]))

    def test_email_chain_preview_hides_resolvable_work_title_artifact(self) -> None:
        campaign_id, _email_id = self._prepare_campaign_with_email_template(
            body_html="<p>на {{вид работ}} для {{company}}.</p>",
        )
        patch_resp = self.client.patch(
            f"/api/v1/campaigns/{campaign_id}",
            json={"work_type": "stp_mo"},
        )
        self.assertEqual(patch_resp.status_code, 200, patch_resp.text)

        preview = self.client.post(f"/api/v1/campaigns/{campaign_id}/email-chain/preview")
        self.assertEqual(preview.status_code, 200, preview.text)
        item = preview.json()["result"]["items"][0]
        self.assertNotIn("{{вид работ}}", item["body_html"])
        self.assertFalse(any(issue.get("kind") == "artifact" for issue in item["issues"]))

    def test_email_chain_preview_substitutes_legacy_name_placeholders(self) -> None:
        campaign_id, _email_id = self._prepare_campaign_with_email_template(
            body_html="<p>Здравствуйте, {{Имя}} {{Отчество}}!</p>",
        )
        recipients = self.client.put(
            f"/api/v1/campaigns/{campaign_id}/recipients",
            json={
                "recipients": [
                    {
                        "company": "ООО Техностар",
                        "contact_name": "Федорова Ирина Александровна",
                        "email": "test@example.com",
                    },
                ]
            },
        )
        self.assertEqual(recipients.status_code, 200, recipients.text)

        preview = self.client.post(f"/api/v1/campaigns/{campaign_id}/email-chain/preview")
        self.assertEqual(preview.status_code, 200, preview.text)
        item = preview.json()["result"]["items"][0]
        self.assertIn("Ирина Александровна", item["body_html"])
        self.assertNotIn("{{Имя}}", item["body_html"])
        self.assertNotIn("{{Отчество}}", item["body_html"])
        self.assertFalse(any(issue.get("kind") == "artifact" for issue in item["issues"]))

    def test_validation_auto_fix_reports_skipped_when_fragment_missing(self) -> None:
        campaign_id, _email_id = self._prepare_campaign_with_email_template(
            body_html="<p>чистый текст без артефактов.</p>",
        )

        with patch(
            "src.campaigns.validation_auto_fix_service.substitution_validation_issues",
            return_value=[
                {
                    "template_id": _email_id,
                    "template_name": "AutoFix email template",
                    "field": "body_html",
                    "kind": "artifact",
                    "severity": "error",
                    "fragment": "{missing fragment}",
                    "message": "artifact",
                }
            ],
        ):
            auto_fix = self.client.post(f"/api/v1/campaigns/{campaign_id}/validation/auto-fix")
        self.assertEqual(auto_fix.status_code, 200, auto_fix.text)
        payload = auto_fix.json()["result"]
        self.assertTrue(
            any("Не удалось определить замену" in item.get("message", "") for item in payload["skipped"]),
            msg=f"skipped={payload['skipped']}",
        )


if __name__ == "__main__":
    unittest.main()

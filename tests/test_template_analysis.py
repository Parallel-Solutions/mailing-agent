from __future__ import annotations

import shutil
import sys
import types
import unittest
import uuid
from contextlib import contextmanager
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document
from openpyxl import Workbook

if "alembic" not in sys.modules:
    alembic_stub = types.ModuleType("alembic")
    alembic_config_stub = types.ModuleType("alembic.config")
    alembic_command_stub = types.ModuleType("alembic.command")
    alembic_config_stub.Config = object
    alembic_stub.command = alembic_command_stub
    sys.modules["alembic"] = alembic_stub
    sys.modules["alembic.config"] = alembic_config_stub
    sys.modules["alembic.command"] = alembic_command_stub

from src.generator.generation.template_analysis import analyze_template_file, build_template_analysis_context

KP_TEMPLATE_FILENAME = "kp_template_source.docx"


@contextmanager
def _workspace_temp_dir():
    root = Path.cwd() / "tmp_test_template_analysis_workspace"
    path = root / f"case_{uuid.uuid4().hex}"
    path.mkdir(parents=True, exist_ok=False)
    try:
        yield str(path)
    finally:
        shutil.rmtree(path, ignore_errors=True)

class TemplateAnalysisTests(unittest.TestCase):
    def test_analyze_docx_extracts_placeholders_and_tables(self) -> None:
        with _workspace_temp_dir() as tmp:
            template_path = Path(tmp) / KP_TEMPLATE_FILENAME
            doc = Document()
            doc.add_paragraph("Уважаемый HEAD_FIO!")
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "КП для ADM_NAME"
            doc.save(template_path)

            result = analyze_template_file(template_path, kind="kp")

            self.assertTrue(result["exists"])
            self.assertEqual(result["format"], "docx")
            self.assertIn("HEAD_FIO", result["placeholders"])
            self.assertIn("ADM_NAME", result["placeholders"])
            self.assertGreaterEqual(result["table_count"], 1)
            self.assertTrue(any(block["kind"] == "table" for block in result["blocks"]))
            self.assertTrue(result["style_profile"]["available"])
            self.assertEqual(result["style_profile"]["normalization"]["images"], "preserve_template_layout")

    def test_build_template_analysis_includes_data_headers(self) -> None:
        with _workspace_temp_dir() as tmp:
            root = Path(tmp)
            templates_dir = root / "templates"
            input_dir = root / "input"
            templates_dir.mkdir()
            input_dir.mkdir()

            template_path = templates_dir / KP_TEMPLATE_FILENAME
            doc = Document()
            doc.add_paragraph("Коммерческое предложение для ADM_NAME и HEAD_FIO")
            doc.save(template_path)

            data_path = input_dir / "data.xlsx"
            workbook = Workbook()
            worksheet = workbook.active
            worksheet.cell(row=2, column=1).value = "ADM_NAME"
            worksheet.cell(row=2, column=2).value = "EMAIL"
            worksheet.cell(row=3, column=1).value = "Администрация тестового района"
            worksheet.cell(row=3, column=2).value = "test@example.com"
            workbook.save(data_path)

            fake_paths = SimpleNamespace(templates_dir=templates_dir, data_xlsx=data_path)
            with patch("src.generator.generation.template_analysis.resolve_job_paths", return_value=fake_paths):
                result = build_template_analysis_context(job_id="job-test", document_mode="kp")

            self.assertEqual(result["data"]["row_count"], 1)
            self.assertIn("ADM_NAME", result["data"]["headers"])
            self.assertIn("HEAD_FIO", result["all_placeholders"])
            self.assertIn("HEAD_FIO", result["placeholders_without_same_named_column"])
            self.assertNotIn("ADM_NAME", result["placeholders_without_same_named_column"])
            self.assertEqual(result["normalization_plan"]["renderer"], "docx_template_pdf_fit")
            self.assertTrue(result["normalization_plan"]["one_page_required"])
            mapping = {item["placeholder"]: item for item in result["field_mapping_suggestions"]}
            self.assertEqual(mapping["ADM_NAME"]["status"], "mapped")


if __name__ == "__main__":
    unittest.main()

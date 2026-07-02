from __future__ import annotations

import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Cm, Pt

from src.generator.delivery import sender_agent
from src.generator.generation.document_builder import (
    build_kp_filename,
    build_kp_replacements,
    normalize_kp_formatting,
    render_docx,
    stabilize_kp_pdf_layout,
)
from src.generator.generation.transforms import build_document_context
from src.generator.generation.work_types import (
    WORK_TYPE_RANDOM_FOREST,
    WORK_TYPE_STP_MO,
    WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES,
    normalize_work_type,
)


class WorkTypeProfileTests(unittest.TestCase):
    def _row(self) -> dict:
        return {
            "ID": "1",
            "SUB_RF": "Иркутская область",
            "MUN_R_NAME": "Усть-Кутский муниципальный район",
            "MUN_NAME": "Нийское сельское поселение",
            "ADM_NAME": "Администрация Нийского сельского поселения",
            "HEAD_FIO": "Иванов Иван Иванович",
            "POPULATION": "123",
            "EMAIL_OSN": "test@example.com",
            "TEL_OSN": "+70000000000",
        }

    def test_stp_mo_context_adds_work_profile_fields(self) -> None:
        context = build_document_context(self._row(), 101, work_type=WORK_TYPE_STP_MO)

        self.assertEqual(context["WORK_TYPE"], WORK_TYPE_STP_MO)
        self.assertEqual(context["WORK_TYPE_LABEL"], "СТП МО")
        self.assertEqual(context["WORK_SHORT_NAME"], "СТП МО")
        self.assertEqual(context["WORK_FILENAME_LABEL"], "СТП")
        self.assertIn("схемы территориального планирования", context["WORK_TITLE"])
        self.assertIn("СТП МО", context["MAIL_SUBJECT_DEFAULT"])
        self.assertNotIn("местных нормативов", context["WORK_TITLE"])

    def test_random_forest_context_adds_work_profile_fields(self) -> None:
        context = build_document_context(self._row(), 101, work_type=WORK_TYPE_RANDOM_FOREST)

        self.assertEqual(normalize_work_type("random_forest"), WORK_TYPE_RANDOM_FOREST)
        self.assertEqual(context["WORK_TYPE"], WORK_TYPE_RANDOM_FOREST)
        self.assertEqual(context["WORK_TYPE_LABEL"], "Случайный лес")
        self.assertEqual(context["WORK_SHORT_NAME"], "Случайный лес")
        self.assertEqual(context["WORK_FILENAME_LABEL"], "Случайный_лес")
        self.assertIn("интеллектуальной автоматизации", context["WORK_TITLE"])
        self.assertIn("интеллектуальной автоматизации", context["MAIL_SUBJECT_DEFAULT"])
        self.assertNotIn("местных нормативов", context["WORK_TITLE"])

    def test_random_forest_kp_replacements_include_head_greeting(self) -> None:
        context = build_document_context(self._row(), 101, work_type=WORK_TYPE_RANDOM_FOREST)
        replacements = dict(build_kp_replacements(context))

        self.assertEqual(replacements["Уважаемый (ая) HEAD_FIO  !"], "Уважаемый Иван Иванович!")
        self.assertEqual(replacements["HEAD_GREETING"], "Уважаемый Иван Иванович!")
        self.assertEqual(replacements["HEAD_FIO_SHORT"], "И.И. Иванов")
        self.assertEqual(replacements["HEAD_FIO"], "И.И. Иванов")

        row = {**self._row(), "HEAD_FIO": "Юнусова Лилия Римовна"}
        female_context = build_document_context(row, 101, work_type=WORK_TYPE_RANDOM_FOREST)
        female_replacements = dict(build_kp_replacements(female_context))
        self.assertEqual(female_replacements["Уважаемый (ая) HEAD_FIO  !"], "Уважаемая Лилия Римовна!")

    def test_random_forest_render_replaces_split_head_greeting(self) -> None:
        Path("tmp").mkdir(exist_ok=True)
        template_path = Path("tmp") / "random_forest_kp_template_source.docx"
        output_path = Path("tmp") / "random_forest_kp_result.docx"
        try:
            template_path.unlink(missing_ok=True)
            output_path.unlink(missing_ok=True)
            doc = Document()
            paragraph = doc.add_paragraph()
            for chunk in ["Уважаемы", "й (", "ая", ")", " ", "HEAD_FIO", " ", " ", "!"]:
                paragraph.add_run(chunk)
            table = doc.add_table(rows=1, cols=1)
            cell_paragraph = table.rows[0].cells[0].paragraphs[0]
            cell_paragraph.add_run("ADM")
            cell_paragraph.add_run("_NAME")
            cell_paragraph.add_run(" ")
            doc.save(template_path)

            context = build_document_context(self._row(), 101, work_type=WORK_TYPE_RANDOM_FOREST)
            render_docx(template_path, build_kp_replacements(context), output_path, context)

            rendered = Document(output_path)
            text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
            cell_text = rendered.tables[0].rows[0].cells[0].text
            self.assertIn("Уважаемый Иван Иванович!", text)
            self.assertNotIn("HEAD_FIO", text)
            self.assertIn("Администрации", cell_text)
            self.assertIn("Нийское сельское поселение", cell_text)
            self.assertNotIn("ADM_NAME", cell_text)
        finally:
            template_path.unlink(missing_ok=True)
            output_path.unlink(missing_ok=True)

    def test_territorial_zone_context_adds_work_profile_fields(self) -> None:
        context = build_document_context(self._row(), 101, work_type=WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES)

        self.assertEqual(context["WORK_TYPE"], WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES)
        self.assertIn("территориальных зон", context["WORK_TITLE"])
        self.assertNotIn("местных нормативов", context["WORK_TITLE"])

    def test_territorial_zone_inflection_context_uses_work_profile(self) -> None:
        context = build_document_context(self._row(), 101, work_type=WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES)
        trace = {item["field"]: item for item in context["INFLECTION_TRACE"]}

        self.assertIn("территориальных зон", trace["MUN_NAME_2"]["context_sentence"])
        self.assertNotIn("местных нормативов", trace["MUN_NAME_2"]["context_sentence"])

    def test_stp_mo_kp_filename_uses_short_uppercase_prefix(self) -> None:
        row = self._row()
        row["MUN_NAME"] = "Жигаловский муниципальный округ"
        row["MUN_R_NAME"] = "Жигаловский муниципальный округ"
        context = build_document_context(row, 101, work_type=WORK_TYPE_STP_MO)

        filename = build_kp_filename(row, context)

        self.assertEqual(filename, "КП_СТП_Жигаловский муниципальный округ.docx")
        self.assertNotIn("СТП_МО", filename)
        self.assertNotIn("Стп", filename)

    def test_territorial_zone_kp_filename_uses_profile_prefix(self) -> None:
        row = self._row()
        context = build_document_context(row, 101, work_type=WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES)

        filename = build_kp_filename(row, context)

        self.assertTrue(filename.startswith("КП_Территориальные_зоны_"))
        self.assertNotIn("МНГП", filename)

    def test_territorial_zone_replacements_include_work_placeholders(self) -> None:
        context = build_document_context(self._row(), 101, work_type=WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES)

        replacements = dict(build_kp_replacements(context))

        self.assertIn("территориальных зон", replacements["WORK_TITLE"])
        self.assertIn("описание местоположения границ территориальных зон", replacements["WORK_RESULT_NAME"])
        self.assertIn("Нийского сельского поселения", replacements["MUN_R_NAME SUB_RF"])
        self.assertIn("Усть-Кутского муниципального района", replacements["MUN_R_NAME SUB_RF"])
        self.assertEqual(replacements["ADM_NAME"], replacements["ADM_NAME_1"])
        self.assertTrue(replacements["ADM_NAME"].startswith("Администрации"))
        self.assertNotIn("местных нормативов", replacements["WORK_TITLE"])

    def test_territorial_zone_consent_body_does_not_use_mngp_text(self) -> None:
        body = sender_agent._build_consent_request_body(
            self._row(),
            consent_url="https://example.test/consent",
            attachment_mode="kp",
            work_type=WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES,
        )

        self.assertIn("территориальных зон", body)
        self.assertNotIn("МНГП", body)
        self.assertNotIn("местных нормативов", body)

    def test_territorial_zone_kp_preserves_signature_contact_row(self) -> None:
        doc = Document()
        doc.add_table(rows=1, cols=2)
        work_table = doc.add_table(rows=3, cols=2)
        work_table.rows[1].cells[0].text = "Выполнение работ по разработке проекта"
        signature_table = doc.add_table(rows=2, cols=3)
        signature_table.rows[0].cells[0].text = "С уважением,\nисполнительный директор"
        signature_table.rows[1].cells[0].text = "Исп. Черкашина Наталья\nтел. +7 963 912-74-25\nks@parresh.ru"
        for _ in range(12):
            doc.add_paragraph("Текст")
        context = build_document_context(self._row(), 101, work_type=WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES)

        normalize_kp_formatting(doc, context)

        self.assertEqual(len(signature_table.rows), 2)
        self.assertIn("Черкашина", signature_table.rows[1].cells[0].text)
        self.assertIn("ks@parresh", signature_table.rows[1].cells[0].text)

    def test_territorial_zone_render_preserves_uploaded_kp_layout(self) -> None:
        Path("tmp").mkdir(exist_ok=True)
        template_path = Path("tmp") / "kp_test_work_type_template_source.docx"
        output_path = Path("tmp") / "test_work_type_kp_result.docx"
        try:
            template_path.unlink(missing_ok=True)
            output_path.unlink(missing_ok=True)
            doc = Document()
            doc.sections[0].top_margin = Cm(0.05)
            doc.sections[0].bottom_margin = Cm(0.57)
            paragraph = doc.add_paragraph(
                "ООО предлагает выполнить работы по разработке файлов описания местоположения границ "
                "территориальных зон для MUN_R_NAME SUB_RF."
            )
            paragraph.runs[0].font.name = "Arial"
            paragraph.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
            paragraph.paragraph_format.space_after = Pt(12)
            table = doc.add_table(rows=1, cols=2)
            admin_run = table.rows[0].cells[1].paragraphs[0].add_run("ADM_NAME")
            admin_run.bold = True
            admin_run.font.name = "Courier New"
            admin_run.font.size = Pt(14)
            doc.save(template_path)

            context = build_document_context(self._row(), 101, work_type=WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES)
            render_docx(template_path, build_kp_replacements(context), output_path, context)

            rendered = Document(output_path)
            self.assertAlmostEqual(rendered.sections[0].top_margin.cm, 0.05, places=2)
            self.assertAlmostEqual(rendered.sections[0].bottom_margin.cm, 0.57, places=2)
            self.assertIn("разработке файлов описания", rendered.paragraphs[0].text)
            self.assertNotIn("MUN_R_NAME SUB_RF", rendered.paragraphs[0].text)
            self.assertEqual(rendered.paragraphs[0].paragraph_format.space_after.pt, 12)
            self.assertEqual(rendered.paragraphs[0].runs[0].font.name, "Arial")
            self.assertEqual(rendered.paragraphs[0].runs[0].font.highlight_color, WD_COLOR_INDEX.YELLOW)
            self.assertNotEqual(rendered.paragraphs[0].runs[0].font.name, "Tahoma")
            self.assertIn("Нийского сельского поселения", rendered.paragraphs[0].text)
            admin_runs = [run for run in rendered.tables[0].rows[0].cells[1].paragraphs[0].runs if run.text.strip()]
            self.assertTrue(admin_runs)
            self.assertTrue(all(run.bold is True for run in admin_runs))
            self.assertTrue(all(run.font.name == "Courier New" for run in admin_runs))
            self.assertTrue(all(run.font.size.pt == 14 for run in admin_runs if run.font.size))
        finally:
            template_path.unlink(missing_ok=True)
            output_path.unlink(missing_ok=True)

    def test_kp_pdf_layout_after_price_note_is_stabilized(self) -> None:
        doc = Document()
        paragraph = doc.add_paragraph("Стоимость выполнения работ составляет ")
        amount_run = paragraph.add_run("99 000")
        amount_run.bold = True
        paragraph.add_run(" рублей 00 копеек, в том числе НДС 5% — 4 714,29 руб.")
        doc.add_paragraph("")
        doc.add_paragraph("ООО «Параллельные Решения» специализируется на комплексной разработке документов.")

        stabilize_kp_pdf_layout(doc, {})

        self.assertEqual(
            doc.paragraphs[0].text,
            "Стоимость выполнения работ составляет 99 000 рублей 00 копеек, в том числе НДС 5% — 4 714,29 руб.",
        )
        self.assertTrue(doc.paragraphs[0].runs[1].bold)
        self.assertEqual(doc.paragraphs[0].paragraph_format.space_after.pt, 0)
        self.assertEqual(doc.paragraphs[1].text, " ")
        self.assertEqual(doc.paragraphs[1].runs[0].font.size.pt, 1)
        self.assertEqual(doc.paragraphs[1].paragraph_format.line_spacing.pt, 4)
        self.assertEqual(
            doc.paragraphs[2].text,
            "ООО «Параллельные Решения» специализируется на комплексной разработке документов.",
        )

    def test_territorial_zone_kp_price_gap_is_compacted(self) -> None:
        doc = Document()
        doc.add_paragraph(
            "Стоимость выполнения работ (за 1 территориальную зону) составляет "
            "10 000 рублей 00 копеек, в том числе НДС 5% — 476,19 руб."
        )
        doc.add_paragraph("")
        doc.add_paragraph("ООО «Параллельные Решения» специализируется на комплексной разработке документов.")

        stabilize_kp_pdf_layout(doc, {"WORK_TYPE": WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES})

        self.assertEqual(len(doc.paragraphs), 3)
        self.assertEqual(doc.paragraphs[0].paragraph_format.space_after.pt, 0)
        self.assertEqual(doc.paragraphs[1].text, " ")
        self.assertEqual(doc.paragraphs[1].runs[0].font.size.pt, 1)
        self.assertEqual(doc.paragraphs[1].paragraph_format.line_spacing.pt, 4)
        self.assertEqual(
            doc.paragraphs[2].text,
            "ООО «Параллельные Решения» специализируется на комплексной разработке документов.",
        )


if __name__ == "__main__":
    unittest.main()

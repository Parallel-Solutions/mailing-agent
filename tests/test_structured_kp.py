from __future__ import annotations

import re
import shutil
import struct
import zlib
import unittest
from pathlib import Path
from unittest.mock import patch
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

from src.generator.generation import document_builder
from src.generator.generation.structured_kp import build_structured_kp_model, render_structured_kp_docx
from src.generator.generation.transforms import build_document_context


SIMPLE_BACKGROUND_SVG = (
    b'<svg viewBox="0 0 10 10" xmlns="http://www.w3.org/2000/svg">'
    b'<path d="M0 0L10 0L10 10L0 10Z" fill="#EAEAEA"/>'
    b'</svg>'
)
COMPOUND_BACKGROUND_SVG = (
    b'<svg viewBox="0 0 10 10" xmlns="http://www.w3.org/2000/svg">'
    b'<style>.MsftOfcResponsive_Fill_eaeaea { fill:#EAEAEA; }</style>'
    b'<path class="MsftOfcResponsive_Fill_eaeaea" fill="#D9D9D9" '
    b'd="M0 0L10 0L10 10L0 10Z M3 3L3 7L7 7L7 3Z"/>'
    b'</svg>'
)


def png_size(payload: bytes) -> tuple[int, int]:
    if not payload.startswith(b"\x89PNG\r\n\x1a\n"):
        raise AssertionError("payload is not a PNG")
    return struct.unpack(">II", payload[16:24])


def png_rgba_pixel(payload: bytes, x: int, y: int) -> tuple[int, int, int, int]:
    width, height = png_size(payload)
    if not (0 <= x < width and 0 <= y < height):
        raise AssertionError("pixel is outside PNG bounds")
    offset = 8
    idat = b""
    while offset < len(payload):
        length = struct.unpack(">I", payload[offset : offset + 4])[0]
        chunk_type = payload[offset + 4 : offset + 8]
        chunk_data = payload[offset + 8 : offset + 8 + length]
        if chunk_type == b"IDAT":
            idat += chunk_data
        offset += 12 + length
    raw = zlib.decompress(idat)
    stride = 1 + width * 4
    row = raw[y * stride : (y + 1) * stride]
    if row[0] != 0:
        raise AssertionError("unexpected PNG filter")
    pixel_offset = 1 + x * 4
    return tuple(row[pixel_offset : pixel_offset + 4])  # type: ignore[return-value]

class StructuredKPTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp_dir = Path("tmp") / "structured_kp_tests"
        shutil.rmtree(self.tmp_dir, ignore_errors=True)
        self.tmp_dir.mkdir(parents=True, exist_ok=True)

    def tearDown(self) -> None:
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def _row(self) -> dict:
        return {
            "ID": "1",
            "SUB_RF": "Республика Адыгея",
            "MUN_R_NAME": "Тахтамукайский муниципальный район",
            "MUN_NAME": "Яблоновское городское поселение",
            "ADM_NAME": "Администрация Яблоновского городского поселения",
            "HEAD_FIO": "Иванов Иван Иванович",
            "EMAIL_OSN": "test@example.com",
            "TEL_OSN": "+70000000000",
        }

    def test_build_structured_kp_model_uses_generation_context(self) -> None:
        context = build_document_context(self._row(), 101)

        model = build_structured_kp_model(context)

        self.assertEqual(model.outgoing_number, "101")
        self.assertIn("Яблоновского городского поселения", model.work_scope)
        self.assertIn("местных нормативов", model.work_title)
        self.assertEqual(model.price.amount_rubles, 99000)
        self.assertEqual(model.price.amount_words, "девяносто девять тысяч")

    def test_render_structured_kp_docx_creates_document_without_placeholders(self) -> None:
        context = build_document_context(self._row(), 101)
        output_path = self.tmp_dir / "structured_kp.docx"

        render_structured_kp_docx(context, output_path)

        doc = Document(output_path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
        full_text = f"{text}\n{table_text}"
        self.assertIn("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ", full_text)
        self.assertIn("Яблоновского городского поселения", full_text)
        self.assertIn("99 000,00", full_text)
        self.assertNotIn("MUN_", full_text)
        self.assertNotIn("ADM_", full_text)

    def test_generate_documents_for_row_can_use_structured_kp_without_template(self) -> None:
        row = self._row()
        context = build_document_context(row, 101)
        output_dir = self.tmp_dir / "output"
        batch_dir = self.tmp_dir / "batch"
        templates_dir = self.tmp_dir / "templates"
        templates_dir.mkdir()

        with patch.object(document_builder, "KP_GENERATION_ENGINE", "structured"):
            generated = document_builder.generate_documents_for_row(
                row,
                context,
                output_dir=output_dir,
                batch_docx_dir=batch_dir,
                templates_dir=templates_dir,
                document_mode="kp",
            )

        self.assertIn("kp", generated)
        self.assertTrue(generated["kp"].exists())
        self.assertTrue(generated["kp_final_docx"].name.startswith("КП_"))
        self.assertIn("Яблоновское городское поселение", generated["kp_final_docx"].name)


    def test_restore_svg_assets_keeps_png_fallback_for_pdf_export(self) -> None:
        template_path = self.tmp_dir / "template.docx"
        output_path = self.tmp_dir / "output.docx"
        svg_fragment = (
            '<a:blip r:embed="rId14"><a:extLst>'
            '<a:ext uri="{96DAC541-7B7A-43D3-8B79-37D633B846F1}">'
            '<asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" '
            'r:embed="rId15"/></a:ext></a:extLst></a:blip>'
        )
        template_document = f"<w:document>{svg_fragment}</w:document>"
        output_document = '<w:document><a:blip r:embed="rId14"/></w:document>'
        template_rels = (
            '<Relationships>'
            '<Relationship Id="rId14" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image8.png"/>'
            '<Relationship Id="rId15" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image9.svg"/>'
            '</Relationships>'
        )
        output_rels = (
            '<Relationships>'
            '<Relationship Id="rId14" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image8.png"/>'
            '</Relationships>'
        )
        with ZipFile(template_path, "w", compression=ZIP_DEFLATED) as archive:
            archive.writestr("word/document.xml", template_document)
            archive.writestr("word/_rels/document.xml.rels", template_rels)
            archive.writestr("word/media/image8.png", b"png")
            archive.writestr("word/media/image9.svg", SIMPLE_BACKGROUND_SVG)
        with ZipFile(output_path, "w", compression=ZIP_DEFLATED) as archive:
            archive.writestr("word/document.xml", output_document)
            archive.writestr("word/_rels/document.xml.rels", output_rels)
            archive.writestr("word/media/image8.png", b"png")

        document_builder.restore_svg_assets_from_template(template_path, output_path)

        with ZipFile(output_path) as archive:
            restored_document = archive.read("word/document.xml").decode("utf-8")
            restored_rels = archive.read("word/_rels/document.xml.rels").decode("utf-8")
            names = set(archive.namelist())
        self.assertIn('<a:blip r:embed="rId14"><a:extLst>', restored_document)
        self.assertIn('asvg:svgBlip', restored_document)
        self.assertIn('Id="rId14"', restored_rels)
        self.assertIn('Target="media/image8.png"', restored_rels)
        self.assertIn('Id="rId15"', restored_rels)
        self.assertIn('Target="media/image9.svg"', restored_rels)
        self.assertIn("word/media/image8.png", names)
        self.assertIn("word/media/image9.svg", names)

    def test_strip_background_svg_extensions_keeps_self_closing_blip_valid(self) -> None:
        anchor = (
            '<wp:anchor behindDoc="1"><pic:blipFill>'
            '<a:blip r:embed="rId14"><a:extLst>'
            '<a:ext uri="{96DAC541-7B7A-43D3-8B79-37D633B846F1}">'
            '<asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" '
            'r:embed="rId15"/></a:ext></a:extLst></a:blip>'
            '<a:stretch><a:fillRect/></a:stretch>'
            '</pic:blipFill></wp:anchor>'
        )

        stripped, _, changed, _ = document_builder.strip_background_svg_extensions_for_pdf(anchor, "<Relationships/>")

        self.assertTrue(changed)
        self.assertIn('<a:blip r:embed="rId14"/><a:stretch>', stripped)
        self.assertNotIn("//>", stripped)
        self.assertNotIn("asvg:svgBlip", stripped)

    def test_malformed_self_closing_blip_is_normalized_before_svg_restore(self) -> None:
        normalized, changed = document_builder.normalize_malformed_self_closing_blips(
            '<pic:blipFill><a:blip r:embed="rId14"//><a:stretch/></pic:blipFill>'
        )

        self.assertTrue(changed)
        self.assertEqual(normalized, '<pic:blipFill><a:blip r:embed="rId14"/><a:stretch/></pic:blipFill>')

    def test_restore_svg_assets_readds_missing_signature_background_with_template_position(self) -> None:
        template_path = self.tmp_dir / "background-template.docx"
        output_path = self.tmp_dir / "background-output.docx"
        background_run = (
            '<w:r><w:drawing><wp:anchor behindDoc="1"><wp:positionH relativeFrom="column"><wp:posOffset>828675</wp:posOffset></wp:positionH><wp:positionV relativeFrom="paragraph"><wp:posOffset>-889000</wp:posOffset></wp:positionV><wp:extent cx="4294505" cy="2900680"/>'
            '<a:blip r:embed="rId14"><a:extLst>'
            '<a:ext uri="{96DAC541-7B7A-43D3-8B79-37D633B846F1}">'
            '<asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" '
            'r:embed="rId15"/></a:ext></a:extLst></a:blip>'
            '</wp:anchor></w:drawing></w:r>'
        )
        template_document = f"<w:document><w:body><w:tbl><w:tr><w:tc><w:p>{background_run}</w:p></w:tc></w:tr></w:tbl></w:body></w:document>"
        output_document = "<w:document><w:body><w:tbl><w:tr><w:tc><w:p></w:p></w:tc></w:tr></w:tbl></w:body></w:document>"
        template_rels = (
            '<Relationships>'
            '<Relationship Id="rId14" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image8.png"/>'
            '<Relationship Id="rId15" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image9.svg"/>'
            '</Relationships>'
        )
        output_rels = "<Relationships></Relationships>"
        with ZipFile(template_path, "w", compression=ZIP_DEFLATED) as archive:
            archive.writestr("word/document.xml", template_document)
            archive.writestr("word/_rels/document.xml.rels", template_rels)
            archive.writestr("word/media/image8.png", b"png")
            archive.writestr("word/media/image9.svg", SIMPLE_BACKGROUND_SVG)
        with ZipFile(output_path, "w", compression=ZIP_DEFLATED) as archive:
            archive.writestr("word/document.xml", output_document)
            archive.writestr("word/_rels/document.xml.rels", output_rels)

        document_builder.restore_svg_assets_from_template(template_path, output_path)

        with ZipFile(output_path) as archive:
            restored_document = archive.read("word/document.xml").decode("utf-8")
            restored_rels = archive.read("word/_rels/document.xml.rels").decode("utf-8")
            names = set(archive.namelist())
            png_payload = archive.read("word/media/image8.png")
        self.assertIn('behindDoc="1"', restored_document)
        self.assertIn('<wp:posOffset>828675</wp:posOffset>', restored_document)
        self.assertIn('<wp:posOffset>-889000</wp:posOffset>', restored_document)
        self.assertNotIn('asvg:svgBlip', restored_document)
        self.assertIn('<a:blip r:embed="rId14"/>', restored_document)
        self.assertIn('Id="rId14"', restored_rels)
        self.assertIn('Id="rId15"', restored_rels)
        self.assertIn("word/media/image8.png", names)
        self.assertIn("word/media/image9.svg", names)
        self.assertEqual(png_size(png_payload), (1409, 952))

    def test_restore_svg_assets_places_signature_background_in_signature_table_not_contact_table(self) -> None:
        template_path = self.tmp_dir / "signature-background-template.docx"
        output_path = self.tmp_dir / "signature-background-output.docx"
        background_run = (
            '<w:r><w:drawing><wp:anchor behindDoc="1"><wp:positionH relativeFrom="column"><wp:posOffset>828675</wp:posOffset></wp:positionH><wp:positionV relativeFrom="paragraph"><wp:posOffset>-889000</wp:posOffset></wp:positionV><wp:extent cx="4294505" cy="2900680"/>'
            '<a:blip r:embed="rId14"><a:extLst>'
            '<a:ext uri="{96DAC541-7B7A-43D3-8B79-37D633B846F1}">'
            '<asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" '
            'r:embed="rId15"/></a:ext></a:extLst></a:blip>'
            '</wp:anchor></w:drawing></w:r>'
        )
        signature_table_template = (
            '<w:tbl>'
            '<w:tr><w:tc><w:p><w:r><w:t>С уважением, исполнительный директор</w:t></w:r></w:p></w:tc><w:tc><w:p/></w:tc><w:tc><w:p><w:r><w:t>К.И. Крашенинников</w:t></w:r></w:p></w:tc></w:tr>'
            f'<w:tr><w:tc><w:p><w:r><w:t>Исп. Черкашина Наталья</w:t></w:r></w:p></w:tc><w:tc><w:p/></w:tc><w:tc><w:p>{background_run}</w:p></w:tc></w:tr>'
            '</w:tbl>'
        )
        signature_table_output = signature_table_template.replace(background_run, "")
        contact_table = '<w:tbl><w:tr><w:tc><w:p><w:r><w:t>Исп. Черкашина Наталья</w:t></w:r></w:p></w:tc></w:tr></w:tbl>'
        template_document = f"<w:document><w:body>{signature_table_template}{contact_table}</w:body></w:document>"
        output_document = f"<w:document><w:body>{signature_table_output}{contact_table}</w:body></w:document>"
        template_rels = (
            '<Relationships>'
            '<Relationship Id="rId14" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image8.png"/>'
            '<Relationship Id="rId15" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image9.svg"/>'
            '</Relationships>'
        )
        with ZipFile(template_path, "w", compression=ZIP_DEFLATED) as archive:
            archive.writestr("word/document.xml", template_document)
            archive.writestr("word/_rels/document.xml.rels", template_rels)
            archive.writestr("word/media/image8.png", b"png")
            archive.writestr("word/media/image9.svg", SIMPLE_BACKGROUND_SVG)
        with ZipFile(output_path, "w", compression=ZIP_DEFLATED) as archive:
            archive.writestr("word/document.xml", output_document)
            archive.writestr("word/_rels/document.xml.rels", "<Relationships></Relationships>")

        document_builder.restore_svg_assets_from_template(template_path, output_path)

        with ZipFile(output_path) as archive:
            restored_document = archive.read("word/document.xml").decode("utf-8")
        tables = re.findall(r"<w:tbl\b[\s\S]*?</w:tbl>", restored_document)
        self.assertEqual(len(tables), 2)
        self.assertIn('behindDoc="1"', tables[0])
        self.assertNotIn('behindDoc="1"', tables[1])

    def test_restore_svg_assets_readds_signature_foreground_drawings_from_template(self) -> None:
        template_path = self.tmp_dir / "signature-foreground-template.docx"
        output_path = self.tmp_dir / "signature-foreground-output.docx"
        stamp_run = (
            '<w:r><w:drawing><wp:anchor behindDoc="0"><wp:extent cx="1638300" cy="1857375"/>'
            '<a:blip r:embed="rId7"/></wp:anchor></w:drawing></w:r>'
        )
        icon_run = (
            '<w:r><w:drawing><wp:anchor behindDoc="0"><wp:extent cx="114935" cy="78105"/>'
            '<a:blip r:embed="rId8"><a:extLst>'
            '<a:ext uri="{96DAC541-7B7A-43D3-8B79-37D633B846F1}">'
            '<asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" '
            'r:embed="rId9"/></a:ext></a:extLst></a:blip>'
            '</wp:anchor></w:drawing></w:r>'
        )
        signature_table_template = (
            '<w:tbl>'
            f'<w:tr><w:tc><w:p><w:r><w:t>С уважением, исполнительный директор</w:t></w:r></w:p></w:tc><w:tc><w:p>{stamp_run}</w:p></w:tc><w:tc><w:p><w:r><w:t>К.И. Крашенинников</w:t></w:r></w:p></w:tc></w:tr>'
            f'<w:tr><w:tc><w:p>{icon_run}<w:r><w:t>Исп. Черкашина Наталья</w:t></w:r></w:p></w:tc><w:tc><w:p/></w:tc><w:tc><w:p/></w:tc></w:tr>'
            '</w:tbl>'
        )
        output_document = (
            signature_table_template
            .replace(stamp_run, "")
            .replace(icon_run, "")
        )
        rels = (
            '<Relationships>'
            '<Relationship Id="rId7" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/stamp.jpeg"/>'
            '<Relationship Id="rId8" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/phone.png"/>'
            '<Relationship Id="rId9" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/phone.svg"/>'
            '</Relationships>'
        )
        with ZipFile(template_path, "w", compression=ZIP_DEFLATED) as archive:
            archive.writestr("word/document.xml", f"<w:document><w:body>{signature_table_template}</w:body></w:document>")
            archive.writestr("word/_rels/document.xml.rels", rels)
            archive.writestr("word/media/stamp.jpeg", b"stamp")
            archive.writestr("word/media/phone.png", b"phone-png")
            archive.writestr("word/media/phone.svg", SIMPLE_BACKGROUND_SVG)
        with ZipFile(output_path, "w", compression=ZIP_DEFLATED) as archive:
            archive.writestr("word/document.xml", f"<w:document><w:body>{output_document}</w:body></w:document>")
            archive.writestr("word/_rels/document.xml.rels", "<Relationships></Relationships>")

        document_builder.restore_svg_assets_from_template(template_path, output_path)

        with ZipFile(output_path) as archive:
            restored_document = archive.read("word/document.xml").decode("utf-8")
            restored_rels = archive.read("word/_rels/document.xml.rels").decode("utf-8")
            names = set(archive.namelist())
        self.assertEqual(restored_document.count('behindDoc="0"'), 2)
        self.assertIn('r:embed="rId7"', restored_document)
        self.assertIn('asvg:svgBlip', restored_document)
        self.assertIn('Id="rId7"', restored_rels)
        self.assertIn('Id="rId8"', restored_rels)
        self.assertIn('Id="rId9"', restored_rels)
        self.assertIn("word/media/stamp.jpeg", names)
        self.assertIn("word/media/phone.svg", names)
    def test_restore_svg_assets_resets_existing_signature_background_to_template_position(self) -> None:
        template_path = self.tmp_dir / "existing-background-template.docx"
        output_path = self.tmp_dir / "existing-background-output.docx"
        background_run = (
            '<w:r><w:drawing><wp:anchor behindDoc="1"><wp:positionH relativeFrom="column"><wp:posOffset>828675</wp:posOffset></wp:positionH><wp:positionV relativeFrom="paragraph"><wp:posOffset>-889000</wp:posOffset></wp:positionV><wp:extent cx="4294505" cy="2900680"/>'
            '<a:blip r:embed="rId14"><a:extLst>'
            '<a:ext uri="{96DAC541-7B7A-43D3-8B79-37D633B846F1}">'
            '<asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" '
            'r:embed="rId15"/></a:ext></a:extLst></a:blip>'
            '</wp:anchor></w:drawing></w:r>'
        )
        document_xml = f"<w:document><w:body><w:p>{background_run}</w:p></w:body></w:document>"
        output_document_xml = document_xml.replace("828675", "1450000").replace("-889000", "-350000")
        rels_xml = (
            '<Relationships>'
            '<Relationship Id="rId14" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image8.png"/>'
            '<Relationship Id="rId15" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image9.svg"/>'
            '</Relationships>'
        )
        with ZipFile(template_path, "w", compression=ZIP_DEFLATED) as archive:
            archive.writestr("word/document.xml", document_xml)
            archive.writestr("word/_rels/document.xml.rels", rels_xml)
            archive.writestr("word/media/image8.png", b"png")
            archive.writestr("word/media/image9.svg", SIMPLE_BACKGROUND_SVG)
        with ZipFile(output_path, "w", compression=ZIP_DEFLATED) as archive:
            archive.writestr("word/document.xml", output_document_xml)
            archive.writestr("word/_rels/document.xml.rels", rels_xml)
            archive.writestr("word/media/image8.png", b"png")
            archive.writestr("word/media/image9.svg", SIMPLE_BACKGROUND_SVG)

        document_builder.restore_svg_assets_from_template(template_path, output_path)

        with ZipFile(output_path) as archive:
            restored_document = archive.read("word/document.xml").decode("utf-8")
        self.assertIn('<wp:posOffset>828675</wp:posOffset>', restored_document)
        self.assertIn('<wp:posOffset>-889000</wp:posOffset>', restored_document)
        self.assertNotIn('<wp:posOffset>1450000</wp:posOffset>', restored_document)
        self.assertNotIn('<wp:posOffset>-350000</wp:posOffset>', restored_document)

    def test_existing_background_png_payload_is_rendered_from_svg_for_pdf(self) -> None:
        payloads = {
            "word/media/image8.png": b"old-png",
            "word/media/image9.svg": SIMPLE_BACKGROUND_SVG,
        }

        document_builder.enhance_background_png_payloads(
            part_name="word/document.xml",
            output_text=(
                '<w:document><wp:anchor behindDoc="1"><wp:extent cx="914400" cy="457200"/>'
                '<a:blip r:embed="rId14"><a:extLst><a:ext uri="{96DAC541-7B7A-43D3-8B79-37D633B846F1}">'
                '<asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" r:embed="rId15"/>'
                '</a:ext></a:extLst></a:blip></wp:anchor></w:document>'
            ),
            output_rels_text=(
                '<Relationships>'
                '<Relationship Id="rId14" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image8.png"/>'
                '<Relationship Id="rId15" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image9.svg"/>'
                '</Relationships>'
            ),
            payloads=payloads,
        )

        self.assertEqual(png_size(payloads["word/media/image8.png"]), (300, 150))

    def test_pdf_background_png_payload_renders_svg_without_pillow(self) -> None:
        result = document_builder.build_pdf_background_png_payload(
            b"old-png",
            svg_payload=SIMPLE_BACKGROUND_SVG,
            width_px=40,
            height_px=20,
        )

        self.assertEqual(png_size(result), (40, 20))

    def test_pdf_background_png_payload_preserves_compound_holes_and_css_fill(self) -> None:
        result = document_builder.build_pdf_background_png_payload(
            b"old-png",
            svg_payload=COMPOUND_BACKGROUND_SVG,
            width_px=20,
            height_px=20,
        )

        self.assertEqual(png_size(result), (20, 20))
        self.assertEqual(png_rgba_pixel(result, 2, 2), (234, 234, 234, 255))
        self.assertEqual(png_rgba_pixel(result, 10, 10)[3], 0)

    def test_kp_recipient_address_line_is_capitalized(self) -> None:
        context = build_document_context(self._row(), 101)
        context.update({"OUTGOING_NUMBER": "101", "DATE": "22.06.2026"})
        context["ADM_NAME_1"] = 'администрации муниципального образования "Энемское городское поселение"'

        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[1].text = "ADM_NAME_1"

        document_builder.normalize_kp_formatting(doc, context)
        replacements = dict(document_builder.build_kp_replacements(context))

        expected = 'Администрации муниципального образования "Энемское городское поселение"'
        self.assertEqual(table.rows[0].cells[1].text, expected)
        self.assertEqual(replacements["ADM_NAME_1"], expected)

    def test_contract_replacements_keep_inflected_admin_case_as_is(self) -> None:
        context = build_document_context(self._row(), 101)
        context.update({"CONTRACT_NUMBER": "101", "DATE": "22.06.2026"})
        context["ADM_NAME_1"] = 'администрации муниципального образования "Энемское городское поселение"'

        replacements = dict(document_builder.build_contract_replacements(context))

        self.assertEqual(
            replacements["ADM_NAME_1"],
            'администрации муниципального образования "Энемское городское поселение"',
        )

    def test_kp_pdf_layout_removes_trailing_blank_body_paragraph(self) -> None:
        doc = Document()
        doc.add_paragraph("\u0422\u0435\u043a\u0441\u0442 \u041a\u041f")
        doc.add_paragraph("")
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = "\u0421 \u0443\u0432\u0430\u0436\u0435\u043d\u0438\u0435\u043c,\n\u0438\u0441\u043f\u043e\u043b\u043d\u0438\u0442\u0435\u043b\u044c\u043d\u044b\u0439 \u0434\u0438\u0440\u0435\u043a\u0442\u043e\u0440"
        table.rows[0].cells[2].text = "\u041a.\u0418. \u041a\u0440\u0430\u0448\u0435\u043d\u0438\u043d\u043d\u0438\u043a\u043e\u0432"
        blank_signature_paragraph = table.rows[1].cells[1].add_paragraph("")
        doc.add_paragraph("").add_run(" ")

        document_builder.stabilize_kp_pdf_layout(doc, {})

        body_children = list(doc._body._element)
        last_content = next(child for child in reversed(body_children) if document_builder._local_name(child) != "sectPr")
        previous = table._tbl.getprevious()
        previous_previous = previous.getprevious()
        previous_text = "".join(node.text or "" for node in previous_previous.iter())
        self.assertEqual(document_builder._local_name(last_content), "tbl")
        self.assertEqual(document_builder._local_name(previous), "p")
        self.assertIn('w:line="60"', previous.xml)
        self.assertIn("\u0422\u0435\u043a\u0441\u0442 \u041a\u041f", previous_text)
        self.assertIn("w:cantSplit", table._tbl.xml)
        self.assertIn("w:keepNext", table._tbl.xml)
        self.assertEqual(blank_signature_paragraph.text, " ")
        self.assertEqual(blank_signature_paragraph.paragraph_format.line_spacing_rule, WD_LINE_SPACING.EXACTLY)

    def test_territorial_zone_layout_preserves_body_spacing_and_signature_font(self) -> None:
        doc = Document()
        paragraph = doc.add_paragraph("\u041e\u043f\u0438\u0441\u0430\u043d\u0438\u0435 \u043c\u0435\u0441\u0442\u043e\u043f\u043e\u043b\u043e\u0436\u0435\u043d\u0438\u044f \u0433\u0440\u0430\u043d\u0438\u0446 \u0442\u0435\u0440\u0440\u0438\u0442\u043e\u0440\u0438\u0430\u043b\u044c\u043d\u044b\u0445 \u0437\u043e\u043d \u0437\u0430\u043d\u0438\u043c\u0430\u0435\u0442 \u043c\u043d\u043e\u0433\u043e \u0441\u0442\u0440\u043e\u043a.")
        paragraph.paragraph_format.space_after = Pt(12)
        body_run = paragraph.runs[0]
        table = doc.add_table(rows=1, cols=3)
        signature_run = table.rows[0].cells[0].paragraphs[0].add_run("\u0421 \u0443\u0432\u0430\u0436\u0435\u043d\u0438\u0435\u043c,")
        table.rows[0].cells[2].text = "\u041a.\u0418. \u041a\u0440\u0430\u0448\u0435\u043d\u0438\u043d\u043d\u0438\u043a\u043e\u0432"

        document_builder.stabilize_kp_pdf_layout(
            doc,
            {"WORK_TYPE": document_builder.WORK_TYPE_TERRITORIAL_ZONE_BOUNDARIES},
        )

        self.assertEqual(paragraph.paragraph_format.space_after, Pt(12))
        self.assertIsNone(body_run.font.size)
        self.assertIsNone(signature_run.font.size)

    def test_signature_contact_spacing_uses_pdf_stable_gap(self) -> None:
        output_text = (
            '<w:document><w:body><w:tbl>'
            '<w:tr><w:tc><w:p><w:r><w:t>\u0421 \u0443\u0432\u0430\u0436\u0435\u043d\u0438\u0435\u043c,</w:t></w:r></w:p></w:tc>'
            '<w:tc><w:p/></w:tc>'
            '<w:tc><w:p><w:r><w:t>\u041a.\u0418. \u041a\u0440\u0430\u0448\u0435\u043d\u0438\u043d\u043d\u0438\u043a\u043e\u0432</w:t></w:r></w:p></w:tc></w:tr>'
            '<w:tr><w:tc><w:tbl><w:tr><w:tc><w:p>'
            '<w:r><w:t>\u0418\u0441\u043f. \u0427\u0435\u0440\u043a\u0430\u0448\u0438\u043d\u0430 \u041d\u0430\u0442\u0430\u043b\u044c\u044f</w:t></w:r>'
            '<w:r><w:br/></w:r><w:r><w:t xml:space="preserve">       </w:t></w:r><w:r><w:t>\u0442\u0435\u043b. +7 963 912-74-25</w:t></w:r>'
            '<w:r><w:br/></w:r><w:r><w:t xml:space="preserve">       </w:t></w:r><w:r><w:t>ks@parresh.ru</w:t></w:r>'
            '</w:p></w:tc></w:tr></w:tbl></w:tc><w:tc/><w:tc/></w:tr>'
            '</w:tbl></w:body></w:document>'
        )

        adjusted, changed = document_builder.normalize_signature_contact_spacing_for_pdf(output_text)

        self.assertTrue(changed)
        self.assertNotIn('>       </w:t>', adjusted)
        self.assertEqual(adjusted.count('>   </w:t>'), 2)

    def test_signature_contact_icons_use_pdf_stable_vertical_offsets(self) -> None:
        phone_anchor = (
            '<wp:anchor behindDoc="0"><wp:positionV relativeFrom="paragraph"><wp:posOffset>250190</wp:posOffset></wp:positionV>'
            '<wp:extent cx="113665" cy="121920"/></wp:anchor>'
        )
        mail_anchor = (
            '<wp:anchor behindDoc="0"><wp:positionV relativeFrom="paragraph"><wp:posOffset>431165</wp:posOffset></wp:positionV>'
            '<wp:extent cx="114935" cy="78105"/></wp:anchor>'
        )
        output_text = (
            '<w:document><w:body><w:tbl>'
            '<w:tr><w:tc><w:p><w:r><w:t>\u0421 \u0443\u0432\u0430\u0436\u0435\u043d\u0438\u0435\u043c,</w:t></w:r></w:p></w:tc>'
            '<w:tc><w:p/></w:tc>'
            '<w:tc><w:p><w:r><w:t>\u041a.\u0418. \u041a\u0440\u0430\u0448\u0435\u043d\u0438\u043d\u043d\u0438\u043a\u043e\u0432</w:t></w:r></w:p></w:tc></w:tr>'
            '<w:tr><w:tc><w:p>'
            f'<w:r><w:drawing>{phone_anchor}</w:drawing></w:r>'
            f'<w:r><w:drawing>{mail_anchor}</w:drawing></w:r>'
            '<w:r><w:t>\u0418\u0441\u043f. \u0427\u0435\u0440\u043a\u0430\u0448\u0438\u043d\u0430 \u041d\u0430\u0442\u0430\u043b\u044c\u044f</w:t></w:r>'
            '<w:r><w:br/></w:r><w:r><w:t>\u0442\u0435\u043b. +7 963 912-74-25</w:t></w:r>'
            '<w:r><w:br/></w:r><w:r><w:t>ks@parresh.ru</w:t></w:r>'
            '</w:p></w:tc><w:tc/><w:tc/></w:tr>'
            '</w:tbl></w:body></w:document>'
        )

        adjusted, changed = document_builder.adjust_signature_contact_icon_positions_for_pdf(output_text)
        adjusted_again, changed_again = document_builder.adjust_signature_contact_icon_positions_for_pdf(adjusted)

        self.assertTrue(changed)
        self.assertFalse(changed_again)
        self.assertEqual(adjusted_again, adjusted)
        self.assertIn('<wp:posOffset>150190</wp:posOffset>', adjusted)
        self.assertIn('<wp:posOffset>351165</wp:posOffset>', adjusted)

    def test_territorial_zone_signature_stamp_is_shifted_below_date_line(self) -> None:
        stamp_anchor = (
            '<wp:anchor behindDoc="0"><wp:positionV relativeFrom="paragraph"><wp:posOffset>-558800</wp:posOffset></wp:positionV>'
            '<wp:extent cx="1638300" cy="1857375"/></wp:anchor>'
        )
        output_text = (
            '<w:document><w:body>'
            '<w:p><w:r><w:t>\u043e\u043f\u0438\u0441\u0430\u043d\u0438\u044f \u043c\u0435\u0441\u0442\u043e\u043f\u043e\u043b\u043e\u0436\u0435\u043d\u0438\u044f \u0433\u0440\u0430\u043d\u0438\u0446 \u0442\u0435\u0440\u0440\u0438\u0442\u043e\u0440\u0438\u0430\u043b\u044c\u043d\u044b\u0445 \u0437\u043e\u043d</w:t></w:r></w:p>'
            '<w:tbl><w:tr><w:tc><w:p><w:r><w:t>\u0421 \u0443\u0432\u0430\u0436\u0435\u043d\u0438\u0435\u043c,</w:t></w:r></w:p></w:tc>'
            f'<w:tc><w:p><w:r><w:drawing>{stamp_anchor}</w:drawing></w:r></w:p></w:tc>'
            '<w:tc><w:p><w:r><w:t>\u041a.\u0418. \u041a\u0440\u0430\u0448\u0435\u043d\u0438\u043d\u043d\u0438\u043a\u043e\u0432</w:t></w:r></w:p></w:tc></w:tr></w:tbl>'
            '</w:body></w:document>'
        )

        adjusted, changed = document_builder.adjust_territorial_signature_stamp_position(output_text)

        self.assertTrue(changed)
        self.assertIn('<wp:posOffset>-350000</wp:posOffset>', adjusted)

if __name__ == "__main__":
    unittest.main()

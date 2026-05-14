import shutil
import unittest
from pathlib import Path

from docx import Document

from src.generator.document_review_agent import review_docx
from src.generator.philologist_agent import _auto_fix_docx
from src.generator.philologist_tools import PhilologistToolRunner


class PhilologistRegressionTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp_dir = Path("tmp_test_philologist_regressions")
        self.tmp_dir.mkdir(exist_ok=True)

    def tearDown(self) -> None:
        if self.tmp_dir.exists():
            shutil.rmtree(self.tmp_dir)

    def _docx(self, name: str, *paragraphs: str) -> Path:
        path = self.tmp_dir / name
        doc = Document()
        for text in paragraphs:
            doc.add_paragraph(text)
        doc.save(path)
        return path

    def _text(self, path: Path) -> str:
        doc = Document(path)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    def test_contract_body_legal_terms_are_lowercased_without_removing_links(self) -> None:
        path = self._docx(
            "legal_terms.docx",
            (
                "1.2. Выполнение Работ осуществляется по месту нахождения Исполнителя "
                "на условиях и в сроки, установленные настоящим Договором, "
                "Техническим заданием (Приложение № 1 к Договору), "
                "Календарным планом выполнения работ (Приложение № 2 к Договору), "
                "которые являются неотъемлемой частью настоящего Договора."
            ),
        )

        review = review_docx(path, ai_enabled=False)
        result = _auto_fix_docx(path, review, client=None, tool_runner=PhilologistToolRunner())
        text = self._text(path)

        self.assertGreaterEqual(result["applied_fix_count"], 8)
        self.assertIn("Выполнение работ", text)
        self.assertIn("нахождения исполнителя", text)
        self.assertIn("настоящим договором", text)
        self.assertIn("техническим заданием (приложение № 1 к договору)", text)
        self.assertIn("календарным планом выполнения работ (приложение № 2 к договору)", text)
        self.assertIn("настоящего договора", text)
        self.assertIn("приложение № 1", text)
        self.assertIn("приложение № 2", text)

    def test_uppercase_commercial_offer_heading_is_not_lowercased(self) -> None:
        path = self._docx("heading.docx", "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ")

        review = review_docx(path, ai_enabled=False)
        result = _auto_fix_docx(path, review, client=None, tool_runner=PhilologistToolRunner())

        self.assertEqual(review["issue_count"], 0)
        self.assertEqual(result["applied_fix_count"], 0)
        self.assertEqual(self._text(path), "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ")

    def test_attachment_heading_keeps_capitalized_label(self) -> None:
        path = self._docx("attachment_heading.docx", "Приложение № 1 к Договору")

        review = review_docx(path, ai_enabled=False)
        result = _auto_fix_docx(path, review, client=None, tool_runner=PhilologistToolRunner())

        self.assertEqual(result["applied_fix_count"], 1)
        self.assertEqual(self._text(path), "Приложение № 1 к договору")


if __name__ == "__main__":
    unittest.main()
